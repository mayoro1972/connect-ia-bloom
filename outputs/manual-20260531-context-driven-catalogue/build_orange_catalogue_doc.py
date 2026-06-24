from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260531-context-driven-catalogue/output"
)
OUT_PATH = OUT_DIR / "Orange_CI_Mini_Catalogue_Automatise_V3_CoSigne_2026-05-31.docx"
LOGO_PATH = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-prospection/assets/transferai-africa-nettelecom-co-brand.png"
)
FOOTER_LINE = (
    "TransferAI | Hub IA de Nettelecom CI | contact@transferai.ci | "
    "www.transferai.ci | +225 07 16 57 39 90 | Riviera 3, carrefour Sainte Famille, Abidjan"
)

ACCENT = RGBColor(0xE7, 0x6F, 0x1D)
INK = RGBColor(0x1A, 0x35, 0x54)
SOFT = RGBColor(0x61, 0x70, 0x86)
TEAL = RGBColor(0x1C, 0x8A, 0x78)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].font.color.rgb = INK
styles["Normal"].paragraph_format.space_after = Pt(5)
styles["Normal"].paragraph_format.line_spacing = 1.15

for name, size, color in [
    ("Heading 1", 16, INK),
    ("Heading 2", 12.5, ACCENT),
    ("Heading 3", 11.5, TEAL),
]:
    style = styles[name]
    style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color

if "Eyebrow" not in styles:
    eyebrow = styles.add_style("Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
    eyebrow.font.name = "Aptos"
    eyebrow.font.size = Pt(8.5)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = ACCENT

hdr = section.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
if LOGO_PATH.exists():
    hdr_run = hdr.add_run()
    hdr_run.add_picture(str(LOGO_PATH), width=Cm(4.7))
hdr.paragraph_format.space_after = Pt(0)

hdr_note = section.header.add_paragraph()
hdr_note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = hdr_note.add_run("Mini-catalogue automatisé | Formation, accompagnement et pilote IA")
run.font.name = "Aptos"
run.font.size = Pt(8)
run.font.color.rgb = SOFT

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_before = Pt(6)
footer_run = footer.add_run(FOOTER_LINE)
footer_run.font.name = "Aptos"
footer_run.font.size = Pt(8)
footer_run.font.color.rgb = SOFT

p = doc.add_paragraph(style="Eyebrow")
p.add_run("PROPOSITION CIBLÉE")

p = doc.add_paragraph()
r = p.add_run("Formation et accompagnement IA pour Orange Côte d'Ivoire")
r.font.name = "Aptos Display"
r.font.size = Pt(23)
r.font.bold = True
r.font.color.rgb = INK

p = doc.add_paragraph()
r = p.add_run(
    "Document modèle automatisé construit à partir de la logique des nœuds Sanitize Prospect Data For LLM et Assemble Prospect Context, aligné avec le courrier exécutif, le deck et le mini-catalogue co-signé."
)
r.font.name = "Aptos"
r.font.size = Pt(11.5)
r.font.color.rgb = SOFT

p = doc.add_paragraph()
shade_paragraph(p, "FBF5EE")
r = p.add_run(
    "Résumé. La proposition recommandée ne consiste pas à dérouler un catalogue large. Elle consiste à choisir un premier terrain crédible, former les bonnes équipes, encadrer la gouvernance et rester engagé sur 90 jours pour transformer la formation en usage réel."
)
r.font.name = "Aptos"
r.font.size = Pt(10.5)
r.font.color.rgb = INK
r.bold = True

doc.add_heading("1. Synthèse executive", level=1)
doc.add_paragraph(
    "Cette proposition vise à aider Orange Côte d'Ivoire à transformer des flux à fort volume en gains visibles pour la direction et les équipes opérationnelles. La logique recommandée n'est pas un catalogue large, mais un enchaînement ciblé entre audit, formation utile, cas d'usage priorisés, accompagnement 90 jours et décision d'extension sur la base de KPI concrets."
)

doc.add_heading("2. Lecture du contexte du prospect", level=1)
doc.add_paragraph(
    "Orange Côte d'Ivoire opère des flux où la qualité de réponse, la coordination entre canaux, la diffusion des procédures et la vitesse de pilotage ont un impact direct sur l'expérience client et l'efficacité opérationnelle. Les signaux disponibles suggèrent une priorité probable autour de la relation client, du reporting managérial, de la standardisation des scripts et de la mise à disposition rapide d'une base de connaissances exploitable par les équipes."
)

doc.add_heading("3. Priorités métier et irritants visibles", level=1)
for item in [
    "Mieux traiter les demandes récurrentes sur plusieurs canaux.",
    "Réduire le temps de préparation des synthèses de pilotage.",
    "Rendre les procédures et scripts plus accessibles aux équipes front et support.",
    "Améliorer la cohérence des réponses et la continuité de traitement.",
    "Disposer d'une lecture plus rapide des écarts de qualité de service.",
]:
    add_bullet(doc, item)

doc.add_heading("4. Cas d'usage et quick wins recommandés", level=1)
cases = [
    ("Copilote service client multicanal", "Demandes récurrentes réparties entre plusieurs canaux avec réponses hétérogènes.", "Copilote de réponse, résumés d'interactions et meilleure continuité de traitement.", "Délai moyen de réponse, taux de résolution au premier contact."),
    ("Synthèse tickets et incidents", "Analyse manuelle des tickets et reporting lent pour les managers.", "Synthèses automatiques, regroupement des motifs récurrents et alertes plus lisibles.", "Temps de préparation du reporting, délai d'escalade."),
    ("Base procédures et scripts terrain", "Réponses et procédures dispersées selon les équipes et les documents.", "Base de connaissances assistée pour scripts, contrôles et bonnes réponses.", "Temps de recherche d'information, temps d'onboarding."),
    ("Commentaire IA du pilotage qualité", "Les indicateurs existent mais sont peu commentés et difficiles à exploiter rapidement.", "Commentaires automatiques des écarts, synthèses managers et meilleurs arbitrages.", "Temps de lecture des KPI, réactivité de pilotage."),
]
for title, before, after, kpi in cases:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Avant : {before}")
    doc.add_paragraph(f"Après : {after}")
    doc.add_paragraph(f"KPI de preuve : {kpi}")

doc.add_heading("5. Porte d'entrée recommandée", level=1)
doc.add_paragraph(
    "La porte d'entrée recommandée est un audit IA gratuit suivi d'un échange expert de 30 minutes. Cette étape doit permettre de confirmer les priorités, préciser les populations à former, valider le premier terrain d'exécution et calibrer les KPI de preuve."
)

doc.add_heading("6. Parcours recommandé", level=1)
for item in [
    "J+0 : audit et cadrage des flux service client et pilotage.",
    "J+15 : premier livrable exploitable, par exemple un copilote ciblé ou une première synthèse automatisée.",
    "J+45 : formation ciblée, gouvernance d'usage et lancement du pilote.",
    "J+90 : mesure des résultats, lecture de l'adoption et décision d'extension.",
]:
    add_bullet(doc, item)

doc.add_heading("7. Formations prioritaires", level=1)
trainings = [
    ("IA appliquée au service client multicanal", "Responsables relation client, superviseurs, équipes front et support.", "Améliorer la cohérence des réponses, accélérer le traitement des demandes répétitives, structurer les usages de résumé et de préparation de réponse."),
    ("IA pour reporting, tickets et synthèses managériales", "Managers de service, responsables qualité, équipes pilotage.", "Accélérer la préparation des synthèses, rendre les signaux de friction plus lisibles, améliorer la vitesse d'arbitrage."),
    ("Gouvernance des procédures, scripts et base de connaissances", "Équipes support, qualité, transformation, référents métier.", "Structurer la diffusion des bonnes réponses, réduire le temps de recherche d'information, sécuriser la mise à jour des procédures."),
    ("Gouvernance IA, confidentialité et supervision humaine", "Direction, managers, fonctions transformation, qualité et conformité.", "Encadrer les usages IA autorisés, clarifier la supervision humaine, sécuriser les premiers pilotes."),
]
for title, audience, goals in trainings:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f"Public concerné : {audience}")
    doc.add_paragraph(f"Objectifs : {goals}")
    doc.add_paragraph("Format / durée : à confirmer selon le périmètre.")

doc.add_heading("8. Livrables attendus", level=1)
for item in [
    "Note de priorités et séquence recommandée.",
    "Supports de formation contextualisés.",
    "Prompts, trames et checklists réutilisables.",
    "Premier cas d'usage encadré ou premier livrable pilote.",
    "Feuille de route d'exécution sur 90 jours.",
    "Points de revue pour mesurer adoption, qualité et impact.",
]:
    add_bullet(doc, item)

doc.add_heading("9. Gouvernance, confidentialité et conduite du changement", level=1)
doc.add_paragraph(
    "Le dispositif recommandé doit intégrer un cadre simple mais explicite : supervision humaine avant généralisation, périmètre pilote défini, accès encadré aux contenus et sorties, validation progressive des usages autorisés et conduite du changement centrée sur les gestes métier."
)

doc.add_heading("10. Accompagnement post-formation", level=1)
doc.add_paragraph(
    "L'accompagnement ne s'arrête pas à la fin des sessions. Les personnes et équipes formées gardent un accès gratuit à la communauté TransferAI afin de partager leurs cas d'entreprise, les blocages rencontrés et leurs questions d'exécution. Cet accompagnement permet de transformer la formation en usage réel, puis en preuve d'impact."
)

doc.add_heading("11. Accompagnement 90 jours", level=1)
doc.add_paragraph(
    "La vraie différence TransferAI par rapport à une structure de formation classique se joue ici. Une structure de formation classique transmet un contenu puis clôture son intervention. TransferAI forme, puis reste engagé sur 90 jours pour aider l'organisation à consolider les usages réellement adoptés, traiter les blocages terrain, ajuster les prompts, scripts ou trames, relire les premiers KPI et arbitrer la suite entre pilote, extension ou réajustement."
)
for item in [
    "Points de suivi.",
    "Revues de cas réels.",
    "Assistance questions/réponses.",
    "Conseils pratiques de mise en œuvre.",
    "Revue de gouvernance et de supervision.",
    "Lecture d'impact orientée direction.",
]:
    add_bullet(doc, item)

doc.add_heading("12. Proposition immédiate", level=1)
doc.add_paragraph(
    "Nous recommandons de démarrer par un audit IA gratuit suivi d'un échange expert de 30 minutes, puis de confirmer un premier terrain de valeur autour du service client, du reporting ou de la base de connaissances. L'objectif n'est pas de tout lancer d'un bloc, mais de choisir un premier pilote crédible, former les bonnes équipes et sécuriser un accompagnement sur 90 jours pour transformer la formation en usage concret."
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(14)
r = p.add_run("Signataires du mini-catalogue")
r.font.name = "Aptos"
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = SOFT

for item in [
    "Soulemane Konate — Directeur IA & Innovation",
    "Medard Sery — Consultant expert · Data engineering, cloud & plateformes IA",
]:
    add_bullet(doc, item)

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
