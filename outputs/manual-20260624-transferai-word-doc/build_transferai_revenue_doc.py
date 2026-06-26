from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260624-transferai-word-doc"
)
DOCX_PATH = OUTPUT_DIR / "TransferAI_Comment_IA_Regenerer_Argent_Par_Secteur_2026-06-24.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
ORANGE = RGBColor(0xE7, 0x6F, 0x1D)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(dxa))


def set_table_width(table, dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(dxa))


def set_table_indent(table, dxa):
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(dxa))


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_borders(table, color="D9E2F3"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_run_font(run, size, color=TEXT, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_styled_paragraph(doc, text="", style=None, before=0, after=6, line=1.25):
    p = doc.add_paragraph(style=style)
    if text:
        p.add_run(text)
    set_paragraph_spacing(p, before=before, after=after, line=line)
    return p


def build_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    if "Card Title" not in styles:
        card_title = styles.add_style("Card Title", WD_STYLE_TYPE.PARAGRAPH)
        card_title.font.name = "Calibri"
        card_title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        card_title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        card_title.font.size = Pt(12)
        card_title.font.bold = True
        card_title.font.color.rgb = DARK_BLUE
        card_title.paragraph_format.space_before = Pt(0)
        card_title.paragraph_format.space_after = Pt(5)
        card_title.paragraph_format.line_spacing = 1.15


def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=2, line=1.0)
    run = p.add_run("DOCUMENT DE TRAVAIL TRANSFERAI")
    set_run_font(run, 10, MUTED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=4, line=1.0)
    run = p.add_run("Comment l’IA régénère de l’argent par secteur")
    set_run_font(run, 24, color=TEXT, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=14, line=1.15)
    run = p.add_run(
        "Version Word prête à partager pour formaliser la slide commerciale, son message central et sa mise en page recommandée."
    )
    set_run_font(run, 12, color=MUTED, bold=False)

    meta_rows = [
        ("Date", "24 juin 2026"),
        ("Usage", "Base de slide / support de vente"),
        ("Angle", "Revenu récupéré, conversion, marge, capacité"),
    ]
    for label, value in meta_rows:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.15)
        run = p.add_run(f"{label} : ")
        set_run_font(run, 11, TEXT, bold=True)
        run = p.add_run(value)
        set_run_font(run, 11, TEXT, bold=False)

    doc.add_paragraph()


def add_message_callout(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=10, line=1.2)
    run = p.add_run(
        "Idée-force : chez TransferAI, l’IA ne crée pas de valeur en théorie. "
        "Elle récupère du revenu perdu, accélère la conversion, protège la marge et augmente la capacité d’exécution."
    )
    set_run_font(run, 12, color=ORANGE, bold=True)


def add_slide_table(doc):
    doc.add_paragraph("Message central de la slide", style="Heading 1")

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=8, line=1.25)
    run = p.add_run(
        "L’IA crée de la valeur quand elle améliore un KPI métier déjà connu : churn, conversion, délai, volume traité, coût par dossier ou marge protégée."
    )
    set_run_font(run, 11, TEXT, bold=False)

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.allow_autofit = False
    set_table_layout_fixed(table)
    set_table_width(table, 9360)
    set_table_indent(table, 120)
    set_table_borders(table, color="D9E2F3")

    cards = [
        ("1. Revenu récupéré", "Churn réduit\nRéclamations mieux traitées\nSignaux clients mieux détectés", "FCE4D6"),
        ("2. Conversion accélérée", "Segmentation plus utile\nRelances plus rapides\nMessages plus performants", "DDEBF7"),
        ("3. Marge protégée", "Moins d’erreurs\nMoins de retards\nMoins de rework", "E2F0D9"),
        ("4. Capacité augmentée", "Plus de volume traité\nMême équipe\nDécisions plus rapides", "FFF2CC"),
    ]

    for idx, (title, lines, fill) in enumerate(cards):
        cell = table.rows[idx // 2].cells[idx % 2]
        set_cell_width(cell, 4680)
        set_cell_shading(cell, fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = doc.styles["Card Title"]
        set_paragraph_spacing(p, before=0, after=4, line=1.15)
        run = p.add_run(title)
        set_run_font(run, 12, DARK_BLUE, bold=True)
        for line in lines.split("\n"):
            p = cell.add_paragraph()
            p.style = "Normal"
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            set_paragraph_spacing(p, before=0, after=3, line=1.18)
            run = p.add_run("• ")
            set_run_font(run, 11, TEXT, bold=True)
            run = p.add_run(line)
            set_run_font(run, 11, TEXT, bold=False)

    doc.add_paragraph()


def add_copy_ready_content(doc):
    doc.add_paragraph("Texte prêt à copier dans la slide", style="Heading 1")

    sections = [
        (
            "Titre",
            "Comment l’IA régénère de l’argent par secteur",
        ),
        (
            "Sous-titre",
            "Chez TransferAI, l’IA ne crée pas de valeur en théorie : elle récupère du revenu perdu, accélère la conversion, protège la marge et augmente la capacité d’exécution.",
        ),
        (
            "Exemples sectoriels",
            "Télécommunications : churn mieux anticipé, upsell mieux ciblé, service client plus rapide.\n"
            "Banque & Assurance : dossiers traités plus vite, meilleure conversion, meilleure rétention.\n"
            "Retail & FMCG : campagnes plus ciblées, ventes mieux commentées, réseau commercial plus réactif.\n"
            "Industrie & Logistique : moins de pertes cachées, arbitrages plus rapides, meilleure continuité opérationnelle.",
        ),
        (
            "Phrase de conclusion",
            "L’IA régénère de l’argent quand elle agit sur un KPI métier clair : conversion, churn, délai, volume traité, coût par dossier ou marge protégée.",
        ),
    ]

    for heading, content in sections:
        doc.add_paragraph(heading, style="Heading 2")
        if "\n" not in content:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=0, after=6, line=1.25)
            run = p.add_run(content)
            set_run_font(run, 11, TEXT, bold=False)
            continue
        for line in content.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.14)
            set_paragraph_spacing(p, before=0, after=3, line=1.22)
            run = p.add_run("• ")
            set_run_font(run, 11, TEXT, bold=True)
            run = p.add_run(line)
            set_run_font(run, 11, TEXT, bold=False)


def add_layout_guidance(doc):
    doc.add_paragraph("Mise en page recommandée", style="Heading 1")

    guidance = [
        ("Structure", "Titre en haut à gauche, sous-titre sur deux lignes, bloc message central au milieu, puis quatre cartes autour ou sous le centre."),
        ("Fond", "Blanc cassé ou gris très clair pour garder une lecture premium et simple."),
        ("Couleurs", "Orange TransferAI pour le message central, bleu marine pour le titre et les sous-titres, accents teal ou dorés pour les cartes."),
        ("Typographie", "Titre très affirmé, sous-titre plus léger, très peu de texte par carte pour préserver la lisibilité."),
        ("Animation", "Si la slide est présentée à l’oral, faire apparaître les quatre cartes l’une après l’autre autour du message central."),
        ("Formule orale", "Notre promesse n’est pas l’IA pour l’IA. Notre promesse est simple : récupérer le revenu qui fuit, convertir plus vite, protéger la marge et augmenter la capacité des équipes, secteur par secteur."),
    ]

    for label, content in guidance:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=4, line=1.2)
        run = p.add_run(f"{label} : ")
        set_run_font(run, 11, TEXT, bold=True)
        run = p.add_run(content)
        set_run_font(run, 11, TEXT, bold=False)

def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    build_styles(doc)
    add_title_block(doc)
    add_message_callout(doc)
    add_slide_table(doc)
    add_copy_ready_content(doc)
    add_layout_guidance(doc)

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
