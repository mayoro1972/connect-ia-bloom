from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260531-context-driven-deck/documents/resume-explicatif-deck-experts/output"
)
OUT_PATH = OUT_DIR / "Note_Explicative_Deck_Automatise_Experts_TransferAI_2026-05-31.docx"

ACCENT = RGBColor(0xE7, 0x6F, 0x1D)
INK = RGBColor(0x1A, 0x35, 0x54)
SOFT = RGBColor(0x61, 0x70, 0x86)
TEAL = RGBColor(0x1C, 0x8A, 0x78)
LINE = "D7CEC3"
PANEL = "FBF5EE"
NAVY_FILL = "163556"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=INK, size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55 * level)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = INK


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21.59)
section.page_height = Cm(27.94)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Normal"].font.color.rgb = INK
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.15

for name, size, color in [
    ("Heading 1", 16, INK),
    ("Heading 2", 12.5, ACCENT),
    ("Heading 3", 11.5, TEAL),
]:
    style = styles[name]
    style.font.name = "Aptos Display" if name == "Heading 1" else "Aptos"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color

if "SmallCapsLabel" not in styles:
    label_style = styles.add_style("SmallCapsLabel", WD_STYLE_TYPE.PARAGRAPH)
    label_style.font.name = "Aptos"
    label_style.font.size = Pt(8.5)
    label_style.font.bold = True
    label_style.font.color.rgb = ACCENT

hdr = section.header.paragraphs[0]
hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hdr_run = hdr.add_run("TransferAI | Note experts | Deck automatisé prospect")
hdr_run.font.name = "Aptos"
hdr_run.font.size = Pt(8)
hdr_run.font.color.rgb = SOFT

ftr = section.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
ftr_run = ftr.add_run("Document de travail interne pour revue experts")
ftr_run.font.name = "Aptos"
ftr_run.font.size = Pt(8)
ftr_run.font.color.rgb = SOFT

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("NOTE EXPLICATIVE")
r.font.name = "Aptos"
r.font.bold = True
r.font.size = Pt(9)
r.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Deck de présentation automatisé TransferAI")
r.font.name = "Aptos Display"
r.font.bold = True
r.font.size = Pt(24)
r.font.color.rgb = INK

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "Résumé à partager aux experts pour expliquer le rôle du deck, sa logique d’alimentation depuis le workflow n8n et les validations encore nécessaires avant production."
)
r.font.name = "Aptos"
r.font.size = Pt(11.5)
r.font.color.rgb = SOFT

callout = doc.add_table(rows=1, cols=1)
callout.autofit = False
callout.columns[0].width = Cm(16.6)
cell = callout.cell(0, 0)
shade_cell(cell, PANEL)
set_cell_text(
    cell,
    "Résumé exécutif. Le deck n’est plus conçu comme une proposition figée inspirée d’ELTON Oil CI. Il devient un support commercial piloté par les informations scrappées, nettoyées, enrichies par OpenAI puis consolidées dans le nœud “Assemble Prospect Context”. L’objectif est de générer automatiquement un PowerPoint crédible, sectoriel et décisionnel, avec des cas d’usage, des hypothèses de ROI et une feuille de route adaptés à chaque prospect.",
    size=10.5,
)

doc.add_paragraph()
doc.add_heading("1. Finalité du deck", level=1)
doc.add_paragraph(
    "Le deck doit servir de document de pré-vente executive. Il a pour rôle d’aider TransferAI à passer d’un simple discours catalogue à une proposition contextualisée, lisible par une direction générale, un comité métier ou un sponsor de transformation."
)
add_bullet(doc, "Montrer que TransferAI comprend le contexte du prospect avant de parler solution.")
add_bullet(doc, "Présenter des cas d’usage directement reliés aux irritants métier visibles.")
add_bullet(doc, "Positionner le ROI comme hypothèse crédible à confirmer, et non comme promesse arbitraire.")
add_bullet(doc, "Conclure sur une prochaine étape simple : audit IA gratuit suivi d’un échange expert de 30 minutes.")

doc.add_heading("2. Ce que démontre la version actuelle du deck", level=1)
doc.add_paragraph(
    "La version de démonstration produite prend Orange Côte d’Ivoire comme exemple de prospect télécom. Elle prouve que les pages les plus sensibles du deck peuvent déjà être alimentées automatiquement par un paquet de données structuré."
)
add_bullet(doc, "Couverture : nom du prospect, angle sectoriel, synthèse de contexte, CTA et premières cartes ROI.")
add_bullet(doc, "Slides de contexte : quatre cas d’usage recommandés avec logique avant/après.")
add_bullet(doc, "Slide ROI : hypothèses chiffrées et KPI à instrumenter par terrain d’usage.")
add_bullet(doc, "Parcours 90 jours : séquence de cadrage, prototype, formation, pilote et décision d’extension.")
add_bullet(doc, "Slide finale : cadre de confiance, message de gouvernance et appel à l’action.")

doc.add_heading("3. Chaîne d’alimentation depuis le workflow n8n", level=1)
doc.add_paragraph(
    "Le fonctionnement cible repose sur une chaîne de transformation déjà présente dans le workflow V3. Les experts doivent lire ce deck comme la couche de restitution d’un contexte déjà calculé, et non comme un livrable rédigé entièrement à la main."
)

flow = doc.add_table(rows=1, cols=3)
flow.autofit = False
flow.columns[0].width = Cm(4.2)
flow.columns[1].width = Cm(5.3)
flow.columns[2].width = Cm(7.1)
headers = ["Étape workflow", "Rôle", "Sortie utile pour le deck"]
for idx, label in enumerate(headers):
    shade_cell(flow.cell(0, idx), NAVY_FILL)
    set_cell_text(flow.cell(0, idx), label, bold=True, color=WHITE, size=9.5)

rows = [
    ("Build Source URLs", "Sélection des pages publiques à lire", "URL de pages utiles par prospect"),
    ("Normalize Public Signals", "Nettoyage et synthèse des contenus publics", "Texte consolidé et signaux métier"),
    ("Sanitize Prospect Data For LLM", "Pseudonymisation et politique de sécurité", "Payload autorisé pour enrichissement IA"),
    ("Call OpenAI Pre-Audit", "Lecture de contexte", "Résumé organisationnel, besoins probables, angle d’entrée"),
    ("Call OpenAI Problems Solutions", "Lecture métier et commerciale", "Cas d’usage, quick wins, offre recommandée"),
    ("Call OpenAI ROI", "Hypothèses d’impact", "ROI, gains attendus, timeline"),
    ("Assemble Prospect Context", "Consolidation finale", "Contexte complet prêt à alimenter le deck"),
    ("Generate Deck Brief", "Orchestration éditoriale", "Profil de deck, messages clés, structure des slides"),
]
for i, row in enumerate(rows, start=1):
    cells = flow.add_row().cells
    for idx, value in enumerate(row):
        if i % 2 == 1:
            shade_cell(cells[idx], PANEL)
        set_cell_text(cells[idx], value, size=9.5 if idx != 2 else 9.3)

doc.add_paragraph()
doc.add_heading("4. Contrat de données du deck", level=1)
doc.add_paragraph(
    "Le deck peut être considéré comme un moteur de rendu qui attend un objet de contexte standardisé. Dans la démonstration actuelle, cet objet est simulé par un “samplePack”. En production, il doit être rempli directement par la sortie de “Assemble Prospect Context” et de “Generate Deck Brief”."
)
add_bullet(doc, "organization_name, organization_type, sector_guess")
add_bullet(doc, "organization_summary et probable_needs")
add_bullet(doc, "single_primary_cta")
add_bullet(doc, "recommended_case_study avec titre, avant, après, KPI, icône et accent")
add_bullet(doc, "roi_hypothesis avec valeur, libellé et note de prudence")
add_bullet(doc, "delivery_timeline et training_focus")

doc.add_heading("5. Ce que les experts doivent valider", level=1)
doc.add_paragraph(
    "La qualité du deck en production dépend moins du design que de la crédibilité métier des blocs générés. La revue experts doit donc porter sur le fond et sur la gouvernance du contenu."
)
add_bullet(doc, "La justesse des cas d’usage par secteur et par type d’organisation.")
add_bullet(doc, "Le niveau de prudence des hypothèses ROI, pour éviter toute surpromesse commerciale.")
add_bullet(doc, "Le vocabulaire de direction, qui doit rester simple, sérieux et compatible avec des comités exécutifs.")
add_bullet(doc, "Le caractère réaliste de la feuille de route J+0 / J+15 / J+45 / J+90.")
add_bullet(doc, "La cohérence entre le deck et les autres sorties du workflow : lettre executive, formulaire d’audit, catalogue ciblé.")
add_bullet(doc, "La présence d’un cadre de confiance explicite : confidentialité, supervision humaine, périmètre pilote, validation avant extension.")

doc.add_heading("6. Recommandations avant passage en production", level=1)
add_bullet(doc, "Figer un schéma JSON minimal et stable pour l’alimentation du deck.")
add_bullet(doc, "Créer une bibliothèque experte de cas d’usage, de KPI et de formules ROI par secteur.")
add_bullet(doc, "Prévoir une couche de validation humaine légère avant envoi externe du deck.")
add_bullet(doc, "Garder deux profils de narration : commercial entreprise et institutionnel.")
add_bullet(doc, "Mesurer sur les prochains pilotes la qualité réelle des recommandations générées pour améliorer le prompt et le modèle de données.")

doc.add_heading("7. Conclusion", level=1)
doc.add_paragraph(
    "Le deck automatisé constitue une évolution importante du projet TransferAI : il transforme un workflow de collecte et d’enrichissement en un support de vente executive cohérent, sectoriel et actionnable. La démonstration actuelle est déjà suffisamment aboutie pour présenter la logique aux experts. La prochaine étape n’est plus de redessiner le deck, mais de sécuriser le contrat de données et la validation métier qui permettront de le produire à grande échelle."
)

doc.add_paragraph()
lab = doc.add_paragraph(style="SmallCapsLabel")
lab.add_run("Pièce liée")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
r = p.add_run(
    "Aperçu PowerPoint de démonstration : Apercu_Deck_Automatise_Orange_CI_V3_2026-05-31.pptx"
)
r.font.name = "Aptos"
r.font.size = Pt(10.5)
r.font.color.rgb = TEAL
r.bold = True

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
