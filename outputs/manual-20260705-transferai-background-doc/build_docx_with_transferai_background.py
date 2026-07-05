from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from lxml import etree
from PIL import Image, ImageDraw, ImageFilter, ImageFont


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
WORK_DIR = ROOT / "outputs/manual-20260705-transferai-background-doc"
SOURCE_DOCX = WORK_DIR / "source.docx"
SHARED_SOURCE_DOCX = WORK_DIR / "source_shared_no_tarifs.docx"
LOGO_PATH = ROOT / "src/assets/logo-transferai-nettelecom.png"
BACKGROUND_PNG = WORK_DIR / "transferai-background-page.png"
INTERMEDIATE_DOCX = WORK_DIR / "with-inline-header-image.docx"
OUTPUT_DOCX = WORK_DIR / "Programme_Formation_IA_Appliquee_Votre_Metier_Session_29_30_juillet_2026_background_transferai.docx"
LOGO_SCALE = 0.64
LOGO_OPACITY = 0.16
SECONDARY_LOGO_OPACITY = 0.08
LOGO_VERTICAL_OFFSET = 0.1
TEXT_BAND_OPACITY = 0.08
BLUE = (30, 102, 178)
ORANGE = (242, 138, 30)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
V_NS = "urn:schemas-microsoft-com:vml"
O_NS = "urn:schemas-microsoft-com:office:office"

NS = {
    "w": W_NS,
    "wp": WP_NS,
    "pic": PIC_NS,
    "a": A_NS,
    "r": R_NS,
    "v": V_NS,
    "o": O_NS,
}


def w(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def v(tag: str) -> str:
    return f"{{{V_NS}}}{tag}"


def o(tag: str) -> str:
    return f"{{{O_NS}}}{tag}"


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            try:
                return ImageFont.truetype(str(path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def build_shared_source() -> None:
    doc = Document(str(SOURCE_DOCX))

    session_text = (
        "Session : 29-30 juillet 2026 & 10-11 Aout 2026\n"
        "Format : 2 jours en présentiel\n"
        "Lieu : Riviera 3, Abidjan\n"
        "Capacité : 20 participants maximum\n"
        "Matériel : ordinateurs fournis"
    )
    session_paragraph = next(
        para for para in doc.paragraphs if para.text.strip().startswith("Session :")
    )
    session_paragraph.clear()
    session_paragraph.add_run(session_text)

    heading_index = next(
        i for i, para in enumerate(doc.paragraphs) if para.text.strip() == "Tarifs"
    )
    next_heading_index = next(
        i
        for i, para in enumerate(doc.paragraphs[heading_index + 1 :], start=heading_index + 1)
        if para.style.name.startswith("Heading")
    )

    replacement = doc.paragraphs[heading_index + 1]
    replacement.text = "Tarifs communiques sur demande."

    for paragraph in doc.paragraphs[heading_index + 2 : next_heading_index]:
        remove_paragraph(paragraph)

    doc.save(str(SHARED_SOURCE_DOCX))


def build_background_image() -> None:
    page = Image.new("RGBA", (2550, 3300), (255, 255, 255, 0))
    accents = Image.new("RGBA", page.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(accents)
    logo = Image.open(LOGO_PATH).convert("RGBA")

    draw.rounded_rectangle(
        (-120, 210, 1520, 820),
        radius=120,
        fill=(*BLUE, 10),
    )
    draw.rounded_rectangle(
        (1240, 150, 2660, 820),
        radius=140,
        fill=(*ORANGE, 8),
    )
    draw.rounded_rectangle(
        (-80, 2050, 1630, 2860),
        radius=140,
        fill=(*ORANGE, 6),
    )
    draw.rounded_rectangle(
        (980, 2220, 2650, 3020),
        radius=160,
        fill=(*BLUE, 8),
    )
    accents = accents.filter(ImageFilter.GaussianBlur(26))
    page.alpha_composite(accents)

    target_width = int(page.width * LOGO_SCALE)
    target_height = int(logo.height * (target_width / logo.width))
    hero_logo = logo.resize((target_width, target_height), Image.LANCZOS)
    hero_alpha = hero_logo.getchannel("A").point(lambda value: int(value * LOGO_OPACITY))
    hero_logo.putalpha(hero_alpha)
    hero_x = (page.width - hero_logo.width) // 2
    hero_y = int((page.height - hero_logo.height) * LOGO_VERTICAL_OFFSET)
    page.alpha_composite(hero_logo, (hero_x, hero_y))

    secondary_width = int(page.width * 0.52)
    secondary_height = int(logo.height * (secondary_width / logo.width))
    secondary_logo = logo.resize((secondary_width, secondary_height), Image.LANCZOS).rotate(
        -3, resample=Image.BICUBIC, expand=True
    )
    secondary_alpha = secondary_logo.getchannel("A").point(
        lambda value: int(value * SECONDARY_LOGO_OPACITY)
    )
    secondary_logo.putalpha(secondary_alpha)
    secondary_x = (page.width - secondary_logo.width) // 2
    secondary_y = int(page.height * 0.74)
    page.alpha_composite(secondary_logo, (secondary_x, secondary_y))

    text_layer = Image.new("RGBA", page.size, (255, 255, 255, 0))
    text_draw = ImageDraw.Draw(text_layer)
    font = load_font(120)
    watermark = "TransferAI x Nettelecom"
    bbox = text_draw.textbbox((0, 0), watermark, font=font)
    text_width = bbox[2] - bbox[0]
    text_x = (page.width - text_width) // 2
    text_y = 1590
    text_draw.text((text_x, text_y), watermark, font=font, fill=(*BLUE, int(255 * TEXT_BAND_OPACITY)))
    text_draw.text((text_x, text_y + 150), watermark, font=font, fill=(*ORANGE, int(255 * 0.05)))
    text_layer = text_layer.filter(ImageFilter.GaussianBlur(1))
    page.alpha_composite(text_layer)

    page.save(BACKGROUND_PNG, dpi=(300, 300))


def add_inline_header_image() -> None:
    doc = Document(str(SHARED_SOURCE_DOCX))

    for section in doc.sections:
        section.different_first_page_header_footer = False
        header = section.header
        if not header.paragraphs:
            header.add_paragraph("")
        paragraph = header.paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run()
        run.add_picture(str(BACKGROUND_PNG), width=section.page_width, height=section.page_height)

    doc.save(str(INTERMEDIATE_DOCX))


def make_vml_picture(rel_id: str, width_pt: float, height_pt: float) -> etree._Element:
    pict = etree.Element(w("pict"), nsmap={"v": V_NS, "o": O_NS, "r": R_NS})
    shape = etree.SubElement(
        pict,
        v("shape"),
        {
            "id": "TransferAIBackground",
            o("spid"): "_x0000_s2049",
            "type": "#_x0000_t75",
            "style": (
                "position:absolute;"
                f"width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;"
                "z-index:-251654144;"
                "mso-wrap-edited:f;"
                "mso-position-horizontal:center;"
                "mso-position-horizontal-relative:page;"
                "mso-position-vertical:center;"
                "mso-position-vertical-relative:page;"
            ),
            "stroked": "f",
        },
    )
    etree.SubElement(shape, v("imagedata"), {f"{{{R_NS}}}id": rel_id, o("title"): "TransferAI background"})
    return pict


def patch_header_to_background() -> None:
    with zipfile.ZipFile(INTERMEDIATE_DOCX, "r") as zin:
        header_names = [name for name in zin.namelist() if name.startswith("word/header") and name.endswith(".xml")]
        overrides: dict[str, bytes] = {}

        for header_name in header_names:
            root = etree.fromstring(zin.read(header_name))
            first_drawing = root.find(".//w:drawing", namespaces=NS)
            if first_drawing is None:
                continue

            extent = first_drawing.find(".//wp:extent", namespaces=NS)
            blip = first_drawing.find(".//a:blip", namespaces=NS)
            if extent is None or blip is None:
                continue

            rel_id = blip.get(f"{{{R_NS}}}embed")
            cx = int(extent.get("cx"))
            cy = int(extent.get("cy"))
            width_pt = cx / 12700
            height_pt = cy / 12700

            run = first_drawing.getparent()
            run.remove(first_drawing)
            run.append(make_vml_picture(rel_id, width_pt, height_pt))

            overrides[header_name] = etree.tostring(
                root,
                xml_declaration=True,
                encoding="UTF-8",
                standalone="yes",
            )

        with zipfile.ZipFile(OUTPUT_DOCX, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = overrides.get(info.filename, zin.read(info.filename))
                zout.writestr(info, data)


def main() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    build_shared_source()
    build_background_image()
    add_inline_header_image()
    patch_header_to_background()
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
