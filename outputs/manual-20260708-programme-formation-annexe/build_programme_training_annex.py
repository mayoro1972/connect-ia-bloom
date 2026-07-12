from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("outputs/manual-20260708-programme-formation-annexe")
DOCX_PATH = OUTPUT_DIR / "Programme_Formation_TransferAI_Annexe_2026-07-08.docx"

BLUE = RGBColor(26, 76, 130)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(95, 99, 104)
LIGHT = RGBColor(242, 245, 248)
BORDER = "DADCE0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, color=DARK, bold=False, italic=False):
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, before=0, after=6, line=1.10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def add_text_paragraph(doc, text, before=0, after=6, size=11, color=DARK, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=before, after=after, line=1.10, alignment=alignment)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, before=0, after=4, line=1.15)
    run = paragraph.add_run(text)
    set_font(run, size=11, color=DARK)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        style_paragraph(paragraph, before=16, after=8, line=1.0)
        size = 16
        color = BLUE
    elif level == 2:
        style_paragraph(paragraph, before=12, after=6, line=1.0)
        size = 13
        color = BLUE
    else:
        style_paragraph(paragraph, before=8, after=4, line=1.0)
        size = 12
        color = RGBColor(31, 77, 120)
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=True)
    return paragraph


def add_rule(doc):
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph, before=0, after=10, line=1.0)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return paragraph


def build_info_table(doc):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(3.2), Inches(3.2)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

    entries = [
        ("Format", "Formation présentielle de 2 jours"),
        ("Public visé", "PME, grandes entreprises et équipes métiers"),
        ("Approche", "Démonstrations concrètes, ateliers pratiques, plan d'action"),
        ("Point fort", "Protection des données sensibles et outils opérationnels"),
    ]
    for index, (label, value) in enumerate(entries):
        row = index // 2
        col = index % 2
        cell = table.rows[row].cells[col]
        set_cell_shading(cell, "F2F5F8")
        p1 = cell.paragraphs[0]
        style_paragraph(p1, before=0, after=2, line=1.0)
        r1 = p1.add_run(label)
        set_font(r1, size=10.5, color=MUTED, bold=True)
        p2 = cell.add_paragraph()
        style_paragraph(p2, before=0, after=0, line=1.10)
        r2 = p2.add_run(value)
        set_font(r2, size=11, color=DARK, bold=False)


def build_expert_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(3.0), Inches(3.5)]
    headers = ["Expert", "Domaine"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_shading(cell, "EAF0F6")
        p = cell.paragraphs[0]
        style_paragraph(p, before=0, after=0, line=1.0)
        run = p.add_run(headers[idx])
        set_font(run, size=11, color=BLUE, bold=True)

    experts = [
        ("Dr GOUHO Bi Néantien Jean Baptiste", "Expert en IA, automatisation et systèmes d'entreprise"),
        ("Roland COFFI", "Consultant expert en diplomatie, IA appliquée et formation"),
        ("Jean Serge ANON", "Expert auditeur et formateur-coach multipotentiel"),
        ("Sery Medard", "Consultant expert en data engineering, cloud et plateformes IA"),
    ]
    for name, domain in experts:
        row = table.add_row()
        for idx, value in enumerate((name, domain)):
            cell = row.cells[idx]
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.10)
            run = p.add_run(value)
            set_font(run, size=10.5, color=DARK, bold=(idx == 0))


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(paragraph, before=0, after=0, line=1.0)
    run = paragraph.add_run("TransferAI Africa | Annexe programme de formation")
    set_font(run, size=9, color=MUTED)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    style_r_pr = normal._element.get_or_add_rPr()
    style_r_fonts = style_r_pr.rFonts
    if style_r_fonts is None:
        style_r_fonts = OxmlElement("w:rFonts")
        style_r_pr.append(style_r_fonts)
    style_r_fonts.set(qn("w:ascii"), "Calibri")
    style_r_fonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    style_paragraph(title, before=0, after=4, line=1.0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Programme de formation IA appliquée")
    set_font(title_run, size=23, color=DARK, bold=True)

    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, before=0, after=8, line=1.0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Annexe prête à joindre à un courrier de proposition")
    set_font(sub_run, size=12, color=MUTED)

    add_rule(doc)
    add_text_paragraph(
        doc,
        "Cette formation TransferAI est conçue pour aider les organisations à comprendre les usages concrets de l'IA, sécuriser l'utilisation des données sensibles et définir un plan d'action réaliste à partir de cas métier.",
        before=0,
        after=10,
        size=11,
        color=DARK,
    )

    build_info_table(doc)

    add_heading(doc, "Objectifs de la session", level=1)
    add_bullet(doc, "Comprendre ce que l'IA générative peut et ne peut pas faire dans une PME ou une grande entreprise.")
    add_bullet(doc, "Découvrir des solutions TransferAI déjà en production et directement transposables à des contextes métier.")
    add_bullet(doc, "Adopter les bons réflexes de confidentialité avant toute utilisation d'un outil d'IA grand public.")
    add_bullet(doc, "Construire un plan d'action IA personnalisé à partir des enjeux réels des participants.")

    add_heading(doc, "Programme détaillé", level=1)

    add_heading(doc, "Jour 1 (matin) | Panorama IA et services TransferAI", level=2)
    add_bullet(doc, "Introduction à l'IA générative : usages pertinents, limites et points de vigilance pour l'entreprise.")
    add_bullet(doc, "Panorama des solutions TransferAI déjà en production : automatisation de la relation prospect/client, audit digital automatisé, assistant conversationnel multicanal, recherche documentaire intelligente et comptes-rendus de réunion par transcription IA.")
    add_bullet(doc, "Mise en perspective des bénéfices mesurables : gain de temps, disponibilité 24/7, réduction des erreurs de saisie et scalabilité sans recrutement proportionnel.")

    add_heading(doc, "Jour 1 (après-midi) | Sensibilisation aux données sensibles", level=2)
    add_bullet(doc, "Pourquoi la question est critique : une donnée partagée dans une IA publique peut être mémorisée, journalisée ou exposée.")
    add_bullet(doc, "Typologie des données sensibles en entreprise : identité, coordonnées bancaires, données RH, informations clients confidentielles et secrets commerciaux.")
    add_bullet(doc, "Démonstration en direct du module PII Guard de TransferAI : anonymisation locale sans appel externe avant traitement IA, puis restitution réservée aux personnes habilitées.")
    add_bullet(doc, "Bonnes pratiques immédiatement applicables : ce qu'il est permis d'utiliser avec l'IA et ce qu'il ne faut jamais copier-coller dans un outil grand public.")
    add_bullet(doc, "Atelier d'identification des données sensibles propres à l'activité de chaque participant.")

    add_heading(doc, "Jour 2 | Mise en pratique et plan d'action", level=2)
    add_bullet(doc, "Ateliers pratiques sur des cas réels apportés par les participants : prospection, audit interne, support client et organisation documentaire.")
    add_bullet(doc, "Construction d'un plan d'action IA personnalisé par participant ou par entreprise.")
    add_bullet(doc, "Session de questions-réponses avec un expert du réseau des 13 experts internationaux TransferAI.")
    add_bullet(doc, "Présentation des modalités d'accompagnement post-formation.")

    add_heading(doc, "Livrables remis aux participants", level=1)
    add_bullet(doc, "Support de formation complet au format PDF.")
    add_bullet(doc, "Grille d'auto-diagnostic IA pour l'entreprise.")
    add_bullet(doc, "Plan d'action IA personnalisé remis individuellement.")
    add_bullet(doc, "Accès de démonstration à un ou plusieurs outils présentés.")
    add_bullet(doc, "Attestation de participation TransferAI Africa.")

    add_heading(doc, "Différenciateurs clés", level=1)
    add_bullet(doc, "Ordinateurs portables fournis pour faciliter la mise en pratique pendant la session.")
    add_bullet(doc, "Accompagnement post-formation pour aider à transformer les apprentissages en actions concrètes.")
    add_bullet(doc, "Focus assumé sur la protection des données sensibles et la maîtrise des risques liés à l'usage de l'IA.")
    add_bullet(doc, "Équipe d'experts de référence incluant notamment Dr GOUHO Bi Néantien Jean Baptiste, Docteur-Ingénieur en Informatique avec plus de 15 ans d'expérience en IA, ainsi qu'Emmanuelle Koffi, certifiée Google AI Professional.")

    add_heading(doc, "Experts mobilisés", level=1)
    build_expert_table(doc)

    add_text_paragraph(
        doc,
        "Les modalités pratiques complémentaires, telles que les horaires détaillés, le lieu exact et les conditions tarifaires, peuvent être précisées dans la version finale ou dans le courrier d'accompagnement.",
        before=10,
        after=0,
        size=10,
        color=MUTED,
        italic=True,
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
