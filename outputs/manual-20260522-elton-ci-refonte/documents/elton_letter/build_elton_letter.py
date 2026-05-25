from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260522-elton-ci-refonte/documents/elton_letter")
OUTPUT_DOCX = BASE_DIR / "TransferAI_ELTON_CI_Courrier_Revu_Final.docx"
ACCENT = RGBColor(0xDE, 0x6A, 0x1F)
INK = RGBColor(0x17, 0x26, 0x3E)
MUTED = RGBColor(0x5C, 0x67, 0x77)


def set_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def clear_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
      tbl_borders = OxmlElement("w:tblBorders")
      tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "nil")
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                left={"val": "nil"},
                top={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
            )


def add_bullet(doc, text_value):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"• {text_value}")
    set_font(run, size=11, color=INK)


doc = Document()
section = doc.sections[0]
section.start_type = WD_SECTION.NEW_PAGE
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = INK

for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

styles["Title"].font.size = Pt(18)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = INK
styles["Heading 1"].font.size = Pt(13)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = INK
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = INK

header = section.header
header_p = header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(4)
run = header_p.add_run("TransferAI Africa")
set_font(run, size=13, bold=True, color=INK)
sub = header.add_paragraph()
sub.paragraph_format.space_after = Pt(2)
sub_run = sub.add_run("Audit IA | Formation ciblée | Accompagnement de mise en œuvre")
set_font(sub_run, size=9.5, color=MUTED)
rule = header.add_paragraph()
rule.paragraph_format.space_after = Pt(8)
rule_run = rule.add_run(" ")
rule_run.font.size = Pt(2)
rule_p = rule._p
p_pr = rule_p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "DE6A1F")
p_bdr.append(bottom)
p_pr.append(p_bdr)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date_p.paragraph_format.space_after = Pt(10)
run = date_p.add_run("Abidjan, le 22 mai 2026")
set_font(run, size=10.5, color=MUTED)

recipient_lines = [
    "À l'attention de Monsieur le Directeur Général",
    "ELTON Oil CI",
    "Abidjan, Côte d'Ivoire",
]
for index, line in enumerate(recipient_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1 if index < len(recipient_lines) - 1 else 10)
    run = p.add_run(line)
    set_font(run, size=11, bold=index < 2, color=INK if index < 2 else MUTED)

subject = doc.add_paragraph()
subject.paragraph_format.space_after = Pt(12)
run = subject.add_run("Objet : Proposition d'audit IA et d'accompagnement opérationnel pour ELTON Oil CI")
set_font(run, size=11.5, bold=True, color=INK)

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(10)
lead_run = lead.add_run("Monsieur le Directeur Général,")
set_font(lead_run, size=11, color=INK)

paragraphs = [
    "Je me permets de vous adresser, en pièce jointe, une offre ciblée conçue pour ELTON Oil CI.",
    "Chez TransferAI Africa, nous accompagnons les entreprises qui souhaitent aborder l'intelligence artificielle de manière utile, progressive et maîtrisée. Notre conviction est qu'une structure comme ELTON peut tirer une valeur concrète de l'IA non pas en ajoutant un outil générique, mais en partant de quelques priorités opérationnelles clairement identifiées.",
    "Au regard de vos activités, nous pensons qu'une première démarche structurée pourrait utilement porter sur :",
]

for value in paragraphs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(value)
    set_font(run, size=11, color=INK)

for bullet in [
    "le pilotage des comptes professionnels et le suivi des relances ;",
    "l'exploitation des données liées à la flotte, aux cartes carburant et aux opérations ;",
    "l'amélioration du suivi client et du service auto ;",
    "la diffusion plus simple des procédures, standards et informations utiles aux équipes.",
]:
    add_bullet(doc, bullet)

transition_1 = doc.add_paragraph()
transition_1.paragraph_format.space_before = Pt(4)
transition_1.paragraph_format.space_after = Pt(8)
run = transition_1.add_run("Notre approche se distingue par quatre principes simples :")
set_font(run, size=11, color=INK)

for bullet in [
    "partir des besoins métier avant de parler d'outil ;",
    "former les bonnes populations au bon niveau ;",
    "intégrer dès le départ les exigences de confidentialité, de gouvernance et de validation humaine ;",
    "n'engager un pilote qu'après un diagnostic court et des priorités partagées.",
]:
    add_bullet(doc, bullet)

transition_2 = doc.add_paragraph()
transition_2.paragraph_format.space_before = Pt(4)
transition_2.paragraph_format.space_after = Pt(8)
run = transition_2.add_run("Dans cet esprit, nous souhaiterions vous proposer un point d'entrée simple et sans engagement :")
set_font(run, size=11, color=INK)

for bullet in [
    "un échange de cadrage de 45 minutes avec vous ou la personne que vous désignerez ;",
    "un audit d'entrée gratuit centré sur quelques sujets prioritaires ;",
    "une note de recommandations présentant les cas d'usage à arbitrer, les besoins de formation et une première feuille de route.",
]:
    add_bullet(doc, bullet)

for value in [
    "Si cette démarche retient votre attention, nous serions heureux de convenir d'un court échange afin de vous présenter l'offre et de vérifier ensemble les sujets sur lesquels TransferAI Africa pourrait vous être utile.",
    "Je vous remercie par avance pour votre attention et vous prie d'agréer, Monsieur le Directeur Général, l'expression de ma haute considération.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(value)
    set_font(run, size=11, color=INK)

closing = doc.add_paragraph()
closing.paragraph_format.space_before = Pt(8)
closing.paragraph_format.space_after = Pt(6)
run = closing.add_run("Pour TransferAI Africa")
set_font(run, size=11, bold=True, color=INK)

signers = [
    ("Casimir Kassi BEDA", "Directeur Général"),
    ("Eric N'Guessan", "Dir. Audit, Pédagogie & Certifications"),
    ("Marius Ayoro", "Dir. Développement & Partenariats"),
]

for name, role in signers:
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_after = Pt(1)
    r1 = p_name.add_run(name)
    set_font(r1, size=10.5, bold=True, color=INK)

    p_role = doc.add_paragraph()
    p_role.paragraph_format.space_after = Pt(5)
    r2 = p_role.add_run(role)
    set_font(r2, size=9.5, color=MUTED)

OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT_DOCX)
print(OUTPUT_DOCX)
