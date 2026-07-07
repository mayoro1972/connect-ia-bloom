from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260706-pre-cloture-guide")
DOCX_PATH = OUTPUT_DIR / "Guide_Utilisateur_Pre_Cloture_Bancaire_n8n_TransferAI_Nettelecom_CI.docx"
JSON_PATH = OUTPUT_DIR / "workflow_precloture_bancaire_multibanques_n8n.json"
LOGO_PATH = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-prospection/assets/transferai-africa-nettelecom-co-brand.png")


BLUE = RGBColor(22, 53, 86)
DARK = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "EDEDED"
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)


def set_run_font(run, size=11, bold=False, color=BLACK, name="Arial", italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_paragraph(paragraph, after=6, before=0, line=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def add_title(doc, title, subtitle):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run()
    run_logo.add_picture(str(LOGO_PATH), width=Inches(2.3))
    style_paragraph(p_logo, after=10)

    p = doc.add_paragraph()
    style_paragraph(p, after=4, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(title)
    set_run_font(run, size=23, bold=True, color=BLUE)

    p2 = doc.add_paragraph()
    style_paragraph(p2, after=16, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run2 = p2.add_run(subtitle)
    set_run_font(run2, size=12, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    widths = [1.6, 4.5]
    rows = [
        ("Document", "Guide utilisateur de mise en place du workflow n8n"),
        ("Fichier n8n", JSON_PATH.name),
        ("Public visé", "Équipes finance, contrôle financier, audit interne et administrateurs n8n"),
    ]
    for r_idx, (label, value) in enumerate(rows):
        row = meta.rows[r_idx]
        for c_idx, content in enumerate((label, value)):
            cell = row.cells[c_idx]
            set_cell_width(cell, widths[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=3, before=3)
            run = p.add_run(content)
            set_run_font(run, size=10.5, bold=(c_idx == 0), color=DARK)
            if c_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)

    doc.add_paragraph("")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(
        p,
        before=16 if level == 1 else 10,
        after=6,
    )
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
    else:
        set_run_font(run, size=12, bold=True, color=DARK)


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    style_paragraph(p, after=6)
    run = p.add_run(text)
    set_run_font(run, size=11, color=DARK, italic=italic)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=4)
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style_paragraph(p, after=4)
        run = p.add_run(item)
        set_run_font(run, size=11, color=DARK)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0]
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_width(cell, widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, LIGHT_FILL)
        p = cell.paragraphs[0]
        style_paragraph(p, after=3, before=3)
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, color=DARK)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, after=3, before=3)
            run = p.add_run(str(value))
            set_run_font(run, size=10.2, color=DARK)
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 6.2)
    set_cell_shading(cell, "F6F8FA")
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, before=0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, size=9.5, color=DARK, name="Courier New")


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, before=0)
    run = p.add_run("TransferAI x Nettelecom CI")
    set_run_font(run, size=9, color=MUTED)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    styles["Normal"].font.size = Pt(11)

    add_title(
        doc,
        "Guide utilisateur du workflow de pré-clôture bancaire dans n8n",
        "Cas d’usage concret, paramétrage, tests et mise en production pour une banque opérant en Côte d’Ivoire",
    )

    add_heading(doc, "1. Objet du document")
    add_body(
        doc,
        "Ce document explique comment déployer, paramétrer et exploiter un workflow n8n de pré-clôture bancaire à partir d’un fichier JSON exportable. Il est rédigé pour un client francophone et peut être utilisé par toute banque opérant en Côte d’Ivoire.",
    )
    add_body(
        doc,
        f"Le workflow de référence est fourni dans le fichier suivant : {JSON_PATH.name}.",
    )

    add_heading(doc, "2. Cas d’usage concret")
    add_body(
        doc,
        "À chaque fin de semaine ou fin de mois, la direction financière d’une banque doit analyser un export comptable ou une balance de pré-clôture avant validation. Cette étape est souvent ralentie par des contrôles manuels, des écarts détectés trop tard et des commentaires produits de manière hétérogène selon les équipes.",
    )
    add_body(
        doc,
        "Le workflow proposé automatise la lecture du fichier, détecte les variations significatives, signale les justificatifs manquants et génère une synthèse managériale en français. Un responsable valide ensuite cette synthèse avant sa diffusion au contrôle financier, à la comptabilité, à l’audit interne ou au comité concerné.",
    )
    add_heading(doc, "3. Résultat attendu", level=2)
    add_bullets(
        doc,
        [
            "Réduire le temps de préparation de la pré-clôture.",
            "Détecter plus tôt les anomalies, écarts et exceptions.",
            "Homogénéiser la qualité des notes de synthèse.",
            "Maintenir une validation humaine avant tout usage officiel.",
        ],
    )

    add_heading(doc, "4. Contenu du workflow n8n")
    add_body(
        doc,
        "Le workflow exportable comprend les étapes suivantes. Elles sont déjà câblées dans le fichier JSON et peuvent être ajustées aux besoins du client.",
    )
    add_table(
        doc,
        ["Bloc", "Nom dans n8n", "Rôle"],
        [
            ("Déclenchement", "Trigger manuel / Webhook - Données pré-clôture", "Lance le workflow en mode test ou via un envoi HTTP."),
            ("Préparation", "Jeu de données démo / Normaliser entrée webhook", "Met en forme les données de pré-clôture et applique les valeurs par défaut."),
            ("Contrôles", "Détecter écarts et priorités", "Calcule les variations, classe les anomalies et prépare les KPI."),
            ("Protection", "Masquer données sensibles pour IA", "Pseudonymise les données sensibles avant tout envoi au moteur IA."),
            ("IA", "Générer synthèse IA", "Produit une synthèse managériale structurée en français à partir de données pseudonymisées."),
            ("Validation", "Construire email validation / Envoyer email validation", "Envoie un e-mail de validation humaine avant diffusion."),
            ("Diffusion", "Webhook - Approbation finale / Envoyer synthèse finale", "Diffuse la synthèse uniquement après approbation."),
        ],
        [1.2, 2.2, 3.1],
    )

    add_heading(doc, "5. Prérequis techniques")
    add_bullets(
        doc,
        [
            "Une instance n8n active et accessible.",
            "Un accès d’administration pour importer un workflow JSON.",
            "Une clé API OpenAI pour la génération de synthèse.",
            "Une clé API Resend pour l’envoi des e-mails de validation et de diffusion.",
            "Une URL publique ou interne stable pour n8n si la validation se fait par lien.",
        ],
    )

    add_heading(doc, "6. Variables d’environnement à configurer")
    add_table(
        doc,
        ["Variable", "Obligatoire", "Usage"],
        [
            ("OPENAI_API_KEY", "Oui", "Authentifie l’appel au modèle IA chargé de rédiger la synthèse."),
            ("RESEND_API_KEY", "Oui", "Permet l’envoi des e-mails de validation et de diffusion."),
            ("N8N_BASE_URL", "Oui", "Construit les liens d’approbation envoyés au valideur."),
            ("OUTREACH_FROM_EMAIL", "Optionnel", "Adresse d’envoi affichée dans les e-mails."),
            ("INTERNAL_REVIEW_EMAIL", "Optionnel", "Adresse de validation par défaut si aucune adresse n’est fournie dans le payload."),
        ],
        [1.8, 0.9, 3.8],
    )

    add_heading(doc, "7. Structure des données attendues dans le webhook")
    add_body(
        doc,
        "Le webhook principal attend un payload JSON contenant les métadonnées du client bancaire et les lignes de pré-clôture à analyser.",
    )
    add_table(
        doc,
        ["Champ", "Type", "Description"],
        [
            ("entity", "Texte", "Nom de l’entité bancaire ou de la direction financière."),
            ("report_period", "Texte", "Période analysée, par exemple : Juin 2026."),
            ("currency", "Texte", "Devise utilisée, par exemple FCFA."),
            ("manager_email", "Texte", "Adresse e-mail du valideur principal."),
            ("recipients", "Liste", "Destinataires de la synthèse finale après validation."),
            ("absolute_variation_threshold", "Nombre", "Seuil absolu à partir duquel une variation est signalée."),
            ("relative_variation_threshold", "Nombre", "Seuil relatif à partir duquel une variation est signalée."),
            ("rows", "Liste", "Ensemble des lignes de balance ou d’export comptable à analyser."),
        ],
        [2.0, 1.0, 3.5],
    )

    add_heading(doc, "8. Protection des données / PII")
    add_body(
        doc,
        "Pour tenir compte des données sensibles en contexte bancaire, le workflow applique une pseudonymisation avant l’étape IA. Cette approche réduit l’exposition des données à caractère personnel tout en conservant les informations strictement nécessaires à l’analyse des écarts.",
    )
    add_bullets(
        doc,
        [
            "Le workflow ne transmet au modèle IA que des anomalies pseudonymisées et des KPI agrégés.",
            "Le nom de l’établissement, les références détaillées, les e-mails, les téléphones et les identifiants longs sont masqués ou généricisés avant l’appel IA.",
            "Les informations utiles à l’analyse restent disponibles pour les équipes internes dans le flux n8n, sous contrôle d’accès approprié.",
            "La validation humaine reste obligatoire avant toute diffusion officielle ou tout arbitrage de gestion.",
        ],
    )
    add_heading(doc, "8.1 Champs à masquer ou à pseudonymiser", level=2)
    add_table(
        doc,
        ["Champ ou catégorie", "Traitement recommandé", "Justification"],
        [
            ("Nom de banque ou d’entité", "Génériciser pour l’IA", "Le modèle n’a pas besoin du nom exact de l’établissement pour commenter les écarts."),
            ("Nom de client ou de collaborateur", "Masquer ou remplacer par un alias", "Réduit l’exposition de données à caractère personnel."),
            ("Numéro de compte ou identifiant long", "Tronquer ou remplacer par une référence interne", "Conserve un repère sans exposer l’identifiant complet."),
            ("E-mail, téléphone, adresse", "Masquer systématiquement", "Ces données ne sont pas nécessaires à la synthèse managériale."),
            ("Libellés libres", "Nettoyer si des données personnelles y figurent", "Évite qu’une donnée sensible remonte dans le prompt IA."),
        ],
        [1.8, 1.8, 2.9],
    )

    add_heading(doc, "9. Exemple de cas concret pour une banque en Côte d’Ivoire")
    add_body(
        doc,
        "Une banque participante réalise sa pré-clôture mensuelle sur un export comptable transmis par la direction financière. Le workflow identifie plusieurs écarts inhabituels, dont un compte d’attente sur opérations cartes, une TVA à régulariser et des charges diverses à justifier. L’IA produit alors une note courte indiquant ce qui a bougé, ce qui doit être revu immédiatement, ce qui peut attendre et ce qui doit être remonté au manager.",
    )
    add_body(
        doc,
        "Le responsable du contrôle financier reçoit un e-mail de validation. Après relecture et approbation, la synthèse finale est envoyée aux équipes concernées. Ce schéma reste compatible avec des usages hebdomadaires, mensuels ou ad hoc selon l’organisation de la banque.",
    )
    add_body(
        doc,
        "Dans cette configuration, la synthèse IA est générée à partir de données pseudonymisées, ce qui permet de limiter la circulation des informations sensibles tout en maintenant la valeur opérationnelle du commentaire.",
    )

    add_heading(doc, "10. Étapes de mise en place dans n8n")
    add_numbered(
        doc,
        [
            "Importer le fichier JSON exportable dans n8n.",
            "Vérifier le nom du workflow et l’activer en mode brouillon.",
            "Configurer les variables d’environnement OPENAI_API_KEY, RESEND_API_KEY et N8N_BASE_URL.",
            "Adapter les adresses e-mail par défaut aux destinataires du client.",
            "Vérifier les règles de pseudonymisation et confirmer les champs sensibles à masquer selon la politique interne du client.",
            "Tester d’abord le workflow avec le nœud « Trigger manuel » et le jeu de données de démonstration.",
            "Tester ensuite le webhook avec un payload réel ou simulé envoyé en HTTP POST.",
            "Vérifier la réception de l’e-mail de validation et le fonctionnement du lien d’approbation.",
            "Confirmer que la synthèse finale n’est envoyée qu’après validation humaine.",
        ],
    )

    add_heading(doc, "11. Paramètres à adapter avant production")
    add_bullets(
        doc,
        [
            "Le nom de l’entité bancaire et les adresses des destinataires.",
            "Les seuils absolus et relatifs de détection des écarts.",
            "La liste des champs sensibles à masquer ou à tronquer avant l’étape IA.",
            "Le contenu du commentaire managérial généré par l’IA si un ton spécifique est attendu.",
            "Le canal de diffusion final : e-mail, espace partagé, ou autre système interne.",
            "Le mode de déclenchement : manuel, webhook, dépôt de fichier ou intégration applicative.",
        ],
    )

    add_heading(doc, "12. Procédure de test recommandée")
    add_bullets(
        doc,
        [
            "Lancer le workflow avec les données de démonstration pour valider la logique générale.",
            "Envoyer un payload webhook contenant des lignes réelles anonymisées ou des données de test proches du réel.",
            "Vérifier que les e-mails, numéros, identifiants longs et libellés sensibles sont bien masqués avant l’étape IA.",
            "Vérifier que les anomalies remontées correspondent bien aux règles métier attendues.",
            "Contrôler la qualité de la synthèse générée en français professionnel.",
            "Tester le refus d’approbation afin de vérifier qu’aucune diffusion finale n’est déclenchée.",
            "Tester l’approbation afin de vérifier la diffusion finale au bon périmètre.",
        ],
    )

    add_heading(doc, "13. Checklist de mise en production")
    add_bullets(
        doc,
        [
            "Les clés API sont présentes et fonctionnelles.",
            "Les destinataires de validation et de diffusion ont été validés.",
            "Le lien N8N_BASE_URL est accessible depuis la messagerie du valideur.",
            "Les seuils d’écart ont été validés par la direction financière.",
            "Les règles de pseudonymisation ont été validées par les équipes sécurité, conformité ou risque si nécessaire.",
            "Le workflow a été testé sur un cas réel ou quasi réel.",
            "La validation humaine a été maintenue comme étape obligatoire.",
        ],
    )

    add_heading(doc, "14. Commande de test webhook", level=2)
    add_body(
        doc,
        "La commande suivante peut être utilisée pour tester le webhook principal après import du workflow :",
    )
    add_code_block(
        doc,
        [
            "curl -X POST \"https://VOTRE_N8N/webhook/pre-cloture-bancaire-upload\" \\",
            "  -H \"Content-Type: application/json\" \\",
            "  -d @payload-precloture.json",
        ],
    )

    add_heading(doc, "15. Références de livraison")
    add_bullets(
        doc,
        [
            "Présentation PowerPoint source : TransferAI_Presentation_Pre_Cloture_Bancaire_Multi_Banques.pptx",
            f"Workflow n8n exportable générique : {JSON_PATH.name}",
            "Guide utilisateur : le présent document.",
        ],
    )

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_document()
