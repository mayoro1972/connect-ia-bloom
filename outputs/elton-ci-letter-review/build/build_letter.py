from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUTPUT_DIR = ROOT / "outputs/elton-ci-letter-review/output"
OUTPUT_DOCX = OUTPUT_DIR / "Courrier_ELTON_Oil_CI_TransferAI_Relecture.docx"
LOGO_PATH = Path("/Users/marius_ayoro/Downloads/image (1).png")


BODY_COLOR = RGBColor(32, 44, 56)
ACCENT_COLOR = RGBColor(180, 137, 76)
FONT_NAME = "Aptos"


def format_run(run, *, size=11, bold=False, color=BODY_COLOR):
    font = run.font
    font.name = FONT_NAME
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = color


def add_text_paragraph(doc: Document, text: str, *, bold=False, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.08
    run = para.add_run(text)
    format_run(run, bold=bold)
    return para


def add_bullet(doc: Document, text: str):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.02
    run = para.add_run(text)
    format_run(run)
    return para


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.8)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(10)
    header_run = header_para.add_run()
    header_run.add_picture(str(LOGO_PATH), width=Cm(5.4))

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(3)
    footer_para.paragraph_format.space_after = Pt(0)
    footer_run = footer_para.add_run(
        "TransferAI | Hub IA de Nettelecom CI | contact@transferai.ci | "
        "www.transferai.ci | +225 07 16 57 39 90 | Riviera 3, carrefour Sainte Famille, Abidjan"
    )
    format_run(footer_run, size=7, color=RGBColor(90, 99, 109))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run("Abidjan, le 26 mai 2026")
    format_run(run, size=11)

    for line in (
        "À l'attention de Monsieur le Directeur Général",
        "ELTON Oil CI",
        "Abidjan, Côte d'Ivoire",
    ):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        format_run(run, size=11)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    obj = doc.add_paragraph()
    obj.paragraph_format.space_after = Pt(10)
    obj_run = obj.add_run(
        "Objet : Proposition d'audit gratuit, d'accompagnement et de formation pour ELTON Oil CI"
    )
    format_run(obj_run, bold=True, color=ACCENT_COLOR)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    run = salutation.add_run("Monsieur le Directeur Général,")
    format_run(run)

    add_text_paragraph(
        doc,
        "TransferAI, hub IA de Nettelecom CI, accompagne les entreprises dans l'identification, "
        "la structuration et la mise en œuvre d'usages concrets de l'intelligence artificielle. "
        "Notre approche associe diagnostic, automatisation, pilotage, formation et gouvernance, "
        "afin de transformer l'IA en levier opérationnel réellement utile pour les équipes.",
    )

    add_text_paragraph(
        doc,
        "TransferAI s'appuie sur une équipe de 13 experts en intelligence artificielle, mobilisable "
        "sur des enjeux de diagnostic, d'automatisation, de pilotage, de formation et de gouvernance. "
        "Cette équipe intervient sur différents secteurs d'activité, ce qui nous permet d'adapter les "
        "solutions proposées aux contraintes métier, aux exigences de confidentialité et aux réalités "
        "d'exécution propres à chaque entreprise.",
    )

    add_text_paragraph(
        doc,
        "En observant les activités visibles d'ELTON CI, notamment autour des comptes professionnels, "
        "de la carte carburant, du suivi de flotte, des services auto, des lubrifiants, ainsi que de la "
        "relation client et des opérations terrain, nous pensons qu'il existe plusieurs leviers sur "
        "lesquels nous pourrions utilement vous accompagner.",
    )

    add_text_paragraph(
        doc,
        "Notre démarche ne consiste pas à proposer une IA générique ou un simple effet d'annonce. "
        "Nous privilégions au contraire une logique de service : cadrer les priorités, qualifier les "
        "cas d'usage, former les bonnes équipes et lancer des pilotes mesurables dans un cadre maîtrisé.",
    )

    intro_cases = doc.add_paragraph()
    intro_cases.paragraph_format.space_after = Pt(4)
    intro_run = intro_cases.add_run(
        "À ce stade, les cas d'usage qui nous paraissent les plus pertinents pour ELTON CI sont les suivants :"
    )
    format_run(intro_run, bold=True)

    add_bullet(
        doc,
        "le pilotage des comptes professionnels et du recouvrement, avec des relances mieux "
        "préparées, des revues de portefeuille plus rapides et une meilleure priorisation des comptes sensibles ;",
    )
    add_bullet(
        doc,
        "le suivi de la flotte, des cartes carburant et des consommations, afin de consolider les "
        "données, détecter plus vite les écarts et mieux piloter les coûts ;",
    )
    add_bullet(
        doc,
        "le service client et les services auto, pour homogénéiser les réponses, améliorer la "
        "réactivité opérationnelle et mieux exploiter l'historique des interactions ;",
    )
    add_bullet(
        doc,
        "la diffusion des procédures internes et la montée en compétence des équipes, grâce à une base "
        "de connaissances plus accessible et à des formations ciblées sur les usages réellement utiles.",
    )

    add_text_paragraph(
        doc,
        "Dans cette logique, nous souhaiterions vous proposer un audit IA gratuit centré sur quelques "
        "priorités métier, complété par un rendez-vous de cadrage de 45 minutes avec vous ou avec la "
        "personne que vous désignerez. À l'issue de cet échange, nous pourrions vous remettre une première "
        "note de recommandations, sans engagement de votre part, avec une priorisation claire des pistes "
        "les plus réalistes pour ELTON CI.",
    )

    add_text_paragraph(
        doc,
        "Pour faciliter la lecture de notre démarche, ce courrier est accompagné, en pièces jointes, "
        "du deck de présentation ELTON CI ainsi que du catalogue de formation ciblé correspondant.",
    )

    add_text_paragraph(
        doc,
        "Nous serions honorés de pouvoir vous présenter cette démarche plus en détail et d'examiner avec "
        "vous les conditions d'un premier pilote utile, gouverné et mesurable.",
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(10)
    closing_run = closing.add_run(
        "Je vous prie d'agréer, Monsieur le Directeur Général, l'expression de ma considération distinguée."
    )
    format_run(closing_run)

    signature = doc.add_paragraph()
    signature.paragraph_format.space_after = Pt(0)
    signature_run = signature.add_run("L'équipe Hub TransferAI")
    format_run(signature_run, bold=True)

    sig2 = doc.add_paragraph()
    sig2.paragraph_format.space_after = Pt(0)
    run = sig2.add_run("Hub IA de Nettelecom CI")
    format_run(run)

    sig3 = doc.add_paragraph()
    sig3.paragraph_format.space_after = Pt(0)
    run = sig3.add_run("Site web : https://www.transferai.ci")
    format_run(run)

    sig4 = doc.add_paragraph()
    sig4.paragraph_format.space_after = Pt(0)
    run = sig4.add_run("Assistant IA : +225 07 16 57 39 90")
    format_run(run)

    sig5 = doc.add_paragraph()
    sig5.paragraph_format.space_after = Pt(0)
    run = sig5.add_run("contact@transferai.ci")
    format_run(run)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_document()
