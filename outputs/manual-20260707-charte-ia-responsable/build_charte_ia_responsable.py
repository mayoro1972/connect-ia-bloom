from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/outputs/manual-20260707-charte-ia-responsable"
)
OUT_PATH = BASE_DIR / "Charte_Utilisation_Responsable_IA_TransferAI_Africa_2026-07.docx"
LOGO_PATH = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/src/assets/logo-transferai-nettelecom.png"
)

NAVY = RGBColor(0x14, 0x2D, 0x4C)
GOLD = RGBColor(0xC2, 0x90, 0x2D)
INK = RGBColor(0x2B, 0x2B, 0x2B)
SOFT = RGBColor(0x6C, 0x74, 0x7F)
PALE_GOLD = "F7F0E2"
PALE_BLUE = "EEF3F8"
FOOTER_LINE_1 = "www.transferai.ci | contact@transferai.ci | WhatsApp et agent IA : +225 07 16 57 39 90"
FOOTER_LINE_2 = "TransferAI Africa - Hub IA de NettelecomCI - Riviera 3, carrefour Sainte Famille, Abidjan, Côte d'Ivoire"


def set_run_font(run, name, size, color, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_paragraph_border(paragraph, edge="bottom", color="C2902D", size="8", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)


def set_page_borders(section, color="C2902D", size=14, space=20):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "double")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), str(space))
        border.set(qn("w:color"), color)
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def configure_section(section):
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.25)
    set_page_borders(section)


def add_footer(section):
    footer = section.footer
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(0)
    set_paragraph_border(p1, edge="top", color="CFC3A0", size="6", space="4")
    r1 = p1.add_run(FOOTER_LINE_1)
    set_run_font(r1, "Aptos", 8.2, NAVY, bold=False)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(FOOTER_LINE_2)
    set_run_font(r2, "Aptos", 7.6, SOFT, bold=False)


def define_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.6)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, font_name, size, color, bold in [
        ("Heading 1", "Cambria", 14.5, NAVY, True),
        ("Heading 2", "Cambria", 12.5, GOLD, True),
        ("Heading 3", "Aptos", 11.0, NAVY, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    custom_styles = [
        ("CoverKicker", WD_STYLE_TYPE.PARAGRAPH, "Aptos", 10.0, GOLD, True),
        ("CoverTitle", WD_STYLE_TYPE.PARAGRAPH, "Cambria", 21.0, NAVY, True),
        ("CoverSubtitle", WD_STYLE_TYPE.PARAGRAPH, "Cambria", 12.5, SOFT, False),
        ("ArticleLabel", WD_STYLE_TYPE.PARAGRAPH, "Aptos", 8.8, GOLD, True),
        ("BodySmall", WD_STYLE_TYPE.PARAGRAPH, "Aptos", 9.4, SOFT, False),
        ("TableText", WD_STYLE_TYPE.PARAGRAPH, "Aptos", 9.3, INK, False),
    ]

    for name, style_type, font_name, size, color, bold in custom_styles:
        if name in doc.styles:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, style_type)
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.1


def add_logo(paragraph, width_cm):
    if LOGO_PATH.exists():
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(width_cm))


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    add_logo(p, 5.6)

    p = doc.add_paragraph(style="CoverKicker")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run("TRANSFERAI AFRICA - HUB IA DE NETTELECOMCI - CÔTE D'IVOIRE")

    p = doc.add_paragraph(style="CoverTitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run("CHARTE D'UTILISATION RESPONSABLE")

    p = doc.add_paragraph(style="CoverTitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.add_run("DE L'INTELLIGENCE ARTIFICIELLE")

    p = doc.add_paragraph(style="CoverSubtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Livret officiel remis aux participants de la formation")

    p = doc.add_paragraph(style="CoverSubtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('"IA appliquée à votre métier"')
    set_run_font(r, "Cambria", 12.5, NAVY, italic=True)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Inches(2.75)
    meta.columns[1].width = Inches(2.75)
    meta.cell(0, 0).width = Inches(2.75)
    meta.cell(0, 1).width = Inches(2.75)
    meta.cell(1, 0).width = Inches(2.75)
    meta.cell(1, 1).width = Inches(2.75)
    values = [
        ("Référence documentaire", "CHARTE-IA-TAI-2026-07"),
        ("Édition", "Juillet 2026"),
        ("Application", "Dès le lendemain de la formation"),
        ("Version", "Participant / entreprise"),
    ]
    for index, (label, value) in enumerate(values):
        row = index // 2
        col = index % 2
        cell = meta.cell(row, col)
        set_cell_shading(cell, PALE_BLUE if row == 0 else PALE_GOLD)
        set_cell_border(
            cell,
            left={"val": "single", "sz": 8, "color": "D8C999"},
            right={"val": "single", "sz": 8, "color": "D8C999"},
            top={"val": "single", "sz": 8, "color": "D8C999"},
            bottom={"val": "single", "sz": 8, "color": "D8C999"},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        label_run = p.add_run(f"{label}\n")
        set_run_font(label_run, "Aptos", 8.6, GOLD, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, "Cambria", 10.5, NAVY, bold=True)

    doc.add_paragraph()

    callout = doc.add_paragraph()
    callout.alignment = WD_ALIGN_PARAGRAPH.CENTER
    callout.paragraph_format.left_indent = Cm(0.4)
    callout.paragraph_format.right_indent = Cm(0.4)
    callout.paragraph_format.space_after = Pt(12)
    set_paragraph_border(callout, edge="top", color="C2902D", size="10", space="5")
    set_paragraph_border(callout, edge="bottom", color="C2902D", size="10", space="5")
    r = callout.add_run(
        "Cette charte fixe les règles minimales à respecter pour utiliser l'IA générative en entreprise "
        "avec prudence, traçabilité et discernement humain."
    )
    set_run_font(r, "Cambria", 11.8, NAVY, italic=True)

    badge = doc.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.paragraph_format.space_before = Pt(8)
    badge.paragraph_format.space_after = Pt(0)
    r = badge.add_run("Confidentialité - Responsabilité - Gouvernance - Confiance")
    set_run_font(r, "Aptos", 9.2, GOLD, bold=True)

    doc.add_page_break()


def add_article_heading(doc, number, title):
    p = doc.add_paragraph(style="ArticleLabel")
    p.paragraph_format.space_after = Pt(1)
    p.add_run(f"ARTICLE {number}")
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 1"]
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(5)
    heading.add_run(title)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_principles_box(doc):
    box = doc.add_table(rows=5, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    box.columns[0].width = Inches(6.2)
    items = [
        ("Minimisation", "ne partager avec un outil d'IA que ce qui est strictement nécessaire à la tâche."),
        ("Anonymisation avant tout", "retirer ou masquer les noms, numéros, emails, identifiants et pièces nominatives avant toute utilisation."),
        ("Traçabilité", "conserver la trace des usages significatifs, des prompts sensibles et des validations importantes."),
        ("Habilitation", "limiter l'accès aux outils, aux données et aux résultats aux personnes effectivement autorisées."),
        ("Réversibilité du doute", "en cas d'incertitude, choisir l'option la plus prudente et repasser par une validation humaine."),
    ]
    for row, (label, detail) in enumerate(items):
        cell = box.cell(row, 0)
        set_cell_shading(cell, PALE_GOLD if row % 2 == 0 else "FCFBF7")
        set_cell_border(
            cell,
            left={"val": "single", "sz": 6, "color": "E2D7B6"},
            right={"val": "single", "sz": 6, "color": "E2D7B6"},
            top={"val": "single", "sz": 6, "color": "E2D7B6"},
            bottom={"val": "single", "sz": 6, "color": "E2D7B6"},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        lead = p.add_run(f"{label} : ")
        set_run_font(lead, "Aptos", 10.0, NAVY, bold=True)
        detail_run = p.add_run(detail)
        set_run_font(detail_run, "Aptos", 10.0, INK)


def make_table(doc, headers, widths_in):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (header, width_in) in enumerate(zip(headers, widths_in)):
        cell = table.cell(0, idx)
        cell.width = Inches(width_in)
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(
            cell,
            left={"val": "single", "sz": 8, "color": "C9D5E2"},
            right={"val": "single", "sz": 8, "color": "C9D5E2"},
            top={"val": "single", "sz": 8, "color": "C9D5E2"},
            bottom={"val": "single", "sz": 8, "color": "C9D5E2"},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(header)
        set_run_font(r, "Aptos", 9.2, NAVY, bold=True)
    return table


def add_table_row(table, values, fill=None, alignments=None):
    row = table.add_row()
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Pt(18)
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        if fill:
            set_cell_shading(cell, fill)
        set_cell_border(
            cell,
            left={"val": "single", "sz": 6, "color": "E4E7EB"},
            right={"val": "single", "sz": 6, "color": "E4E7EB"},
            top={"val": "single", "sz": 6, "color": "E4E7EB"},
            bottom={"val": "single", "sz": 6, "color": "E4E7EB"},
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.style = "TableText"
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        set_run_font(run, "Aptos", 9.2, INK)


def add_articles(doc):
    add_article_heading(doc, "1", "Préambule")
    add_body_paragraph(
        doc,
        "La présente charte s'adresse à toute personne utilisant l'intelligence artificielle générative dans le cadre de son activité professionnelle. "
        "Elle fixe un socle de règles simples, immédiatement applicables, afin de protéger les données, la réputation de l'organisation et la qualité des décisions prises avec l'appui de l'IA. "
        "Sauf instruction contraire de la direction, cette charte s'applique dès le lendemain de la formation."
    )

    add_article_heading(doc, "2", "Cadre légal et périmètre")
    add_body_paragraph(
        doc,
        "La présente charte s'inscrit dans le respect du cadre légal ivoirien applicable à la protection des données, notamment la référence communiquée pendant la formation relative à la loi ivoirienne n° 2013-450 et au cadre ARTCI. "
        "Elle s'applique à tout salarié, manager, dirigeant, stagiaire, consultant, prestataire ou partenaire amené à utiliser un outil d'IA pour produire, analyser, synthétiser ou transformer de l'information liée à l'entreprise."
    )

    add_article_heading(doc, "3", "Principes fondamentaux")
    add_principles_box(doc)

    add_article_heading(doc, "4", "Classification des données sensibles par secteur")
    add_body_paragraph(
        doc,
        "Les données ci-dessous ne doivent jamais être copiées telles quelles dans un outil d'IA ouvert sans anonymisation préalable, validation adaptée et respect du périmètre d'habilitation."
    )
    sector_table = make_table(doc, ["Secteur", "Exemples de données sensibles"], [2.0, 4.2])
    sector_rows = [
        ("Banque / Finance", "Numéros de compte, relevés, informations KYC, profils de risque, documents de crédit, données de transaction."),
        ("Assurance", "Contrats, sinistres, pièces d'identité, primes, informations médicales ou patrimoniales."),
        ("Santé", "Dossiers patients, examens, comptes rendus, ordonnances, données biologiques ou biométrie."),
        ("Industrie / Énergie", "Incidents HSE nominatifs, badges, accès, plans sensibles, journaux d'intervention, données sous secret industriel."),
        ("Port / Logistique", "Manifestes, identités chauffeurs, mouvements de cargaison, géolocalisation, données de sûreté portuaire."),
        ("Administration / Diplomatie", "Données citoyens, correspondances officielles, notes internes, passeports, documents confidentiels de mission."),
        ("Formation / Éducation", "Dossiers apprenants, notes, évaluations, pièces d'inscription, données sociales ou disciplinaires."),
        ("Institutionnel", "Budgets non publics, délibérations, contacts sensibles, projets de décision, documents de gouvernance."),
        ("Distribution / Commerce", "Fichiers clients, historiques d'achat, paiements, fidélité, réclamations nominatives."),
        ("Transversal RH", "CV, paie, évaluations, santé au travail, sanctions, données disciplinaires ou d'identification."),
    ]
    for idx, row in enumerate(sector_rows):
        fill = "FCFBF8" if idx % 2 == 0 else None
        add_table_row(sector_table, row, fill=fill)

    add_article_heading(doc, "5", "Jamais / Possible copier-coller")
    do_table = make_table(doc, ["Jamais", "Possible sous conditions"], [3.05, 3.15])
    do_rows = [
        (
            "Copier un contrat, un dossier client ou un document RH avec noms, numéros et pièces réelles.",
            "Coller un contenu après anonymisation ou pseudonymisation complète.",
        ),
        (
            "Copier mots de passe, codes MFA, secrets commerciaux, clés API ou documents classés.",
            "Utiliser l'IA pour générer un modèle vierge, une trame, une check-list ou une structure.",
        ),
        (
            "Envoyer un relevé financier, un bulletin de salaire ou un dossier médical identifiant.",
            "Faire reformuler, résumer ou corriger un texte non sensible déjà autorisé.",
        ),
        (
            "Demander à l'IA de prendre seule une décision juridique, RH, financière ou médicale.",
            "Utiliser l'IA comme aide à l'analyse avant validation humaine compétente.",
        ),
        (
            "Contourner la voie de validation interne parce que l'outil semble rapide ou convaincant.",
            "Tester un cas d'usage sur données factices ou de démonstration avant extension.",
        ),
    ]
    for idx, row in enumerate(do_rows):
        fill = "FDF9F0" if idx % 2 == 0 else None
        add_table_row(do_table, row, fill=fill)

    add_article_heading(doc, "6", "Rôle du module d'anonymisation")
    add_body_paragraph(
        doc,
        "Le module d'anonymisation, illustré pendant la démonstration PII Guard, constitue un filtre de prudence avant l'envoi d'un contenu vers un moteur d'IA. "
        "Il a pour rôle de détecter puis de masquer les noms, numéros, emails, téléphones, références documentaires, identifiants et autres marqueurs personnels ou confidentiels."
    )
    add_body_paragraph(
        doc,
        "Ce module ne remplace pas le jugement humain. Il diminue le risque, mais l'utilisateur reste responsable de vérifier le résultat, de retirer les informations inutiles et de solliciter une validation supplémentaire lorsque le contexte l'exige."
    )

    doc.add_page_break()

    add_article_heading(doc, "7", "Gouvernance")
    gov_table = make_table(doc, ["Acteur", "Responsabilités minimales"], [1.85, 4.35])
    gov_rows = [
        ("Collaborateur", "Respecte la charte, anonymise avant usage, documente les cas sensibles et remonte tout doute sans attendre."),
        ("Référent IA interne", "Diffuse les bonnes pratiques, tient à jour les cas d'usage autorisés, accompagne les équipes et escalade les incidents."),
        ("Manager / validateur", "Contrôle les livrables à risque, arbitre les exceptions, valide les usages sensibles avant diffusion ou décision."),
        ("Direction générale", "Fixe l'orientation, nomme les responsables, soutient la revue annuelle et décide des mesures correctives majeures."),
    ]
    for idx, row in enumerate(gov_rows):
        fill = "F7FAFD" if idx % 2 == 0 else None
        add_table_row(gov_table, row, fill=fill)

    add_article_heading(doc, "8", "Procédure en cas de doute")
    steps = [
        ("1. Suspendre", "Ne pas copier-coller la donnée ou diffuser la sortie tant que le doute existe."),
        ("2. Nettoyer", "Retirer, anonymiser ou remplacer les éléments sensibles par des données factices."),
        ("3. Vérifier", "Consulter le référent IA interne ou le manager/validateur avec un contexte suffisant mais non sensible."),
        ("4. Tracer", "Conserver la décision prise et choisir l'option la plus prudente si l'incertitude persiste."),
    ]
    steps_table = make_table(doc, ["Étape", "Action attendue"], [1.2, 5.0])
    for idx, row in enumerate(steps):
        fill = "FCFBF8" if idx % 2 == 0 else None
        add_table_row(steps_table, row, fill=fill, alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

    add_article_heading(doc, "9", "Manquements et révision")
    add_body_paragraph(
        doc,
        "Tout manquement à cette charte, toute fuite présumée de données, tout usage non autorisé ou toute sortie IA diffusée sans validation adéquate doit être signalé sans délai. "
        "Selon la gravité, des mesures correctives, un retrait temporaire d'accès ou des suites disciplinaires internes peuvent être engagés conformément aux règles de l'organisation."
    )
    add_body_paragraph(
        doc,
        "La présente charte fait l'objet d'une révision annuelle minimum, et plus tôt si le cadre légal, les outils utilisés ou les risques de l'organisation évoluent de manière significative."
    )

    add_article_heading(doc, "10", "Engagement du participant")
    add_body_paragraph(
        doc,
        "En signant ce document, le participant reconnaît avoir pris connaissance des principes de cette charte et s'engage à les appliquer dans le cadre de son activité professionnelle."
    )

    signature = doc.add_table(rows=4, cols=1)
    signature.alignment = WD_TABLE_ALIGNMENT.CENTER
    signature.autofit = False
    signature.columns[0].width = Inches(6.15)
    fields = [
        "Nom et prénom : _______________________________________",
        "Entreprise / Fonction : _______________________________________",
        "Date : _______________________________________",
        "Signature : _______________________________________",
    ]
    for idx, field in enumerate(fields):
        cell = signature.cell(idx, 0)
        set_cell_shading(cell, "FFFDF8")
        set_cell_border(
            cell,
            left={"val": "single", "sz": 6, "color": "E3D3A4"},
            right={"val": "single", "sz": 6, "color": "E3D3A4"},
            top={"val": "single", "sz": 6, "color": "E3D3A4"},
            bottom={"val": "single", "sz": 6, "color": "E3D3A4"},
        )
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(field)
        set_run_font(run, "Cambria", 11.0, NAVY)

    closing = doc.add_paragraph(style="BodySmall")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(8)
    closing.paragraph_format.space_after = Pt(0)
    run = closing.add_run("TransferAI Africa - www.transferai.ci")
    set_run_font(run, "Aptos", 9.3, GOLD, bold=True)


doc = Document()
define_styles(doc)
configure_section(doc.sections[0])
add_footer(doc.sections[0])
add_cover(doc)
add_articles(doc)

BASE_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
