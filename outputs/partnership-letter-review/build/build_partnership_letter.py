from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUTPUT_DIR = ROOT / "outputs/partnership-letter-review/output"
OUTPUT_DOCX = OUTPUT_DIR / "Proposition_partenariat_strategique_TransferAI_signee.docx"
LOGO_PATH = Path("/Users/marius_ayoro/Downloads/image (1).png")
SIGNATURE_PATH = Path("/Users/marius_ayoro/Downloads/signature_transparent.png")

BODY_COLOR = RGBColor(32, 44, 56)
ACCENT_COLOR = RGBColor(180, 137, 76)
FOOTER_COLOR = RGBColor(90, 99, 109)
FONT_NAME = "Aptos"


def format_run(run, *, size=11, bold=False, color=BODY_COLOR, italic=False):
    font = run.font
    font.name = FONT_NAME
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    font.color.rgb = color


def add_para(doc: Document, text: str, *, after=6, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    format_run(r, bold=bold)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    format_run(r)
    return p


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.8)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(10)
    hr = hp.add_run()
    hr.add_picture(str(LOGO_PATH), width=Cm(5.4))

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run(
        "TransferAI | contact@transferai.ci | "
        "www.transferai.ci | +225 07 16 57 39 90 | Riviera 3, carrefour Sainte Famille, Abidjan"
    )
    format_run(fr, size=7, color=FOOTER_COLOR)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(10)
    r = meta.add_run("Abidjan, le 26 mai 2026")
    format_run(r, size=11)

    recipient_lines = (
        "À l'attention de Madame Mbarka Zineddine",
        "Administratrice générale",
        "Académie des Métiers de la Diplomatie",
    )
    for line in recipient_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        format_run(r, size=11)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    obj = doc.add_paragraph()
    obj.paragraph_format.space_after = Pt(10)
    r = obj.add_run("Objet : Proposition de partenariat stratégique")
    format_run(r, bold=True, color=ACCENT_COLOR)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    r = salutation.add_run("Madame l'Administratrice générale,")
    format_run(r)

    add_para(
        doc,
        "Permettez-moi de vous soumettre une proposition de partenariat stratégique entre TransferAI "
        "et l'Académie des Métiers de la Diplomatie, autour d'un enjeu désormais central : "
        "l'intégration de l'intelligence artificielle dans la pratique diplomatique, la gouvernance "
        "publique et la formation des cadres appelés à évoluer dans des environnements internationaux "
        "de plus en plus numériques.",
    )

    add_para(
        doc,
        "TransferAI accompagne les organisations dans la formation à l'intelligence artificielle appliquée, "
        "la structuration de parcours exécutifs et le développement de dispositifs pédagogiques orientés vers "
        "la pratique, l'employabilité et l'impact opérationnel. Notre approche vise à rendre ces usages "
        "accessibles, concrets et directement mobilisables dans les environnements institutionnels et professionnels.",
    )

    add_para(
        doc,
        "Dans un contexte marqué par la montée en puissance de la diplomatie numérique, de la donnée, de la "
        "cybersécurité et des enjeux technologiques dans les relations internationales, il existe une "
        "opportunité forte de positionner votre Académie comme un acteur de référence sur ces nouveaux champs "
        "d'expertise en Afrique.",
    )

    add_para(
        doc,
        "Nous pensons qu'un partenariat entre votre institution et TransferAI pourrait contribuer à renforcer "
        "ce positionnement autour de plusieurs axes complémentaires :",
        after=4,
        bold=True,
    )

    add_bullet(
        doc,
        "la co-construction de programmes spécialisés en intelligence artificielle, diplomatie numérique, "
        "data, cybersécurité et gouvernance technologique ;",
    )
    add_bullet(
        doc,
        "le développement de formations exécutives à destination des diplomates, cadres publics et "
        "responsables institutionnels ;",
    )
    add_bullet(
        doc,
        "la mise en place de certifications co-brandées à vocation régionale et internationale ;",
    )
    add_bullet(
        doc,
        "un déploiement progressif en Afrique de l'Ouest, appuyé sur notre présence à Abidjan et notre "
        "capacité d'exécution locale.",
    )

    add_para(
        doc,
        "Au-delà des contenus, notre proposition vise à inscrire ce partenariat dans une logique d'impact : "
        "renforcer l'attractivité de votre Académie, former des profils hybrides capables d'articuler "
        "technologie et affaires internationales, et créer une offre différenciante à forte valeur ajoutée "
        "pour les décideurs publics et les partenaires internationaux.",
    )

    add_para(
        doc,
        "Pour faciliter l'appréciation de notre démarche, ce courrier est accompagné d'une présentation de "
        "TransferAI ainsi que des éléments de proposition permettant de mieux visualiser notre approche, "
        "nos capacités et les premières pistes de collaboration envisageables.",
    )

    add_para(
        doc,
        "Je serais honoré de pouvoir échanger avec vous afin d'examiner les modalités concrètes de cette "
        "collaboration et d'identifier les formes de partenariat les plus pertinentes pour votre Académie.",
    )

    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(5)
    r = closing.add_run(
        "Je vous prie d'agréer, Madame l'Administratrice générale, l'expression de ma très haute considération."
    )
    format_run(r)

    sig_img = doc.add_paragraph()
    sig_img.paragraph_format.space_before = Pt(2)
    sig_img.paragraph_format.space_after = Pt(1)
    sig_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sig_img.add_run()
    r.add_picture(str(SIGNATURE_PATH), width=Cm(3.2))

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_after = Pt(0)
    r = sig_name.add_run("M. Marius AYORO")
    format_run(r, bold=True)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    r = sig_title.add_run("Directeur Développement & Partenariats")
    format_run(r)

    sig_org = doc.add_paragraph()
    sig_org.paragraph_format.space_after = Pt(0)
    r = sig_org.add_run("TransferAI")
    format_run(r)

    sig_contact = doc.add_paragraph()
    sig_contact.paragraph_format.space_after = Pt(0)
    r = sig_contact.add_run("www.transferai.ci | +225 07 16 57 39 90 | contact@transferai.ci")
    format_run(r)

    doc.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build_document()
