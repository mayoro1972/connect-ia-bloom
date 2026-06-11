from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUT_DIR = ROOT / "outputs/manual-20260610-transferai-letterhead/output"
OUT_PATH = OUT_DIR / "Papier_Entete_TransferAI_x_NettelecomCI_2026-06-10.docx"
SAMPLE_OUT_PATH = OUT_DIR / "Courrier_Entreprise_TransferAI_Test_2026-06-10.docx"
FIPME_OUT_PATH = OUT_DIR / "Courrier_FIPME_Serge_Jean_ANON_2026-06-10.docx"
LOGO_PATH = ROOT / "src/assets/logo-transferai-nettelecom.png"

BLUE = RGBColor(0x1E, 0x66, 0xB2)
ORANGE = RGBColor(0xF2, 0x8A, 0x1E)
INK = RGBColor(0x1F, 0x2D, 0x3D)
MUTED = RGBColor(0x6B, 0x78, 0x86)
LIGHT = RGBColor(0xE9, 0xEF, 0xF7)
FONT_BODY = "Aptos"
FONT_DISPLAY = "Aptos Display"


def set_cell_border(cell, *, bottom=None):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    if bottom:
        edge = OxmlElement("w:bottom")
        for key, value in bottom.items():
            edge.set(qn(f"w:{key}"), str(value))
        tc_borders.append(edge)


def set_paragraph_bottom_border(paragraph, color_hex: str, size: int = 12):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)


def format_run(run, *, font=FONT_BODY, size=11, bold=False, color=INK, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_text(
    doc: Document,
    text: str,
    *,
    size=11,
    bold=False,
    color=INK,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    after=6,
    italic=False,
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    format_run(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    format_run(run, size=10.5)
    return paragraph


def configure_page(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)
    return section


def build_header(section):
    header = section.header
    table = header.add_table(rows=2, cols=1, width=Cm(16.6))
    table.autofit = False

    top_cell = table.cell(0, 0)
    top_para = top_cell.paragraphs[0]
    top_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_para.paragraph_format.space_after = Pt(0)
    logo_run = top_para.add_run()
    logo_run.add_picture(str(LOGO_PATH), width=Cm(5.9))
    set_cell_border(
        top_cell,
        bottom={"val": "single", "sz": "12", "space": "0", "color": "D7E3F2"},
    )

    bottom_cell = table.cell(1, 0)
    descriptor = bottom_cell.paragraphs[0]
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(2)
    descriptor.paragraph_format.space_after = Pt(0)
    r1 = descriptor.add_run("TransferAI | Hub IA de Nettelecom CI")
    format_run(r1, size=9.5, bold=True, color=BLUE)

    descriptor2 = header.add_paragraph()
    descriptor2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor2.paragraph_format.space_before = Pt(1)
    descriptor2.paragraph_format.space_after = Pt(1)
    r2 = descriptor2.add_run(
        "Adoption concrète, responsable et performante de l'Intelligence Artificielle"
    )
    format_run(r2, size=8.2, color=MUTED)


def build_footer(section):
    footer = section.footer
    p1 = footer.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(
        "www.transferai.ci | contact@transferai.ci | WhatsApp UK : +44 7595 715838 | "
        "IA Assistant / WhatsApp CI : +225 07 16 57 39 90"
    )
    format_run(r1, size=7.8, color=MUTED)

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(
        "Riviera 3, carrefour Sainte Famille, Abidjan, Côte d'Ivoire | Coordination : Londres"
    )
    format_run(r2, size=7.6, color=MUTED)


def build_body(doc: Document):
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run("Abidjan, le [date]")
    format_run(meta_run, size=10.5, color=INK)

    for line in (
        "À l'attention de [Titre] [Nom]",
        "[Fonction]",
        "[Entreprise / Institution]",
        "[Ville, Pays]",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        format_run(r, size=10.8, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(10)
    subject_run = subject.add_run("Objet : [Objet du courrier]")
    format_run(subject_run, size=11, bold=True, color=ORANGE)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(6)
    salutation_run = salutation.add_run("Madame, Monsieur,")
    format_run(salutation_run, size=11)

    intro = (
        "TransferAI est une organisation spécialisée dans l’accompagnement des entreprises, "
        "institutions et organisations professionnelles dans l’adoption concrète, responsable "
        "et performante de l’Intelligence Artificielle. Basée à Londres avec une présence "
        "opérationnelle en Côte d’Ivoire, TransferAI aide les dirigeants et les équipes à "
        "transformer l’IA en gains mesurables de productivité, de qualité de service, de "
        "coordination et de compétitivité, grâce à des actions combinées de sensibilisation, "
        "formation, conseil stratégique, gouvernance et déploiement progressif de solutions "
        "adaptées aux réalités du terrain africain."
    )
    add_text(doc, intro, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)

    add_text(
        doc,
        "[Insérer ici le contenu détaillé du courrier : contexte, proposition de valeur, bénéfices "
        "attendus, pièces jointes, demande de rendez-vous.]",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        color=MUTED,
        italic=True,
        after=12,
    )

    add_text(doc, "Veuillez agréer, Madame, Monsieur, l'expression de ma considération distinguée.", after=10)

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_before = Pt(8)
    sig_name.paragraph_format.space_after = Pt(0)
    r_name = sig_name.add_run("Marius AYORO")
    format_run(r_name, size=11, bold=True)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    r_title = sig_title.add_run("Directeur Développement & Partenariats")
    format_run(r_title, size=10.5)

    sig_org = doc.add_paragraph()
    sig_org.paragraph_format.space_after = Pt(0)
    r_org = sig_org.add_run("TransferAI")
    format_run(r_org, size=10.5)


def build_sample_letter(doc: Document):
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run("Abidjan, le [date]")
    format_run(meta_run, size=10.2, color=INK)

    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(7)
    subject_run = subject.add_run(
        "Objet : Proposition d’échange sur l’adoption concrète de l’Intelligence Artificielle en entreprise"
    )
    format_run(subject_run, size=10.8, bold=True, color=ORANGE)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(5)
    salutation_run = salutation.add_run("Madame, Monsieur,")
    format_run(salutation_run, size=10.6)

    paragraphs = [
        "Je me permets de vous contacter au nom de TransferAI.",
        "TransferAI est une organisation spécialisée dans l’accompagnement des entreprises, institutions et organisations professionnelles dans l’adoption concrète, responsable et performante de l’Intelligence Artificielle. Basée à Londres avec une présence opérationnelle en Côte d’Ivoire, TransferAI aide les dirigeants et les équipes à transformer l’IA en gains mesurables de productivité, de qualité de service, de coordination et de compétitivité, grâce à des actions combinées de sensibilisation, formation, conseil stratégique, gouvernance et déploiement progressif de solutions adaptées aux réalités du terrain africain.",
        "Dans le cadre du lancement de nos activités en Côte d’Ivoire, nous souhaitons établir des échanges avec des entreprises désireuses de mieux comprendre comment l’Intelligence Artificielle peut créer de la valeur de manière concrète et mesurable dans leurs opérations, leurs fonctions support et leur développement stratégique.",
        "Notre approche consiste à aider les organisations à identifier les usages les plus pertinents, à renforcer les compétences de leurs équipes, à structurer un cadre de gouvernance adapté et à déployer, de façon progressive, des solutions utiles, réalistes et alignées sur leurs priorités métier.",
        "C’est dans cette perspective que nous souhaiterions solliciter un rendez-vous avec vous afin de vous présenter notre structure, nos offres de formation et d’accompagnement, ainsi que les opportunités de collaboration susceptibles de générer un retour sur investissement clair pour votre entreprise.",
        "Vous trouverez ci-joints un deck de présentation de TransferAI ainsi qu’un mini-catalogue de formations pouvant être adapté aux enjeux spécifiques de votre secteur et de vos équipes.",
        "Je vous remercie par avance pour l’attention portée à cette démarche et reste à votre entière disposition pour toute information complémentaire ou pour convenir d’un échange à votre meilleure convenance.",
        "Veuillez agréer, Madame, Monsieur, l’expression de ma considération distinguée.",
    ]
    for text in paragraphs:
        add_text(
            doc,
            text,
            size=10.2,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            after=5,
        )

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_before = Pt(10)
    sig_name.paragraph_format.space_after = Pt(0)
    r_name = sig_name.add_run("Marius AYORO")
    format_run(r_name, size=10.8, bold=True)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    r_title = sig_title.add_run("Directeur Développement & Partenariats")
    format_run(r_title, size=10.2)

    sig_org = doc.add_paragraph()
    sig_org.paragraph_format.space_after = Pt(0)
    r_org = sig_org.add_run("TransferAI")
    format_run(r_org, size=10.2)


def build_fipme_letter(doc: Document):
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.paragraph_format.space_after = Pt(8)
    meta_run = meta.add_run("Abidjan, le [date]")
    format_run(meta_run, size=10.2, color=INK)

    for line in (
        "À l'attention de Monsieur Serge Jean ANON",
        "Directeur",
        "Fédération Ivoirienne des Petites et Moyennes Entreprises (FIPME)",
        "Abidjan, Côte d'Ivoire",
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line)
        format_run(r, size=9.8, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)

    subject = doc.add_paragraph()
    subject.paragraph_format.space_after = Pt(5)
    subject_run = subject.add_run(
        "Objet : Proposition d’échange sur l’adoption concrète de l’Intelligence Artificielle au service des PME ivoiriennes"
    )
    format_run(subject_run, size=10, bold=True, color=ORANGE)

    salutation = doc.add_paragraph()
    salutation.paragraph_format.space_after = Pt(3)
    salutation_run = salutation.add_run("Monsieur le Directeur,")
    format_run(salutation_run, size=10)

    paragraphs = [
        "Je me permets de vous contacter au nom de TransferAI.",
        "TransferAI est une organisation spécialisée dans l’accompagnement des entreprises, institutions et organisations professionnelles dans l’adoption concrète, responsable et performante de l’Intelligence Artificielle. Basée à Londres avec une présence opérationnelle en Côte d’Ivoire, TransferAI aide les dirigeants et les équipes à transformer l’IA en gains mesurables de productivité, de qualité de service, de coordination et de compétitivité, grâce à des actions combinées de sensibilisation, formation, conseil stratégique, gouvernance et déploiement progressif de solutions adaptées aux réalités du terrain africain.",
        "Dans le cadre du lancement de nos activités en Côte d’Ivoire, nous souhaitons établir des échanges avec les acteurs qui contribuent directement à la dynamisation et à la compétitivité du tissu entrepreneurial ivoirien, en particulier au bénéfice des PME.",
        "À ce titre, la Fédération Ivoirienne des Petites et Moyennes Entreprises occupe une place stratégique. Nous sommes convaincus qu’une réflexion commune sur les usages concrets de l’Intelligence Artificielle pourrait aider les PME ivoiriennes à gagner en productivité, en organisation interne, en qualité de service et en montée en compétence.",
        "Notre approche consiste à identifier les usages les plus pertinents, renforcer les compétences des équipes, structurer un cadre de gouvernance adapté et déployer progressivement des solutions utiles et réalistes, alignées sur les priorités métier.",
        "C’est dans cette perspective que nous souhaiterions solliciter un rendez-vous avec vous afin de vous présenter TransferAI, notre approche d’accompagnement et les pistes de collaboration envisageables avec la FIPME au bénéfice de ses membres.",
        "Vous trouverez ci-joints un deck de présentation de TransferAI ainsi qu’un mini-catalogue de formations pouvant être adapté aux besoins des PME, des fédérations professionnelles et des structures d’appui à l’entrepreneuriat.",
        "Je vous remercie par avance pour l’attention portée à cette démarche et reste à votre entière disposition pour toute information complémentaire ou pour convenir d’un échange à votre meilleure convenance.",
        "Je vous prie d’agréer, Monsieur le Directeur, l’expression de ma considération distinguée.",
    ]

    for text in paragraphs:
        add_text(
            doc,
            text,
            size=9.5,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            after=2,
        )

    sig_name = doc.add_paragraph()
    sig_name.paragraph_format.space_before = Pt(4)
    sig_name.paragraph_format.space_after = Pt(0)
    r_name = sig_name.add_run("Marius AYORO")
    format_run(r_name, size=10.1, bold=True)

    sig_title = doc.add_paragraph()
    sig_title.paragraph_format.space_after = Pt(0)
    r_title = sig_title.add_run("Directeur Développement & Partenariats")
    format_run(r_title, size=9.6)

    sig_org = doc.add_paragraph()
    sig_org.paragraph_format.space_after = Pt(0)
    r_org = sig_org.add_run("TransferAI")
    format_run(r_org, size=9.6)


def make_document(body_builder, output_path: Path):
    doc = Document()
    configure_page(doc)
    build_header(doc.sections[0])
    build_footer(doc.sections[0])
    body_builder(doc)
    doc.save(str(output_path))


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_document(build_body, OUT_PATH)
    make_document(build_sample_letter, SAMPLE_OUT_PATH)
    make_document(build_fipme_letter, FIPME_OUT_PATH)
    print(OUT_PATH)
    print(SAMPLE_OUT_PATH)
    print(FIPME_OUT_PATH)


if __name__ == "__main__":
    main()
