from pathlib import Path

from docx import Document


SOURCE = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/output/"
    "Proposition_Professionnelle_TransferAI_ELTON_CI_revisee.docx"
)
OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/output/"
    "Proposition_Professionnelle_TransferAI_ELTON_CI_hub_ia.docx"
)


def main() -> None:
    doc = Document(str(SOURCE))

    doc.paragraphs[3].text = (
        "TransferAI Africa accompagne les entreprises africaines dans l’amélioration "
        "de leur performance grâce à l’intelligence artificielle, à la structuration "
        "des processus métier et à la montée en compétences des équipes. En Côte "
        "d’Ivoire, TransferAI se positionne comme le Hub IA de Nettelecom CI, avec "
        "pour mission de transformer les usages de l’IA en solutions concrètes, "
        "utiles et immédiatement activables pour les organisations."
    )

    doc.paragraphs[4].text = (
        "« En tant que Hub IA de Nettelecom CI, nous ne proposons pas de l’IA pour "
        "l’IA. Nous proposons des solutions concrètes de performance, de formation "
        "et de transformation adaptées aux réalités terrain. »"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    main()
