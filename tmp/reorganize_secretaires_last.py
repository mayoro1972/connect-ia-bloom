from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/output/"
    "Proposition_Professionnelle_TransferAI_ELTON_CI_hub_ia_v2.docx"
)
OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/output/"
    "Proposition_Professionnelle_TransferAI_ELTON_CI_hub_ia_v3.docx"
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style_name: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
    if style_name:
        new_para.style = style_name
    return new_para


def remove_matching_paragraphs(doc: Document, exact_texts: set[str]) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in exact_texts:
            element = paragraph._element
            element.getparent().remove(element)


def main() -> None:
    doc = Document(str(SOURCE))

    # Replace the current intro so the section starts with the core business use cases.
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Dans la perspective du lancement, au mois de juillet"):
            paragraph.text = (
                "Les formations ci-dessous répondent en priorité aux cas d’usage métier "
                "identifiés pour ELTON CI, en lien avec la direction, les opérations, la "
                "finance, l’audit, la relation client, l’optimisation des processus et la "
                "gouvernance."
            )
            break

    texts_to_remove = {
        "Ces fonctions occupent une place stratégique dans la circulation de l’information, la préparation des décisions et la qualité d’exécution administrative. Lorsqu’ils sont bien formés à l’IA, les secrétaires et assistants de direction deviennent un levier direct de productivité, de réactivité et de fiabilité pour l’ensemble de l’organisation.",
        "IA pour Secrétaires & Assistants de Direction : Structurer les comptes rendus, préparer les dossiers, accélérer les courriers et fiabiliser le suivi des échéances de direction.",
        "Rédiger et synthétiser les réunions de direction avec des comptes rendus, plans d’action et relances produits plus rapidement.",
        "Préparer les notes, courriers, présentations et dossiers décisionnels avec un meilleur niveau de qualité rédactionnelle et moins de temps de production.",
        "Suivre les agendas, échéances, validations et demandes transverses afin de réduire les oublis et de fluidifier la coordination entre directions.",
    }
    remove_matching_paragraphs(doc, texts_to_remove)

    anchor = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Protection des Données & Gouvernance IA : Maîtriser les bonnes pratiques de sécurité et conformité.":
            anchor = paragraph
            break

    if anchor is None:
        raise RuntimeError("Anchor paragraph not found")

    p = insert_paragraph_after(
        anchor,
        (
            "IA pour Secrétaires & Assistants de Direction : Renforcer l’accueil, la "
            "coordination administrative, la préparation des dossiers et la qualité de "
            "circulation de l’information."
        ),
        "List Number",
    )
    p = insert_paragraph_after(
        p,
        (
            "Bien que ce corps de métier ne constitue pas le coeur des cas d’usage "
            "opérationnels présentés dans cette proposition, cette formation demeure "
            "essentielle. La secrétaire reste la porte d’entrée de toute structure : elle "
            "oriente l’information, organise les priorités, prépare les interactions de "
            "direction et contribue directement à l’image, à la réactivité et à la qualité "
            "d’exécution de l’entreprise."
        ),
        "Normal",
    )
    p = insert_paragraph_after(
        p,
        "Dans la perspective du lancement prévu en juillet, cette montée en compétences permettra notamment de :",
        "Normal",
    )
    p = insert_paragraph_after(
        p,
        (
            "Rédiger et synthétiser les réunions de direction avec des comptes rendus, "
            "plans d’action et relances produits plus rapidement."
        ),
        "List Bullet",
    )
    p = insert_paragraph_after(
        p,
        (
            "Préparer les notes, courriers, présentations et dossiers décisionnels avec "
            "un meilleur niveau de qualité rédactionnelle et moins de temps de production."
        ),
        "List Bullet",
    )
    insert_paragraph_after(
        p,
        (
            "Suivre les agendas, échéances, validations et demandes transverses afin de "
            "réduire les oublis et de fluidifier la coordination entre directions."
        ),
        "List Bullet",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    main()
