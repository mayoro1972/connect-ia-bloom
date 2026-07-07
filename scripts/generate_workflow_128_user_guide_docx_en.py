from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "139_User_Guide_Workflow_128_International_Meeting_Report_Training_July_August_2026.docx"
)


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(document, text, subtitle=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.0)
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color="10263F")
    if subtitle:
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=0, after=14, line=1.0)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=11, color="666666")


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6 if level == 1 else 4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=15, bold=True, color="2E74B5")
    elif level == 2:
        set_font(r, size=12.5, bold=True, color="2E74B5")
    else:
        set_font(r, size=11.5, bold=True, color="1F4D78")


def add_body(document, text):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.15)
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_two_col_table(document, rows, widths=(2.0, 4.5), header=("Item", "Configuration")):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])
    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            set_paragraph_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="10263F")
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before=0, after=0, line=1.08)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0))
    document.add_paragraph()


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("Workflow 128 User Guide - TransferAI")
    set_font(r, size=9, color="777777")


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("TransferAI - Sensitive meeting report workflow")
    set_font(r, size=9, bold=True, color="777777")


def build_doc():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header(section)
    add_footer(section)

    add_title(
        document,
        "User Guide - Workflow 128",
        "International meeting report, local pseudonymization, and sanitized deliverables - July / August 2026 training session",
    )

    add_heading(document, "1. Purpose of this document", level=1)
    add_body(
        document,
        "This guide explains how to use and configure workflow 128 to process sensitive meeting recordings, produce structured meeting reports, and generate sanitized external deliverables."
    )
    add_body(
        document,
        "It is intended for the TransferAI training session use case in July / August 2026, in an international context involving multiple recording sources and confidentiality requirements."
    )

    add_heading(document, "2. Recommended use case", level=1)
    add_body(
        document,
        "The client records its scoping, follow-up, or debrief meetings through Zoom, Webex, Microsoft Teams, a dictaphone, a smartphone, or a meeting room device."
    )
    add_body(
        document,
        "The goal is to obtain quickly a usable meeting report, a client follow-up email, and a sanitized version, while protecting sensitive data before any AI call."
    )
    add_bullets(
        document,
        [
            "Scoping meeting: AI Training Programme Executive Secretariat - August 2026 session",
            "Recording source: Teams, Zoom, Webex, dictaphone, or smartphone",
            "Expected output: meeting report, client email, sanitized archive",
        ],
    )

    add_heading(document, "3. Workflow architecture", level=1)
    add_numbered(
        document,
        [
            "Recording intake",
            "Metadata normalization",
            "Audio retrieval or reception",
            "Audio splitting",
            "OpenAI transcription",
            "Local pseudonymization",
            "Sanitized report generation",
            "Human validation and final delivery",
        ],
    )

    add_heading(document, "4. Supported input modes", level=1)
    add_heading(document, "4.1 Manual upload", level=2)
    add_body(document, "Use this for .m4a files, .mp3 files, smartphone exports, and dictaphone exports.")
    add_heading(document, "4.2 Google Drive", level=2)
    add_body(document, "Use this when files are uploaded into a watched folder or when a Drive link is provided in the form.")
    add_heading(document, "4.3 Meeting platform webhook", level=2)
    add_body(
        document,
        "Use this when an upstream Zoom, Teams, or Webex workflow retrieves the recording and then calls the ingestion webhook of workflow 128."
    )

    add_heading(document, "5. Technical prerequisites", level=1)
    add_two_col_table(
        document,
        [
            ("OpenAI key", "Environment variable OPENAI_API_KEY"),
            ("Resend key", "Environment variable RESEND_API_KEY"),
            ("Report model", "gpt-5 recommended; otherwise gpt-4.1 or gpt-4o"),
            ("Google Drive", "Active OAuth credential for reading and archiving"),
            ("Audio splitting", "Service available at http://n8n-pxlk-audio-splitter-1:8000/split"),
            ("Webhook domain", "https://n8n-pxlk.srv1480638.hstgr.cloud/webhook"),
        ],
    )

    add_heading(document, "6. Step-by-step configuration", level=1)
    steps = [
        ("Step 1 - Import the workflow", "In n8n, open Workflows, choose Import from file, then select the JSON file for workflow 128."),
        ("Step 2 - Check the entry points", "Review the nodes Formulaire - Audio sensible, Google Drive - Nouveau fichier audio, and Webhook - Ingestion plateforme meeting."),
        ("Step 3 - Configure the form", "Fill in the fields for subject, participants, date, recording source, ingestion mode, final recipient, validator, audio language, output language, requested deliverable, sensitive entities to mask, and masking level."),
        ("Step 4 - Configure the source folder", "Check the watched folder ID in the node Google Drive - Nouveau fichier audio."),
        ("Step 5 - Configure the archive folder", "Check the archive folder ID in the node Archiver version sanitisee."),
        ("Step 6 - Configure Resend", "Validate the nodes Envoyer au validateur and Envoyer livrable final sanitise, as well as the authorized sender address."),
        ("Step 7 - Configure OpenAI", "Check the nodes OpenAI - Transcrire segment and OpenAI - Rapport structure sanitise."),
        ("Step 8 - Verify the validation webhooks", "Review the nodes Webhook - Approuver and Webhook - Rejeter, together with the production domain."),
        ("Step 9 - Run a pilot test", "Test with a short recording, approve the deliverable, then verify the archived HTML output."),
    ]
    add_two_col_table(document, steps)

    add_heading(document, "7. Recommended settings for July / August 2026", level=1)
    add_bullets(
        document,
        [
            "Requested deliverable: Both",
            "Masking level: Strict",
            "Output language: French",
            "Source folder: Recordings - AI Training - July August 2026",
            "Archive folder: Sanitized Archives - AI Training - July August 2026",
        ],
    )

    add_heading(document, "8. Meeting source configuration", level=1)
    add_heading(document, "8.1 Dictaphone", level=2)
    add_bullets(
        document,
        [
            "Simple mode: export the file and upload it through the form.",
            "Semi-automated mode: sync the files to a watched Google Drive folder.",
        ],
    )
    add_heading(document, "8.2 Zoom", level=2)
    add_bullets(
        document,
        [
            "Enable Cloud Recording in Zoom.",
            "Use an upstream workflow that receives the recording.completed event.",
            "Download the recording, then call the cr-intl-source-ingest webhook.",
        ],
    )
    add_heading(document, "8.3 Microsoft Teams", level=2)
    add_bullets(
        document,
        [
            "Retrieve recordings from OneDrive or SharePoint.",
            "Use Microsoft Graph or an upstream collection workflow.",
            "Then pass the binary file or the metadata to workflow 128.",
        ],
    )
    add_heading(document, "8.4 Webex", level=2)
    add_bullets(
        document,
        [
            "Enable Webex recordings.",
            "Use the Recordings API to retrieve the file.",
            "Forward the recording to the ingestion webhook of workflow 128.",
        ],
    )

    add_heading(document, "9. Expected ingestion webhook structure", level=1)
    add_body(
        document,
        "The cr-intl-source-ingest webhook should ideally receive: source_system, source_label, meeting_title, meeting_date, participants, recipient email, validator email, language, masking level, meeting reference, source object, and recording_url. When possible, send the audio binary directly."
    )

    add_heading(document, "10. Go-live checklist", level=1)
    add_bullets(
        document,
        [
            "The workflow imports without error.",
            "Google Drive is connected.",
            "Resend is connected.",
            "OpenAI is connected.",
            "The audio splitting service responds.",
            "The validator receives the preview.",
            "The HTML archive is created.",
            "No re-identified version is sent to the client.",
        ],
    )

    add_heading(document, "11. Final recommendation", level=1)
    add_body(
        document,
        "For the July / August 2026 training session, the recommended starting point is manual upload and Google Drive. Once the business process and report quality are validated, add dedicated ingestion workflows for Zoom, Teams, and Webex."
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
