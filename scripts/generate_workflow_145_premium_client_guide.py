from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MD_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "145_Guide_Premium_Client_Workflow_143_Traduction_Internationale_Securisee_06juillet2026.md"
)
DOCX_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "145_Guide_Premium_Client_Workflow_143_Traduction_Internationale_Securisee_06juillet2026.docx"
)


TITLE = "Guide premium client — Workflow 143"
SUBTITLE = (
    "Traduction internationale sécurisée, installation pas à pas, gouvernance des "
    "données sensibles et argumentaire de présentation"
)


SUMMARY_ROWS = [
    ("Nom du workflow", "Workflow 143 — traduction officielle internationale sécurisée"),
    ("Public cible", "Organisations internationales, sommets, ambassades, institutions, cabinets, ONG"),
    ("Objectif", "Traduire des documents importants sans exposer inutilement les données sensibles"),
    ("Décision clé", "Validation humaine obligatoire ou fortement recommandée selon la sensibilité du document"),
]

SECTIONS = [
    (
        "1. Résumé exécutif",
        "paragraphs",
        [
            "Le workflow 143 a été conçu pour permettre à TransferAI de traiter des traductions internationales sensibles dans un cadre beaucoup plus sûr qu’un flux classique envoyé en clair à un modèle d’IA.",
            "Il répond à un besoin très concret : traduire vite, bien et de manière professionnelle, tout en réduisant l’exposition des noms, numéros de dossier, identifiants, adresses e-mail, références de visa, données personnelles et informations confidentielles.",
            "Ce document a été pensé pour servir à la fois de support d’installation, de support de démonstration et de support commercial lors des visites clients.",
        ],
    ),
    (
        "2. Pourquoi ce workflow est important",
        "bullets",
        [
            "Parce que les documents de visa, les pièces d’identité, les actes, les notes verbales, les courriers officiels et certains comptes rendus contiennent souvent des données sensibles.",
            "Parce qu’un client institutionnel doit comprendre non seulement que l’IA produit un résultat utile, mais aussi que le traitement respecte une logique de confidentialité.",
            "Parce qu’en contexte international, la confiance, la conformité, la traçabilité et la validation humaine comptent autant que la qualité de traduction.",
        ],
    ),
    (
        "3. Définitions à expliquer à toute audience",
        "definitions",
        [
            (
                "PII",
                "PII signifie Personal Identifiable Information, c’est-à-dire les informations permettant d’identifier directement ou indirectement une personne. Exemples : nom complet, e-mail, numéro de téléphone, numéro de passeport, référence de dossier, date de naissance, adresse.",
            ),
            (
                "Anonymisation",
                "L’anonymisation consiste à retirer ou à transformer des informations de manière à ce qu’il ne soit plus possible de retrouver l’identité de la personne à partir des données traitées. Dans un workflow, cela revient à supprimer définitivement certaines données ou à les remplacer par une forme non réversible.",
            ),
            (
                "Pseudonymisation",
                "La pseudonymisation consiste à remplacer les données identifiantes par des jetons techniques, par exemple PERSON_001 ou PASSPORT_001. La différence avec l’anonymisation est qu’une table de correspondance peut exister localement pour rétablir les vraies valeurs à la fin du traitement, sans jamais les exposer au fournisseur d’IA.",
            ),
        ],
    ),
    (
        "4. Pourquoi appliquer PII, anonymisation et pseudonymisation",
        "bullets",
        [
            "Pour réduire le risque d’exposition de données sensibles pendant le traitement.",
            "Pour limiter ce que le modèle externe voit réellement.",
            "Pour rassurer les organisations internationales, les cabinets, les directions générales et les équipes conformité.",
            "Pour prouver que le workflow ne se contente pas de traduire, mais applique une gouvernance sérieuse des données.",
            "Pour rendre le service plus crédible lors des visites clients, démonstrations et discussions de cadrage.",
        ],
    ),
    (
        "5. Ce que le workflow 143 fait concrètement",
        "numbered",
        [
            "Reçoit une demande de traduction via webhook n8n.",
            "Analyse les métadonnées du document et classe son niveau de sensibilité.",
            "Détecte les éléments PII présents dans les champs structurés et dans le texte.",
            "Pseudonymise localement les éléments sensibles avant tout appel au modèle.",
            "Envoie au modèle uniquement la version pseudonymisée du contenu.",
            "Récupère la traduction et lance une relecture qualité également sur une version pseudonymisée.",
            "Ré-identifie localement les jetons après traitement.",
            "Assemble un document bilingue propre avec score de qualité.",
            "Route soit vers le client final, soit vers un validateur humain interne si le document est sensible ou si le score est insuffisant.",
            "Purge la table de correspondance sensible en fin de traitement.",
        ],
    ),
    (
        "6. Installation pas à pas dans n8n",
        "steps",
        [
            (
                "Étape 1 — Préparer l’environnement",
                "Vérifiez que votre instance n8n est accessible, que vous pouvez créer ou importer des workflows et que vous disposez d’un espace de test séparé avant la mise en production.",
            ),
            (
                "Étape 2 — Importer le JSON",
                "Dans n8n, ouvrez Workflows, choisissez l’option d’import, puis collez ou importez le fichier JSON du workflow 143 prêt à l’emploi.",
            ),
            (
                "Étape 3 — Configurer la clé OpenAI",
                "Créez ou mettez à jour la variable d’environnement OPENAI_API_KEY. Elle sera utilisée pour la traduction pseudonymisée et la relecture qualité pseudonymisée.",
            ),
            (
                "Étape 4 — Configurer la clé Resend",
                "Créez ou mettez à jour la variable d’environnement RESEND_API_KEY pour les envois e-mail sortants.",
            ),
            (
                "Étape 5 — Définir l’adresse de revue interne",
                "Ajoutez INTERNAL_REVIEW_EMAIL si vous souhaitez qu’un validateur reçoive automatiquement les documents nécessitant une vérification humaine.",
            ),
            (
                "Étape 6 — Définir l’adresse expéditeur",
                "Ajoutez OUTREACH_FROM_EMAIL pour maîtriser l’adresse utilisée lors de l’envoi des documents ou des demandes de validation.",
            ),
            (
                "Étape 7 — Vérifier les nœuds sensibles",
                "Contrôlez les nœuds de détection PII, de pseudonymisation, de traduction, de relecture, d’envoi validateur et d’envoi client afin de confirmer que les variables sont bien référencées et qu’aucune clé n’est inscrite en dur.",
            ),
            (
                "Étape 8 — Tester avec un document fictif",
                "Envoyez un document de test contenant des données fictives, par exemple un faux dossier de visa, pour vérifier la bonne détection des champs sensibles et le bon routage du workflow.",
            ),
            (
                "Étape 9 — Vérifier le comportement de validation",
                "Confirmez qu’un document déclaré très sensible, ou un document au score insuffisant, est bien envoyé au validateur humain et non directement au client final.",
            ),
            (
                "Étape 10 — Passer en production",
                "Une fois les tests validés, activez le workflow, documentez les rôles de validation et limitez l’accès aux exécutions n8n contenant encore temporairement des données réelles.",
            ),
        ],
    ),
    (
        "7. Paramètres attendus à l’entrée du workflow",
        "bullets",
        [
            "`texte` ou `content` : contenu à traduire.",
            "`langue_source` : langue d’origine, par exemple `fr`.",
            "`langue_cible` : langue de destination, par exemple `en`.",
            "`type_document` : visa, note verbale, courrier officiel, contrat, compte rendu, etc.",
            "`demandeur` : nom du demandeur, si nécessaire à l’orchestration interne.",
            "`organisation` : structure cliente ou institution concernée.",
            "`reference_dossier` : identifiant de suivi interne.",
            "`email_retour` : destinataire final du document.",
            "`email_validateur` : destinataire interne de la validation humaine.",
            "`contexte_mission` : sommet international, mission terrain, protocole, coopération, immigration, etc.",
        ],
    ),
    (
        "8. Ce que l’IA voit et ce qu’elle ne voit pas",
        "table",
        [
            ("Nom réel du demandeur", "Remplacé par un jeton comme PERSON_001"),
            ("Numéro de passeport", "Remplacé par un jeton comme PASSPORT_001"),
            ("E-mail réel", "Remplacé par un jeton comme EMAIL_001"),
            ("Référence de dossier", "Remplacée par un jeton comme REF_001"),
            ("Texte métier utile à la traduction", "Conservé si nécessaire au résultat"),
            ("Table de correspondance", "Conservée localement dans n8n puis purgée"),
        ],
    ),
    (
        "9. Cas d’usage recommandés",
        "bullets",
        [
            "Traduction de documents de visa et d’immigration.",
            "Traduction de lettres administratives ou institutionnelles.",
            "Traduction de notes verbales et de correspondances diplomatiques.",
            "Traduction de contrats, avenants, projets d’accord et documents juridiques à valider.",
            "Traduction de documents de sommet, ateliers régionaux, réunions bilatérales et événements internationaux.",
        ],
    ),
    (
        "10. Messages clés à utiliser lors des visites clients",
        "bullets",
        [
            "Notre workflow ne se limite pas à traduire : il applique une couche de protection des données avant même l’appel à l’IA.",
            "Le modèle ne voit pas les identifiants sensibles en clair lorsque la pseudonymisation est activée.",
            "La ré-identification n’a lieu que localement dans le workflow, sous votre contrôle opérationnel.",
            "Les documents les plus critiques peuvent être bloqués pour validation humaine avant diffusion.",
            "Cette approche aide l’organisation à concilier productivité, qualité linguistique et confidentialité.",
        ],
    ),
    (
        "11. Checklist de mise en production",
        "bullets",
        [
            "Vérifier que les clés API sont stockées hors workflow.",
            "Tester le flux avec des documents fictifs réalistes.",
            "Vérifier que les documents sensibles sont bien routés vers la validation humaine.",
            "Contrôler les journaux d’exécution et la politique de conservation.",
            "Définir qui approuve, qui corrige et qui diffuse les livrables finaux.",
            "Effectuer immédiatement une rotation des secrets si un ancien JSON contenait des clés en clair.",
        ],
    ),
    (
        "12. Référence du JSON prêt à coller dans n8n",
        "paragraphs",
        [
            "La dernière version locale du workflow prête à importer ou à coller dans n8n est : 143_n8n_Traduction_Officielle_Internationale_PII_Revamp_06juillet2026.json.",
            "Cette version est la version sécurisée à utiliser de préférence à l’ancien fichier de démonstration présent dans Downloads, lequel contenait des secrets en clair et ne mettait pas en œuvre de vraie pseudonymisation locale.",
        ],
    ),
]


def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=8, after=4, line=1.0)
    r = p.add_run(TITLE)
    set_font(r, size=24, bold=True, color="163556")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p2, before=0, after=14, line=1.15)
    r2 = p2.add_run(SUBTITLE)
    set_font(r2, size=12, color="4B5563")


def add_intro_band(doc):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    shade_cell(cell, "EEF4FB")
    p = cell.paragraphs[0]
    set_spacing(p, before=0, after=0, line=1.2)
    r = p.add_run(
        "Message clé : ce workflow protège les données sensibles avant l’appel au modèle, "
        "puis rétablit localement les informations utiles uniquement à la fin du traitement."
    )
    set_font(r, size=11, bold=True, color="163556")
    doc.add_paragraph()


def add_summary_table(doc):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.6)

    hdr = table.rows[0].cells
    hdr[0].text = "Élément"
    hdr[1].text = "Information essentielle"
    for cell in hdr:
        shade_cell(cell, "DCE7F3")
        for p in cell.paragraphs:
            set_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="163556")

    for left, right in SUMMARY_ROWS:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        shade_cell(cells[0], "EEF4FB")
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_spacing(p, before=0, after=0, line=1.2)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0), color="163556" if idx == 0 else None)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_spacing(p, before=18, after=8, line=1.0)
        r = p.add_run(text)
        set_font(r, size=16, bold=True, color="2E74B5")
    else:
        set_spacing(p, before=12, after=6, line=1.0)
        r = p.add_run(text)
        set_font(r, size=12.5, bold=True, color="1F4D78")


def add_paragraph(doc, text, muted=False):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6, line=1.25)
    r = p.add_run(text)
    set_font(r, size=11, color="555555" if muted else None)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, before=0, after=4, line=1.25)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_spacing(p, before=0, after=4, line=1.25)
        r = p.add_run(item)
        set_font(r, size=11)


def add_steps_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)

    hdr = table.rows[0].cells
    hdr[0].text = "Étape"
    hdr[1].text = "Action à réaliser"
    for cell in hdr:
        shade_cell(cell, "DCE7F3")
        for p in cell.paragraphs:
            set_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="163556")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        shade_cell(cells[0], "EEF4FB")
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_spacing(p, before=0, after=0, line=1.2)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0), color="163556" if idx == 0 else None)
    doc.add_paragraph()


def add_definition_block(doc, term, definition):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=3, line=1.2)
    r1 = p.add_run(f"{term} : ")
    set_font(r1, size=11, bold=True, color="163556")
    r2 = p.add_run(definition)
    set_font(r2, size=11)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.35)
    table.columns[1].width = Inches(4.15)

    hdr = table.rows[0].cells
    hdr[0].text = "Élément"
    hdr[1].text = "Traitement appliqué"
    for cell in hdr:
        shade_cell(cell, "DCE7F3")
        for p in cell.paragraphs:
            set_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="163556")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_spacing(p, before=0, after=0, line=1.2)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0))
    doc.add_paragraph()


def add_json_block(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6, line=1.15)
    r = p.add_run(text)
    set_font(r, name="Courier New", size=9.5)


def build_markdown():
    lines = [f"# {TITLE}", "", SUBTITLE, ""]
    lines.extend(
        [
            "| Élément | Information essentielle |",
            "|---|---|",
        ]
    )
    for left, right in SUMMARY_ROWS:
        lines.append(f"| {left} | {right} |")
    lines.append("")

    for title, kind, content in SECTIONS:
        lines.append(f"## {title}")
        lines.append("")
        if kind == "paragraphs":
            for paragraph in content:
                lines.append(paragraph)
                lines.append("")
        elif kind == "bullets":
            for item in content:
                lines.append(f"- {item}")
            lines.append("")
        elif kind == "numbered":
            for idx, item in enumerate(content, start=1):
                lines.append(f"{idx}. {item}")
            lines.append("")
        elif kind == "definitions":
            for term, definition in content:
                lines.append(f"**{term}** : {definition}")
                lines.append("")
        elif kind == "steps":
            for title_step, description in content:
                lines.append(f"### {title_step}")
                lines.append("")
                lines.append(description)
                lines.append("")
        elif kind == "table":
            lines.extend(["| Élément | Traitement appliqué |", "|---|---|"])
            for left, right in content:
                lines.append(f"| {left} | {right} |")
            lines.append("")

    json_example = """{
  "texte": "Nom : Fatou Diallo\\nPassport No: AB1234567\\nEmail: fatou.diallo@example.org\\nObjet : demande de visa officiel pour le sommet régional.",
  "langue_source": "fr",
  "langue_cible": "en",
  "type_document": "Demande de visa officiel",
  "demandeur": "Fatou Diallo",
  "organisation": "Organisation régionale fictive",
  "reference_dossier": "VISA-2026-0041",
  "email_retour": "translation.office@example.org",
  "email_validateur": "compliance@example.org",
  "contexte_mission": "Sommet international"
}"""
    lines.extend(["## Exemple de charge utile", "", "```json", json_example, "```", ""])
    MD_OUTPUT.write_text("\n".join(lines), encoding="utf-8")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc)
    add_intro_band(doc)
    add_summary_table(doc)

    for title, kind, content in SECTIONS:
        add_heading(doc, title, level=1)
        if kind == "paragraphs":
            for paragraph in content:
                add_paragraph(doc, paragraph)
        elif kind == "bullets":
            add_bullets(doc, content)
        elif kind == "numbered":
            add_numbered(doc, content)
        elif kind == "definitions":
            for term, definition in content:
                add_definition_block(doc, term, definition)
        elif kind == "steps":
            add_steps_table(doc, content)
        elif kind == "table":
            add_two_col_table(doc, content)

    add_heading(doc, "13. Exemple de charge utile", level=1)
    add_paragraph(
        doc,
        "Exemple de charge utile à envoyer au workflow pour un cas de test ou de démonstration.",
    )
    add_json_block(
        doc,
        '{\n'
        '  "texte": "Nom : Fatou Diallo\\nPassport No: AB1234567\\nEmail: fatou.diallo@example.org\\nObjet : demande de visa officiel pour le sommet régional.",\n'
        '  "langue_source": "fr",\n'
        '  "langue_cible": "en",\n'
        '  "type_document": "Demande de visa officiel",\n'
        '  "demandeur": "Fatou Diallo",\n'
        '  "organisation": "Organisation régionale fictive",\n'
        '  "reference_dossier": "VISA-2026-0041",\n'
        '  "email_retour": "translation.office@example.org",\n'
        '  "email_validateur": "compliance@example.org",\n'
        '  "contexte_mission": "Sommet international"\n'
        '}'
    )

    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_spacing(p, before=0, after=0, line=1.0)
        r = p.add_run("Guide premium client — Workflow 143 — TransferAI")
        set_font(r, size=9, color="777777")

    doc.save(DOCX_OUTPUT)


def main():
    build_markdown()
    build_docx()


if __name__ == "__main__":
    main()
