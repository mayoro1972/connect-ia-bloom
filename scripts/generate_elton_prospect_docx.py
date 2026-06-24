from __future__ import annotations

import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE_BLUE = RGBColor(18, 53, 91)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(107, 114, 128)
ACCENT = RGBColor(201, 91, 37)
LIGHT_FILL = "F4F6F8"
CALLOUT_FILL = "FCEDE4"
BORDER = "D9DDE3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color: str = BORDER, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pbdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pbdr")
        p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def ensure_style(document: Document, name: str, base: str, size: int, *, bold: bool = False, color: RGBColor | None = None):
    styles = document.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color
    return style


def configure_document(document: Document, header_text: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    ensure_style(document, "TA Title", "Title", 24, bold=True, color=TITLE_BLUE)
    ensure_style(document, "TA Subtitle", "Subtitle", 12, color=MUTED)
    ensure_style(document, "TA Heading 1", "Heading 1", 16, bold=True, color=TITLE_BLUE)
    ensure_style(document, "TA Heading 2", "Heading 2", 13, bold=True, color=TITLE_BLUE)
    ensure_style(document, "TA Heading 3", "Heading 3", 12, bold=True, color=ACCENT)
    ensure_style(document, "TA Small", "Normal", 9, color=MUTED)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run(header_text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_paragraph_bottom_border(header_p)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    add_page_number(footer_p)


def add_cover(document: Document, title: str, subtitle: str, callout: str | None) -> None:
    title_p = document.add_paragraph(style="TA Title")
    title_p.paragraph_format.space_after = Pt(8)
    title_p.add_run(title)

    subtitle_p = document.add_paragraph(style="TA Subtitle")
    subtitle_p.paragraph_format.space_after = Pt(12)
    subtitle_p.add_run(subtitle)

    meta = document.add_paragraph(style="TA Small")
    meta.add_run(f"Document de prospection | Généré le {datetime.now().strftime('%d/%m/%Y')}")

    if callout:
        table = document.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(callout)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = TITLE_BLUE
        set_cell_shading(cell, CALLOUT_FILL)

    document.add_paragraph()


def clean_inline(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("`", "")
    return text.strip()


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(clean_inline(text))
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(clean_inline(text))
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_callout(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean_inline(text))
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT
    set_cell_shading(cell, LIGHT_FILL)
    document.add_paragraph()


def add_paragraph_with_bold_label(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    if ":" in text and not text.startswith("http"):
        label, rest = text.split(":", 1)
        label_run = paragraph.add_run(clean_inline(label) + " :")
        label_run.bold = True
        label_run.font.color.rgb = TITLE_BLUE
        rest_value = clean_inline(rest)
        if rest_value:
            paragraph.add_run(" " + rest_value)
    else:
        paragraph.add_run(clean_inline(text))


def add_source_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(clean_inline(text))
    label_run.font.size = Pt(10.5)


def add_markdown_content(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    skip_first_title = True
    in_sources = False

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            document.add_paragraph(clean_inline(stripped[2:]), style="TA Heading 1")
            in_sources = clean_inline(stripped[2:]).lower() == "sources"
            continue

        if stripped.startswith("## "):
            document.add_paragraph(clean_inline(stripped[3:]), style="TA Heading 1")
            in_sources = clean_inline(stripped[3:]).lower() == "sources"
            continue

        if stripped.startswith("### "):
            document.add_paragraph(clean_inline(stripped[4:]), style="TA Heading 2")
            continue

        if stripped.startswith("#### "):
            document.add_paragraph(clean_inline(stripped[5:]), style="TA Heading 3")
            continue

        if stripped.startswith("- "):
            content = stripped[2:].strip()
            if in_sources:
                add_source_line(document, "• " + content)
            else:
                add_bullet(document, content)
            continue

        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            add_number(document, number_match.group(1))
            continue

        if stripped.lower().startswith("message clé") or stripped.lower().startswith("message cle"):
            add_callout(document, stripped)
            continue

        add_paragraph_with_bold_label(document, stripped)


def build_docx(markdown_path: Path, output_path: Path, title: str, subtitle: str, header_text: str, callout: str | None) -> None:
    document = Document()
    configure_document(document, header_text)
    add_cover(document, title, subtitle, callout)
    add_markdown_content(document, markdown_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) < 6:
        print(
            "Usage: python generate_elton_prospect_docx.py <input.md> <output.docx> <title> <subtitle> <header_text> [callout]"
        )
        return 1

    markdown_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()
    title = sys.argv[3]
    subtitle = sys.argv[4]
    header_text = sys.argv[5]
    callout = sys.argv[6] if len(sys.argv) > 6 else None

    if not markdown_path.exists():
        print(f"Input file not found: {markdown_path}")
        return 1

    build_docx(markdown_path, output_path, title, subtitle, header_text, callout)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
