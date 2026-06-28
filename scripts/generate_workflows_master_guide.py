"""
Generate Master Guide — All TransferAI Workflows
Covers all 12 active workflows with purpose, nodes, triggers, flow, and interactions
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

OUTPUT_DIR = "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin"


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_title(doc, text):
    h = doc.add_heading(text, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return h


def h1(doc, text):
    h = doc.add_heading(text, 1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return h


def h2(doc, text):
    return doc.add_heading(text, 2)


def h3(doc, text):
    return doc.add_heading(text, 3)


def body(doc, text):
    return doc.add_paragraph(text)


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_col_width(table, 0, 5)
    set_col_width(table, 1, 11)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
    doc.add_paragraph()
    return table


def separator(doc):
    doc.add_paragraph("─" * 90)


# ─────────────────────────────────────────────────────────────────────────────
def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER ────────────────────────────────────────────────────────────────
    add_title(doc, "TransferAI — Guide Maître des Workflows n8n")
    doc.add_paragraph(
        f"Pipeline de Prospection Automatisé | Version 1.0 | {datetime.today().strftime('%d/%m/%Y')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Ce document décrit l'ensemble des workflows actifs dans l'environnement n8n TransferAI, "
        "leur rôle, leurs déclencheurs, la séquence de traitement et leurs interactions."
    )
    separator(doc)

    # ── TABLE DES MATIÈRES (sommaire manuel) ─────────────────────────────────
    h1(doc, "Sommaire")
    toc = [
        "1. Architecture globale du pipeline",
        "2. Workflow 1 — Admin Prospection Controller V1",
        "3. Workflow 2 & 3 — V3 CRM Enhanced [FINAL] et [V20-FIXED]",
        "4. Workflow 4 — Post-Audit Expert Routing V6 MVP",
        "5. Workflow 5 — Post-Audit Expert Routing V2 (legacy)",
        "6. Workflow 6 — Google Forms Social Lead Sequence V2",
        "7. Workflow 7 — Multi-Prospect V4 Batch Orchestrator",
        "8. Workflow 8 — Zoho Reply Intelligence V1",
        "9. Workflow 9 & 10 — Chatwoot Auto Reply V5.5 / V5.5.2",
        "10. Workflow 11 — V3 Follow-Up Sequence",
        "11. Workflow 12 — Calendly Meeting Booked Webhook V1",
        "12. Points de vigilance et clarifications",
        "13. Schéma d'interaction entre workflows",
    ]
    for item in toc:
        bullet(doc, item)

    doc.add_page_break()

    # ── 1. ARCHITECTURE GLOBALE ───────────────────────────────────────────────
    h1(doc, "1. Architecture globale du pipeline TransferAI")
    body(doc,
        "Le pipeline TransferAI est composé de 12 workflows n8n interconnectés couvrant l'ensemble "
        "du cycle de vie prospect : découverte → enrichissement → génération de pack → envoi → "
        "suivi → audit post-soumission → notification expert → relance."
    )

    h2(doc, "Phases du pipeline")
    phases = [
        ("Phase 1 — Découverte & Enrichissement",
         "V4 Batch Orchestrator identifie les prospects éligibles depuis Supabase/Airtable/Google Sheets. "
         "V3 récupère et analyse les signaux publics du prospect (site web, secteur, signaux IA)."),
        ("Phase 2 — Génération du Pack",
         "V3 génère via GPT-4 : Executive Letter personnalisée, Mini-Catalogue sectoriel, "
         "Deck Brief, Formulaire d'audit adapté. Le pack est stocké dans ai_prospecting_packs."),
        ("Phase 3 — Validation & Envoi",
         "Admin Controller orchestre la validation humaine (approve/reject). "
         "V3 Follow-Up assure les relances automatiques J+1, J+3, J+7."),
        ("Phase 4 — Audit Prospect",
         "Le prospect remplit le formulaire d'audit (transferai.ci/questionnaire-audit). "
         "La Edge Function resolve-invitation gère l'accès et la personnalisation."),
        ("Phase 5 — Traitement Post-Audit",
         "V6 détecte la soumission, génère une fiche pré-RDV, route vers l'expert, "
         "envoie les notifications email et met à jour le CRM."),
        ("Phase 6 — Leads entrants & Chat",
         "Google Forms V2 traite les leads issus des réseaux sociaux. "
         "Chatwoot V5.5/V5.5.2 répond automatiquement aux messages entrants. "
         "Zoho Reply Intelligence classe les réponses emails."),
        ("Phase 7 — RDV",
         "Calendly Webhook met à jour le CRM et notifie l'équipe lors d'une prise de RDV."),
    ]
    for phase, desc in phases:
        h3(doc, phase)
        body(doc, desc)

    doc.add_page_break()

    # ── 2. ADMIN CONTROLLER ───────────────────────────────────────────────────
    h1(doc, "2. Workflow 1 — TransferAI Admin Prospection Controller V1")
    info_table(doc, [
        ("Nom complet", "TransferAI Admin Prospection Controller V1 [CONFIGURED]"),
        ("Noeuds", "77"),
        ("Déclencheurs", "Webhook HTTP (Admin Controller Webhook) + Manual Trigger"),
        ("Rôle", "Orchestre toutes les actions admin : approve, reject, send, re-run, health check"),
        ("Statut", "Actif — noeud central de contrôle"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "Ce workflow est le cerveau administratif du pipeline. Il reçoit toutes les commandes "
        "de validation via un webhook sécurisé (Authorization header requis) et les dispatche "
        "vers les workflows appropriés (V3, V6, etc.)."
    )

    h2(doc, "Actions supportées")
    actions = [
        ("approve", "Approuver un pack prospect → déclenche l'envoi de l'email"),
        ("reject", "Rejeter un pack → met à jour le statut dans Supabase"),
        ("send", "Forcer l'envoi d'un pack déjà approuvé"),
        ("re-run", "Relancer la génération V3 pour un prospect"),
        ("health_check", "Vérifier l'état du pipeline"),
        ("batch_run", "Déclencher V4 Batch Orchestrator"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Action"
    r.cells[1].text = "Description"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for action, desc in actions:
        row = t.add_row()
        row.cells[0].text = action
        row.cells[1].text = desc
    doc.add_paragraph()

    h2(doc, "Flux de traitement")
    steps = [
        "Réception de la requête HTTP (webhook ou manual)",
        "Validation du header Authorization",
        "Validation de l'action demandée",
        "Routage vers le sous-workflow correspondant",
        "Exécution via Call Legacy Approve/Reject/Send Webhook",
        "Retour de la réponse JSON au demandeur",
    ]
    for s in steps:
        numbered(doc, s)

    h2(doc, "Sécurité")
    body(doc, "Chaque requête doit contenir un header Authorization valide. Sans ce header, "
         "le workflow retourne une réponse 401 Unauthorized et s'arrête.")

    doc.add_page_break()

    # ── 3. V3 ────────────────────────────────────────────────────────────────
    h1(doc, "3. Workflows 2 & 3 — V3 CRM Enhanced [FINAL] et [V20-FIXED]")
    info_table(doc, [
        ("Noms", "TransferAI Prospecting V3 CRM Enhanced [FINAL] (51 nœuds)\nTransferAI Prospecting V3 CRM Enhanced [V20-FIXED] (52 nœuds)"),
        ("Déclencheurs", "Manual Trigger + Execute Workflow Trigger (appelé par Admin Controller ou V4)"),
        ("Rôle", "Analyser un prospect et générer l'intégralité du pack de prospection"),
        ("Statut", "FINAL = version principale | V20-FIXED = variante de secours"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "V3 est le workflow de génération de contenu. Pour chaque prospect, il récupère les "
        "signaux publics (site web, LinkedIn, secteur), les analyse avec GPT-4, et génère "
        "un pack complet personnalisé : Executive Letter, Mini-Catalogue, Deck Brief, "
        "Formulaire d'audit, et stocke le tout dans Supabase."
    )

    h2(doc, "Flux détaillé")
    v3_steps = [
        ("Set Target", "Définit le prospect cible : URL, nom, email, commercial_priority_default"),
        ("Build Source URLs", "Construit les URLs à scraper (site, LinkedIn, Crunchbase...)"),
        ("Fetch Public Page 1-5", "Scrape jusqu'à 5 pages web du prospect (timeout géré)"),
        ("Normalize Public Signals", "Normalise et fusionne les données scrapées"),
        ("Sanitize Prospect Data For LLM", "Anonymise les données (ORG_TARGET) avant envoi au LLM"),
        ("Call OpenAI Pre-Audit", "GPT-4 analyse le secteur, détecte les signaux IA, identifie les besoins"),
        ("Call OpenAI Problems Solutions", "GPT-4 génère les problèmes probables et solutions recommandées"),
        ("Call OpenAI ROI", "GPT-4 calcule l'hypothèse ROI et les gains attendus"),
        ("Assemble Prospect Context", "Fusionne tous les outputs en un contexte unifié"),
        ("Generate Executive Letter", "GPT-4 génère la lettre executive personnalisée"),
        ("Generate Tailored Catalogue", "GPT-4 génère le mini-catalogue sectoriel"),
        ("Generate Tailored Audit Form", "GPT-4 génère le formulaire d'audit adapté"),
        ("Generate Deck Brief", "GPT-4 génère le brief pour le deck de présentation"),
        ("Assemble Prospect Pack", "Assemble le pack final avec tous les éléments générés"),
        ("Upsert To Supabase", "Stocke ou met à jour le pack dans ai_prospecting_packs"),
        ("Approval Webhook", "Notifie l'Admin Controller que le pack est prêt pour validation"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Nœud"
    r.cells[1].text = "Action"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for node, action in v3_steps:
        row = t.add_row()
        row.cells[0].text = node
        row.cells[1].text = action
    doc.add_paragraph()

    h2(doc, "Correction importante appliquée (session 28/06/2026)")
    body(doc, "Dans le nœud 'Assemble Prospect Pack', la propagation du commercial_priority_tier "
         "a été corrigée :")
    code(doc, "// AVANT : commercial_priority_tier: ctx.commercial_priority_tier || ''")
    code(doc, "// APRÈS : commercial_priority_tier: ctx.commercial_priority_tier || ctx.commercial_priority_default || ''")
    body(doc, "Cette valeur (tier1/tier2/tier3) est essentielle pour que V6 détermine la priorité "
         "de routage et déclenche l'envoi d'email expert.")

    h2(doc, "Données produites dans ai_prospecting_packs")
    fields = [
        "pack_id, organization_name, target_email",
        "executive_letter, tailored_catalogue, tailored_audit_form",
        "deck_brief (JSON structuré)",
        "audit_form_url (lien direct vers le formulaire d'audit)",
        "commercial_priority_tier (tier1/tier2/tier3)",
        "probable_needs, probable_problems, probable_quick_wins",
        "recommended_offer, recommended_use_case",
        "roi_hypothesis, expected_time_savings",
    ]
    for f in fields:
        bullet(doc, f)

    doc.add_page_break()

    # ── 4. V6 ────────────────────────────────────────────────────────────────
    h1(doc, "4. Workflow 4 — Post-Audit Expert Routing V6 MVP")
    info_table(doc, [
        ("Nom complet", "TransferAI Post-Audit Expert Routing V6 MVP"),
        ("Noeuds", "37"),
        ("Déclencheurs", "Schedule toutes les 30 min + Webhook 'Audit Completed' + Manual Trigger"),
        ("Rôle", "Traiter les formulaires d'audit soumis et router vers l'expert avec notifications"),
        ("Statut", "Actif — VERSION EN PRODUCTION"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "V6 est le workflow post-audit. Il détecte les nouvelles soumissions du formulaire d'audit, "
        "génère une fiche pré-RDV avec GPT-4, détermine la priorité de routage, et envoie "
        "deux emails de notification à l'équipe TransferAI."
    )

    h2(doc, "Flux détaillé")
    v6_steps = [
        ("Reconciliation Schedule Trigger", "S'exécute toutes les 30 minutes"),
        ("Fetch Recent Completed Responses", "Récupère les form_responses is_completed=true des 30 dernières min"),
        ("Extract Next Completed Response", "Traite une réponse à la fois"),
        ("Normalize Post-Audit Request", "Normalise les données de la requête"),
        ("Get Pack Row / Extract Pack Row", "Récupère le pack correspondant dans ai_prospecting_packs"),
        ("If Pack Found", "Vérifie que le pack existe ($json.found is true)"),
        ("Get Latest Form Response", "Récupère les données du formulaire soumis"),
        ("Extract Latest Form Response", "Extrait form_data, completion_percentage, audit_completed"),
        ("If Audit Completed", "Vérifie que l'audit est bien complété ($json.audit_completed is true)"),
        ("Get Invitation Row / Extract Invitation Row", "Récupère l'invitation liée"),
        ("Get Prospect Target Row / Extract Prospect Target Row", "Récupère le prospect dans prospect_targets"),
        ("If Already Post-Audit Processed", "Évite le double traitement ($json.already_processed is true → skip)"),
        ("Build Expert Routing", "Calcule routing_priority (high/medium/normal) selon tier et completion"),
        ("Build Post-Audit CRM Context", "Assemble le contexte complet pour le CRM et les emails"),
        ("Generate Internal Pre-RDV Brief", "GPT-4 génère la fiche pré-RDV complète"),
        ("Send Internal Brief Email", "Envoie la fiche pré-RDV via Resend à contact@transferai.ci"),
        ("If Expert Notification Enabled", "Envoie alerte expert si routing_priority = high OU medium"),
        ("Send Expert Notification Email", "Email [PRIORITÉ HAUTE] à l'expert assigné"),
        ("Upsert Follow-Up Tracking", "Met à jour le suivi dans follow_up_tracking"),
        ("Patch Pack Post-Audit Status", "Met à jour le statut du pack dans ai_prospecting_packs"),
        ("Update Prospect Target Post-Audit", "Met à jour prospect_targets avec résultat audit"),
        ("Sync to Google Sheets Dashboard", "Synchronise avec le tableau de bord Google Sheets"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Nœud"
    r.cells[1].text = "Action"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for node, action in v6_steps:
        row = t.add_row()
        row.cells[0].text = node
        row.cells[1].text = action
    doc.add_paragraph()

    h2(doc, "Logique de priorité")
    body(doc, "Le nœud 'Build Expert Routing' calcule routing_priority :")
    priority_rules = [
        "completion_percentage >= 95 → HIGH",
        "maturity_label contient 'avance' ou 'advanced' → HIGH",
        "commercial_priority_tier = 'tier1' ou 'A' → HIGH",
        "commercial_priority_tier = 'tier2' ou 'B' → MEDIUM (si pas déjà HIGH)",
        "Sinon → NORMAL",
    ]
    for r in priority_rules:
        bullet(doc, r)
    body(doc, "Email expert envoyé si routing_priority = HIGH ou MEDIUM.")

    h2(doc, "Emails générés")
    body(doc, "Email 1 — [PRIORITÉ HAUTE] Nouveau dossier post-audit à traiter :")
    bullet(doc, "Organisation, contact, secteur, maturité, service recommandé, pack_id, lien Calendly")
    body(doc, "Email 2 — Fiche pré-RDV post-audit :")
    bullet(doc, "Fiche GPT-4 : identification, synthèse, maturité IA, cas d'usage prioritaires, quick wins, contraintes, orientation TransferAI, prochaine étape, documents à consulter")

    doc.add_page_break()

    # ── 5. V2 LEGACY ─────────────────────────────────────────────────────────
    h1(doc, "5. Workflow 5 — Post-Audit Expert Routing V2 (Legacy)")
    info_table(doc, [
        ("Nom complet", "TransferAI Post-Audit Expert Routing V2"),
        ("Noeuds", "33"),
        ("Déclencheurs", "Schedule + Webhook + Manual Trigger (mêmes que V6)"),
        ("Rôle", "Ancienne version du routing post-audit"),
        ("Statut", "⚠️ ACTIF — mais devrait être DÉSACTIVÉ car V6 le remplace"),
    ])

    h2(doc, "⚠️ Alerte — Double exécution")
    body(doc,
        "V2 et V6 ont les mêmes déclencheurs (schedule 30 min + webhook). "
        "Si les deux sont actifs, chaque soumission de formulaire sera traitée DEUX FOIS, "
        "générant des emails en double et des mises à jour CRM contradictoires."
    )
    body(doc, "ACTION REQUISE : Désactiver V2 dans n8n (toggle Active → Off).")

    h2(doc, "Différences V2 vs V6")
    diff = [
        ("Noeuds", "33 (V2)", "37 (V6)"),
        ("Fiche pré-RDV", "Basique", "GPT-4 complète (V6)"),
        ("Conditions If", "Ancien format", "Nouveau format v2"),
        ("Google Sheets sync", "Absent", "Présent"),
        ("Support tier1/tier2", "Absent", "Présent"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Critère"
    r.cells[1].text = "V2 (legacy)"
    r.cells[2].text = "V6 (production)"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for label, v2, v6 in diff:
        row = t.add_row()
        row.cells[0].text = label
        row.cells[1].text = v2
        row.cells[2].text = v6
    doc.add_paragraph()

    doc.add_page_break()

    # ── 6. GOOGLE FORMS ───────────────────────────────────────────────────────
    h1(doc, "6. Workflow 6 — Google Forms Social Lead Sequence V2")
    info_table(doc, [
        ("Nom complet", "TransferAI Google Forms Social Lead Sequence V2"),
        ("Noeuds", "19 (Clean Importable) / 30 (version complète)"),
        ("Déclencheurs", "Webhook Google Forms + Schedule horaire"),
        ("Rôle", "Traiter les leads entrants depuis les formulaires Google (réseaux sociaux)"),
        ("Statut", "Actif"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "Ce workflow capture les leads générés via les campagnes de réseaux sociaux "
        "(LinkedIn, Facebook, etc.) qui utilisent Google Forms. Il normalise les données, "
        "insère le lead dans le CRM Supabase, et déclenche une séquence d'emails de suivi."
    )

    h2(doc, "Flux de traitement")
    gf_steps = [
        "Réception du webhook Google Forms (nouveau lead)",
        "Normalisation des données du lead",
        "Insertion dans Supabase (table social_prospects ou contact_requests)",
        "Vérification si le lead est prêt pour l'envoi",
        "Envoi du premier email de la séquence sociale",
        "Planification des follow-ups suivants",
    ]
    for s in gf_steps:
        numbered(doc, s)

    h2(doc, "Schedule horaire")
    body(doc, "En plus du webhook, un schedule horaire vérifie les leads en attente de suivi "
         "et envoie les emails de relance selon la séquence définie.")

    doc.add_page_break()

    # ── 7. V4 BATCH ───────────────────────────────────────────────────────────
    h1(doc, "7. Workflow 7 — Multi-Prospect V4 Batch Orchestrator")
    info_table(doc, [
        ("Nom complet", "TransferAI Prospecting Multi-Prospect V4 Batch Orchestrator"),
        ("Noeuds", "23"),
        ("Déclencheurs", "Daily Schedule Trigger + Manual Trigger"),
        ("Rôle", "Orchestrer le traitement en batch de multiples prospects vers V3"),
        ("Statut", "Actif — ⚠️ 53 prospects tier1 jamais contactés"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "V4 est l'orchestrateur batch. Chaque jour, il identifie les prospects éligibles "
        "depuis plusieurs sources (Supabase, Airtable, Google Sheets), filtre ceux qui n'ont "
        "pas encore été traités, et les envoie un par un vers V3 pour génération de pack."
    )

    h2(doc, "Sources de prospects")
    sources = [
        ("Supabase", "Table prospect_targets — source principale avec scoring et tiering"),
        ("Airtable", "Base de données prospects externes"),
        ("Google Sheets CSV", "Import manuel de listes de prospects"),
    ]
    for source, desc in sources:
        h3(doc, source)
        body(doc, desc)

    h2(doc, "Logique d'éligibilité")
    body(doc, "Un prospect est éligible si :")
    eligibility = [
        "status ≠ 'sent', 'completed', 'rejected', 'do_not_contact'",
        "paused = false",
        "do_not_contact = false",
        "Quota journalier non atteint (Set Batch Config)",
    ]
    for e in eligibility:
        bullet(doc, e)

    h2(doc, "⚠️ Note importante")
    body(doc, "53 prospects tier1 dans Supabase n'ont jamais été contactés. "
         "V4 Batch Orchestrator doit être déclenché manuellement ou son schedule activé "
         "pour lancer la prospection sur ces cibles prioritaires.")

    doc.add_page_break()

    # ── 8. ZOHO REPLY ─────────────────────────────────────────────────────────
    h1(doc, "8. Workflow 8 — Zoho Reply Intelligence V1")
    info_table(doc, [
        ("Nom complet", "TransferAI Zoho Reply Intelligence V1"),
        ("Noeuds", "15"),
        ("Déclencheurs", "Schedule toutes les 15 minutes"),
        ("Rôle", "Classifier les réponses emails reçues dans Zoho Mail"),
        ("Statut", "Actif"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "Ce workflow surveille la boîte email Zoho Mail (contact@transferai.ci) toutes les "
        "15 minutes et classifie les nouvelles réponses avec GPT-4. Il détecte si une réponse "
        "nécessite une intervention humaine ou peut être traitée automatiquement."
    )

    h2(doc, "Flux de traitement")
    zoho_steps = [
        ("Read Unread Zoho Emails", "Récupère les emails non lus depuis Zoho Mail"),
        ("Parse Incoming Reply", "Extrait le corps, l'expéditeur, le sujet"),
        ("Find Prospect By Email", "Cherche le prospect correspondant dans Supabase"),
        ("Load Prospect Context", "Charge le contexte du prospect (pack, historique)"),
        ("OpenAI Classify Reply", "GPT-4 classe la réponse : intérêt positif / négatif / désinscription / hors sujet"),
        ("Parse AI Classification", "Extrait la classification et le niveau de confiance"),
        ("IF Needs Human Handoff", "Si réponse complexe → alerte équipe"),
        ("Log to Google Sheets", "Enregistre la réponse dans le dashboard"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Nœud"
    r.cells[1].text = "Action"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for node, action in zoho_steps:
        row = t.add_row()
        row.cells[0].text = node
        row.cells[1].text = action
    doc.add_paragraph()

    doc.add_page_break()

    # ── 9. CHATWOOT ───────────────────────────────────────────────────────────
    h1(doc, "9. Workflows 9 & 10 — Chatwoot Auto Reply V5.5 / V5.5.2")
    info_table(doc, [
        ("Noms", "TransferAI Chatwoot Auto Reply V5.5 Calendly\nTransferAI Chatwoot Auto Reply V5.5.2 Calendly"),
        ("Noeuds", "12 chacun"),
        ("Déclencheurs", "Chatwoot Webhook (nouveau message entrant)"),
        ("Rôle", "Répondre automatiquement aux messages entrants dans Chatwoot avec GPT-4"),
        ("Statut", "⚠️ Les DEUX sont actifs — risque de double réponse"),
    ])

    h2(doc, "⚠️ Alerte — Double webhook")
    body(doc,
        "V5.5 et V5.5.2 ont le même déclencheur Chatwoot. Si les deux sont actifs sur le "
        "même compte Chatwoot, chaque message reçu génèrera DEUX réponses automatiques. "
        "V5.5.2 est la version corrigée — V5.5 doit être désactivé."
    )
    body(doc, "ACTION REQUISE : Vérifier si V5.5 et V5.5.2 pointent vers des comptes Chatwoot "
         "différents. Si même compte → désactiver V5.5.")

    h2(doc, "Flux de traitement (commun aux deux)")
    chatwoot_steps = [
        ("Chatwoot Webhook", "Reçoit les nouveaux messages clients"),
        ("Filter Chatwoot Event", "Filtre les événements pertinents (message_created, agent non-bot)"),
        ("Get Conversation History", "Récupère l'historique de la conversation"),
        ("Build OpenAI Reply Payload", "Prépare le contexte pour GPT-4 (historique + prompt TransferAI)"),
        ("Call OpenAI Reply", "GPT-4 génère une réponse adaptée"),
        ("Format AI Reply", "Formate la réponse (ton, longueur, Calendly si besoin)"),
        ("Send Reply to Chatwoot", "Poste la réponse dans la conversation Chatwoot"),
        ("Build Qualification Payload", "Prépare les données pour qualifier le lead"),
        ("Call OpenAI Qualification", "GPT-4 évalue le niveau d'intérêt et la maturité"),
        ("Parse Qualification + Compute Labels", "Extrait score et labels (chaud/tiède/froid)"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Nœud"
    r.cells[1].text = "Action"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for node, action in chatwoot_steps:
        row = t.add_row()
        row.cells[0].text = node
        row.cells[1].text = action
    doc.add_paragraph()

    doc.add_page_break()

    # ── 10. V3 FOLLOW-UP ─────────────────────────────────────────────────────
    h1(doc, "10. Workflow 11 — V3 Follow-Up Sequence")
    info_table(doc, [
        ("Nom complet", "TransferAI Prospecting V3 Follow-Up Sequence [V1]"),
        ("Noeuds", "12"),
        ("Déclencheurs", "Daily V3 Follow-Up Schedule + Manual Trigger"),
        ("Rôle", "Envoyer les emails de relance aux prospects V3 (J+1, J+3, J+7)"),
        ("Statut", "Actif"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "Après l'envoi du pack initial par V3, ce workflow assure les relances automatiques "
        "aux prospects qui n'ont pas répondu. Il respecte des délais configurables et "
        "s'arrête dès qu'une réponse positive est détectée."
    )

    h2(doc, "Séquence de relance")
    followup_steps = [
        ("J+1", "Premier suivi — confirmation de réception + invitation formulaire audit"),
        ("J+3", "Deuxième suivi — partage d'un cas d'usage concret du secteur"),
        ("J+7", "Troisième suivi — dernière tentative + offre de RDV direct"),
    ]
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    r = t.add_row()
    r.cells[0].text = "Délai"
    r.cells[1].text = "Email envoyé"
    for cell in r.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for delay, email in followup_steps:
        row = t.add_row()
        row.cells[0].text = delay
        row.cells[1].text = email
    doc.add_paragraph()

    h2(doc, "Flux de traitement")
    fu_steps = [
        ("Fetch V3 Sent Prospects", "Récupère les prospects avec status='sent' depuis Supabase"),
        ("Build Due V3 Follow-Ups", "Calcule quels prospects sont dus pour une relance"),
        ("If V3 Follow-Up Ready", "Filtre les prospects éligibles"),
        ("Build V3 Follow-Up Email", "Génère le contenu de la relance"),
        ("Send V3 Follow-Up Email", "Envoie via Resend"),
        ("Log V3 Follow-Up Attempt", "Enregistre la tentative dans Supabase"),
        ("Update Prospect V3 Follow-Up CRM", "Met à jour le statut du prospect"),
    ]
    for node, action in fu_steps:
        bullet(doc, f"{node} → {action}")

    doc.add_page_break()

    # ── 11. CALENDLY ─────────────────────────────────────────────────────────
    h1(doc, "11. Workflow 12 — Calendly Meeting Booked Webhook V1")
    info_table(doc, [
        ("Nom complet", "TransferAI Calendly Meeting Booked Webhook V1"),
        ("Noeuds", "6"),
        ("Déclencheurs", "Webhook Calendly (invitee.created)"),
        ("Rôle", "Mettre à jour le CRM et notifier l'équipe lors d'une prise de RDV"),
        ("Statut", "Actif"),
    ])

    h2(doc, "Rôle")
    body(doc,
        "Dès qu'un prospect prend un RDV via Calendly, ce webhook déclenche la mise à jour "
        "du CRM Supabase et envoie une notification à l'équipe TransferAI."
    )

    h2(doc, "Flux")
    cal_steps = [
        ("Calendly Webhook", "Reçoit l'événement invitee.created"),
        ("Parse Calendly Event", "Extrait email, nom, date/heure du RDV, type de meeting"),
        ("Find Prospect By Email", "Cherche le prospect dans Supabase par email"),
        ("Build Meeting Booked Context", "Construit le contexte de la notification"),
        ("Update Prospect Meeting Booked", "Met à jour le statut CRM (meeting_booked=true, date)"),
        ("Send Admin Meeting Alert", "Email à l'équipe avec les détails du RDV"),
    ]
    for node, action in cal_steps:
        bullet(doc, f"{node} → {action}")

    doc.add_page_break()

    # ── 12. POINTS DE VIGILANCE ───────────────────────────────────────────────
    h1(doc, "12. Points de vigilance et clarifications")

    h2(doc, "⚠️ Workflows en doublon actif")
    alerts = [
        ("Post-Audit V2 + V6",
         "Les deux ont le même schedule (30 min) et le même webhook. "
         "V2 doit être DÉSACTIVÉ dans n8n. V6 est la version en production.",
         "ACTION : Désactiver V2 (toggle Off)"),
        ("Chatwoot V5.5 + V5.5.2",
         "Si les deux pointent vers le même compte Chatwoot, chaque message reçoit 2 réponses. "
         "V5.5.2 est la version corrigée (probablement fix sur la détection bot).",
         "ACTION : Vérifier les comptes Chatwoot. Si même compte → désactiver V5.5"),
    ]
    for title_text, desc, action in alerts:
        h3(doc, title_text)
        body(doc, desc)
        body(doc, action)

    h2(doc, "📌 Points d'attention opérationnels")
    ops = [
        "V3 peut échouer sur 'Fetch Public Page' si le site du prospect est inaccessible (timeout) — normal, géré",
        "V6 s'exécute toutes les 30 min — une soumission peut attendre jusqu'à 30 min avant traitement",
        "Le schedule V4 doit être activé manuellement pour lancer la prospection des 53 prospects tier1",
        "Les Edge Functions Supabase doivent être redéployées après chaque modification du code",
        "Les noeuds 'If' en ancien format n8n (boolean v1) doivent être migrés vers le format v2",
    ]
    for op in ops:
        bullet(doc, op)

    doc.add_page_break()

    # ── 13. SCHÉMA D'INTERACTION ──────────────────────────────────────────────
    h1(doc, "13. Schéma d'interaction entre workflows")
    body(doc, "Représentation textuelle du flux global :")

    schema = """
┌─────────────────────────────────────────────────────────────────────┐
│                    ENTRÉES PROSPECTS                                 │
│  Google Forms Social  │  Supabase/Airtable/Sheets  │  Manual        │
└──────────┬────────────┴──────────┬─────────────────┴────────┬───────┘
           │                       │                           │
           ▼                       ▼                           │
  [Google Forms V2]      [V4 Batch Orchestrator] ─────────────┘
  (leads réseaux sociaux)  (sélection quotidienne)
           │                       │
           ▼                       ▼
  [CRM Supabase]         [V3 CRM Enhanced FINAL]
                         (enrichissement + génération pack)
                                   │
                    ┌──────────────┼──────────────┐
                    ▼              ▼               ▼
              Executive        Mini-Catalogue    Deck Brief
              Letter           Sectoriel        + Audit Form URL
                    └──────────────┬──────────────┘
                                   ▼
                    [Admin Controller V1]
                    (validation humaine approve/reject)
                                   │
                                   ▼ (si approved)
                    [Email prospect via Resend]
                    + [V3 Follow-Up Sequence]
                    (J+1, J+3, J+7 si pas de réponse)
                                   │
                                   ▼ (prospect remplit formulaire)
                    [Edge Function resolve-invitation]
                    (création/résolution invitation)
                    [Formulaire audit transferai.ci]
                                   │
                                   ▼ (soumission)
                    [V6 Post-Audit Expert Routing]
                    (schedule 30 min ou webhook)
                                   │
                    ┌──────────────┼──────────────┐
                    ▼              ▼               ▼
              Fiche pré-RDV   Email expert    CRM Update
              (GPT-4)         [PRIORITÉ HAUTE] Supabase
                    └──────────────┬──────────────┘
                                   │
           ┌───────────────────────┤
           │                       │
           ▼                       ▼
  [Calendly Webhook]    [Chatwoot Auto Reply V5.5.2]
  (RDV pris → CRM)     (messages entrants → réponse GPT-4)

  [Zoho Reply V1]
  (réponses email → classification GPT-4)
"""
    p = doc.add_paragraph()
    run = p.add_run(schema)
    run.font.name = "Courier New"
    run.font.size = Pt(8)

    # Save
    out_path = os.path.join(OUTPUT_DIR, "107_Guide_Maitre_Workflows_TransferAI_V1.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    p = build_doc()
    print(f"\n✓ Document généré : {p}")
