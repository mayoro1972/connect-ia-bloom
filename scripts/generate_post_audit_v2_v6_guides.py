from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUT_DIR = ROOT / "docs" / "transferai-admin"
USER_GUIDE_PATH = OUT_DIR / "Guide_Utilisateur_Workflow_Post_Audit_V2_V6.docx"
TROUBLESHOOTING_PATH = OUT_DIR / "Troubleshooting_Guide_Workflow_Post_Audit_V2_V6.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, RGBColor(15, 46, 75)),
        ("Heading 1", 15, RGBColor(15, 46, 75)),
        ("Heading 2", 12.5, RGBColor(36, 82, 120)),
        ("Heading 3", 11, RGBColor(58, 86, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title_block(document, title, subtitle, project_scope):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 46, 75)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.name = "Arial"
    run2.font.size = Pt(10.5)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(90, 90, 90)

    rows = [
        ("Projet", "Workflow Post-Audit CRM-Aware V2 + V6 Dashboard"),
        ("Date", "07 juin 2026"),
        ("Portée", project_scope),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(5.5)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "DDEBF7")
    document.add_paragraph()


def add_section(document, title, paragraphs=None, bullets=None, numbered=None):
    document.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(5)
    for item in bullets or []:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)
    for item in numbered or []:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_subsection(document, title, paragraphs=None, bullets=None, numbered=None):
    document.add_heading(title, level=2)
    for paragraph in paragraphs or []:
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(4)
    for item in bullets or []:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)
    for item in numbered or []:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_issue(document, title, symptome, cause, resolution, verification, statut="Résolu"):
    document.add_heading(title, level=2)
    rows = [
        ("Symptôme", symptome),
        ("Cause", cause),
        ("Résolution", resolution),
        ("Vérification", verification),
        ("Statut", statut),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.5)
        cells[1].width = Inches(5.6)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "EAF2F8")
    document.add_paragraph()


def build_user_guide():
    doc = Document()
    set_document_defaults(doc)
    add_title_block(
        doc,
        "Guide Utilisateur V2 + V6",
        "Workflow post-audit TransferAI avec tableau de bord Google Sheets",
        "Guide opératoire pour équipes audit, experts et commerciaux non-tech.",
    )

    add_section(
        doc,
        "1. Objet du guide",
        paragraphs=[
            "Ce guide décrit la chaîne post-audit V2 enrichie par V6. Il explique comment les réponses du formulaire d'audit sont récupérées, consolidées, transformées en fiche expert, puis synchronisées dans un tableau Google Sheets utilisable par les commerciaux et les experts non-tech.",
            "Le guide suppose que la V3 de prospection, d'envoi du pack et d'accès au formulaire dynamique est déjà opérationnelle.",
        ],
    )

    add_section(
        doc,
        "2. Ce que fait V2",
        bullets=[
            "Récupère la réponse d'audit complétée depuis Supabase.",
            "Relie la réponse au bon pack, à l'invitation et au prospect CRM.",
            "Construit un contexte CRM-aware consolidé.",
            "Génère une fiche interne Pré-RDV par IA.",
            "Met à jour le suivi post-audit dans Supabase.",
            "Envoie une fiche interne à l'équipe experte.",
        ],
    )

    add_section(
        doc,
        "3. Ce que V6 ajoute",
        bullets=[
            "Transforme la sortie métier de V2 en ligne de dashboard lisible.",
            "Synchronise cette ligne dans Google Sheets.",
            "Déclenche une notification expert si le dossier est suffisamment complet.",
            "Produit un résumé d'exécution V6 pour contrôle et audit.",
        ],
    )

    add_section(
        doc,
        "4. Flux global V2 + V6",
        numbered=[
            "Le prospect complète le formulaire d'audit lié à un pack_id.",
            "La réponse est enregistrée dans Supabase via save-form-response.",
            "Le workflow V2 post-audit se déclenche.",
            "V2 relit le pack, la réponse, l'invitation et le prospect CRM.",
            "V2 construit le contexte consolidé dans Build Post-Audit CRM Context.",
            "V2 génère la fiche Pré-RDV et met à jour post_audit_follow_ups.",
            "Build Post-Audit Result construit la sortie métier finale.",
            "V6 transforme cette sortie en ligne commerciale de dashboard.",
            "La ligne est synchronisée dans Google Sheets.",
            "Si le score de complétion atteint le seuil défini, l'expert reçoit une notification.",
        ],
    )

    add_section(
        doc,
        "5. Nœuds clés à connaître",
        bullets=[
            "Build Post-Audit CRM Context : fusion métier pack + formulaire + CRM.",
            "Generate Internal Pre-RDV Brief : fiche expert générée par IA.",
            "Upsert Follow-Up Tracking : suivi structuré dans post_audit_follow_ups.",
            "Build Post-Audit Result : sortie consolidée V2.",
            "Build V6 Dashboard Row : formatage des données pour le tableau non-tech.",
            "Sync Google Sheets Dashboard : écriture ou mise à jour de la ligne dashboard.",
            "If Expert Notification Enabled : règle de déclenchement d'alerte.",
            "Send Expert Notification : email expert avec résumé du dossier.",
            "Build V6 Summary : résumé de succès de la chaîne V6.",
        ],
    )

    add_section(
        doc,
        "6. Pré-requis de configuration",
        bullets=[
            "Le workflow V3 doit déjà être opérationnel.",
            "Le workflow V2 post-audit doit être validé de bout en bout.",
            "Une feuille Google Sheets doit exister avec l'onglet responses_dashboard.",
            "Le credential n8n Google Sheets OAuth2 doit être créé.",
            "La clé RESEND_API_KEY doit être disponible dans n8n.",
            "Le domaine d'envoi contact@transferai.ci doit rester vérifié dans Resend.",
        ],
    )

    add_subsection(
        doc,
        "6.1 Colonnes recommandées du dashboard Google Sheets",
        paragraphs=[
            "Pour le MVP V6 retenu, l'onglet responses_dashboard doit contenir une ligne d'en-tête avec les colonnes suivantes :",
        ],
        bullets=[
            "response_id",
            "pack_id",
            "submitted_at",
            "organization_name",
            "decision_maker_name",
            "target_email",
            "sector_guess",
            "maturity_level",
            "recommended_offer",
            "completion_percentage",
            "assigned_expert_email",
            "workflow_status",
            "next_action_at",
            "booking_link",
            "follow_up_status",
            "commercial_notes",
        ],
    )

    add_section(
        doc,
        "7. Ce que voit un commercial non-tech",
        paragraphs=[
            "Le commercial n'a pas besoin d'ouvrir Supabase. Le tableau Google Sheets devient sa vue de pilotage. Chaque ligne correspond à une réponse d'audit traitée.",
        ],
        bullets=[
            "Organisation et email du prospect",
            "Secteur et niveau de maturité IA",
            "Service recommandé",
            "Expert assigné",
            "Date de prochaine action",
            "Statut du suivi",
            "Champ libre de notes commerciales",
        ],
    )

    add_section(
        doc,
        "8. Ce que voit un expert",
        paragraphs=[
            "L'expert continue de recevoir la fiche Pré-RDV issue de V2. V6 lui ajoute une notification lisible et un point d'entrée vers le tableau de suivi.",
        ],
        bullets=[
            "Organisation",
            "Secteur",
            "Maturité IA",
            "Service recommandé",
            "Niveau de complétion",
            "Email prospect",
            "Date de prochaine action",
        ],
    )

    add_section(
        doc,
        "9. Procédure de test recommandée",
        numbered=[
            "Soumettre un formulaire d'audit complet ou quasi complet.",
            "Vérifier que V2 se termine avec status = post_audit_processed.",
            "Vérifier que Build V6 Dashboard Row contient les champs attendus.",
            "Vérifier que la ligne apparaît ou se met à jour dans Google Sheets.",
            "Vérifier que la notification expert part si la règle est satisfaite.",
            "Contrôler que Build V6 Summary retourne status = v6_sync_completed.",
        ],
    )

    add_section(
        doc,
        "10. Critères d'acceptation V2 + V6",
        bullets=[
            "Une réponse d'audit traitée crée ou met à jour une ligne dashboard.",
            "Le dashboard est lisible par un commercial non-tech.",
            "L'expert reçoit l'alerte selon la règle métier retenue.",
            "Le suivi Supabase et le dashboard restent cohérents.",
            "Aucune étape ne dépend d'une lecture manuelle dans Supabase pour l'usage courant.",
        ],
    )

    return doc


def build_troubleshooting_guide():
    doc = Document()
    set_document_defaults(doc)
    add_title_block(
        doc,
        "Troubleshooting Guide V2 + V6",
        "Incidents, diagnostics et résolutions pour le workflow post-audit et le dashboard Google Sheets",
        "Guide de résolution des incidents pour équipes techniques et référents n8n.",
    )

    add_section(
        doc,
        "1. Objet du guide",
        paragraphs=[
            "Ce guide recense les incidents les plus probables sur la chaîne V2 + V6 et la méthode de résolution associée. Il part du principe que le formulaire d'audit dynamique et la V3 ont déjà été stabilisés.",
        ],
    )

    add_section(
        doc,
        "2. Points de contrôle rapides",
        bullets=[
            "Le formulaire prospect soumet bien ses réponses dans form_responses.",
            "Le workflow V2 passe par Build Post-Audit Result sans erreur.",
            "Build V6 Dashboard Row produit un payload complet.",
            "Google Sheets reçoit bien la ligne response_id / pack_id.",
            "Le seuil de notification expert est conforme à la règle choisie.",
        ],
    )

    add_issue(
        doc,
        "3.1 V2 ne se déclenche pas après soumission du formulaire",
        symptome="Le prospect soumet le formulaire mais aucun traitement post-audit n'apparaît dans n8n.",
        cause="Webhook ou déclencheur post-audit non appelé, ou payload de save-form-response incomplet.",
        resolution="Vérifier la configuration du webhook/trigger post-audit, confirmer que save-form-response envoie bien l'événement attendu, puis relancer un test avec un pack_id réel.",
        verification="Une nouvelle exécution V2 apparaît immédiatement après une soumission de formulaire.",
    )

    add_issue(
        doc,
        "3.2 Build Post-Audit CRM Context est incomplet",
        symptome="Le workflow V2 démarre mais organisation, email, secteur ou maturité sont vides dans la suite du traitement.",
        cause="Jointure incomplète entre pack, réponse formulaire, invitation ou prospect CRM.",
        resolution="Vérifier les sorties de Extract Pack Row, Extract Latest Form Response, Extract Invitation Row et Extract Prospect Target Row. Corriger les fallbacks dans Build Post-Audit CRM Context si un champ a changé de nom.",
        verification="Build Post-Audit CRM Context contient bien organization_name, target_email, sector_guess, maturity_level et pack_id.",
    )

    add_issue(
        doc,
        "3.3 La fiche Pré-RDV n'est pas envoyée",
        symptome="Le contexte est construit mais l'expert ne reçoit pas la fiche interne.",
        cause="Échec OpenAI, échec Resend, ou internal_email_to vide.",
        resolution="Vérifier Generate Internal Pre-RDV Brief, Send Internal Brief Email, la clé RESEND_API_KEY et la valeur internal_email_to. Tester avec un destinataire interne en dur pour isoler l'incident.",
        verification="Build Post-Audit Result contient un resend_id et status = post_audit_processed.",
    )

    add_issue(
        doc,
        "3.4 La ligne n'apparaît pas dans Google Sheets",
        symptome="V2 se termine correctement mais aucun enregistrement n'apparaît dans le dashboard.",
        cause="Credential Google Sheets absent, mauvais document ID, mauvais nom d'onglet, ou mapping cassé.",
        resolution="Vérifier le credential Google Sheets, l'ID du fichier, l'onglet responses_dashboard et les permissions du compte Google. Rejouer Build V6 Dashboard Row puis Sync Google Sheets Dashboard.",
        verification="La ligne est créée ou mise à jour dans l'onglet responses_dashboard avec response_id comme clé.",
    )

    add_issue(
        doc,
        "3.5 La ligne est créée mais les colonnes sont décalées",
        symptome="Les données arrivent dans Google Sheets mais dans les mauvaises colonnes ou de manière illisible.",
        cause="Entêtes non conformes au mapping V6 ou changement manuel dans la feuille.",
        resolution="Réinitialiser la ligne d'en-tête du dashboard avec les colonnes validées, puis revalider le mapping du nœud Google Sheets.",
        verification="Chaque champ se range dans la bonne colonne après un nouvel envoi de test.",
    )

    add_issue(
        doc,
        "3.6 La notification expert ne part pas",
        symptome="Le dashboard est mis à jour mais aucun email expert n'est reçu.",
        cause="La règle du nœud If n'est pas satisfaite, assigned_expert_email est vide, ou la requête Resend échoue.",
        resolution="Vérifier la règle métier choisie pour If Expert Notification Enabled, contrôler assigned_expert_email et inspecter la réponse du nœud Send Expert Notification.",
        verification="Quand la règle est vraie, Send Expert Notification retourne un id Resend.",
    )

    add_issue(
        doc,
        "3.7 Un commercial non-tech ne comprend pas le tableau",
        symptome="Le dashboard est techniquement rempli mais difficilement exploitable par les équipes commerciales.",
        cause="Colonnes trop techniques, absence de statuts lisibles, ou trop d'informations brutes.",
        resolution="Réduire la vue à une liste métier courte, ajouter filtres et couleurs dans Google Sheets, et documenter les statuts follow_up_status et workflow_status.",
        verification="Un commercial peut comprendre qui a répondu, quel service est recommandé et quelle est la prochaine action sans ouvrir Supabase.",
        statut="À surveiller"
    )

    add_issue(
        doc,
        "3.8 La cohérence entre Supabase et Google Sheets dérive",
        symptome="Le tableau montre un état différent du suivi dans post_audit_follow_ups ou prospect_targets.",
        cause="Le dashboard est alimenté mais les règles de mise à jour ou de réexécution ne sont pas clairement définies.",
        resolution="Définir response_id comme clé primaire du dashboard, figer les statuts métier, et documenter les cas de rerun manuel ou de reconciliation.",
        verification="Une nouvelle réponse ou un rerun met à jour la bonne ligne sans dupliquer les dossiers.",
        statut="À surveiller"
    )

    add_section(
        doc,
        "4. Checklist de reprise d'incident",
        numbered=[
            "Vérifier que le prospect a bien soumis une réponse exploitable.",
            "Confirmer que V2 a démarré et atteint Build Post-Audit Result.",
            "Vérifier les champs de sortie de Build V6 Dashboard Row.",
            "Tester Sync Google Sheets Dashboard isolément si nécessaire.",
            "Tester Send Expert Notification séparément si la ligne dashboard existe déjà.",
            "Contrôler le résumé Build V6 Summary pour confirmer l'état final.",
        ],
    )

    add_section(
        doc,
        "5. Résultat attendu après stabilisation",
        bullets=[
            "V2 traite systématiquement chaque audit utile.",
            "V6 reflète les réponses dans un tableau simple et actionnable.",
            "Les experts sont prévenus à temps.",
            "Les commerciaux non-tech travaillent sans dépendre de Supabase.",
            "Le système est prêt pour une future V7 de dashboard natif Supabase.",
        ],
    )

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_user_guide().save(USER_GUIDE_PATH)
    build_troubleshooting_guide().save(TROUBLESHOOTING_PATH)
    print(USER_GUIDE_PATH)
    print(TROUBLESHOOTING_PATH)


if __name__ == "__main__":
    main()
