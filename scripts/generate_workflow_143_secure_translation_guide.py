from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


MD_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "144_Guide_Utilisateur_Workflow_143_Traduction_Internationale_Securisee_06juillet2026.md"
)
DOCX_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "144_Guide_Utilisateur_Workflow_143_Traduction_Internationale_Securisee_06juillet2026.docx"
)


TITLE = "Guide utilisateur - Workflow 143"
SUBTITLE = (
    "Traduction internationale securisee avec PII, pseudonymisation locale, "
    "relecture qualite et gouvernance humaine"
)


SECTIONS = [
    (
        "1. Objet du guide",
        [
            "Ce document presente le workflow 143 de TransferAI pour la traduction internationale de documents sensibles. Il a ete concu pour les organisations internationales, institutions publiques, cabinets, directions generales, ambassades, ONG, agences onusiennes et equipes evenementielles qui manipulent des contenus officiels ou confidentiels.",
            "Le workflow met en oeuvre une logique security-by-design : detection PII, anonymisation selective, pseudonymisation locale, envoi minimal au modele, reidentification uniquement dans n8n, puis validation humaine pour les cas a risque eleve.",
        ],
    ),
    (
        "2. Cas d'usage cibles",
        [
            "Traduction de dossiers de visa, passeports, actes, notes verbales, invitations officielles, lettres administratives et documents de conformite.",
            "Traduction de documents de sommet international, reunion bilaterale, atelier regional, conference de presse, protocole et compte rendu sensible.",
            "Traduction d'accords, projets de contrat, memos, notes diplomatiques, supports de mission et pieces documentaires devant circuler entre plusieurs langues.",
        ],
    ),
    (
        "3. Principes de protection appliques",
        [
            "Minimisation des donnees : seules les informations strictement necessaires a la traduction sont transmises au modele.",
            "Pseudonymisation locale : les identifiants directs sont remplaces par des jetons comme PERSON_001, ORG_001, EMAIL_001, REF_001 ou PASSPORT_001.",
            "Anonymisation contextuelle : les metadonnees accessoires ou non necessaires peuvent etre retirees ou generalisees.",
            "Reidentification locale uniquement : le fournisseur IA ne voit jamais les donnees reelles lorsque la pseudonymisation est active.",
            "Purge post-traitement : la table de correspondance sensible est detruite a la fin du traitement operationnel.",
            "Validation humaine : pour les visas, documents d'identite, actes d'etat civil, contenus diplomatiques et toute traduction critique, le workflow bascule en revue humaine avant diffusion.",
        ],
    ),
    (
        "4. Architecture du workflow",
        [
            "Reception de la demande via webhook securise.",
            "Extraction des parametres et classification de sensibilite.",
            "Detection PII dans le texte et les champs structures.",
            "Pseudonymisation locale avant appel au modele.",
            "Traduction pseudonymisee via OpenAI.",
            "Relecture qualite pseudonymisee avec score et risques.",
            "Reidentification locale dans n8n.",
            "Routage automatique : envoi client direct ou envoi au validateur interne.",
            "Purge de la table sensible.",
        ],
    ),
    (
        "5. Champs d'entree recommandes",
        [
            "`texte` ou `content` : texte a traduire.",
            "`langue_source` et `langue_cible` : codes courts comme `fr`, `en`, `es`, `ar`, `pt`.",
            "`type_document` : visa, note verbale, courrier officiel, contrat, compte rendu, etc.",
            "`demandeur`, `organisation`, `reference_dossier` : seulement si necessaire a l'orchestration interne.",
            "`email_retour` : destinataire final.",
            "`email_validateur` : validateur interne ou compliance officer pour les documents sensibles.",
            "`contexte_mission` : sommet, reunion internationale, mission terrain, immigration, protocole, cooperation, etc.",
        ],
    ),
    (
        "6. Exemple de charge utile",
        [
            "Exemple JSON :",
            "```json\n{\n  \"texte\": \"Nom : Fatou Diallo\\nPassport No: AB1234567\\nEmail: fatou.diallo@example.org\\nObjet: demande de visa officiel pour le sommet regional.\",\n  \"langue_source\": \"fr\",\n  \"langue_cible\": \"en\",\n  \"type_document\": \"Demande de visa officiel\",\n  \"demandeur\": \"Fatou Diallo\",\n  \"organisation\": \"Organisation regionale fictive\",\n  \"reference_dossier\": \"VISA-2026-0041\",\n  \"email_retour\": \"translation.office@example.org\",\n  \"email_validateur\": \"compliance@example.org\",\n  \"contexte_mission\": \"Sommet international\"\n}\n```",
        ],
    ),
    (
        "7. Comportement de validation",
        [
            "Le workflow marque `review_required = true` dans trois cas principaux : document de sensibilite elevee, score qualite insuffisant, ou validation automatique negative.",
            "Quand la revue est requise, le document n'est pas diffuse au client final. Il est adresse au validateur interne defini dans `email_validateur`.",
            "Quand la revue n'est pas requise, le document bilingue est envoye directement au destinataire final via Resend.",
        ],
    ),
    (
        "8. Prerequis techniques",
        [
            "Variable d'environnement `OPENAI_API_KEY` pour la traduction et la relecture.",
            "Variable d'environnement `RESEND_API_KEY` pour les envois email.",
            "Optionnel : `INTERNAL_REVIEW_EMAIL` pour definir un email de revue par defaut.",
            "Optionnel : `OUTREACH_FROM_EMAIL` pour l'adresse expediteur.",
            "Import du fichier JSON dans n8n, puis verification des credentials et tests sur un document fictif avant mise en production.",
        ],
    ),
    (
        "9. Recommandations operationnelles pour organisations internationales",
        [
            "Toujours tester avec un document factice avant premier usage en production.",
            "Ne pas reutiliser des cles API en clair dans les noeuds ; utiliser exclusivement variables d'environnement ou credentials n8n.",
            "Definir une politique par type de document : visa, identite, diplomatique, legal, RH, medical, financier, media.",
            "Maintenir une validation humaine pour toute piece qui peut avoir un impact legal, migratoire, reputionnel ou diplomatique.",
            "Limiter les personnes pouvant consulter les executions n8n contenant encore des donnees reelles pendant la fenetre de traitement.",
        ],
    ),
    (
        "10. Checklist de mise en production",
        [
            "Verifier que les cles API ont ete configurees hors workflow.",
            "Verifier que les executions n8n et les logs sont conformes a votre politique interne.",
            "Tester au moins un cas visa, un cas reunion internationale et un cas courrier officiel.",
            "Verifier la logique de routage vers `email_validateur`.",
            "Documenter qui peut approuver, corriger et diffuser le livrable final.",
            "Planifier une rotation immediate des secrets si un ancien workflow contenait des cles en dur.",
        ],
    ),
]


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=4, line=1.0)
    r = p.add_run(TITLE)
    set_font(r, size=20, bold=True, color="2E74B5")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p2, before=0, after=14, line=1.15)
    r2 = p2.add_run(SUBTITLE)
    set_font(r2, size=11, color="555555")


def add_meta_table(doc):
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    widths = (1.88, 4.62)
    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)

    rows = [
        ("Public cible", "Organisations internationales, sommets, reunions sensibles, cabinets, institutions"),
        ("Workflow", "Workflow 143 importable pour n8n (fichier JSON securise fourni dans le depot)"),
        ("Mode securite", "Detection PII, pseudonymisation locale, payload minimal LLM, purge"),
        ("Decision critique", "Validation humaine requise pour documents a sensibilite elevee"),
    ]

    for i, (left, right) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].text = left
        cells[1].text = right
        shade_cell(cells[0], "E8EEF5")
        for c_index, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, before=0, after=0, line=1.25)
                if paragraph.runs:
                    set_font(paragraph.runs[0], size=10.5, bold=(c_index == 0), color="10263F" if c_index == 0 else None)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        set_spacing(p, before=18, after=10, line=1.0)
        r = p.add_run(text)
        set_font(r, size=16, bold=True, color="2E74B5")
    else:
        set_spacing(p, before=14, after=7, line=1.0)
        r = p.add_run(text)
        set_font(r, size=13, bold=True, color="2E74B5")


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=6, line=1.25)
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_spacing(p, before=0, after=4, line=1.25)
        r = p.add_run(item)
        set_font(r, size=11)


def build_markdown():
    lines = [f"# {TITLE}", "", SUBTITLE, ""]
    lines.extend(
        [
            "| Element | Configuration |",
            "|---|---|",
            "| Public cible | Organisations internationales, sommets, reunions sensibles, cabinets, institutions |",
            "| Workflow | `143_n8n_Traduction_Officielle_Internationale_PII_Revamp_06juillet2026.json` |",
            "| Mode securite | Detection PII, pseudonymisation locale, payload minimal LLM, purge |",
            "| Decision critique | Validation humaine requise pour documents a sensibilite elevee |",
            "",
        ]
    )

    for title, paragraphs in SECTIONS:
        lines.append(f"## {title}")
        lines.append("")
        for text in paragraphs:
            if text.startswith("```json"):
                lines.append(text)
            elif text.startswith("`"):
                lines.append(f"- {text}")
            elif "```json" in text:
                lines.append(text)
            else:
                if title in {
                    "2. Cas d'usage cibles",
                    "3. Principes de protection appliques",
                    "4. Architecture du workflow",
                    "5. Champs d'entree recommandes",
                    "8. Prerequis techniques",
                    "9. Recommandations operationnelles pour organisations internationales",
                    "10. Checklist de mise en production",
                }:
                    lines.append(f"- {text}")
                else:
                    lines.append(text)
            lines.append("")

    MD_OUTPUT.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    add_title(doc)
    add_meta_table(doc)

    for title, paragraphs in SECTIONS:
        add_heading(doc, title, level=1)
        if title in {
            "2. Cas d'usage cibles",
            "3. Principes de protection appliques",
            "4. Architecture du workflow",
            "5. Champs d'entree recommandes",
            "8. Prerequis techniques",
            "9. Recommandations operationnelles pour organisations internationales",
            "10. Checklist de mise en production",
        }:
            add_bullets(doc, paragraphs)
            continue

        for text in paragraphs:
            if text.startswith("```json"):
                add_heading(doc, "Exemple JSON", level=2)
                add_paragraph(doc, text.replace("```json\n", "").replace("\n```", ""))
            else:
                add_paragraph(doc, text)

    doc.save(DOCX_OUTPUT)


def main():
    build_markdown()
    build_docx()


if __name__ == "__main__":
    main()
