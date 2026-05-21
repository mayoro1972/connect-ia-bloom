from __future__ import annotations

from datetime import datetime
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(201, 91, 37)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(107, 114, 128)
BORDER = "D9DDE3"
HEADER_TEXT = "TransferAI Africa | Veille réglementaire IA banque CI"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color: str = BORDER, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = p_pr.find(qn("w:pbdr"))
    if pbdr is None:
      pbdr = OxmlElement("w:pbdr")
      p_pr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def ensure_style(document: Document, name: str, base: str, size: int, bold: bool = False, color: RGBColor | None = None):
    styles = document.styles
    if name in styles:
        style = styles[name]
    else:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles[base]
    font = style.font
    font.name = "Arial"
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color
    return style


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    ensure_style(document, "TA Title", "Title", 22, bold=True, color=TEXT)
    ensure_style(document, "TA Subtitle", "Subtitle", 12, color=MUTED)
    ensure_style(document, "TA Heading 1", "Heading 1", 16, bold=True, color=TEXT)
    ensure_style(document, "TA Heading 2", "Heading 2", 14, bold=True, color=TEXT)
    ensure_style(document, "TA Heading 3", "Heading 3", 12, bold=True, color=TEXT)
    ensure_style(document, "TA Small", "Normal", 10, color=MUTED)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run(HEADER_TEXT)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_paragraph_bottom_border(header_p)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Page ")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = MUTED
    add_page_number(footer_p)


def add_cover(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph(style="TA Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.color.rgb = TEXT

    sub = document.add_paragraph(style="TA Subtitle")
    sub.add_run(subtitle)

    meta = document.add_paragraph(style="TA Small")
    meta.add_run(f"Document opérationnel | Généré le {datetime.now().strftime('%d/%m/%Y')}")

    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    cell.text = (
        "Ce document sert de mode opératoire court pour une équipe conformité, veille ou transformation "
        "afin de capter, qualifier et publier rapidement les signaux réglementaires IA utiles à une banque "
        "en Côte d'Ivoire."
    )
    set_cell_shading(cell, "F7F1EA")
    document.add_paragraph()


def add_markdown_content(document: Document, markdown_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    skip_first_title = True

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith("# "):
            if skip_first_title:
                skip_first_title = False
                continue
            document.add_paragraph(stripped[2:].strip(), style="TA Heading 1")
            continue

        if stripped.startswith("## "):
            document.add_paragraph(stripped[3:].strip(), style="TA Heading 1")
            continue

        if stripped.startswith("### "):
            document.add_paragraph(stripped[4:].strip(), style="TA Heading 2")
            continue

        if stripped.startswith("#### "):
            document.add_paragraph(stripped[5:].strip(), style="TA Heading 3")
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="Normal")
            paragraph.style = document.styles["List Bullet"]
            run = paragraph.add_run(stripped[2:].strip())
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.color.rgb = TEXT
            continue

        paragraph = document.add_paragraph(style="Normal")
        paragraph.add_run(stripped)


def build_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    add_cover(
        document,
        "Guide opératoire | Veille réglementaire IA",
        "Banque - Côte d'Ivoire",
    )
    add_markdown_content(document, markdown_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: python generate_regulatory_watch_guide_docx.py <input.md> <output.docx>")
        return 1

    markdown_path = Path(sys.argv[1]).resolve()
    output_path = Path(sys.argv[2]).resolve()

    if not markdown_path.exists():
        print(f"Input file not found: {markdown_path}")
        return 1

    build_docx(markdown_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
