from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
SOURCE_MD = ROOT / "outputs" / "proposition-openclaw-paradis-enfants.md"
OUTPUT_DIR = ROOT / "outputs" / "openclaw-paradis-enfants-doc"
OUTPUT_DOCX = OUTPUT_DIR / "Proposition_Commerciale_OpenClaw_Paradis_Enfants.docx"

TITLE_BLUE = RGBColor(0x17, 0x2B, 0x4D)
HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
BODY = RGBColor(0x10, 0x19, 0x28)
MUTED = RGBColor(0x47, 0x55, 0x67)
PINK = RGBColor(0xC8, 0x1E, 0x78)
LIGHT_FILL = "F8F5FF"
SOFT_FILL = "F4F6F9"
GRID = "D9E2F1"

INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE = re.compile(r"`(.+?)`")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
      p_bdr = OxmlElement("w:pBdr")
      p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
      bottom = OxmlElement("w:bottom")
      p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GRID)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    h1 = document.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_BLUE
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h2 = document.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = HEADING_BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h3 = document.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = HEADING_DARK
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.15
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Projet piloté par TransferAI, notre hub IA, en collaboration avec Nettelecom CI")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED

    settings = document.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def add_runs(paragraph, text: str, *, size: int = 11, color: RGBColor = BODY) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        bold_match = INLINE_BOLD.fullmatch(part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = True
            run.font.color.rgb = color
            continue
        code_match = INLINE_CODE.fullmatch(part)
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Consolas"
            run.font.size = Pt(max(size - 1, 9))
            run.font.color.rgb = BODY
            continue
        run = paragraph.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = color


def add_title_block(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.05
    run = title.add_run("Proposition commerciale OpenClaw pour Paradis D'enfants")
    run.font.name = "Calibri"
    run.font.size = Pt(23)
    run.font.bold = True
    run.font.color.rgb = TITLE_BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.line_spacing = 1.15
    srun = subtitle.add_run(
        "Adaptation du modèle OpenClaw Boutique Abidjan à un shop bébé et enfants piloté par TikTok, Facebook, WhatsApp, CRM et Mobile Money."
    )
    srun.font.name = "Calibri"
    srun.font.size = Pt(11)
    srun.font.color.rgb = MUTED
    add_bottom_rule(subtitle)

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    meta.autofit = False
    meta.columns[0].width = Inches(1.45)
    meta.columns[1].width = Inches(4.85)

    rows = [
        ("Client visé", "Paradis D'enfants"),
        ("Positionnement", "Boutique bébé et enfants portée par les réseaux sociaux"),
        ("Canaux couverts", "TikTok, Facebook, WhatsApp, tableau de bord CRM, Telegram"),
        ("Périmètre clé", "Conversion, relance, CRM, Mobile Money, livraison et pilotage"),
    ]
    for idx, (label, value) in enumerate(rows):
        left = meta.rows[idx].cells[0]
        right = meta.rows[idx].cells[1]
        left.width = Inches(1.45)
        right.width = Inches(4.85)
        set_cell_shading(left, SOFT_FILL)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for cell, text, bold, color in [
            (left, label, True, TITLE_BLUE),
            (right, value, False, BODY),
        ]:
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(10.5)
            run.font.bold = bold
            run.font.color.rgb = color

    document.add_paragraph()

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    callout.style = "Table Grid"
    cell = callout.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    lead = p.add_run("Promesse commerciale : ")
    lead.font.name = "Calibri"
    lead.font.size = Pt(11)
    lead.font.bold = True
    lead.font.color.rgb = PINK
    body = p.add_run(
        "transformer l'audience sociale de la boutique en conversations suivies, commandes mieux gérées, paiements Mobile Money tracés et livraisons pilotées."
    )
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = BODY
    document.add_paragraph("Projet piloté par TransferAI, notre hub IA, en collaboration avec Nettelecom CI.")
    document.add_paragraph()


def add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(p, text)


def add_numbered(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(p, text)


def add_heading(document: Document, level: int, text: str) -> None:
    style_name = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
    p = document.add_paragraph(style=style_name)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs(p, text)


def export_proposal() -> Path:
    document = Document()
    configure_document(document)
    add_title_block(document)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    start_index = next(
        (idx for idx, line in enumerate(lines) if line.startswith("## 1. ")),
        0,
    )

    for raw in lines[start_index:]:
        line = raw.strip()
        if not line or line == "---":
            continue

        if line.startswith("## "):
            add_heading(document, 1, line[3:].strip())
            continue
        if line.startswith("### "):
            add_heading(document, 2, line[4:].strip())
            continue
        if line.startswith("#### "):
            add_heading(document, 3, line[5:].strip())
            continue

        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            add_bullet(document, bullet_match.group(1))
            continue

        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            add_numbered(document, number_match.group(1))
            continue

        add_body_paragraph(document, line)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = export_proposal()
    print(path)
