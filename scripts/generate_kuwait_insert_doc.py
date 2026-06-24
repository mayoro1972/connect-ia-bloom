from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUT_DIR = ROOT / "outputs" / "manual-20260610-kuwait-insert" / "output"
ASSET_DIR = OUT_DIR / "assets"

SOURCE_IMAGES = {
    "config": Path("/private/tmp/kuwait-doc-images/IMG_3486.HEIC.png"),
    "diagram": Path("/private/tmp/kuwait-doc-images/IMG_3496.HEIC.png"),
}

OUTPUT_DOCX = OUT_DIR / "Kuwait_ABMS_Section_3_1_1_Insert.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_landscape(section):
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def crop_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    diagram_image = Image.open(SOURCE_IMAGES["diagram"])
    # Focus on the PDF content and drop most of the laptop bezel.
    diagram_crop = diagram_image.crop((40, 260, 1310, 1405))
    diagram_crop.save(ASSET_DIR / "kuwait-network-diagram-crop.png")

    config_image = Image.open(SOURCE_IMAGES["config"])
    # Keep the DC/DR table area from the spreadsheet view.
    config_crop = config_image.crop((40, 360, 1735, 1005))
    config_crop.save(ASSET_DIR / "kuwait-infinity-config-crop.png")

    return {
        "diagram": ASSET_DIR / "kuwait-network-diagram-crop.png",
        "config": ASSET_DIR / "kuwait-infinity-config-crop.png",
    }


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)
        style.font.size = Pt(size)


def add_bullets(doc, title, items):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

    for item in items:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent = Inches(0.2)
        bp.paragraph_format.space_after = Pt(0)
        run = bp.add_run(item)
        run.font.size = Pt(10)


def add_note_table(doc):
    table = doc.add_table(rows=2, cols=1)
    table.style = "Table Grid"
    table.autofit = True

    header = table.cell(0, 0)
    set_cell_shading(header, "D87A2A")
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("Replication Note")
    header_run.bold = True
    header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_run.font.name = "Arial"
    header_run.font.size = Pt(10)

    body = table.cell(1, 0).paragraphs[0]
    body.paragraph_format.space_before = Pt(4)
    body.paragraph_format.space_after = Pt(4)
    run = body.add_run(
        "SyncIQ replication should be enabled between the DC and DR NAS platforms "
        "to support resilient data replication across both sites."
    )
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_image_page(doc, section_title, caption, image_path, width_inches):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run(section_title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(6)
    caption_p.paragraph_format.space_after = Pt(0)
    cap_run = caption_p.add_run(caption)
    cap_run.italic = True
    cap_run.font.name = "Arial"
    cap_run.font.size = Pt(9)


def build_document(cropped_paths):
    doc = Document()
    set_landscape(doc.sections[0])
    style_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("3.1.1 DC and DR Infrastructure Sizing / IP Address Requirements")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.add_run(
        "The following section provides the proposed Data Centre (DC) and Disaster "
        "Recovery (DR) infrastructure details for the Kuwait deployment, including "
        "server roles, IP address requirements, management interfaces, private "
        "cluster communication, storage partitions, RAID configuration, and NAS "
        "replication requirements."
    )

    add_bullets(
        doc,
        "DC Environment",
        [
            "OS LAN IP requirements",
            "iDRAC / management IP requirements",
            "Cluster IPs",
            "Server/NAS table",
            "SyncIQ replication note",
        ],
    )
    doc.add_paragraph()
    add_bullets(
        doc,
        "DR Environment",
        [
            "OS LAN IP requirements",
            "iDRAC / management IP requirements",
            "Cluster IPs",
            "Server/NAS table",
            "SyncIQ replication note",
        ],
    )
    doc.add_paragraph()
    add_note_table(doc)

    doc.add_page_break()
    add_image_page(
        doc,
        "Kuwait Network Diagram",
        "Figure 1. Kuwait DC/DR network architecture overview.",
        cropped_paths["diagram"],
        9.7,
    )

    doc.add_page_break()
    add_image_page(
        doc,
        "Kuwait Infinity Server Configuration",
        "Figure 2. DC and DR server / NAS configuration capture.",
        cropped_paths["config"],
        9.7,
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


def main():
    cropped_paths = crop_assets()
    build_document(cropped_paths)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
