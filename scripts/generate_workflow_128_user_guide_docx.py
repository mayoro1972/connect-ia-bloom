from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "137_Guide_Utilisateur_Workflow_128_CR_Reunion_Formation_Juillet_Aout_2026.docx"
)


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(document, text, subtitle=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.0)
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color="10263F")
    if subtitle:
        p2 = document.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p2, before=0, after=14, line=1.0)
        r2 = p2.add_run(subtitle)
        set_font(r2, size=11, color="666666")


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6 if level == 1 else 4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=15, bold=True, color="2E74B5")
    elif level == 2:
        set_font(r, size=12.5, bold=True, color="2E74B5")
    else:
        set_font(r, size=11.5, bold=True, color="1F4D78")


def add_body(document, text):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.15)
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_two_col_table(document, rows, widths=(1.9, 4.6)):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])
    hdr = table.rows[0].cells
    hdr[0].text = "Élément"
    hdr[1].text = "Configuration"
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            set_paragraph_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="10263F")
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before=0, after=0, line=1.08)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0))
    document.add_paragraph()


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("Guide utilisateur Workflow 128 - TransferAI")
    set_font(r, size=9, color="777777")


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run("TransferAI - Compte rendu de réunion sensible")
    set_font(r, size=9, bold=True, color="777777")


def build_doc():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header(section)
    add_footer(section)

    add_title(
        document,
        "Guide utilisateur - Workflow 128",
        "Compte rendu de réunion international, pseudonymisation locale et livrables assainis - Session juillet / août 2026",
    )

    add_heading(document, "1. Objet du document", level=1)
    add_body(
        document,
        "Ce guide explique comment utiliser et configurer le workflow 128 pour traiter des enregistrements de réunions sensibles, produire des comptes rendus structurés et générer des livrables externes assainis."
    )
    add_body(
        document,
        "Il est destiné au cas d’usage de la session de formation TransferAI de juillet / août 2026, dans un contexte international impliquant plusieurs sources d’enregistrement et des exigences de confidentialité."
    )

    add_heading(document, "2. Cas d’usage recommandé", level=1)
    add_body(
        document,
        "Le client enregistre ses réunions de cadrage, de suivi ou de restitution via Zoom, Webex, Microsoft Teams, un dictaphone, un smartphone ou un dispositif de salle."
    )
    add_body(
        document,
        "L’objectif est d’obtenir rapidement un compte rendu exploitable, un e-mail de suivi client et une version assainie, tout en protégeant les données sensibles avant appel à l’IA."
    )
    add_bullets(
        document,
        [
            "Réunion de cadrage : Programme Formation IA Secrétariat Direction - session août 2026",
            "Source d’enregistrement : Teams, Zoom, Webex, dictaphone ou smartphone",
            "Sortie attendue : compte rendu, e-mail client, archive assainie",
        ],
    )

    add_heading(document, "3. Architecture du workflow", level=1)
    add_numbered(
        document,
        [
            "Entrée du recording",
            "Normalisation des métadonnées",
            "Récupération ou réception de l’audio",
            "Découpage audio",
            "Transcription OpenAI",
            "Pseudonymisation locale",
            "Génération du rapport assaini",
            "Validation humaine puis envoi final",
        ],
    )

    add_heading(document, "4. Modes d’entrée supportés", level=1)
    add_heading(document, "4.1 Upload manuel", level=2)
    add_body(document, "À utiliser pour les fichiers .m4a, .mp3, les exports smartphone et les exports dictaphone.")
    add_heading(document, "4.2 Google Drive", level=2)
    add_body(document, "À utiliser si les fichiers sont déposés dans un dossier surveillé ou si un lien Drive est fourni dans le formulaire.")
    add_heading(document, "4.3 Webhook plateforme meeting", level=2)
    add_body(
        document,
        "À utiliser si un workflow amont Zoom, Teams ou Webex récupère le recording puis appelle le webhook d’ingestion du workflow 128."
    )

    add_heading(document, "5. Pré-requis techniques", level=1)
    add_two_col_table(
        document,
        [
            ("Clé OpenAI", "Variable d’environnement OPENAI_API_KEY"),
            ("Clé Resend", "Variable d’environnement RESEND_API_KEY"),
            ("Modèle rapport", "gpt-5 recommandé ; sinon gpt-4.1 ou gpt-4o"),
            ("Google Drive", "Credential OAuth actif pour lecture et archivage"),
            ("Découpage audio", "Service disponible sur http://n8n-pxlk-audio-splitter-1:8000/split"),
            ("Domaine webhook", "https://n8n-pxlk.srv1480638.hstgr.cloud/webhook"),
        ],
    )

    add_heading(document, "6. Configuration pas à pas", level=1)
    steps = [
        ("Étape 1 - Importer le workflow", "Dans n8n, ouvrir Workflows, choisir Import from file, puis sélectionner le fichier JSON du workflow 128."),
        ("Étape 2 - Vérifier les points d’entrée", "Contrôler les nœuds Formulaire - Audio sensible, Google Drive - Nouveau fichier audio et Webhook - Ingestion plateforme meeting."),
        ("Étape 3 - Configurer le formulaire", "Renseigner les champs sujet, participants, date, source d’enregistrement, mode d’ingestion, destinataire final, validateur, langue audio, langue de sortie, livrable demandé, entités sensibles à masquer et niveau de masquage."),
        ("Étape 4 - Configurer le dossier source", "Vérifier l’ID du dossier surveillé dans le nœud Google Drive - Nouveau fichier audio."),
        ("Étape 5 - Configurer le dossier d’archive", "Vérifier l’ID du dossier d’archivage dans le nœud Archiver version assainie."),
        ("Étape 6 - Configurer Resend", "Valider les nœuds Envoyer au validateur et Envoyer le livrable final assaini, ainsi que l’adresse expéditeur autorisée."),
        ("Étape 7 - Configurer OpenAI", "Vérifier les nœuds OpenAI - Transcrire segment et OpenAI - Rapport structuré assaini."),
        ("Étape 8 - Vérifier les webhooks de validation", "Contrôler les nœuds Webhook - Approuver et Webhook - Rejeter, ainsi que le domaine de production."),
        ("Étape 9 - Lancer un test pilote", "Faire un test avec un petit enregistrement, approuver le livrable, puis vérifier l’archive HTML."),
    ]
    add_two_col_table(document, steps)

    add_heading(document, "7. Paramétrage conseillé pour la session juillet / août 2026", level=1)
    add_bullets(
        document,
        [
            "Livrable demandé : Les deux",
            "Niveau de masquage : Strict",
            "Langue de sortie : Français",
            "Dossier source : Recordings - Formation IA - Juillet Août 2026",
            "Dossier archive : Archives assainies - Formation IA - Juillet Août 2026",
        ],
    )

    add_heading(document, "8. Configuration des sources de réunion", level=1)
    add_heading(document, "8.1 Dictaphone", level=2)
    add_bullets(
        document,
        [
            "Mode simple : exporter le fichier puis l’envoyer via le formulaire.",
            "Mode semi-automatique : synchroniser les fichiers vers un dossier Google Drive surveillé.",
        ],
    )
    add_heading(document, "8.2 Zoom", level=2)
    add_bullets(
        document,
        [
            "Activer Cloud Recording dans Zoom.",
            "Utiliser un workflow amont recevant l’événement recording.completed.",
            "Télécharger le recording puis appeler le webhook cr-intl-source-ingest.",
        ],
    )
    add_heading(document, "8.3 Microsoft Teams", level=2)
    add_bullets(
        document,
        [
            "Récupérer les enregistrements dans OneDrive ou SharePoint.",
            "Passer par Microsoft Graph ou un workflow de collecte amont.",
            "Transmettre ensuite le binaire ou les métadonnées au workflow 128.",
        ],
    )
    add_heading(document, "8.4 Webex", level=2)
    add_bullets(
        document,
        [
            "Activer les recordings Webex.",
            "Utiliser l’API Recordings pour récupérer le fichier.",
            "Faire suivre le recording vers le webhook d’ingestion du workflow 128.",
        ],
    )

    add_heading(document, "9. Structure attendue pour le webhook d’ingestion", level=1)
    add_body(
        document,
        "Le webhook cr-intl-source-ingest doit idéalement recevoir : source_system, source_label, meeting_title, meeting_date, participants, e-mail destinataire, e-mail validateur, langue, niveau de masquage, référence de réunion, objet source et recording_url. Si possible, transmettre directement le binaire audio."
    )

    add_heading(document, "10. Checklist de mise en production", level=1)
    add_bullets(
        document,
        [
            "Le workflow s’importe sans erreur.",
            "Google Drive est connecté.",
            "Resend est connecté.",
            "OpenAI est connecté.",
            "Le service de découpage audio répond.",
            "Le validateur reçoit l’aperçu.",
            "L’archive HTML est créée.",
            "Aucune version réidentifiée n’est envoyée au client.",
        ],
    )

    add_heading(document, "11. Recommandation finale", level=1)
    add_body(
        document,
        "Pour la session de formation juillet / août 2026, il est recommandé de démarrer avec l’upload manuel et Google Drive, de valider le processus métier et la qualité des rapports, puis d’ajouter dans un second temps des workflows d’ingestion dédiés à Zoom, Teams et Webex."
    )

    document.save(OUTPUT)


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
