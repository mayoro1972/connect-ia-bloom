from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DOCX = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/word/"
    "102_Tableau_Complet_Workflow_Google_Forms_Prospects_Version_Operationnelle_2026-06-23.docx"
)

SOURCE_MD = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "102_Tableau_Complet_Workflow_Google_Forms_Prospects_Version_Operationnelle_2026-06-23.md"
)


ROWS = [
    ("1", "Manual Trigger", "Trigger", "Test", "Lance des essais manuels du workflow", "Exécution manuelle", "Déclenche le flux pour contrôle", "Actif pour tests"),
    ("2", "Hourly Social Follow-Up Schedule", "Schedule Trigger", "Relance", "Déclenche les relances programmées", "Heure planifiée", "Envoie les prospects à relancer", "Actif"),
    ("3", "Google Forms Social Lead Webhook", "Webhook", "Nouveau prospect", "Reçoit le payload Apps Script venant de Google Sheets", "Réponse Google Forms transmise par Apps Script", "Payload brut du prospect", "Actif en production"),
    ("4", "Set Social Sequence Config", "Set", "Configuration", "Charge les variables métier du workflow", "Déclencheur entrant", "Paramètres normalisés de séquence", "Actif"),
    ("5", "If New Lead Payload", "If", "Routage", "Distingue un nouveau prospect d’une relance horaire", "Payload entrant", "Aiguillage vers nouveau lead ou suivi", "Actif"),
    ("6", "Fetch Social Prospect Snapshot", "HTTP Request", "Relance", "Récupère l’état du prospect dans Supabase", "Déclenchement horaire", "Données CRM actuelles du prospect", "Actif"),
    ("7", "Build Due Social Follow-Ups", "Code", "Relance", "Construit la liste des prospects à relancer", "Snapshot Supabase", "Prospect(s) éligibles à une relance", "Actif"),
    ("8", "Normalize Google Forms Lead", "Code", "Nouveau prospect", "Nettoie et normalise les réponses Google Forms", "Payload webhook", "Structure de données homogène", "Actif"),
    ("9", "Prepare Social Prospect Record", "Code", "Nouveau prospect", "Prépare l’enregistrement CRM à écrire dans Supabase", "Lead normalisé", "Objet prospect prêt pour upsert", "Actif"),
    ("10", "Upsert Social Prospect Into CRM", "HTTP Request", "Nouveau prospect", "Crée ou met à jour le prospect dans Supabase", "Record prospect préparé", "Prospect stocké ou mis à jour", "Actif"),
    ("11", "Build Immediate Social Send Context", "Code", "Nouveau prospect", "Prépare le contexte d’envoi immédiat du mail prospect", "Prospect CRM", "Contexte d’email 1 immédiat", "Actif"),
    ("12", "If Social Lead Ready To Send", "If", "Prospect", "Vérifie si le prospect est prêt à recevoir un email", "Contexte immédiat ou relance", "Prospect autorisé à entrer dans la séquence", "Actif"),
    ("13", "Build Social Sequence Email", "Code", "Prospect", "Génère le contenu du mail selon le type de formulaire et l’étape", "Prospect prêt à envoyer", "Sujet, HTML et métadonnées du mail", "Actif"),
    ("14", "Build Social Sent Guard", "Code", "Prospect", "Détermine si l’étape email concernée a déjà été envoyée", "Données prospect + état CRM", "Drapeau anti-doublon et champ ciblé", "Actif"),
    ("15", "If Social Send Allowed", "If", "Prospect", "Bloque l’envoi si le mail de cette étape a déjà été envoyé", "Guard prospect", "Aiguillage vrai ou faux", "Actif"),
    ("16", "Send Social Sequence Email", "HTTP Request", "Prospect", "Envoie l’email prospect via Resend", "Email généré et autorisé", "Identifiant Resend de l’envoi", "Actif"),
    ("17", "Parse Social Send Result", "Code", "Prospect", "Analyse la réponse de l’envoi email", "Réponse Resend", "Résultat exploitable pour la suite", "Actif"),
    ("18", "Build Social Send Result", "Code", "Prospect", "Construit les champs CRM après envoi réussi", "Résultat parse + état du prospect", "Horodatage du bon email et statut de séquence", "Actif"),
    ("19", "If Social Email Sent", "If", "Prospect", "Vérifie que l’envoi prospect est bien confirmé", "Résultat d’envoi prospect", "Branche succès ou échec", "Actif"),
    ("20", "Log Social Outreach Attempt", "HTTP Request", "Prospect", "Enregistre la tentative d’outreach dans Supabase", "Envoi confirmé", "Log d’activité CRM", "Actif"),
    ("21", "Update Prospect After Social Send", "HTTP Request", "Prospect", "Met à jour le prospect après succès d’envoi", "Résultat enrichi", "CRM mis à jour avec statut et horodatage", "Actif"),
    ("22", "Update Prospect Social Failure", "HTTP Request", "Prospect", "Met à jour le prospect en cas d’échec d’envoi", "Branche échec", "CRM mis à jour avec signal d’échec", "Actif"),
    ("23", "Build Admin Lead Alert Context", "Code", "Admin", "Prépare les informations à envoyer à l’équipe interne", "Prospect CRM", "Contexte complet d’alerte admin", "Actif"),
    ("24", "If Admin Alert Eligible", "If", "Admin", "Vérifie qu’une alerte interne doit être générée", "Contexte admin", "Branche éligible ou non", "Actif"),
    ("25", "Build Admin Lead Alert Email", "Code", "Admin", "Génère le sujet et le contenu du mail admin", "Prospect éligible", "Email interne prêt à envoyer", "Actif"),
    ("26", "Build Admin Alert Guard", "Code", "Admin", "Vérifie si une alerte admin a déjà été envoyée", "Prospect + CRM", "Drapeau anti-doublon admin", "Actif"),
    ("27", "If Admin Alert Allowed", "If", "Admin", "Autorise ou bloque l’envoi admin selon l’historique", "Guard admin", "Branche vraie ou fausse", "Actif"),
    ("28", "Send Admin Lead Alert", "HTTP Request", "Admin", "Envoie l’alerte à contact@transferai.ci via Resend", "Email admin autorisé", "Identifiant Resend admin", "Actif"),
    ("29", "Build Admin Alert Send Result", "Code", "Admin", "Construit les champs de retour après envoi admin", "Réponse Resend admin", "Horodatage admin et identifiant message", "Actif"),
    ("30", "Update Prospect After Admin Alert", "HTTP Request", "Admin", "Met à jour le prospect après alerte admin envoyée", "Résultat admin", "admin_alert_sent_at mis à jour", "Actif"),
]


def set_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)


def set_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x17, 0x2B, 0x4D)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(10)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7A)

    for style_name, size in [("Heading 1", 13), ("Heading 2", 11)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F2")
    p_bdr.append(bottom)
    p_pr.insert_element_before(
        p_bdr,
        "w:shd",
        "w:tabs",
        "w:spacing",
        "w:ind",
        "w:jc",
        "w:rPr",
        "w:sectPr",
    )


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(font_size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_docx() -> Path:
    document = Document()
    set_page_geometry(document)
    set_styles(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Tableau complet du workflow Google Forms prospects TransferAI")
    add_bottom_border(title)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Version opérationnelle validée au 23 juin 2026")

    intro = document.add_paragraph(style="Normal")
    intro.paragraph_format.space_after = Pt(8)
    intro.add_run(
        "Ce document synthétise l’architecture réellement exploitée du workflow n8n, "
        "y compris la branche prospect, la branche relance et la branche alerte interne."
    )

    info = document.add_paragraph(style="Normal")
    info.paragraph_format.space_after = Pt(10)
    info.add_run("Source de référence : ")
    info.add_run(str(SOURCE_MD)).italic = True

    headers = [
        "#",
        "Nœud",
        "Type",
        "Branche",
        "Rôle opérationnel",
        "Entrée principale",
        "Sortie attendue",
        "Statut",
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    widths = [Cm(0.8), Cm(4.2), Cm(2.5), Cm(2.4), Cm(6.2), Cm(4.5), Cm(4.8), Cm(2.4)]

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].width = widths[idx]
        set_cell_text(hdr_cells[idx], header, bold=True, font_size=8)
        shade_cell(hdr_cells[idx], "DCE6F1")

    for row_data in ROWS:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].width = widths[idx]
            set_cell_text(cells[idx], value, bold=False, font_size=8)
            if idx in (0, 2, 3, 7):
                shade_cell(cells[idx], "F7FBFF")

    document.add_paragraph("")

    h1 = document.add_paragraph(style="Heading 1")
    h1.add_run("Colonnes de contrôle déjà utilisées en base")

    for item in [
        "admin_alert_sent_at",
        "social_email_1_sent_at",
        "social_email_2_sent_at",
        "social_email_3_sent_at",
        "last_sequence_result",
        "last_response_status",
        "next_action_at",
        "niche_status",
    ]:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)

    h2 = document.add_paragraph(style="Heading 1")
    h2.add_run("Lecture opérationnelle rapide")

    for item in [
        "Un seul point d’entrée Apps Script doit alimenter le webhook n8n.",
        "Le mail prospect et l’alerte admin doivent rester sur deux circuits séparés.",
        "Le lien de rendez-vous actif est https://calendly.com/contact-transferai/30min.",
        "Les guards prospect et admin empêchent les doublons d’envoi.",
        "Le mail 1 prospect et l’alerte admin sont opérationnels au 23/06/2026.",
    ]:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_docx()
    print(path)
