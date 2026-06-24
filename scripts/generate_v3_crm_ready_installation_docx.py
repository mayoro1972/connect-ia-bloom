from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
V3_JSON = DOCS_DIR / "62_n8n_Prospection_Multi_Prospect_V3_CRM_Ready_Export.json"
TODAY = date.today().isoformat()
OUTPUT = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V3_CRM_Ready_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B)),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Installation complète du workflow V3 CRM ready dans n8n")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Approbation, envoi automatique et mise à jour du CRM")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version générée le {TODAY}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.add_run(
        "Ce guide décrit l’installation de la V3 finale prête à exporter. "
        "Elle génère le pack prospect, gère l’approbation interne, envoie automatiquement "
        "le message au prospect et met à jour le CRM opérationnel dans Supabase."
    ).font.size = Pt(10.5)
    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=90, start=110, bottom=90, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def load_workflow() -> dict:
    return json.loads(V3_JSON.read_text(encoding="utf-8"))


def list_nodes(data: dict) -> list[list[str]]:
    return [[str(i), node["name"], node["type"].split(".")[-1]] for i, node in enumerate(data["nodes"], 1)]


def main() -> None:
    data = load_workflow()
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    add_heading(doc, "1. Finalité de cette V3 CRM ready", 1)
    add_bullets(doc, [
        "générer un pack prospect à partir de signaux publics assainis",
        "faire approuver le pack en interne avant envoi",
        "envoyer automatiquement le courrier au prospect après approbation",
        "journaliser l’envoi dans outreach_attempts",
        "mettre à jour prospect_targets selon le résultat métier",
    ])

    add_heading(doc, "2. Variables d’environnement à prévoir", 1)
    add_table(
        doc,
        ["Variable", "Utilité"],
        [
            ["OPENAI_API_KEY", "clé API OpenAI utilisée par les nœuds d’analyse et de génération"],
            ["OPENAI_MODEL", "modèle par défaut, par exemple gpt-4.1-mini"],
            ["BOOKING_LINK_45MIN", "lien de rendez-vous gratuit de 45 minutes"],
            ["SUPABASE_URL", "URL du projet Supabase"],
            ["SUPABASE_SERVICE_ROLE_KEY", "clé service_role pour ai_prospecting_packs, outreach_attempts et prospect_targets"],
            ["RESEND_API_KEY", "clé Resend pour les e-mails internes et externes"],
            ["OUTREACH_FROM_EMAIL", "adresse d’expédition des e-mails"],
            ["INTERNAL_REVIEW_EMAIL", "adresse e-mail qui reçoit les demandes d’approbation et les confirmations"],
            ["N8N_BASE_URL", "URL publique n8n utilisée dans les liens de webhook"],
        ],
        widths=[2.15, 4.15],
    )

    add_heading(doc, "3. Champs prospect attendus à l’entrée", 1)
    add_bullets(doc, [
        "prospect_id",
        "organization_name",
        "website",
        "country",
        "organization_type",
        "sector_guess",
        "decision_maker_name",
        "target_email",
        "custom_page_paths_csv",
        "booking_link_45min",
        "commercial_priority_default",
        "research_scope",
        "source_backend",
        "source_label",
        "raw_source_id",
        "niche_status",
        "last_response_status",
        "last_sequence_result",
    ])

    add_heading(doc, "4. Ordre réel des nœuds", 1)
    add_table(doc, ["Ordre", "Nœud", "Type"], list_nodes(data), widths=[0.65, 3.5, 1.95])

    add_heading(doc, "5. Ce que cette V3 ajoute au pipeline", 1)
    add_table(
        doc,
        ["Bloc", "Rôle"],
        [
            ["Store Pack In Supabase", "enregistre le pack généré dans ai_prospecting_packs"],
            ["Approval Webhook", "reçoit l’action Approuver ou Rejeter"],
            ["Build Send Context", "vérifie que le pack peut vraiment partir"],
            ["If Ready To Send", "sépare l’envoi réel d’un cas approval_error"],
            ["Log Outreach Attempt", "journalise l’action commerciale dans outreach_attempts"],
            ["Update Prospect Target Sent", "met prospect_targets à active / outreach_started"],
            ["Update Prospect Target Approval Error", "met prospect_targets à paused / internal_fix_required"],
            ["Update Prospect Target Rejected", "met prospect_targets à ready / needs_manual_revision"],
        ],
        widths=[2.25, 4.0],
    )

    add_heading(doc, "6. Installation étape par étape", 1)
    add_numbered(doc, [
        "Importer le fichier JSON 62 dans n8n.",
        "Vérifier que les variables d’environnement OpenAI, Supabase, Resend et N8N_BASE_URL existent.",
        "Contrôler le nœud Set Target : il doit conserver includeOtherFields = true pour accepter les champs CRM.",
        "Configurer Supabase pour les nœuds Store Pack In Supabase, Mark Pack Approved, Mark Pack Sent, Mark Pack Approval Error, Mark Pack Rejected et les 3 nœuds Update Prospect Target.",
        "Configurer Resend pour Send Internal Approval Email, Send External Prospect Email et Send Internal Sent Confirmation.",
        "Activer le workflow pour rendre le webhook approve-prospect-pack-v3 accessible.",
        "Lancer un test avec un prospect de démonstration.",
        "Vérifier l’approbation, l’envoi, le log outreach_attempts et la mise à jour de prospect_targets.",
    ])

    add_heading(doc, "7. Configuration des nœuds critiques", 1)
    add_table(
        doc,
        ["Nœud", "Cible", "Point à vérifier"],
        [
            ["Store Pack In Supabase", "ai_prospecting_packs", "status initial = pending_approval, payload complet, llm_redaction_summary présent"],
            ["Approval Webhook", "/webhook/approve-prospect-pack-v3", "workflow activé et N8N_BASE_URL cohérent"],
            ["Build Send Context", "n/a", "can_send doit être true uniquement si target_email et executive_letter existent"],
            ["Send External Prospect Email", "Resend", "to = target_email, html = executive_letter_html"],
            ["Log Outreach Attempt", "outreach_attempts", "métadonnées CRM incluses dans metadata"],
            ["Update Prospect Target Sent", "prospect_targets", "status = active, niche_status = outreach_started"],
            ["Update Prospect Target Approval Error", "prospect_targets", "status = paused, niche_status = internal_fix_required"],
            ["Update Prospect Target Rejected", "prospect_targets", "status = ready, niche_status = needs_manual_revision"],
        ],
        widths=[1.95, 1.6, 2.7],
    )

    add_heading(doc, "8. Logique CRM à retenir", 1)
    add_bullets(doc, [
        "V5 fait entrer les leads publics dans prospect_targets",
        "V4 lit prospect_targets et décide qui traiter aujourd’hui",
        "cette V3 traite un prospect puis réécrit son état CRM",
        "le LLM ne voit que des signaux publics assainis, jamais les données sensibles brutes",
    ])

    add_heading(doc, "9. Tests recommandés", 1)
    add_bullets(doc, [
        "cas 1 : pack approuvé et e-mail envoyé",
        "cas 2 : pack rejeté",
        "cas 3 : approval_error en supprimant target_email ou executive_letter",
        "contrôle des tables ai_prospecting_packs, outreach_attempts et prospect_targets après chaque cas",
    ])

    add_heading(doc, "10. Résultat attendu", 1)
    add_paragraph(
        doc,
        "À la fin, tu disposes d’une V3 exportable qui fonctionne comme moteur de prospection sortante connecté au CRM. "
        "Elle génère, fait approuver, envoie, journalise et met à jour la base prospects pour les étapes suivantes du pipeline."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
