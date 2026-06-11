from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
TODAY = date.today().isoformat()

V2_JSON = DOCS_DIR / "43_n8n_Prospection_Modele_Elton_V2_Supabase_Approval_Email.json"
V3_JSON = DOCS_DIR / "44_n8n_Prospection_Modele_Elton_V3_Approval_AutoSend.json"
V1_WORD = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V1_n8n_{TODAY}.docx"

V2_OUTPUT = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V2_n8n_{TODAY}.docx"
V3_OUTPUT = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V3_n8n_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color, bold in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43), True),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43), True),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B), True),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B), True),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold


def add_footer(document: Document, label: str) -> None:
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_title_page(document: Document, title: str, subtitle: str, intro: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    m = meta.add_run(f"Version générée le {TODAY}")
    m.font.size = Pt(10.5)
    m.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(intro)
    run.font.size = Pt(10.5)
    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=90, start=110, bottom=90, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def load_workflow(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def list_nodes(data: dict) -> list[list[str]]:
    rows = []
    for idx, node in enumerate(data["nodes"], 1):
        rows.append([str(idx), node["name"], node["type"].split(".")[-1]])
    return rows


def find_node(data: dict, name: str) -> dict:
    for node in data["nodes"]:
        if node["name"] == name:
            return node
    raise KeyError(name)


def step_rows_v2() -> list[list[str]]:
    return [
        ["21", "Store Pack In Supabase", "enregistre le prospect pack avec pack_id, payload et statut pending_approval"],
        ["22", "Build Approval Email", "construit les liens Approuver et Rejeter à partir de N8N_BASE_URL"],
        ["23", "Send Internal Approval Email", "envoie l’e-mail de validation interne via Resend"],
        ["24", "Approval Webhook", "reçoit la décision de validation via une URL publique"],
        ["25", "Parse Approval Query", "lit pack_id et decision depuis les paramètres du webhook"],
        ["26", "Get Pack From Supabase", "récupère le pack complet à partir du pack_id"],
        ["27", "Extract Pack Payload", "reconstitue le contexte nécessaire à l’envoi ou au rejet"],
        ["28", "If Approved", "sépare la branche approuvée et la branche rejetée"],
        ["29", "Mark Pack Approved", "met le pack au statut approved dans Supabase"],
        ["30", "Send External Prospect Email", "envoie le courrier au prospect via Resend"],
        ["31", "Log Outreach Attempt", "journalise l’envoi dans outreach_attempts"],
        ["32", "Approval Success Response", "retourne un message de confirmation au navigateur"],
        ["33", "Mark Pack Rejected", "met le pack au statut rejected dans Supabase"],
        ["34", "Rejection Response", "retourne un message de rejet au navigateur"],
    ]


def step_rows_v3() -> list[list[str]]:
    return [
        ["21", "Store Pack In Supabase", "enregistre le prospect pack avec statut pending_approval"],
        ["22", "Build Approval Email", "génère des liens d’approbation spécifiques à la V3"],
        ["23", "Send Internal Approval Email", "transmet la demande d’approbation interne"],
        ["24", "Approval Webhook", "reçoit la décision approved ou rejected"],
        ["25", "Parse Approval Query", "lit pack_id et decision"],
        ["26", "Get Pack From Supabase", "récupère le pack à envoyer"],
        ["27", "Extract Pack Payload", "prépare executive_letter_html et les données d’envoi"],
        ["28", "If Approved", "oriente vers approbation ou rejet"],
        ["29", "Build Send Context", "vérifie qu’il existe un e-mail cible et un contenu de courrier"],
        ["30", "If Ready To Send", "sépare la branche envoi valide et la branche erreur d’approbation"],
        ["31", "Mark Pack Approved", "marque le pack comme approved"],
        ["32", "Send External Prospect Email", "envoie automatiquement le message au prospect"],
        ["33", "Parse Send Result", "capture resend_id et sent_at"],
        ["34", "Mark Pack Sent", "met à jour le pack au statut sent avec horodatage"],
        ["35", "Log Outreach Attempt", "journalise l’envoi dans outreach_attempts"],
        ["36", "Send Internal Sent Confirmation", "notifie l’équipe qu’un envoi a réellement eu lieu"],
        ["37", "Approval Success Response", "confirme dans le navigateur que l’envoi est terminé"],
        ["38", "Mark Pack Approval Error", "marque approval_error si le pack approuvé ne peut pas partir"],
        ["39", "Approval Error Response", "retourne le message d’erreur fonctionnelle"],
        ["40", "Mark Pack Rejected", "met le pack au statut rejected"],
        ["41", "Rejection Response", "retourne la réponse de rejet"],
    ]


def add_env_table(document: Document, extra: list[list[str]] | None = None) -> None:
    rows = [
        ["OPENAI_API_KEY", "clé API OpenAI utilisée par les appels d’analyse et de génération"],
        ["OPENAI_MODEL", "modèle OpenAI par défaut, par exemple gpt-4.1-mini"],
        ["BOOKING_LINK_45MIN", "lien de rendez-vous gratuit de 45 minutes"],
        ["SUPABASE_URL", "URL du projet Supabase"],
        ["SUPABASE_SERVICE_ROLE_KEY", "clé service_role pour les tables REST Supabase"],
        ["RESEND_API_KEY", "clé Resend utilisée pour les e-mails internes et externes"],
        ["OUTREACH_FROM_EMAIL", "adresse d’expédition des e-mails"],
        ["INTERNAL_REVIEW_EMAIL", "boîte e-mail qui reçoit les approbations internes"],
        ["N8N_BASE_URL", "URL publique n8n utilisée pour générer les liens webhook"],
    ]
    if extra:
        rows.extend(extra)
    add_table(document, ["Variable", "Utilité"], rows, widths=[2.2, 4.1])


def build_v2_document() -> Path:
    data = load_workflow(V2_JSON)
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc, "TransferAI Africa - Installation workflow V2 n8n")
    add_title_page(
        doc,
        "Installation complète du workflow V2 dans n8n",
        "Supabase, approbation interne et envoi prospect",
        "Ce guide Word permet d’installer la V2 de bout en bout. Il reprend la base V1, ajoute le stockage Supabase, l’e-mail d’approbation interne, le webhook de décision et l’envoi du courrier au prospect après validation."
    )

    add_heading(doc, "1. Ce que la V2 ajoute à la V1", 1)
    add_bullets(doc, [
        "sauvegarde du prospect pack dans la table ai_prospecting_packs",
        "construction d’un e-mail interne de validation avec liens Approuver et Rejeter",
        "réception de la décision via un webhook public n8n",
        "récupération du pack depuis Supabase après la décision",
        "envoi de l’e-mail prospect après approbation",
        "journalisation de la tentative d’envoi dans outreach_attempts",
    ])

    add_heading(doc, "2. Pré-requis avant installation", 1)
    add_env_table(doc)
    add_paragraph(doc, f"La V2 suppose que la V1 est déjà comprise ou installée. Le guide Word V1 est disponible ici : {V1_WORD.name}")
    add_bullets(doc, [
        "créer ou avoir un projet Supabase actif",
        "créer les tables ai_prospecting_packs et outreach_attempts",
        "disposer d’un compte Resend ou d’un équivalent SMTP/API si tu adaptes le nœud",
        "disposer d’une URL publique n8n, car les liens d’approbation pointent vers le webhook",
    ])

    add_heading(doc, "3. Ordre réel des nœuds V2", 1)
    add_table(doc, ["Ordre", "Nœud", "Type"], list_nodes(data), widths=[0.7, 3.5, 2.0])

    add_heading(doc, "4. Nœuds supplémentaires introduits par la V2", 1)
    add_table(doc, ["Ordre", "Nœud", "Rôle"], step_rows_v2(), widths=[0.65, 2.2, 3.35])

    add_heading(doc, "5. Installation pas à pas", 1)
    add_numbered(doc, [
        "Importer d’abord le workflow JSON V2 dans n8n.",
        "Vérifier que les 20 premiers nœuds issus de la V1 sont présents et correctement configurés.",
        "Créer ou renseigner les variables d’environnement Supabase, Resend et N8N_BASE_URL.",
        "Configurer Store Pack In Supabase avec SUPABASE_URL et SUPABASE_SERVICE_ROLE_KEY.",
        "Configurer Build Approval Email pour générer les liens de validation à partir de N8N_BASE_URL.",
        "Configurer Send Internal Approval Email avec RESEND_API_KEY, OUTREACH_FROM_EMAIL et INTERNAL_REVIEW_EMAIL.",
        "Activer le workflow pour rendre le webhook d’approbation accessible.",
        "Configurer Get Pack From Supabase, Mark Pack Approved, Mark Pack Rejected et Log Outreach Attempt avec les variables Supabase.",
        "Configurer Send External Prospect Email pour utiliser RESEND_API_KEY et l’adresse d’expédition définie.",
        "Tester la chaîne complète avec un prospect de test et une adresse e-mail contrôlée.",
    ])

    add_heading(doc, "6. Configuration nœud par nœud", 1)
    node_rows = [
        ["Store Pack In Supabase", "POST", "{{$env.SUPABASE_URL}}/rest/v1/ai_prospecting_packs", "utiliser apikey et Authorization avec SUPABASE_SERVICE_ROLE_KEY ; ne pas modifier le body JSON déjà prêt"],
        ["Build Approval Email", "Code", "n/a", "laisser le code, mais vérifier N8N_BASE_URL pour générer les liens approve-prospect-pack"],
        ["Send Internal Approval Email", "POST", "https://api.resend.com/emails", "Authorization = Bearer {{$env.RESEND_API_KEY}} ; from = {{$env.OUTREACH_FROM_EMAIL}} ; to = {{$env.INTERNAL_REVIEW_EMAIL}}"],
        ["Approval Webhook", "Webhook GET", "/webhook/approve-prospect-pack", "le workflow doit être activé pour que le lien fonctionne"],
        ["Get Pack From Supabase", "GET", "rest/v1/ai_prospecting_packs?select=*&pack_id=eq.<pack_id>&limit=1", "utiliser les en-têtes apikey et Authorization Supabase"],
        ["Mark Pack Approved", "PATCH", "rest/v1/ai_prospecting_packs?pack_id=eq.<pack_id>", "écrit status = approved et approved_at"],
        ["Send External Prospect Email", "POST", "https://api.resend.com/emails", "envoie executive_letter converti en HTML vers target_email"],
        ["Log Outreach Attempt", "POST", "rest/v1/outreach_attempts", "journaliser pack_id, prospect_id, organization_name, target_email et statut submitted"],
        ["Mark Pack Rejected", "PATCH", "rest/v1/ai_prospecting_packs?pack_id=eq.<pack_id>", "écrit status = rejected et rejected_at"],
    ]
    add_table(doc, ["Nœud", "Méthode", "Cible", "Points à vérifier"], node_rows, widths=[1.75, 0.8, 1.95, 2.1])

    add_heading(doc, "7. Flux fonctionnel de validation", 1)
    add_bullets(doc, [
        "le workflow V2 génère le pack comme en V1 puis le stocke dans Supabase",
        "un e-mail interne de validation est envoyé à l’équipe",
        "au clic sur Approuver, n8n lit pack_id et decision dans le webhook",
        "le pack est relu depuis Supabase pour éviter toute perte de contexte",
        "si decision = approved, le pack passe à approved puis l’e-mail prospect part",
        "si decision = rejected, le pack est simplement marqué rejected",
    ])

    add_heading(doc, "8. Test minimal recommandé", 1)
    add_bullets(doc, [
        "tester d’abord avec une adresse prospect de test",
        "vérifier qu’un enregistrement est bien créé dans ai_prospecting_packs",
        "vérifier la réception de l’e-mail interne de validation",
        "cliquer sur Approuver et contrôler que le webhook répond",
        "vérifier que le statut passe à approved dans Supabase",
        "vérifier que l’e-mail externe est reçu et qu’un log existe dans outreach_attempts",
    ])

    add_heading(doc, "9. Résultat attendu", 1)
    add_paragraph(doc, "À la fin, la V2 permet de générer un prospect pack, de le faire valider par l’équipe et d’envoyer ensuite le courrier au prospect avec une traçabilité minimale dans Supabase.")

    doc.save(V2_OUTPUT)
    return V2_OUTPUT


def build_v3_document() -> Path:
    data = load_workflow(V3_JSON)
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_footer(doc, "TransferAI Africa - Installation workflow V3 n8n")
    add_title_page(
        doc,
        "Installation complète du workflow V3 dans n8n",
        "Approbation, garde-fou d’envoi et auto-envoi prospect",
        "Ce guide Word décrit l’installation complète de la V3. La V3 reprend toute la V2, puis ajoute un contrôle de sendabilité, un statut sent, une confirmation interne après envoi et une branche d’erreur si le pack approuvé ne peut pas partir."
    )

    add_heading(doc, "1. Ce que la V3 ajoute à la V2", 1)
    add_bullets(doc, [
        "contrôle intermédiaire Build Send Context pour vérifier que l’e-mail cible et le courrier existent",
        "branche If Ready To Send pour distinguer un pack approuvé mais incomplet",
        "capture du résultat Resend avec resend_id et sent_at",
        "mise à jour du pack au statut sent dans Supabase",
        "notification interne après envoi effectif",
        "branche approval_error quand le pack est approuvé mais ne peut pas être expédié",
    ])

    add_heading(doc, "2. Pré-requis avant installation", 1)
    add_env_table(doc)
    add_bullets(doc, [
        "avoir déjà une logique V2 comprise ou opérationnelle",
        "conserver le même schéma Supabase que la V2",
        "utiliser une URL n8n publique différente uniquement si tu veux séparer les webhooks V2 et V3",
        "prévoir des adresses de test pour la validation et l’envoi final",
    ])

    add_heading(doc, "3. Ordre réel des nœuds V3", 1)
    add_table(doc, ["Ordre", "Nœud", "Type"], list_nodes(data), widths=[0.7, 3.5, 2.0])

    add_heading(doc, "4. Nœuds supplémentaires introduits ou modifiés dans la V3", 1)
    add_table(doc, ["Ordre", "Nœud", "Rôle"], step_rows_v3(), widths=[0.65, 2.25, 3.3])

    add_heading(doc, "5. Installation pas à pas", 1)
    add_numbered(doc, [
        "Importer le workflow JSON V3 dans n8n.",
        "Vérifier que les 28 premiers nœuds communs à V2 sont présents.",
        "Configurer Store Pack In Supabase, Send Internal Approval Email et Get Pack From Supabase comme en V2.",
        "Vérifier que le webhook V3 utilise bien le chemin approve-prospect-pack-v3.",
        "Configurer Send External Prospect Email avec Resend et une adresse d’expédition autorisée.",
        "Contrôler Build Send Context et If Ready To Send pour que la branche erreur soit claire en cas de cible ou contenu manquant.",
        "Configurer Mark Pack Sent, Mark Pack Approval Error et Send Internal Sent Confirmation avec Supabase et Resend.",
        "Activer le workflow puis tester un cycle complet approved, un cycle rejected et un cycle approval_error.",
    ])

    add_heading(doc, "6. Configuration nœud par nœud", 1)
    node_rows = [
        ["Approval Webhook", "Webhook GET", "/webhook/approve-prospect-pack-v3", "chemin distinct de la V2 ; nécessite workflow activé"],
        ["Build Send Context", "Code", "n/a", "vérifie target_email et executive_letter ; définit can_send et send_failure_reason"],
        ["If Ready To Send", "If", "n/a", "branche true si can_send = true ; branche false vers approval_error"],
        ["Send External Prospect Email", "POST", "https://api.resend.com/emails", "utilise executive_letter_html ; to = target_email"],
        ["Parse Send Result", "Code", "n/a", "capture resend_id et sent_at après la réponse Resend"],
        ["Mark Pack Sent", "PATCH", "rest/v1/ai_prospecting_packs?pack_id=eq.<pack_id>", "écrit status = sent, approved_at, sent_at, resend_message_id"],
        ["Send Internal Sent Confirmation", "POST", "https://api.resend.com/emails", "envoie un e-mail interne de confirmation après expédition réelle"],
        ["Mark Pack Approval Error", "PATCH", "rest/v1/ai_prospecting_packs?pack_id=eq.<pack_id>", "écrit status = approval_error et error_reason"],
        ["Approval Error Response", "Respond to Webhook", "n/a", "retourne un texte du type Pack approuvé mais non envoyé"],
    ]
    add_table(doc, ["Nœud", "Type", "Cible", "Points à vérifier"], node_rows, widths=[1.8, 0.95, 1.75, 2.1])

    add_heading(doc, "7. Logique fonctionnelle de la V3", 1)
    add_bullets(doc, [
        "comme la V2, la V3 crée le pack et l’envoie en validation interne",
        "si approved, le workflow relit le pack depuis Supabase puis construit le contexte d’envoi",
        "si can_send = true, l’e-mail part, la réponse Resend est analysée et le pack passe à sent",
        "si can_send = false, le workflow ne tente pas l’envoi et marque approval_error",
        "une notification interne supplémentaire confirme qu’un prospect a réellement été contacté",
    ])

    add_heading(doc, "8. Tests recommandés", 1)
    add_bullets(doc, [
        "test 1 : approbation complète avec envoi réussi",
        "test 2 : rejet simple depuis le lien Rejeter",
        "test 3 : suppression volontaire de target_email pour vérifier la branche approval_error",
        "contrôle final dans Supabase : pending_approval, approved, sent, rejected, approval_error",
        "contrôle final dans Resend : e-mail interne de validation, e-mail prospect, e-mail interne de confirmation",
    ])

    add_heading(doc, "9. Résultat attendu", 1)
    add_paragraph(doc, "À la fin, la V3 fournit une chaîne complète : pack généré, approbation interne, contrôle d’envoi, auto-envoi prospect, mise à jour du statut sent et confirmation interne.")

    doc.save(V3_OUTPUT)
    return V3_OUTPUT


if __name__ == "__main__":
    v2 = build_v2_document()
    v3 = build_v3_document()
    print(v2)
    print(v3)
