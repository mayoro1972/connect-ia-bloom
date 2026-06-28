"""
Générateur de mini-catalogue personnalisé prospect TransferAI
Nouveau modèle : ancré sur le pré-audit, formations ciblées avant/après, ROI concret.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── Couleurs marque ───────────────────────────────────────────────────
BLEU_FONCE  = RGBColor(0x10, 0x26, 0x3F)
ORANGE_CI   = RGBColor(0xE7, 0x6F, 0x1D)
OR_AMBRE    = RGBColor(0xC8, 0x8C, 0x3A)
TEAL        = RGBColor(0x1C, 0x8A, 0x78)
BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
CREME_BG    = RGBColor(0xF5, 0xEF, 0xE7)
BLEU_TEXTE  = RGBColor(0x61, 0x70, 0x86)
GRIS_CLAIR  = RGBColor(0xF1, 0xEF, 0xE8)


# ════════════════════════════════════════════════════════════════════
# ▶  DONNÉES PROSPECT — modifier pour chaque nouveau catalogue
# ════════════════════════════════════════════════════════════════════

PROSPECT     = "Orange Côte d'Ivoire"
SECTEUR      = "Télécommunications"
PAYS         = "Côte d'Ivoire"
PACK_ID      = "pack-1780665567601-jp48u6z4"
DATE_DOC     = "Juin 2026"

AUDIT_URL    = f"https://www.transferai.ci/questionnaire-audit?pack_id={PACK_ID}"
CALENDLY_URL = "https://calendly.com/contact-transferai/30min"

# 3 enjeux détectés lors du pré-audit
ENJEUX_AUDIT = [
    {
        "num": "01",
        "titre": "Qualité et cohérence des réponses service client",
        "constat": (
            "Les canaux de réponse (téléphone, email, WhatsApp, agences) produisent "
            "des réponses hétérogènes selon les agents et les équipes. Les réclamations "
            "récurrentes manquent de scripts unifiés, entraînant des délais de traitement "
            "allongés et une insatisfaction mesurable."
        ),
        "signal": "Temps moyen de réponse élevé · Taux de réclamations répétées · Qualité inégale entre canaux",
    },
    {
        "num": "02",
        "titre": "Préparation et exploitation des reportings managériaux",
        "constat": (
            "La consolidation des indicateurs de performance (taux de churn, NPS, "
            "incidents réseau, ventes mobile money) reste manuelle et chronophage. "
            "Les comités de direction disposent de données mais sans synthèses "
            "exploitables, ce qui ralentit les prises de décision."
        ),
        "signal": "Temps de préparation des comités · Délai d'analyse des verbatims · Décisions ralenties",
    },
    {
        "num": "03",
        "titre": "Structuration et diffusion des procédures terrain",
        "constat": (
            "Les procédures opérationnelles, scripts d'intervention et bonnes pratiques "
            "sont dispersés dans des fichiers non centralisés. L'onboarding de nouveaux "
            "agents est long, la montée en compétences inégale selon les équipes et les régions."
        ),
        "signal": "Durée d'onboarding · Écarts de qualité entre régions · Procédures non à jour",
    },
]

# 3 formations ciblées tirées du pré-audit
FORMATIONS = [
    {
        "num": "F1",
        "intitule": "Copilote service client multicanal avec l'IA",
        "public": "Superviseurs service client, chefs d'équipe front office, responsables relation client",
        "duree": "2 jours (présentiel ou hybride)",
        "niveau": "Débutant à intermédiaire",
        "probleme": (
            "Les agents répondent différemment selon leur expérience. "
            "Les scripts existants sont éparpillés, rarement mis à jour et peu utilisés en situation réelle."
        ),
        "apres": (
            "Chaque agent dispose d'un copilote IA pour formuler, compléter et valider "
            "ses réponses selon le canal (téléphone, email, WhatsApp). "
            "Les réponses sont plus homogènes, traçables et conformes à la charte qualité."
        ),
        "objectifs": [
            "Utiliser un assistant IA pour structurer et améliorer les réponses client en temps réel",
            "Adapter le ton, le format et la précision selon le canal et le motif de contact",
            "Réduire les escalades inutiles grâce à une meilleure qualification initiale",
            "Mettre en place un système de supervision humaine des réponses générées",
        ],
        "livrables": [
            "Bibliothèque de scripts IA validés pour les 10 motifs de contact les plus fréquents",
            "Canevas de contrôle qualité des réponses assistées par IA",
            "Plan de déploiement du copilote sur les équipes ciblées",
        ],
        "kpi": "Délai moyen de réponse · Taux de résolution au 1er contact · Score satisfaction client",
        "gain": "− 30 à 50 % sur le délai moyen · + homogénéité mesurable dès J+30",
    },
    {
        "num": "F2",
        "intitule": "Synthèse IA des reportings et pilotage par les données",
        "public": "Directeurs, managers opérationnels, responsables qualité de service et contrôle de gestion",
        "duree": "1 jour (présentiel, format décideurs)",
        "niveau": "Intermédiaire",
        "probleme": (
            "Les données existent mais les synthèses sont réalisées manuellement, "
            "souvent par les mêmes 2 ou 3 personnes. Les comités consomment l'information "
            "sans pouvoir l'interroger en profondeur. Les signaux de churn et "
            "d'insatisfaction sont détectés trop tard."
        ),
        "apres": (
            "Les managers reçoivent des synthèses automatisées des indicateurs clés avant "
            "chaque comité. Les verbatims clients sont analysés en quelques minutes. "
            "Les signaux de churn sont détectés automatiquement et remontés aux équipes concernées."
        ),
        "objectifs": [
            "Produire des synthèses managériales à partir de données existantes (CRM, tickets, NPS)",
            "Analyser automatiquement les verbatims et classer les motifs d'insatisfaction",
            "Détecter les signaux de churn et construire un tableau de bord IA simple",
            "Définir un cadre de validation humaine avant diffusion des synthèses",
        ],
        "livrables": [
            "Modèle de synthèse IA pour comité de direction (format hebdomadaire)",
            "Tableau de bord des signaux clients priorisés",
            "Procédure de validation et de diffusion des analyses IA",
        ],
        "kpi": "Temps de préparation des comités · Délai de détection des signaux · Qualité des décisions",
        "gain": "− 60 % sur le temps de préparation reporting · Décisions prises sur données fraîches",
    },
    {
        "num": "F3",
        "intitule": "Base de connaissances IA et onboarding accéléré",
        "public": "Responsables formation, chefs de service terrain, RH et encadrement opérationnel",
        "duree": "1 jour (présentiel) + 1 session en ligne de suivi à J+30",
        "niveau": "Débutant",
        "probleme": (
            "Les procédures sont dispersées dans des fichiers Word, emails et drives partagés "
            "rarement à jour. L'onboarding d'un nouvel agent prend entre 3 et 6 semaines. "
            "Les écarts de pratique entre agences et régions sont importants."
        ),
        "apres": (
            "Une base de connaissances IA centralise toutes les procédures, scripts et "
            "bonnes pratiques. Les nouveaux agents la consultent directement pendant "
            "leur activité. L'onboarding est réduit à 1 à 2 semaines. "
            "Les mises à jour sont propagées instantanément à toutes les équipes."
        ),
        "objectifs": [
            "Structurer et centraliser les procédures opérationnelles dans une base IA interrogeable",
            "Permettre aux agents de trouver la bonne réponse en moins de 30 secondes",
            "Réduire le temps d'onboarding des nouveaux agents grâce à un guide IA interactif",
            "Définir un processus de mise à jour et de validation des procédures",
        ],
        "livrables": [
            "Architecture de la base de connaissances IA (structure, catégories, droits d'accès)",
            "Guide de migration des procédures existantes vers la base IA",
            "Protocole de mise à jour mensuelle validé par les responsables terrain",
        ],
        "kpi": "Durée d'onboarding · Nombre de consultations de la base · Écarts de pratique inter-agences",
        "gain": "Onboarding − 50 % · Procédures accessibles en temps réel · Qualité homogène entre agences",
    },
]

OUT_PATH = f"/Users/marius_ayoro/Downloads/MiniCatalogue_TransferAI_{PROSPECT.replace(' ', '_').replace(chr(39), '')}_V2.docx"


# ════════════════════════════════════════════════════════════════════
# HELPERS DOCX
# ════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in kwargs.items():
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def p_text(doc, text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
           space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def p_label(doc, text, color=ORANGE_CI):
    """Étiquette de section avec barre colorée gauche simulée."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(f"▌ {text}")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color
    return p

def p_bullet(doc, text, size=10.5, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def add_hr(doc, color="E76F1D"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def banner(doc, text, bg="10263F", fg=BLANC, size=13):
    """Bloc coloré pleine largeur simulé via tableau 1 cellule."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    cell.width = Cm(17)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = fg
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def avant_apres_table(doc, avant_text, apres_text):
    """Tableau 2 colonnes Avant / Après."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    cell_av = tbl.rows[0].cells[0]
    cell_ap = tbl.rows[0].cells[1]
    set_cell_bg(cell_av, "FAECE7")
    set_cell_bg(cell_ap, "E1F5EE")

    for cell, label, label_color, text, text_color in [
        (cell_av, "AVANT", "993C1D", avant_text, "712B13"),
        (cell_ap, "APRÈS", "0F6E56", apres_text, "085041"),
    ]:
        p_lbl = cell.paragraphs[0]
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after  = Pt(2)
        r = p_lbl.add_run(label)
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(label_color)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        r2 = p2.add_run(text)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor.from_string(text_color)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def kpi_row(doc, gain_text):
    """Ligne mise en valeur gain / KPI."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "163556")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run(f"  Gain attendu  ·  {gain_text}")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLANC
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════════════════════════
# CONSTRUCTION DU DOCUMENT
# ════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── EN-TÊTE ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TransferAI Africa  ·  www.transferai.ci  ·  IA Assistant / WhatsApp : +225 07 16 57 39 90")
    r.font.size = Pt(8)
    r.font.color.rgb = BLEU_TEXTE
    add_hr(doc)

    # ── COUVERTURE ────────────────────────────────────────────────────
    p_text(doc, "MINI-CATALOGUE PERSONNALISÉ", size=8, bold=True,
           color=ORANGE_CI, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=2)
    p_text(doc, f"Formations IA prioritaires recommandées pour", size=22, bold=True,
           color=BLEU_FONCE, space_before=0, space_after=4)
    p_text(doc, PROSPECT, size=22, bold=True, color=ORANGE_CI, space_before=0, space_after=6)
    p_text(doc, f"{SECTEUR}  ·  {PAYS}  ·  {DATE_DOC}",
           size=10, color=BLEU_TEXTE, space_before=0, space_after=8)

    banner(doc,
           "Ce document a été construit à partir des enjeux identifiés lors de votre pré-audit TransferAI. "
           "Il ne présente pas un catalogue général — il sélectionne les 3 formations les plus utiles "
           "pour vos équipes, avec les gains attendus, les livrables remis et le plan d'action associé.",
           bg="10263F", size=10)

    add_hr(doc)

    # ── SECTION 00 — Ce document vous est adressé ────────────────────
    p_label(doc, f"00 — Ce document vous est adressé, {PROSPECT}")
    p_text(doc,
           f"À la suite de votre pré-audit TransferAI, nous avons identifié trois enjeux "
           f"prioritaires dans vos opérations {SECTEUR.lower()}. Ce mini-catalogue ne présente pas "
           f"toutes nos formations — il sélectionne uniquement celles qui correspondent à ce que "
           f"vos équipes vivent réellement aujourd'hui.",
           size=10.5, space_after=4)
    p_text(doc,
           "Chaque formation présentée inclut : la situation constatée, ce qui change après la "
           "formation, les objectifs pédagogiques, les livrables remis et le gain attendu mesurable.",
           size=10.5, italic=True, color=BLEU_TEXTE, space_after=10)

    # ── SECTION 01 — Enjeux du pré-audit ────────────────────────────
    p_label(doc, "01 — Ce que le pré-audit révèle")

    for enjeu in ENJEUX_AUDIT:
        banner(doc, f"{enjeu['num']}  ·  {enjeu['titre']}", bg="163556", size=11)
        p_text(doc, enjeu["constat"], size=10.5, space_after=4)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("Signaux observés  ·  ")
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = ORANGE_CI
        r2 = p.add_run(enjeu["signal"])
        r2.font.size = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = BLEU_TEXTE

    add_hr(doc)

    # ── SECTION 02 — Formations prioritaires ────────────────────────
    p_label(doc, "02 — Formations prioritaires recommandées")
    p_text(doc,
           "Ces 3 formations ont été sélectionnées spécifiquement pour Orange Côte d'Ivoire "
           "sur la base des enjeux détectés. Elles sont conçues pour être opérationnelles dès "
           "la fin de la formation et non théoriques.",
           size=10.5, space_after=10)

    for f in FORMATIONS:
        # En-tête formation
        banner(doc, f"{f['num']}  ·  {f['intitule']}", bg="E76F1D", fg=BLANC, size=12)

        # Méta
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col_idx, (label, val) in enumerate([
            ("Public cible", f["public"]),
            ("Durée", f["duree"]),
            ("Niveau", f["niveau"]),
        ]):
            cell = tbl.rows[0].cells[col_idx]
            set_cell_bg(cell, "F5EFE7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            r1 = p.add_run(label + "\n")
            r1.font.size = Pt(8)
            r1.font.bold = True
            r1.font.color.rgb = ORANGE_CI
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = BLEU_FONCE
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Avant / Après
        p_text(doc, "Situation actuelle → Ce qui change", size=9, bold=True,
               color=BLEU_FONCE, space_before=4, space_after=2)
        avant_apres_table(doc, f["probleme"], f["apres"])

        # Objectifs
        p_text(doc, "Objectifs pédagogiques", size=10, bold=True,
               color=BLEU_FONCE, space_before=4, space_after=2)
        for obj in f["objectifs"]:
            p_bullet(doc, obj, size=10)

        # Livrables
        p_text(doc, "Livrables remis à la fin de la formation", size=10, bold=True,
               color=TEAL, space_before=6, space_after=2)
        for liv in f["livrables"]:
            p_bullet(doc, liv, size=10, color=TEAL)

        # Gain
        kpi_row(doc, f["gain"])

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_hr(doc)

    # ── SECTION 03 — Livrables globaux ──────────────────────────────
    p_label(doc, "03 — Livrables remis à l'issue du dispositif")

    livrables_globaux = [
        f"Note de cadrage {SECTEUR} : enjeux confirmés, périmètre du pilote, rôles et premières décisions.",
        "Bibliothèque de scripts IA validés pour les cas de contact les plus fréquents.",
        "Modèle de synthèse IA hebdomadaire pour les comités de direction.",
        "Architecture de la base de connaissances terrain avec procédures centralisées.",
        "Plan d'action à 90 jours : KPI cibles, équipes formées, supervision humaine et critères d'extension.",
    ]
    for liv in livrables_globaux:
        p_bullet(doc, liv, size=10.5)

    add_hr(doc)

    # ── SECTION 04 — Parcours 90 jours ──────────────────────────────
    p_label(doc, "04 — Parcours recommandé sur 90 jours")

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    etapes = [
        ("J+0", "Audit & cadrage", "Confirmation des 3 enjeux, sélection des 2 formations prioritaires, définition des KPI."),
        ("J+15", "Premières formations", "Formation F1 (service client) et F2 (reporting) avec équipes ciblées."),
        ("J+45", "Pilote & formation F3", "Déploiement du copilote, lancement de la base de connaissances, formation terrain."),
        ("J+90", "Mesure & extension", "Revue des KPI, décision d'extension sur les autres équipes et usages."),
    ]
    for i, (delai, titre, desc) in enumerate(etapes):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, "F5EFE7")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(2)
        r1 = p.add_run(delai + "\n")
        r1.font.size = Pt(14)
        r1.font.bold = True
        r1.font.color.rgb = ORANGE_CI
        r2 = p.add_run(titre + "\n")
        r2.font.size = Pt(10)
        r2.font.bold = True
        r2.font.color.rgb = BLEU_FONCE
        r3 = p.add_run(desc)
        r3.font.size = Pt(9)
        r3.font.color.rgb = BLEU_TEXTE

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_hr(doc)

    # ── SECTION 05 — CTA ─────────────────────────────────────────────
    p_label(doc, "05 — Prochaine étape")

    banner(doc,
           "Planifier un audit de cadrage gratuit (30 minutes) pour confirmer les priorités "
           "et définir ensemble le périmètre des premières formations.",
           bg="10263F", size=11)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell1, cell2 = tbl.rows[0].cells
    set_cell_bg(cell1, "244667")
    set_cell_bg(cell2, "244667")
    for cell, label, url in [
        (cell1, "Rendez-vous expert →", CALENDLY_URL),
        (cell2, "Formulaire d'audit →", AUDIT_URL),
    ]:
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r1 = p.add_run(label + "\n")
        r1.font.size = Pt(9)
        r1.font.bold = True
        r1.font.color.rgb = OR_AMBRE
        r2 = p.add_run(url)
        r2.font.size = Pt(8.5)
        r2.font.italic = True
        r2.font.color.rgb = BLANC

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── PIED DE PAGE ─────────────────────────────────────────────────
    add_hr(doc, color="C9D4E1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(
        "TransferAI Africa  ·  Hub IA de NettelecomCI  ·  www.transferai.ci  "
        "·  contact@transferai.ci  ·  +225 07 16 57 39 90"
    )
    r.font.size = Pt(7.5)
    r.font.color.rgb = BLEU_TEXTE

    doc.save(OUT_PATH)
    print(f"✓ Catalogue : {OUT_PATH}")
    return OUT_PATH


if __name__ == "__main__":
    build()
