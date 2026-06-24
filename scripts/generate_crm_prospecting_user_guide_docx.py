from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
TODAY = date.today().isoformat()

ATTACHED_V3 = Path("/Users/marius_ayoro/Downloads/TransferAI Prospecting V3 CRM Enhanced [FINAL]-10.json")
FALLBACK_V3 = DOCS_DIR / "62_n8n_Prospection_Multi_Prospect_V3_CRM_Ready_Export.json"

OUTPUT = WORD_DIR / f"TransferAI_Africa_Guide_Utilisateur_CRM_Prospecting_V3_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B)),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title_page(document: Document, workflow_name: str, node_count: int, source_file: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Guide utilisateur - CRM Prospecting V3")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Version opérationnelle du workflow de prospection intelligente CRM Enhanced")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(f"Version générée le {TODAY}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = (
        f"Workflow de référence : {workflow_name}. "
        f"Fichier source : {source_file}. "
        f"État réel : {node_count} nœuds. "
        "Ce guide utilisateur reprend le modèle du guide d’installation V3 CRM ready, "
        "mais l’adapte à la version CRM prospecting enrichie avec génération de pièces jointes, "
        "approbation, envoi, mises à jour CRM et rendu catalogue / deck."
    )
    run = para.add_run(text)
    run.font.size = Pt(10.5)
    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=90, start=110, bottom=90, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def load_v3() -> tuple[dict, Path]:
    path = ATTACHED_V3 if ATTACHED_V3.exists() else FALLBACK_V3
    return json.loads(path.read_text(encoding="utf-8")), path


def list_nodes(data: dict) -> list[list[str]]:
    return [[str(i), node["name"], node["type"].split(".")[-1]] for i, node in enumerate(data["nodes"], 1)]


def find_node_names(data: dict, names: list[str]) -> list[list[str]]:
    out: list[list[str]] = []
    index = {node["name"]: i + 1 for i, node in enumerate(data["nodes"])}
    for name in names:
        if name in index:
            node = next(n for n in data["nodes"] if n["name"] == name)
            out.append([str(index[name]), name, node["type"].split(".")[-1]])
    return out


def add_footer(document: Document, label: str) -> None:
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main() -> None:
    data, source_path = load_v3()
    WORD_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style_document(doc)
    add_title_page(doc, data.get("name", "Workflow V3 CRM"), len(data.get("nodes", [])), source_path.name)

    add_heading(doc, "1. Finalité du guide", 1)
    add_paragraph(
        doc,
        "Ce document sert de guide utilisateur opérationnel pour la version CRM prospecting du workflow V3. "
        "Il est conçu pour un usage quotidien par l’équipe TransferAI : lecture rapide, contrôle des nœuds critiques, "
        "compréhension de la logique d’approbation, suivi CRM, rendu des pièces jointes et dépannage des cas les plus fréquents."
    )

    add_heading(doc, "2. Ce que fait réellement cette version CRM prospecting", 1)
    add_bullets(doc, [
        "scrape et normalise les signaux publics d’un prospect",
        "protège les données sensibles avant exposition au LLM",
        "génère un courrier, un catalogue, une forme d’audit et un brief de deck",
        "stocke un pack prospect dans Supabase",
        "envoie un e-mail d’approbation interne",
        "sur approbation, envoie le courrier prospect avec pièces jointes",
        "met à jour ai_prospecting_packs, outreach_attempts et prospect_targets",
        "rend dynamiquement un catalogue PDF et un deck PPTX avant envoi",
    ])

    add_heading(doc, "3. Variables d’environnement à prévoir", 1)
    add_table(
        doc,
        ["Variable", "Utilité"],
        [
            ["OPENAI_API_KEY", "clé API OpenAI utilisée par les nœuds d’analyse et de génération"],
            ["OPENAI_MODEL", "modèle OpenAI, par exemple gpt-4.1-mini"],
            ["SUPABASE_URL", "URL du projet Supabase"],
            ["SUPABASE_SERVICE_ROLE_KEY", "clé service_role utilisée pour packs, logs, CRM et rendu"],
            ["RESEND_API_KEY", "clé utilisée pour les e-mails internes et prospect"],
            ["OUTREACH_FROM_EMAIL", "adresse d’envoi du courrier prospect"],
            ["INTERNAL_REVIEW_EMAIL", "adresse de réception de l’e-mail d’approbation"],
            ["N8N_BASE_URL", "URL publique n8n utilisée pour le webhook d’approbation"],
            ["BOOKING_LINK_45MIN", "lien de rendez-vous de 45 minutes"],
        ],
        widths=[2.1, 4.2],
    )

    add_heading(doc, "4. Tables et objets utilisés", 1)
    add_bullets(doc, [
        "ai_prospecting_packs : stocke le pack complet et ses statuts",
        "outreach_attempts : journalise les envois effectués",
        "prospect_targets : stocke l’état CRM du prospect",
        "prospecting-artifacts : bucket ou mécanisme de rendu utilisé pour le catalogue et le deck",
        "fonctions de rendu catalogue-renderer et deck-renderer si elles sont déployées côté Supabase",
    ])

    add_heading(doc, "5. Ordre réel des nœuds", 1)
    add_table(doc, ["Ordre", "Nœud", "Type"], list_nodes(data), widths=[0.7, 3.55, 1.9])

    add_heading(doc, "6. Nœuds critiques par bloc fonctionnel", 1)
    add_table(
        doc,
        ["Bloc", "Nœuds principaux", "But"],
        [
            ["Qualification", "Set Target, Build Source URLs, Normalize Public Signals", "préparer le contexte prospect"],
            ["Protection LLM", "Sanitize Prospect Data For LLM", "empêcher l’exposition des données sensibles"],
            ["Génération", "Call OpenAI..., Generate Executive Letter, Generate Tailored Catalogue", "produire le pack commercial"],
            ["Stockage", "Store Pack In Supabase", "sauvegarder le pack avant toute approbation"],
            ["Validation", "Build Approval Email, Send Internal Approval Email, Approval Webhook", "faire approuver le pack en interne"],
            ["Envoi", "Build Send Context, If Ready To Send, Send External Prospect Email", "expédier le message si le pack est complet"],
            ["CRM", "Update Prospect Target Sent, Update Prospect Target Approval Error, Update Prospect Target Rejected", "réécrire le statut métier du prospect"],
            ["Pièces jointes", "Resolve Domain Catalogue, Download Catalogue PDF, Build Deck Render Payload, Render Deck Artifact", "préparer le catalogue PDF et le deck PPTX"],
        ],
        widths=[1.3, 2.7, 2.0],
    )

    add_heading(doc, "7. Séquence utilisateur à retenir", 1)
    add_numbered(doc, [
        "Lancer le workflow avec un prospect ou le recevoir depuis la V4.",
        "Vérifier que Set Target contient bien organization_name, website et target_email.",
        "Laisser la chaîne de génération produire le pack prospect.",
        "Contrôler que Store Pack In Supabase enregistre bien un pack_id.",
        "Recevoir l’e-mail d’approbation interne et cliquer sur Approuver ou Rejeter.",
        "En cas d’approbation, laisser Build Send Context et If Ready To Send contrôler la validité du départ.",
        "Confirmer que le courrier est envoyé et que les pièces jointes sont bien assemblées.",
        "Contrôler enfin la mise à jour des tables Supabase et du statut CRM du prospect.",
    ])

    add_heading(doc, "8. Branchement d’approbation et d’envoi", 1)
    add_table(
        doc,
        ["Nœud", "Contrôle utilisateur à faire"],
        [
            ["Approval Webhook", "workflow activé et URL accessible publiquement"],
            ["Parse Approval Query", "pack_id et decision correctement lus dans la query string"],
            ["Get Pack From Supabase", "headers apikey et Authorization avec la service role key"],
            ["Extract Pack Payload", "reconstruction correcte du payload et de executive_letter_html"],
            ["Build Send Context", "can_send = true seulement si e-mail cible et courrier existent"],
            ["If Ready To Send", "branche true vers Mark Pack Approved, branche false vers Mark Pack Approval Error"],
        ],
        widths=[2.0, 4.3],
    )

    add_heading(doc, "9. Bloc pièces jointes et rendu", 1)
    attachment_nodes = find_node_names(
        data,
        [
            "Resolve Domain Catalogue",
            "Download Catalogue PDF",
            "Assemble Mail Attachments",
            "Build Catalogue Render Payload",
            "render Catalogue Artifact",
            "Merge Catalogue Artifact",
            "Build Deck Render Payload",
            "Render Deck Artifact",
            "Merge Deck Artifact",
        ],
    )
    if attachment_nodes:
        add_table(doc, ["Ordre", "Nœud", "Type"], attachment_nodes, widths=[0.7, 3.7, 1.75])
    add_paragraph(
        doc,
        "Cette version CRM prospecting n’envoie plus seulement un courrier texte. Elle tente aussi de joindre un catalogue PDF "
        "et un deck PPTX ciblés. L’utilisateur doit donc vérifier non seulement la génération du message, mais aussi la réussite "
        "des nœuds de rendu et l’existence finale des pièces jointes dans Assemble Mail Attachments."
    )

    add_heading(doc, "10. Cas d’usage quotidiens", 1)
    add_bullets(doc, [
        "tester un nouveau prospect en mode manuel",
        "contrôler un pack bloqué dans la branche approval_error",
        "valider qu’un prospect rejeté est bien mis à jour dans prospect_targets",
        "vérifier qu’un catalogue PDF existe bien avant l’envoi au DG",
        "contrôler les statuts sent, rejected, approval_error dans ai_prospecting_packs",
    ])

    add_heading(doc, "11. Troubleshooting rapide", 1)
    add_table(
        doc,
        ["Symptôme", "Cause probable", "Action recommandée"],
        [
            ["Get Pack From Supabase renvoie []", "headers ou select incorrect, RLS, mauvaise clé", "vérifier apikey, Authorization, select=* et pack_id exact"],
            ["If Ready To Send bloque", "target_email ou executive_letter manquant", "ouvrir Build Send Context et confirmer can_send"],
            ["Le webhook ne répond pas", "workflow inactif ou mauvaise URL publique", "activer le workflow et tester N8N_BASE_URL"],
            ["Aucune pièce jointe", "renderer catalogue ou deck en échec", "tester les nœuds de rendu séparément"],
            ["Le CRM ne se met pas à jour", "PATCH Supabase mal configuré", "vérifier Update Prospect Target ... et la service role key"],
        ],
        widths=[2.0, 2.1, 2.2],
    )

    add_heading(doc, "12. Résultat attendu", 1)
    add_paragraph(
        doc,
        "À la fin, l’utilisateur doit disposer d’un workflow CRM prospecting capable de qualifier un prospect, générer son pack, "
        "faire approuver l’approche, envoyer un courrier enrichi de pièces jointes, journaliser l’action et mettre à jour la base CRM. "
        "Ce guide doit permettre de reprendre le workflow rapidement sans repartir de zéro."
    )

    add_footer(doc, "TransferAI - Guide utilisateur CRM Prospecting V3")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
