from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
DATA_PATH = ROOT / "docs/transferai-prospection/cofounders-pack-data.json"
OUTPUT_PATH = ROOT / "docs/transferai-prospection/TransferAI_Reunion_CoFondateurs_2026-05-16.docx"


TITLE_BLUE = RGBColor(18, 53, 91)
ACCENT_ORANGE = RGBColor(214, 103, 36)
SOFT_GRAY = RGBColor(90, 96, 104)
LIGHT_FILL = "F4F6F8"
ORANGE_FILL = "FCE6D8"


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
      run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_number(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].width = Inches(widths[idx])
        set_cell_shading(header_cells[idx], LIGHT_FILL)
        set_cell_text(header_cells[idx], header, bold=True, color=TITLE_BLUE, size=10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value, size=10)

    document.add_paragraph()


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 24, TITLE_BLUE),
        ("Heading 1", 16, TITLE_BLUE),
        ("Heading 2", 13, TITLE_BLUE),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_cover(document: Document, data: dict) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.space_before = Pt(0)
    run = title.add_run(data["meta"]["title"])
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = TITLE_BLUE

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(data["meta"]["subtitle"])
    run.font.size = Pt(14)
    run.font.color.rgb = SOFT_GRAY

    meta_rows = [
        ["Date", data["meta"]["date"]],
        ["Objet", data["meta"]["meetingPurpose"]],
        ["Source", "Catalogue TransferAI Africa, catalogue de domaines et playbook de prospection Côte d'Ivoire"],
    ]
    add_table(document, ["Champ", "Valeur"], meta_rows, [1.15, 5.2])

    callout = document.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = callout.rows[0].cells[0]
    set_cell_shading(cell, ORANGE_FILL)
    set_cell_text(
        cell,
        "Message clé : le catalogue TransferAI Africa donne une base large, mais la vente court terme doit se concentrer sur quelques domaines transverses à ROI démontrable.",
        bold=True,
        color=TITLE_BLUE,
        size=11,
    )
    document.add_paragraph()


def add_summary(document: Document, data: dict) -> None:
    document.add_heading("Résumé exécutif", level=1)
    for bullet in data["executiveSummary"]:
        add_bullet(document, bullet)

    document.add_heading("Critères de priorisation", level=2)
    for item in data["selectionCriteria"]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        label = p.add_run(f'{item["label"]} : ')
        label.bold = True
        label.font.color.rgb = TITLE_BLUE
        p.add_run(item["detail"])


def add_priority_matrix(document: Document, data: dict) -> None:
    document.add_heading("Priorisation recommandée", level=1)
    document.add_paragraph(
        "L'objectif n'est pas de tout pousser en même temps. Il faut concentrer la prospection sur les domaines les plus démontrables, les plus transverses et les plus rapides à transformer en missions."
    )

    tier_rows = [
        [item["tier"], item["label"], item["rationale"]] for item in data["priorityTiers"]
    ]
    add_table(document, ["Tier", "Positionnement", "Logique"], tier_rows, [0.9, 1.8, 3.8])

    document.add_heading("Top 5 domaines à pousser immédiatement", level=2)
    rows = []
    for item in data["topPriorities"]:
        rows.append(
            [
                item["domain"],
                item["flagshipUseCase"],
                item["roiStory"],
                ", ".join(item["commercialTargets"][:3]),
            ]
        )
    add_table(document, ["Domaine", "Cas d'usage phare", "ROI à vendre", "Entrées cibles"], rows, [1.55, 2.0, 1.9, 1.05])


def add_priority_detail(document: Document, item: dict) -> None:
    document.add_heading(item["domain"], level=2)
    document.add_paragraph(item["whyNow"])
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(4)
    a = intro.add_run("Cas d'usage phare : ")
    a.bold = True
    a.font.color.rgb = ACCENT_ORANGE
    intro.add_run(item["flagshipUseCase"])

    roi = document.add_paragraph()
    roi.paragraph_format.space_after = Pt(4)
    b = roi.add_run("Promesse ROI : ")
    b.bold = True
    b.font.color.rgb = ACCENT_ORANGE
    roi.add_run(item["roiStory"])

    targets = document.add_paragraph()
    c = targets.add_run("Prospects prioritaires : ")
    c.bold = True
    c.font.color.rgb = ACCENT_ORANGE
    targets.add_run(", ".join(item["commercialTargets"]))

    for proof in item["proofPoints"]:
        add_bullet(document, proof)


def add_market_recommendation(document: Document, data: dict) -> None:
    document.add_heading("Recommandation go-to-market", level=1)
    document.add_paragraph(
        "La meilleure séquence commerciale pour TransferAI Africa est simple : ouvrir avec un audit court, convertir en formation métier, puis transformer les premiers quick wins en pilote concret et mesurable."
    )

    rows = [[offer["name"], offer["promise"], offer["value"]] for offer in data["offers"]]
    add_table(document, ["Offre", "Promesse", "Pourquoi elle vend"], rows, [1.7, 2.55, 2.25])

    document.add_heading("Domaines à développer en deuxième vague", level=2)
    for item in data["tier2Domains"]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'{item["domain"]} - ')
        r.bold = True
        r.font.color.rgb = TITLE_BLUE
        p.add_run(f'{item["focus"]}. {item["reason"]}')

    document.add_heading("Domaines à adresser de façon opportuniste", level=2)
    for item in data["tier3Domains"]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f'{item["domain"]} - ')
        r.bold = True
        r.font.color.rgb = TITLE_BLUE
        p.add_run(f'{item["focus"]}. {item["reason"]}')


def add_action_plan(document: Document, data: dict) -> None:
    document.add_heading("Plan d'action 90 jours", level=1)
    for step in data["actionPlan90Days"]:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f'{step["phase"]} : ')
        run.bold = True
        run.font.color.rgb = TITLE_BLUE
        p.add_run(f'{step["goal"]} ')
        owner = p.add_run(f'({step["ownerHint"]})')
        owner.italic = True
        owner.font.color.rgb = SOFT_GRAY

    document.add_heading("Décisions à valider en réunion", level=2)
    for point in data["decisionPoints"]:
        add_number(document, point)


def add_appendix(document: Document, data: dict) -> None:
    document.add_page_break()
    document.add_heading("Annexe - vue synthétique des 13 domaines", level=1)
    rows = [
        [item["domain"], item["tier"], item["flagshipUseCase"], item["roi"]]
        for item in data["fullDomainList"]
    ]
    add_table(document, ["Domaine", "Tier", "Cas d'usage phare", "ROI"], rows, [1.7, 0.7, 2.45, 1.65])


def main() -> None:
    data = load_data()
    document = Document()
    configure_styles(document)
    add_cover(document, data)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_summary(document, data)
    add_priority_matrix(document, data)
    for item in data["topPriorities"]:
        add_priority_detail(document, item)
    add_market_recommendation(document, data)
    add_action_plan(document, data)
    add_appendix(document, data)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    main()
