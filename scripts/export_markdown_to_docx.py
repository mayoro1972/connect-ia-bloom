from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


INLINE_BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_PATTERN = re.compile(r"`(.+?)`")


def set_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)


def set_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x17, 0x2B, 0x4D)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(11)
    subtitle.font.italic = True
    subtitle.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7A)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(0x17, 0x2B, 0x4D)),
        ("Heading 2", 13, RGBColor(0x1F, 0x4E, 0x79)),
        ("Heading 3", 11, RGBColor(0x2E, 0x5D, 0x86)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_bottom_border(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2F2")
    p_bdr.append(bottom)
    # <w:pBdr> must respect the CT_PPr child order: it has to appear before
    # spacing/ind/jc/rPr and the other later elements, otherwise Word rejects
    # the document as having unreadable content.
    p_pr.insert_element_before(
        p_bdr,
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    )


def add_runs(paragraph, text: str) -> None:
    if not text:
        return

    parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for part in parts:
        if not part:
            continue
        bold_match = INLINE_BOLD_PATTERN.fullmatch(part)
        if bold_match:
            run = paragraph.add_run(bold_match.group(1))
            run.bold = True
            continue
        code_match = INLINE_CODE_PATTERN.fullmatch(part)
        if code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue
        paragraph.add_run(part)


def add_title(document: Document, title_text: str, subtitle_text: str | None) -> None:
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(title, title_text)
    add_bottom_border(title)

    if subtitle_text:
        subtitle = document.add_paragraph(style="Subtitle")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs(subtitle, subtitle_text)
        subtitle.paragraph_format.space_after = Pt(14)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    add_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(2)


def add_numbered(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    add_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(2)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Normal")
    add_runs(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(6)


def add_heading(document: Document, level: int, text: str) -> None:
    heading = document.add_paragraph(style=f"Heading {min(level, 3)}")
    add_runs(heading, text)
    heading.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    heading.paragraph_format.space_after = Pt(3)


def fix_document_settings(document: Document) -> None:
    # python-docx ships a default settings.xml whose <w:zoom> only carries a
    # w:val attribute. The OOXML schema requires w:percent on <w:zoom>; without
    # it Word reports the file as containing unreadable content. Ensure the
    # attribute is present so the document opens cleanly.
    settings = document.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def markdown_to_docx(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    set_page_geometry(document)
    set_base_styles(document)
    fix_document_settings(document)

    title_text = source.stem.replace("_", " ")
    subtitle_text = None
    body_lines = lines

    if lines and lines[0].startswith("# "):
        title_text = lines[0][2:].strip()
        body_lines = lines[1:]

    subtitle_candidates = []
    while body_lines and len(subtitle_candidates) < 2:
        candidate = body_lines[0].strip()
        if not candidate:
            body_lines = body_lines[1:]
            continue
        if candidate.startswith("## ") or candidate.startswith("### "):
            break
        if candidate == "---":
            body_lines = body_lines[1:]
            break
        subtitle_candidates.append(candidate)
        body_lines = body_lines[1:]
    if subtitle_candidates:
        subtitle_text = " · ".join(subtitle_candidates)

    add_title(document, title_text, subtitle_text)

    for line in body_lines:
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue

        heading_match = re.match(r"^(#{2,6})\s+(.+)$", stripped)
        if heading_match:
            add_heading(document, len(heading_match.group(1)) - 1, heading_match.group(2).strip())
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            add_bullet(document, bullet_match.group(1).strip())
            continue

        numbered_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered_match:
            add_numbered(document, numbered_match.group(1).strip())
            continue

        add_body_paragraph(document, stripped)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export a Markdown file to DOCX.")
    parser.add_argument("source", type=Path, help="Markdown source file")
    parser.add_argument("output", type=Path, help="DOCX output path")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    markdown_to_docx(args.source, args.output)


if __name__ == "__main__":
    main()
