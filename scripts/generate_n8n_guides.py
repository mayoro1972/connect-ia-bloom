from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUT_DIR = ROOT / "docs" / "transferai-admin"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, RGBColor(15, 46, 75)),
        ("Heading 1", 15, RGBColor(15, 46, 75)),
        ("Heading 2", 12.5, RGBColor(36, 82, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title_block(document, title, subtitle):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 46, 75)
    p.space_after = Pt(4)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.name = "Arial"
    run2.font.size = Pt(10.5)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(90, 90, 90)
    p2.space_after = Pt(12)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.9)
    rows = [
        ("Projet", "Workflow n8n TransferAI Prospecting V3 CRM Enhanced"),
        ("Date", "05 juin 2026"),
        ("Portee", "Guide operatoire et guide de resolution des incidents"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.7)
        cells[1].width = Inches(5.4)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "DDEBF7")
    document.add_paragraph()


def add_section(document, title, paragraphs=None, bullets=None, numbered=None):
    document.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(5)
    for item in bullets or []:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)
    for item in numbered or []:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_issue(document, title, cause, resolution, status="Corrige"):
    document.add_heading(title, level=2)
    rows = [
        ("Cause", cause),
        ("Resolution", resolution),
        ("Statut", status),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.4)
        cells[1].width = Inches(5.7)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "EAF2F8")
    document.add_paragraph()


def build_user_guide():
    doc = Document()
    set_document_defaults(doc)
    add_title_block(
        doc,
        "Guide Utilisateur",
        "Configuration et processus d'envoi du workflow n8n TransferAI",
    )

    add_section(
        doc,
        "1. Objectif du workflow",
        paragraphs=[
            "Le workflow sert a construire un pack prospect personnalise, a demander une validation interne, puis a envoyer au prospect un courrier accompagne d'un lien de formulaire d'audit pre-RDV et de deux pieces jointes.",
        ],
        bullets=[
            "Courrier executif personnalise",
            "Mini-catalogue PDF",
            "Brief deck PPTX",
            "Lien du formulaire d'audit pre-RDV",
        ],
    )

    add_section(
        doc,
        "2. Configuration realisee",
        bullets=[
            "Suppression des dependances fragiles a $env dans les noeuds OpenAI.",
            "Suppression des anciens fallbacks locaux 127.0.0.1 et _reload=7.",
            "Standardisation du lien de formulaire sur https://www.transferai.ci/questionnaire-audit?pack_id=...",
            "Correction des noeuds de generation de contenu, d'assemblage du pack et de la branche d'approbation.",
            "Configuration d'un expediteur Resend valide sur le domaine transferai.ci.",
        ],
    )

    add_section(
        doc,
        "3. Parcours complet d'envoi",
        numbered=[
            "Le workflow demarre depuis Manual Trigger.",
            "Les informations du prospect sont preparees dans Set Target puis enrichies par les pages publiques.",
            "Le contexte prospect est assemble dans Assemble Prospect Context.",
            "Les noeuds Generate Executive Letter, Generate Tailored Catalogue, Generate Tailored Audit Form et Generate Deck Brief produisent les contenus IA.",
            "Assemble Prospect Pack cree le pack final avec le pack_id et le lien audit_form_url.",
            "Les artefacts documentaires sont rendus puis stockes avec le pack dans Supabase.",
            "Build Approval Email envoie un mail interne de validation avec les liens Approuver et envoyer / Rejeter.",
            "Au clic sur Approuver et envoyer, Approval Webhook declenche la branche d'approbation.",
            "Extract Pack Payload relit le pack stocke dans Supabase.",
            "Build Send Context prepare le mail prospect final, reconstruit les pieces jointes et ajoute le bloc du formulaire.",
            "If Ready To Send verifie que l'email, le courrier, le PDF et le PPTX sont bien disponibles.",
            "Send External Prospect Email envoie le mail final au prospect.",
        ],
    )

    add_section(
        doc,
        "4. Ce que recoit le prospect",
        paragraphs=[
            "Le workflow actuel envoie le pack des l'approbation interne. Il n'attend pas la soumission du formulaire pour envoyer les documents.",
        ],
        bullets=[
            "Le courrier personnalise",
            "Le lien du formulaire d'audit pre-RDV",
            "Le mini-catalogue PDF en piece jointe",
            "Le deck PPTX en piece jointe",
        ],
    )

    add_section(
        doc,
        "5. Verifications de fonctionnement",
        bullets=[
            "Dans Build Send Context, verifier attachments_count = 2.",
            "Verifier can_send = true et send_failure_reason = null.",
            "Dans Send External Prospect Email, verifier la presence d'un id Resend dans la reponse.",
            "Dans Gmail, verifier les deux pieces jointes sur le message prospect.",
            "Dans Resend, verifier que le domaine transferai.ci est bien marque Verified.",
        ],
    )

    add_section(
        doc,
        "6. Limite restante",
        paragraphs=[
            "Le lien du formulaire est maintenant correctement genere dans n8n. Si la page publique retourne encore une 404, le probleme ne vient plus du workflow mais du deploiement du frontend public de transferai.ci.",
        ],
    )

    return doc


def build_troubleshooting_guide():
    doc = Document()
    set_document_defaults(doc)
    add_title_block(
        doc,
        "Troubleshooting Guide",
        "Incidents rencontres et resolutions appliquees sur le workflow n8n TransferAI",
    )

    add_section(
        doc,
        "Vue d'ensemble",
        paragraphs=[
            "Ce document recense les incidents rencontres pendant la reconstruction du workflow et explique, point par point, la cause identifiee et la resolution appliquee.",
        ],
    )

    add_issue(
        doc,
        "1. Erreur Unexpected token 'return' dans Assemble Prospect Pack",
        "Le noeud contenait un collage incomplet ou un return mal place au debut du script.",
        "Le code du noeud a ete remplace en entier par une version propre avec un seul return final.",
    )

    add_issue(
        doc,
        "2. Erreurs missing / et Invalid regular expression",
        "Certaines regex etaient mal collees ou mal echappees dans n8n.",
        "Les blocs fragiles ont ete remplaces par des versions plus robustes, parfois via split(...).join(...).",
    )

    add_issue(
        doc,
        "3. access to env vars denied",
        "Plusieurs noeuds OpenAI utilisaient $env alors que l'instance n8n refusait cet acces.",
        "Les dependances a $env ont ete supprimees et remplacees par des valeurs directes ou deja presentes dans le workflow.",
    )

    add_issue(
        doc,
        "4. Mauvais lien de formulaire",
        "Le workflow utilisait encore des fallbacks locaux comme 127.0.0.1 ou _reload=7.",
        "Le lien a ete standardise vers https://www.transferai.ci/questionnaire-audit?pack_id=... dans les noeuds de generation et d'assemblage.",
    )

    add_issue(
        doc,
        "5. Erreurs sur la branche d'approbation",
        "Extract Pack Payload et Build Send Context contenaient des versions fragiles qui bloquaient le traitement apres clic sur Approuver et envoyer.",
        "Les scripts ont ete simplifies, stabilises et reconnectes a la structure stockee dans Supabase.",
    )

    add_issue(
        doc,
        "6. Mail prospect non envoye",
        "Le webhook d'approbation fonctionnait, mais le workflow pouvait s'arreter avant ou pendant l'envoi effectif.",
        "La chaine If Approved -> Build Send Context -> If Ready To Send -> Send External Prospect Email a ete revue et testee dans les executions n8n.",
    )

    add_issue(
        doc,
        "7. Mail envoye sans pieces jointes",
        "Une version simplifiee de Build Send Context avait perdu la reconstruction des attachments.",
        "La logique de pieces jointes a ete fusionnee avec la nouvelle logique d'injection du lien audit_form_url. Le noeud reconstruit maintenant le PDF et le PPTX et bloque l'envoi si l'un manque.",
    )

    add_issue(
        doc,
        "8. Probleme d'expediteur Resend",
        "Le workflow utilisait onboarding@resend.dev, non ideal pour un envoi reel au prospect.",
        "Le domaine transferai.ci a ete verifie dans Resend et l'expediteur a ete remplace par TransferAI <contact@transferai.ci>.",
    )

    add_issue(
        doc,
        "9. Erreur JSON parameter needs to be valid JSON",
        "Le noeud Send Internal Sent Confirmation utilisait un body JSON trop fragile avec interpolation complexe.",
        "Le body a ete re-ecrit avec des concatenations simples pour garantir un JSON valide.",
    )

    add_issue(
        doc,
        "10. Formulaire public en 404",
        "Le workflow genere maintenant le bon lien, mais le frontend public de transferai.ci ne sert pas encore la route /questionnaire-audit dans la version deployee.",
        "Aucune correction n8n supplementaire n'est necessaire. Il faut redelivrer le frontend public avec la route en production.",
        status="A traiter cote frontend",
    )

    add_section(
        doc,
        "Checklist de diagnostic rapide",
        bullets=[
            "Verifier Build Send Context : attachments_count, can_send, send_failure_reason.",
            "Verifier Send External Prospect Email : presence d'un id Resend.",
            "Verifier le domaine transferai.ci dans Resend : statut Verified.",
            "Verifier le lien formulaire en production : si 404, controler le deploiement frontend et non n8n.",
        ],
    )

    return doc


def save_documents():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    user_path = OUT_DIR / "Guide_Utilisateur_Workflow_n8n_TransferAI.docx"
    trouble_path = OUT_DIR / "Troubleshooting_Guide_Workflow_n8n_TransferAI.docx"

    build_user_guide().save(user_path)
    build_troubleshooting_guide().save(trouble_path)
    return user_path, trouble_path


if __name__ == "__main__":
    paths = save_documents()
    for path in paths:
        print(path)
