from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
SOURCE_GUIDE = DOCS_DIR / "46_Guide_Reconstruction_n8n_V1.md"
CHECKLIST_GUIDE = DOCS_DIR / "55_Checklist_Installation_n8n_V1.md"
CODE_GUIDE = DOCS_DIR / "56_Guide_Copier_Coller_Blocs_Code_n8n_V1.md"
WORKFLOW_JSON = DOCS_DIR / "42_n8n_Prospection_Modele_Elton_V1.json"
TODAY = date.today().isoformat()
OUTPUT_PATH = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V1_n8n_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color, bold in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43), True),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43), True),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B), True),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B), True),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Installation complète du workflow V1 dans n8n")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    r = sub.add_run("Guide Word opérationnel pour le workflow de prospection multi-prospects")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(16)
    m = meta.add_run(f"Version générée le {TODAY}")
    m.font.size = Pt(10.5)
    m.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(
        "Ce document permet de reconstruire, configurer et tester la V1 du workflow n8n "
        "sans revenir au JSON source. Il couvre les prérequis, les nœuds, les connexions, "
        "les appels OpenAI, les blocs de code à copier-coller et le contrôle final."
    )
    run.font.size = Pt(10.5)

    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str, bold_prefix: str | None = None) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_simple_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=90, start=110, bottom=90, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in hdr[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def add_code_block(document: Document, title: str, code: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.35)
    set_cell_shading(cell, "F4F6F8")
    set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1
    run = para.add_run(code.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(8.2)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


def read_code_blocks() -> dict[str, str]:
    text = CODE_GUIDE.read_text(encoding="utf-8")
    pattern = re.compile(r"## Nœud `([^`]+)`.*?```javascript\n(.*?)```", re.S)
    return {name: code.strip() for name, code in pattern.findall(text)}


def load_workflow() -> dict:
    return json.loads(WORKFLOW_JSON.read_text(encoding="utf-8"))


def find_node(data: dict, name: str) -> dict:
    for node in data["nodes"]:
        if node["name"] == name:
            return node
    raise KeyError(name)


def extract_set_target_rows(node: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in node["parameters"]["fields"]["values"]:
        rows.append([item["name"], item["stringValue"].replace("={{", "").replace("}}", "")])
    return rows


def http_summary(node: dict) -> tuple[str, str, str, str]:
    params = node["parameters"]
    system_prompt = ""
    body = params.get("jsonBody", "")
    prompt_match = re.search(r'role: \'system\', content: "([^"]+)"', body)
    if prompt_match:
        system_prompt = prompt_match.group(1)
    body_hint = "Utilise llm_allowed_payload" if "llm_allowed_payload" in body else "Utilise llm_generation_payload"
    return (
        node["name"],
        params["method"],
        params["url"],
        f"{body_hint}. {system_prompt[:210]}..."
    )


def add_footer(section, label: str) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_document() -> Path:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    data = load_workflow()
    code_blocks = read_code_blocks()
    set_target = find_node(data, "Set Target")
    openai_nodes = [
        find_node(data, "Call OpenAI Pre-Audit"),
        find_node(data, "Call OpenAI Problems Solutions"),
        find_node(data, "Call OpenAI ROI"),
        find_node(data, "Generate Executive Letter"),
        find_node(data, "Generate Tailored Catalogue"),
        find_node(data, "Generate Tailored Audit Form"),
        find_node(data, "Generate Deck Brief"),
    ]

    document = Document()
    style_document(document)
    add_footer(document.sections[0], "TransferAI Africa - Installation workflow V1 n8n")
    add_title_page(document)

    add_heading(document, "1. Finalité du workflow", 1)
    add_paragraph(
        document,
        "La V1 est un workflow de prospection multi-prospects. Il part d'informations publiques, "
        "assainit les données avant exposition au LLM, génère un pré-audit commercial et prépare "
        "quatre livrables sur mesure avant revue manuelle.",
    )
    add_bullets(
        document,
        [
            "scraper jusqu'à 5 pages publiques d'un site prospect",
            "normaliser les signaux métier visibles",
            "retirer ou pseudonymiser les données sensibles avant l'appel au LLM",
            "générer un courrier, un mini-catalogue, une forme d'audit et un brief de deck",
            "terminer sur un statut de revue manuelle avant tout envoi",
        ],
    )

    add_heading(document, "2. Pré-requis à préparer dans n8n", 1)
    add_simple_table(
        document,
        ["Variable", "Utilité"],
        [
            ["OPENAI_API_KEY", "clé API utilisée par les 7 nœuds HTTP vers OpenAI"],
            ["OPENAI_MODEL", "modèle par défaut, par exemple gpt-4.1-mini"],
            ["BOOKING_LINK_45MIN", "lien du rendez-vous gratuit de 45 minutes"],
        ],
        widths=[2.1, 4.2],
    )

    add_paragraph(
        document,
        "Créer d'abord un workflow vide dans n8n, puis ajouter les nœuds dans l'ordre réel ci-dessous.",
    )

    add_heading(document, "3. Ordre réel des nœuds du workflow V1", 1)
    node_rows = [
        ["1", "Manual Trigger", "déclenchement manuel"],
        ["2", "Execute Workflow Trigger", "appel depuis un autre workflow"],
        ["3", "Set Target", "définition de la cible prospect"],
        ["4", "Build Source URLs", "construction des pages publiques à explorer"],
        ["5", "Fetch Public Page 1", "collecte page 1"],
        ["6", "Fetch Public Page 2", "collecte page 2"],
        ["7", "Fetch Public Page 3", "collecte page 3"],
        ["8", "Fetch Public Page 4", "collecte page 4"],
        ["9", "Fetch Public Page 5", "collecte page 5"],
        ["10", "Normalize Public Signals", "extraction de texte et de signaux métier"],
        ["11", "Sanitize Prospect Data For LLM", "protection RGPD avant LLM"],
        ["12", "Call OpenAI Pre-Audit", "pré-audit pseudonymisé"],
        ["13", "Call OpenAI Problems Solutions", "problèmes probables et solutions"],
        ["14", "Call OpenAI ROI", "hypothèse de gains et jalons"],
        ["15", "Assemble Prospect Context", "fusion des sorties d'analyse"],
        ["16", "Generate Executive Letter", "courrier de premier contact"],
        ["17", "Generate Tailored Catalogue", "mini-catalogue sur mesure"],
        ["18", "Generate Tailored Audit Form", "forme d'audit pré-appel"],
        ["19", "Generate Deck Brief", "brief de deck taillé sur mesure"],
        ["20", "Assemble Prospect Pack", "réinjection des vraies valeurs et pack final"],
        ["21", "Mark For Review", "sortie ready_for_manual_review"],
    ]
    add_simple_table(document, ["Ordre", "Nœud", "Rôle"], node_rows, widths=[0.65, 2.4, 3.35])

    add_heading(document, "4. Installation pas à pas", 1)
    step_paragraphs = [
        "Créer les deux nœuds de départ: Manual Trigger et Execute Workflow Trigger.",
        "Créer ensuite Set Target de type Set, puis renseigner tous les champs d'entrée prospect.",
        "Ajouter Build Source URLs de type Code pour générer les pages à scraper.",
        "Créer 5 nœuds HTTP Request pour les pages publiques 1 à 5 et les chaîner.",
        "Ajouter Normalize Public Signals puis Sanitize Prospect Data For LLM.",
        "Créer les 3 nœuds OpenAI d'analyse: Pre-Audit, Problems Solutions, ROI.",
        "Créer Assemble Prospect Context pour fusionner les sorties d'analyse.",
        "Créer les 4 nœuds OpenAI de génération: Executive Letter, Tailored Catalogue, Tailored Audit Form, Deck Brief.",
        "Créer Assemble Prospect Pack pour réinjecter les vrais identifiants hors LLM.",
        "Créer enfin Mark For Review pour arrêter le workflow en revue manuelle.",
    ]
    add_numbered(document, step_paragraphs)

    add_heading(document, "5. Configuration exacte du nœud Set Target", 1)
    add_paragraph(
        document,
        "Le nœud Set Target accepte les valeurs entrantes d'un autre workflow, mais fournit aussi des valeurs de secours pour un test manuel.",
    )
    add_simple_table(document, ["Champ", "Expression ou valeur"], extract_set_target_rows(set_target), widths=[2.15, 4.15])

    add_heading(document, "6. Connexions à créer", 1)
    add_bullets(
        document,
        [
            "Manual Trigger -> Set Target",
            "Execute Workflow Trigger -> Set Target",
            "Set Target -> Build Source URLs",
            "Build Source URLs -> Fetch Public Page 1 -> Fetch Public Page 2 -> Fetch Public Page 3 -> Fetch Public Page 4 -> Fetch Public Page 5",
            "Fetch Public Page 5 -> Normalize Public Signals -> Sanitize Prospect Data For LLM",
            "Sanitize Prospect Data For LLM -> Call OpenAI Pre-Audit",
            "Sanitize Prospect Data For LLM -> Call OpenAI Problems Solutions",
            "Sanitize Prospect Data For LLM -> Call OpenAI ROI",
            "Call OpenAI Pre-Audit -> Assemble Prospect Context",
            "Assemble Prospect Context -> Generate Executive Letter",
            "Assemble Prospect Context -> Generate Tailored Catalogue",
            "Assemble Prospect Context -> Generate Tailored Audit Form",
            "Assemble Prospect Context -> Generate Deck Brief",
            "Generate Executive Letter -> Assemble Prospect Pack -> Mark For Review",
        ],
    )

    add_heading(document, "7. Configuration des nœuds HTTP Request OpenAI", 1)
    add_paragraph(
        document,
        "Tous les appels OpenAI utilisent la même base technique: méthode POST, URL https://api.openai.com/v1/chat/completions, en-têtes Authorization et Content-Type, puis un corps JSON construit dynamiquement.",
    )
    http_rows = [list(http_summary(node)) for node in openai_nodes]
    add_simple_table(
        document,
        ["Nœud", "Méthode", "URL", "Résumé du corps JSON"],
        http_rows,
        widths=[1.8, 0.7, 1.45, 2.45],
    )
    add_paragraph(
        document,
        "Important: les 3 nœuds d'analyse utilisent llm_allowed_payload. Les 4 nœuds de génération utilisent llm_generation_payload. Cela évite d'exposer les identifiants réels au modèle.",
    )

    add_heading(document, "8. Protection des données avant LLM", 1)
    add_bullets(
        document,
        [
            "politique deny-by-default",
            "suppression des URLs, e-mails et téléphones",
            "pseudonymisation de l'organisation en ORG_TARGET",
            "pseudonymisation du décideur en DECISION_MAKER_TARGET",
            "transmission au LLM d'un extrait public assaini uniquement",
            "réinjection locale des vrais identifiants après génération des livrables",
        ],
    )

    add_heading(document, "9. Blocs de code exacts à copier-coller", 1)
    add_paragraph(
        document,
        "Les cinq nœuds Code ci-dessous doivent être collés tels quels dans n8n. Respecter les noms de nœuds, car certains scripts relisent d'autres nœuds par leur nom.",
    )
    for node_name in [
        "Build Source URLs",
        "Normalize Public Signals",
        "Sanitize Prospect Data For LLM",
        "Assemble Prospect Context",
        "Assemble Prospect Pack",
    ]:
        add_code_block(document, node_name, code_blocks[node_name])

    add_heading(document, "10. Contrôle final avant usage", 1)
    add_bullets(
        document,
        [
            "tester avec un prospect fictif ou un site de test",
            "vérifier que les 5 URLs sont bien construites",
            "vérifier que page_texts, public_text et roi_clues sont présents",
            "vérifier que llm_allowed_payload n'inclut ni e-mail, ni téléphone, ni website brut",
            "vérifier que llm_redaction_summary est bien généré",
            "vérifier que le courrier, le mini-catalogue, la forme d'audit et le deck brief sont produits",
            "vérifier que la sortie finale passe en ready_for_manual_review",
        ],
    )

    add_heading(document, "11. Résultat attendu", 1)
    add_paragraph(
        document,
        "Une fois la V1 installée, tu disposes d'un workflow n8n capable de produire un pack prospect complet sans envoyer automatiquement d'e-mail. Le workflow est donc exploitable pour la qualification, la personnalisation commerciale et la revue humaine avant passage à la V2 ou à la V3.",
    )

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
