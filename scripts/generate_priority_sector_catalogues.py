from __future__ import annotations

import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_OUTPUT_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-prospection/catalogues-sectoriels-premium-2026-06-12/catalogues-prioritaires-standard"
)
CATALOGUE_DATE = "2026-06-12"

NAVY = RGBColor(0x1C, 0x39, 0x63)
ORANGE = RGBColor(0xEF, 0x6C, 0x12)
GREEN = RGBColor(0x1C, 0x8A, 0x78)
GRAY = RGBColor(0x6B, 0x7C, 0x93)
LIGHT_FILL = "F8F3EC"


SECTORS = [
    {
        "slug": "01_finance-comptabilite",
        "title": "Finance & Comptabilité",
        "cover_line": "pour les directions finance et comptabilité",
        "meta_1": "DAF · Contrôle de gestion · Comptabilité · Audit interne · Trésorerie",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour accélérer le reporting, renforcer les contrôles, fiabiliser les analyses et mieux préparer les arbitrages financiers.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les processus les plus répétitifs, renforcer les capacités d’analyse, sécuriser le cadre de confiance et lancer un premier pilote mesurable autour du reporting, des contrôles ou de la pré-clôture.",
        "benefits": [
            "Temps de préparation des reportings, notes et supports de comité réduit.",
            "Écarts, anomalies et points de vigilance mieux identifiés avant arbitrage humain.",
            "Synthèses financières plus lisibles pour les décideurs et les comités.",
            "Contrôles plus homogènes entre comptabilité, contrôle de gestion, audit et trésorerie.",
        ],
        "summary": "Dans la finance et la comptabilité, la valeur de l’IA se lit d’abord dans la réduction des délais de lecture, de consolidation et de préparation des arbitrages, sans compromis sur la fiabilité, la traçabilité et la supervision humaine.",
        "priorities": [
            "Reporting financier encore lourd à produire, commenter et reformater.",
            "Revues de cohérence et contrôles encore très artisanaux sur de nombreux périmètres.",
            "Décisions retardées faute de lecture rapide et exploitable des KPI.",
            "Demandes ad hoc répétitives qui perturbent la cadence des équipes finance.",
        ],
        "kpis": [
            "Temps de production du reporting mensuel et délai de diffusion.",
            "Temps de revue avant clôture et nombre d’anomalies détectées plus tôt.",
            "Temps de préparation des comités et délai d’arbitrage.",
            "Volume de demandes récurrentes traitées à l’aide de trames réutilisables.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux processus prioritaires, bootcamp ciblé pour décideurs et équipes finance, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Reporting plus rapide et mieux structuré pour la direction.",
            "Revue plus prévisible des écarts, anomalies et exceptions.",
            "Arbitrages mieux documentés grâce à des notes et commentaires plus lisibles.",
            "Première feuille de route claire sur les usages IA à lancer en finance.",
        ],
        "offer_title": "Bootcamp IA Finance & Comptabilité : de l’acculturation utile au premier pilote mesurable.",
        "public": "Directions financières, contrôle de gestion, comptabilité, audit interne, trésorerie.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA en finance et comptabilité.",
            "Structurer les usages liés au reporting, à la pré-clôture, aux contrôles et à l’audit interne.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Reporting financier automatisé et commentaires de performance.",
            "Pré-clôture et détection d’anomalies assistées par l’IA.",
            "Audit interne augmenté et préparation documentaire.",
            "Prévisions budgétaires, trésorerie et lecture des hypothèses.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux finance, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur le reporting financier automatisé et la pré-clôture assistée.",
            "Méthodes, prompts, routines et gestes métier autour des synthèses financières et des revues de cohérence.",
            "Cadre de confiance : confidentialité, validation managériale, traçabilité et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur l’audit interne augmenté, la lecture des écarts et la trésorerie.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : prévisions budgétaires, contrôles récurrents, routines d’aide à l’arbitrage.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage finance avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable au reporting, aux contrôles, à la pré-clôture et à l’audit interne.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs et les équipes finance ciblées.",
        ],
        "domains": [
            "Direction financière",
            "Contrôle de gestion",
            "Comptabilité générale et analytique",
            "Audit interne",
            "Trésorerie et prévisions budgétaires",
        ],
    },
    {
        "slug": "02_service-client",
        "title": "Service Client",
        "cover_line": "pour les directions service client et relation usager",
        "meta_1": "Relation client · Centres de contact · Support · Qualité de service · Supervision",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour accélérer la réponse, stabiliser la qualité de traitement, mieux exploiter les retours clients et renforcer le pilotage omnicanal.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les demandes les plus récurrentes, fiabiliser les réponses, renforcer la lecture des irritants clients et lancer un premier pilote utile autour de la qualité de service et de la résolution.",
        "benefits": [
            "Réponses plus rapides, plus cohérentes et mieux traçables sur plusieurs canaux.",
            "Premier niveau de qualification plus fluide pour les demandes récurrentes.",
            "Verbatims, réclamations et signaux faibles mieux regroupés pour les managers.",
            "KPI de service mieux commentés et plus directement reliés à des plans d’action.",
        ],
        "summary": "Dans le service client, la valeur de l’IA se mesure à travers la réactivité, la qualité de réponse, la résolution, l’exploitation du feedback client et la capacité des managers à piloter rapidement les priorités d’amélioration.",
        "priorities": [
            "Réponses hétérogènes selon les canaux, les personnes ou le niveau de pression.",
            "Traitement lent des demandes simples mais très nombreuses.",
            "Réclamations et verbatims encore peu exploités pour corriger les parcours.",
            "Pilotage difficile faute de lecture rapide et exploitable des indicateurs de service.",
        ],
        "kpis": [
            "Temps moyen de réponse et taux de résolution au premier niveau.",
            "Temps de qualification et délai d’escalade.",
            "Délai d’analyse des verbatims et volume de motifs consolidés.",
            "CSAT / NPS et vitesse de réaction des managers sur les irritants.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux parcours prioritaires, bootcamp ciblé pour décideurs, managers et équipes support, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Temps de réponse réduit sur les demandes simples et récurrentes.",
            "Meilleure cohérence sur plusieurs canaux grâce à des trames et assistants encadrés.",
            "Feedback client plus exploitable pour les superviseurs et responsables qualité.",
            "Première feuille de route claire sur les usages IA à lancer en service client.",
        ],
        "offer_title": "Bootcamp IA Service Client : de l’acculturation utile au premier pilote mesurable.",
        "public": "Directions relation client, centres de contact, support, qualité de service, supervision.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA en service client.",
            "Structurer les usages liés à la réponse, à la qualification, aux verbatims et au pilotage de service.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Assistant de réponse et FAQ métier pour centres de contact.",
            "Qualification et tri des demandes avec supervision humaine.",
            "Analyse des réclamations, verbatims et signaux faibles.",
            "Pilotage omnicanal, satisfaction et commentaires de performance.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux service client, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur l’assistance à la réponse et la qualification des demandes.",
            "Méthodes, prompts, routines et gestes métier autour de la relation client augmentée par l’IA.",
            "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur les verbatims, la gestion des réclamations et le pilotage de la satisfaction.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : scripts omnicanaux, plans d’action qualité, routines de supervision.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage service client avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable à la réponse, au tri des demandes, aux verbatims et au pilotage de service.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs, managers et équipes support ciblées.",
        ],
        "domains": [
            "Centres de contact",
            "Support client",
            "Qualité de service",
            "Pilotage omnicanal",
            "Supervision et expérience client",
        ],
    },
    {
        "slug": "03_ressources-humaines",
        "title": "Ressources Humaines",
        "cover_line": "pour les directions ressources humaines",
        "meta_1": "RH · Recrutement · Formation · HRBP · Managers · Expérience collaborateur",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour accélérer les processus RH, fiabiliser les parcours, mieux standardiser les documents et renforcer la lecture des signaux managériaux.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les processus RH à plus forte valeur de temps, fiabiliser les routines récurrentes, structurer un cadre de confiance et lancer un premier pilote utile sur le recrutement, l’onboarding ou l’analyse RH.",
        "benefits": [
            "Tri, synthèse et préparation d’entretiens plus rapides et plus comparables.",
            "Onboarding plus homogène selon les sites, managers et équipes.",
            "Documents RH, communications internes et supports managériaux mieux standardisés.",
            "Signaux RH et priorités de formation plus visibles pour la direction et les managers.",
        ],
        "summary": "En ressources humaines, la valeur de l’IA ne consiste pas à remplacer le jugement humain. Elle se mesure plutôt dans la réduction des délais, la standardisation des processus, la qualité des synthèses et la capacité à mieux accompagner les équipes et les managers.",
        "priorities": [
            "Recrutement encore chronophage sur la lecture, la présélection et les comptes rendus.",
            "Onboarding inégal selon les équipes, les sites ou la disponibilité des référents.",
            "Documents RH produits sans méthode stable ni trames harmonisées.",
            "Signaux RH, besoins de formation et irritants managériaux repérés trop tard.",
        ],
        "kpis": [
            "Temps de présélection et délai avant premier entretien.",
            "Temps d’intégration et satisfaction à 30 jours.",
            "Temps de production des documents et communications RH.",
            "Temps de reporting RH et nombre de signaux détectés plus tôt.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux processus prioritaires, bootcamp ciblé pour décideurs RH, managers et équipes, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Cycle de recrutement plus fluide sur le tri, les synthèses et la préparation.",
            "Onboarding plus stable sur plusieurs sites ou équipes.",
            "Moins de redites pour les RH et les managers sur les demandes fréquentes.",
            "Première feuille de route claire sur les usages IA à lancer en ressources humaines.",
        ],
        "offer_title": "Bootcamp IA Ressources Humaines : de l’acculturation utile au premier pilote mesurable.",
        "public": "Directions RH, talent acquisition, formation, HRBP, managers de proximité.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA en ressources humaines.",
            "Structurer les usages liés au recrutement, à l’onboarding, aux documents RH et aux signaux managériaux.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Sourcing, tri et synthèse de candidatures.",
            "Onboarding structuré et assistant RH interne.",
            "Documents RH, fiches de poste et communications internes assistés par l’IA.",
            "People analytics et lecture des signaux RH.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux RH, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur le recrutement augmenté et l’onboarding structuré.",
            "Méthodes, prompts, routines et gestes métier autour des synthèses RH et des supports managériaux.",
            "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur les signaux RH, les documents internes et les routines d’accompagnement des managers.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : people analytics, marque employeur, supports de formation et d’intégration.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage RH avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable au recrutement, à l’onboarding, aux documents RH et aux signaux managériaux.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs RH, managers et équipes ciblées.",
        ],
        "domains": [
            "Recrutement et talent acquisition",
            "Onboarding et intégration",
            "HRBP et accompagnement des managers",
            "Formation et expérience collaborateur",
            "People analytics et pilotage RH",
        ],
    },
    {
        "slug": "04_juridique-conformite",
        "title": "Juridique & Conformité",
        "cover_line": "pour les directions juridiques et conformité",
        "meta_1": "Juridique · Conformité · Contentieux · Contrôle interne · Risques",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour accélérer la recherche, la revue documentaire, la préparation des dossiers et la conformité, sans fragiliser le contrôle humain.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les tâches de préparation les plus répétitives, structurer un cadre de confiance robuste et lancer un premier pilote mesurable autour de la veille, de l’analyse contractuelle, de la conformité documentaire ou de la préparation des comités.",
        "benefits": [
            "Recherche juridique et veille plus rapides avant revue experte.",
            "Revue de contrats et repérage d’écarts plus homogènes entre équipes.",
            "Notes, dossiers et synthèses préparatoires plus lisibles pour les décideurs.",
            "Cadre documentaire plus stable sur les sujets de conformité, données et procédures.",
        ],
        "summary": "Dans le juridique et la conformité, la valeur de l’IA ne réside pas dans une automatisation aveugle du jugement. Elle se mesure plutôt dans la vitesse de préparation, la qualité des synthèses, la cohérence des revues et la traçabilité du travail avant validation humaine.",
        "priorities": [
            "Recherche et veille encore longues sur les textes, clauses, repères sectoriels et jurisprudences utiles.",
            "Analyse contractuelle répétitive sur des points récurrents, malgré des motifs connus.",
            "Réponses et procédures de conformité encore peu homogènes selon les équipes et les dossiers.",
            "Préparation de comités, contentieux ou due diligence encore lourde en collecte et reformulation.",
        ],
        "kpis": [
            "Temps de préparation d’une note juridique et délai de réponse interne.",
            "Temps de revue contractuelle et nombre d’écarts repérés plus tôt.",
            "Temps de réponse sur conformité, données et procédures.",
            "Temps de préparation de dossier et délai avant comité.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux processus prioritaires, bootcamp ciblé pour décideurs et équipes juridiques, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Préparation documentaire plus rapide sur la veille, la revue et la synthèse.",
            "Analyse contractuelle plus cohérente grâce à des trames, checklists et routines encadrées.",
            "Conformité plus lisible sur les données, les procédures et les obligations internes.",
            "Première feuille de route sur les usages IA à lancer en juridique et conformité.",
        ],
        "offer_title": "Bootcamp IA Juridique & Conformité : de l’acculturation utile au premier pilote mesurable.",
        "public": "Directions juridiques, conformité, secrétariat général, audit, risques.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA dans le juridique et la conformité.",
            "Structurer les usages liés à la recherche, à la revue contractuelle, à la conformité documentaire et à la préparation des dossiers.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Recherche juridique assistée par l’IA.",
            "Analyse de contrats et repérage d’écarts.",
            "Rédaction juridique et synthèses préparatoires assistées.",
            "Conformité des données, procédures et obligations internes.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux du juridique et de la conformité, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur la recherche juridique assistée et l’analyse de contrats.",
            "Méthodes, prompts, routines et gestes métier autour des synthèses, notes et revues documentaires.",
            "Cadre de confiance : confidentialité, validation experte, traçabilité et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur la conformité documentaire, la préparation des dossiers et les routines de revue.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : données personnelles, anti-blanchiment, due diligence et comités.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage juridique avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable à la recherche, aux contrats, à la conformité documentaire et à la préparation des dossiers.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs et équipes juridiques ciblées.",
        ],
        "domains": [
            "Recherche juridique et veille réglementaire",
            "Analyse contractuelle et clauses sensibles",
            "Conformité documentaire et procédures internes",
            "Protection des données et obligations réglementaires",
            "Préparation des comités, contentieux et due diligence",
        ],
    },
    {
        "slug": "05_transport-logistique",
        "title": "Transport & Logistique",
        "cover_line": "pour les directions transport, logistique et opérations",
        "meta_1": "Exploitation · Supply chain · Flotte · Hubs · Service client · Direction opérations",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour améliorer la coordination terrain, réduire les retards, mieux lire les anomalies et renforcer la qualité de service en environnement distribué.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les flux les plus sensibles, renforcer la lecture opérationnelle des incidents et lancer un premier pilote mesurable autour du cockpit d’exploitation, du suivi d’anomalies, de l’interface service client ou de la maintenance assistée.",
        "benefits": [
            "Lecture plus rapide des flux, retards et incidents pour l’exploitation.",
            "Meilleure coordination entre terrain, management et relation client.",
            "KPI plus commentés, plus lisibles et plus utiles aux arbitrages quotidiens.",
            "Réponses client plus homogènes en cas de perturbation ou d’anomalie.",
        ],
        "summary": "Dans le transport et la logistique, la valeur de l’IA se lit dans la rapidité de coordination, la lecture plus fine des retards, la qualité du pilotage opérationnel et la capacité à relier exploitation, maintenance et service client autour d’une information plus fiable.",
        "priorities": [
            "Informations terrain encore dispersées entre plusieurs canaux, équipes et outils.",
            "Pilotage des flux encore peu commenté malgré des KPI déjà disponibles.",
            "Opérations et service client encore trop séparés en cas de retard ou d’incident.",
            "Maintenance et anomalies encore traitées trop tard faute de lecture consolidée.",
        ],
        "kpis": [
            "Temps de lecture des incidents et réactivité de coordination.",
            "Temps de détection et nombre d’anomalies qualifiées plus tôt.",
            "Temps de réponse client et taux de résolution au premier contact.",
            "Délai d’escalade et temps d’immobilisation évité.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux processus prioritaires, bootcamp ciblé pour décideurs et équipes opérations, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Moins de retards subis grâce à une meilleure lecture des points de friction.",
            "Coordination plus rapide entre le terrain et les managers.",
            "Service plus homogène côté client grâce à une information plus claire en cas d’incident.",
            "Première feuille de route claire sur les usages IA à lancer en transport et logistique.",
        ],
        "offer_title": "Bootcamp IA Transport & Logistique : de l’acculturation utile au premier pilote mesurable.",
        "public": "Exploitation, supply chain, flotte, hubs, service client, direction opérations.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA dans le transport et la logistique.",
            "Structurer les usages liés au pilotage des opérations, au suivi d’anomalies, à l’interface service client et à la maintenance assistée.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Cockpit synthétique des opérations terrain.",
            "Gestion de flotte et suivi d’anomalies.",
            "Interface opérations et service client.",
            "Maintenance et support prédictif.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux transport et logistique, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur le cockpit des opérations et le suivi d’anomalies.",
            "Méthodes, prompts, routines et gestes métier autour de la lecture des incidents, flux et priorités terrain.",
            "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur l’interface service client, la maintenance assistée et les routines d’exploitation.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : flotte, approvisionnement, relation client et support prédictif.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage transport et logistique avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable au pilotage des flux, aux anomalies, au service client et à la maintenance.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs et équipes opérations ciblées.",
        ],
        "domains": [
            "Pilotage des opérations terrain",
            "Gestion de flotte et suivi d’exceptions",
            "Interface exploitation et service client",
            "Maintenance et support prédictif",
            "Supply chain, hubs et coordination opérationnelle",
        ],
    },
    {
        "slug": "06_energie-petrole",
        "title": "Énergie & Pétrole",
        "cover_line": "pour les directions énergie, pétrole et opérations sensibles",
        "meta_1": "Opérations · Maintenance · HSE · Finance · Direction de site · Réseau et distribution",
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour mieux piloter les opérations, les comptes, la maintenance et la qualité de service dans des environnements sensibles.",
        "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les flux critiques, structurer la lecture des alertes et lancer un premier pilote mesurable autour du pilotage des comptes, de la maintenance assistée, du service client réseau ou de la conformité HSE.",
        "benefits": [
            "Lecture plus rapide des comptes, alertes et opérations pour la direction et l’exploitation.",
            "Meilleure réactivité sur les anomalies, incidents et historiques de maintenance.",
            "Réponses plus homogènes côté service client et réseau de distribution.",
            "Cadre documentaire plus solide sur la conformité HSE, les standards et les audits.",
        ],
        "summary": "Dans l’énergie et le pétrole, la valeur de l’IA se mesure dans la capacité à mieux relier opérations, maintenance, HSE, finance et réseau de distribution autour d’une lecture consolidée, commentée et utile à la décision, sans compromis sur la supervision humaine.",
        "priorities": [
            "Multi-activités encore difficiles à relier entre comptes, distribution, maintenance, réseau et service.",
            "Pilotage des comptes et opérations encore trop lent pour les responsables de terrain et de direction.",
            "Maintenance et anomalies encore coûteuses faute de lecture plus précoce des signaux utiles.",
            "Cadre HSE et conformité documentaire encore exigeant en temps, rigueur et coordination.",
        ],
        "kpis": [
            "Temps de préparation de comité et délai de relance ou de décision.",
            "Temps de détection et temps de réaction sur incidents et écarts critiques.",
            "Temps moyen de réponse et taux de résolution au premier niveau.",
            "Temps de préparation documentaire et nombre de points de contrôle couverts.",
        ],
        "approach": "Audit IA gratuit, cadrage de deux processus prioritaires, bootcamp ciblé pour décideurs et équipes opérations, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
        "results": [
            "Lecture plus rapide des comptes et opérations grâce à des synthèses et alertes mieux structurées.",
            "Réactivité accrue sur les anomalies et la maintenance grâce à une lecture plus exploitable des signaux.",
            "Cadre mieux maîtrisé sur HSE, documentation, standards et conformité.",
            "Première feuille de route claire sur les usages IA à lancer en énergie et pétrole.",
        ],
        "offer_title": "Bootcamp IA Énergie & Pétrole : de l’acculturation utile au premier pilote mesurable.",
        "public": "Opérations, maintenance, HSE, finance, direction de site, coordination réseau et distribution.",
        "objectives": [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA dans l’énergie et le pétrole.",
            "Structurer les usages liés au pilotage des comptes, à la maintenance assistée, au service client réseau et à la conformité HSE.",
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [
            "Pilotage des comptes professionnels et de la performance.",
            "Maintenance, support prédictif et incidents.",
            "Service client et réseau de distribution.",
            "Conformité HSE et gouvernance documentaire.",
            "Cadre de confiance, gouvernance et supervision humaine des usages IA.",
        ],
        "j1": [
            "Lecture des enjeux énergie et pétrole, cartographie des irritants prioritaires et sélection des premiers gains visés.",
            "Études de cas guidées sur le pilotage des comptes et la maintenance assistée.",
            "Méthodes, prompts, routines et gestes métier autour des alertes, synthèses, dossiers et flux critiques.",
            "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
        ],
        "j2": [
            "Ateliers de conception sur le service client réseau, la conformité HSE et la gouvernance documentaire.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : performance, maintenance, qualité documentaire et lecture des KPI.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            "Une note de cadrage énergie et pétrole avec priorités confirmées, risques, points de vigilance et premières décisions.",
            "Une matrice des usages prioritaires applicable au pilotage des comptes, à la maintenance, au service client réseau et à la conformité HSE.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs et équipes opérations ciblées.",
        ],
        "domains": [
            "Pilotage des comptes et de la performance",
            "Maintenance et support prédictif",
            "Service client et réseau de distribution",
            "Conformité HSE et gouvernance documentaire",
            "Opérations critiques, réseau et coordination de site",
        ],
    },
]


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")


AUTO_SECTOR_META = {
    "assistanat-et-secretariat": {
        "cover_line": "pour les directions générales, assistanats et secrétariats exécutifs",
        "meta_1": "Direction générale · Assistanat · Secrétariat exécutif · Coordination · Support",
        "public": "Directions générales, assistantes de direction, secrétariats exécutifs, coordinations administratives.",
        "sector_phrase": "dans l’assistanat et le secrétariat",
        "team_phrase": "assistanat et secrétariat",
    },
    "marketing-et-communication": {
        "cover_line": "pour les directions marketing et communication",
        "meta_1": "Marketing · Communication · Digital · Brand · Acquisition · Contenu",
        "public": "Directions marketing, communication, brand, digital, contenu et acquisition.",
        "sector_phrase": "dans le marketing et la communication",
        "team_phrase": "marketing et communication",
    },
    "data-analyse": {
        "cover_line": "pour les directions data, BI et performance",
        "meta_1": "Data · BI · Reporting · Performance · Contrôle · Décision",
        "public": "Directions data, BI, performance, contrôle, équipes d’analystes et managers utilisateurs.",
        "sector_phrase": "dans la data et l’analyse",
        "team_phrase": "data et analyse",
    },
    "administration-et-gestion": {
        "cover_line": "pour les directions administration et gestion",
        "meta_1": "Administration · Back office · Coordination · Opérations support · Gestion",
        "public": "Directions administratives, opérations support, back office, coordination multi-sites.",
        "sector_phrase": "dans l’administration et la gestion",
        "team_phrase": "administration et gestion",
    },
    "management-et-leadership": {
        "cover_line": "pour les comités de direction et managers",
        "meta_1": "Direction générale · Leadership · Management · Transformation · Pilotage",
        "public": "Comités de direction, top management, managers de BU, responsables de transformation.",
        "sector_phrase": "dans le management et le leadership",
        "team_phrase": "management et leadership",
    },
    "it-et-transformation-digitale": {
        "cover_line": "pour les directions IT et transformation digitale",
        "meta_1": "DSI · PMO · Produit · Support · Transformation · Gouvernance",
        "public": "DSI, transformation digitale, PMO, équipes produit, support et gouvernance.",
        "sector_phrase": "dans l’IT et la transformation digitale",
        "team_phrase": "IT et transformation digitale",
    },
    "formation-et-pedagogie": {
        "cover_line": "pour les directions pédagogiques et formation",
        "meta_1": "Pédagogie · Formation · Ingénierie · Cohortes · LMS · Contenus",
        "public": "Directions pédagogiques, responsables académiques, formateurs, ingénierie de formation.",
        "sector_phrase": "dans la formation et la pédagogie",
        "team_phrase": "formation et pédagogie",
    },
    "sante-et-bien-etre": {
        "cover_line": "pour les directions santé, prévention et bien-être",
        "meta_1": "Santé · Prévention · QVT · HSE · Qualité · Programmes",
        "public": "Directions d’établissements, coordination soignante, accueil, expérience patient et support.",
        "sector_phrase": "dans la santé et le bien-être",
        "team_phrase": "santé et bien-être",
    },
    "diplomatie-et-affaires-internationales": {
        "cover_line": "pour les directions diplomatie et affaires internationales",
        "meta_1": "Diplomatie · Coopération · Affaires publiques · Veille · Missions · Multilingue",
        "public": "Affaires publiques, coopération internationale, organisations régionales, cellules diplomatiques, direction pays.",
        "sector_phrase": "dans la diplomatie et les affaires internationales",
        "team_phrase": "diplomatie et affaires internationales",
    },
    "telecommunications": {
        "cover_line": "pour les directions télécommunications et services digitaux",
        "meta_1": "Réseau · Service client · Mobile money · Distribution · Marketing · Performance",
        "public": "Direction générale, service client, opérations réseau, marketing, vente B2B/B2C, qualité de service, mobile money.",
    },
}


def slugify(value: str) -> str:
    value = value.lower()
    replacements = {
        "é": "e",
        "è": "e",
        "ê": "e",
        "ë": "e",
        "à": "a",
        "â": "a",
        "ä": "a",
        "î": "i",
        "ï": "i",
        "ô": "o",
        "ö": "o",
        "ù": "u",
        "û": "u",
        "ü": "u",
        "ç": "c",
        "œ": "oe",
        "’": "-",
        "'": "-",
        "&": "-",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def display_slug(slug: str) -> str:
    slug_value = slug.split("_", 1)[-1]
    return slug_value.replace("-", "_")


def commercial_sector_label(sector: dict) -> str:
    custom_labels = {
        "01_finance-comptabilite": "Finance_Comptabilite",
        "02_service-client": "Service_Client",
        "03_ressources-humaines": "Ressources_Humaines",
        "04_juridique-conformite": "Juridique_Conformite",
        "05_transport-logistique": "Transport_Logistique",
        "06_energie-petrole": "Energie_Petrole",
        "07_assistanat-secretariat": "Assistanat_Secretariat",
        "08_marketing-communication": "Marketing_Communication",
        "09_data-analyse": "Data_Analyse",
        "10_administration-gestion": "Administration_Gestion",
        "11_management-leadership": "Management_Leadership",
        "12_it-transformation-digitale": "IT_Transformation_Digitale",
        "13_formation-pedagogie": "Formation_Pedagogie",
        "14_sante-bien-etre": "Sante_Bien_Etre",
        "15_diplomatie-affaires-internationales": "Diplomatie_Affaires_Internationales",
        "16_telecommunications": "Telecommunications",
    }
    if sector["slug"] in custom_labels:
        return custom_labels[sector["slug"]]
    return display_slug(slugify(sector["title"])).title().replace("_", "")


def commercial_pdf_filename(sector: dict) -> str:
    return f"TransferAI_Mini_Catalogue_{commercial_sector_label(sector)}.pdf"


def standard_docx_filename(sector: dict) -> str:
    return f"Mini_Catalogue_TransferAI_{sector['slug']}_Standard_{CATALOGUE_DATE}.docx"


def load_js_specs() -> list[dict]:
    node_code = """
import('./scripts/presentations/transferai-sector-decks/domain-specs.mjs').then(async (m) => {
  const specs = await m.getSectorDeckSpecs();
  console.log(JSON.stringify(specs));
});
"""
    result = subprocess.run(
        ["node", "-e", node_code],
        cwd=str(ROOT),
        capture_output=True,
        text=True,
        check=True,
    )
    return json.loads(result.stdout)


def normalize_sentence(text: str) -> str:
    return text.strip().rstrip(".") + "."


def to_mid_sentence(text: str) -> str:
    preserved_prefixes = ("IA ", "IT ", "BI ", "QVT", "HSE", "RGPD", "SQL", "LMS", "Power BI")
    for prefix in preserved_prefixes:
        if text.startswith(prefix):
            return text
    return text[0].lower() + text[1:] if text else text


def make_auto_sector(spec: dict, index: int) -> dict:
    meta = AUTO_SECTOR_META[spec["slug"]]
    priorities = [normalize_sentence(item["before"]) for item in spec["priorities"][:4]]
    kpis = [normalize_sentence(item["kpi"]) for item in spec["useCases"][:4]]
    benefits = [normalize_sentence(item["after"]) for item in spec["useCases"][:4]]
    domains = [item["title"] for item in spec["useCases"][:4]] + spec["trainingFocus"][:1]
    title = spec["title"]
    recommended = ", ".join(spec["trainingFocus"][2:5]).replace("'", "’")
    sector_phrase = meta["sector_phrase"]
    team_phrase = meta["team_phrase"]

    return {
        "slug": f"{index:02d}_{slugify(spec['title'])}",
        "title": title,
        "cover_line": meta["cover_line"],
        "meta_1": meta["meta_1"],
        "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
        "promise": normalize_sentence(
            "Transformer l’IA en usages concrets, activables et gouvernés pour fluidifier les processus clés, renforcer la coordination et générer des gains de temps mesurables"
        ),
        "positioning": normalize_sentence(
            f"TransferAI propose une trajectoire simple et pilotable : identifier les usages prioritaires et lancer un premier pilote mesurable autour de {to_mid_sentence(spec['useCases'][0]['title'])}, de {to_mid_sentence(spec['useCases'][1]['title'])} ou de {to_mid_sentence(spec['useCases'][2]['title'])}"
        ),
        "benefits": benefits,
        "summary": normalize_sentence(
            f"{sector_phrase[0].upper() + sector_phrase[1:]}, la valeur de l’IA se mesure d’abord dans la réduction des délais, la qualité des synthèses et une meilleure coordination des décisions"
        ),
        "priorities": priorities,
        "kpis": kpis,
        "approach": normalize_sentence(
            f"Audit IA gratuit, cadrage de deux priorités, bootcamp ciblé pour décideurs et équipes {team_phrase}, puis accompagnement sur 90 jours avec KPI simples et supervision humaine"
        ),
        "results": [
            normalize_sentence(f"{item['value']} {item['label']} ({item['note']})") for item in spec["roiMetrics"][:2]
        ]
        + ["Feuille de route initiale des usages IA prioritaires du secteur."],
        "offer_title": f"Bootcamp IA {title} : de l’acculturation utile au premier pilote mesurable.",
        "public": meta["public"],
        "objectives": [
            normalize_sentence(f"Identifier les priorités métier les plus crédibles à traiter par l’IA {sector_phrase}"),
            normalize_sentence(f"Structurer les usages liés à {to_mid_sentence(spec['useCases'][0]['title'])}, à {to_mid_sentence(spec['useCases'][1]['title'])} et à {to_mid_sentence(spec['useCases'][2]['title'])}"),
            "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
        "trainings": [normalize_sentence(name) for name in spec["trainingFocus"][:5]],
        "j1": [
            normalize_sentence(f"Lecture des enjeux {team_phrase}, cartographie des irritants prioritaires et sélection des premiers gains visés"),
            normalize_sentence(f"Études de cas guidées sur {to_mid_sentence(spec['useCases'][0]['title'])} et {to_mid_sentence(spec['useCases'][1]['title'])}"),
            normalize_sentence(f"Méthodes, prompts, routines et gestes métier autour de {to_mid_sentence(spec['useCases'][2]['title'])} et des synthèses opérationnelles"),
            "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
        ],
        "j2": [
            normalize_sentence(f"Ateliers de conception sur {to_mid_sentence(spec['useCases'][2]['title'])} et {to_mid_sentence(spec['useCases'][3]['title'])}"),
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            normalize_sentence(f"Approfondissements recommandés : {recommended}"),
            "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
        ],
        "deliverables": [
            normalize_sentence(f"Une note de cadrage {team_phrase} avec priorités confirmées, risques, points de vigilance et premières décisions"),
            normalize_sentence(f"Une matrice des usages prioritaires applicable à {to_mid_sentence(spec['useCases'][0]['title'])}, à {to_mid_sentence(spec['useCases'][1]['title'])} et à {to_mid_sentence(spec['useCases'][2]['title'])}"),
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            normalize_sentence(f"Une recommandation de parcours de montée en compétences pour les décideurs et équipes {team_phrase} ciblées"),
        ],
        "domains": domains[:5],
    }


def build_all_sectors() -> list[dict]:
    existing_titles = {sector["title"] for sector in SECTORS}
    js_specs = load_js_specs()
    auto_specs = [spec for spec in js_specs if spec["title"] not in existing_titles and spec["slug"] in AUTO_SECTOR_META]
    start_index = len(SECTORS) + 1
    built = list(SECTORS)
    for offset, spec in enumerate(auto_specs):
        built.append(make_auto_sector(spec, start_index + offset))
    built.append(
        {
            "slug": "16_telecommunications",
            "title": "Télécommunications",
            "cover_line": "pour les directions télécommunications et services digitaux",
            "meta_1": "Réseau · Service client · Mobile money · Distribution · Marketing · Performance",
            "meta_2": "Document de présentation sectoriel | Présentiel · Hybride · En ligne",
            "promise": "Transformer l’IA en usages concrets, activables et gouvernés pour améliorer la qualité de service, la coordination opérationnelle, la performance commerciale et la vitesse de décision dans les télécommunications.",
            "positioning": "TransferAI propose une trajectoire simple et pilotable : identifier les usages prioritaires, former les équipes concernées, sécuriser le cadre de confiance et lancer un premier pilote utile sur la base de KPI concrets.",
            "benefits": [
                "Réponses plus rapides, plus homogènes et mieux traçables sur les canaux clients.",
                "Incidents réseau mieux relayés entre exploitation, relation client, communication interne et management.",
                "Lecture plus rapide des verbatims, signaux de churn, irritants service et opportunités d’amélioration.",
                "Meilleure qualité de préparation des comités commerciaux et opérationnels grâce à des synthèses plus exploitables.",
            ],
            "summary": "Dans les télécommunications, la valeur de l’IA se mesure rapidement à travers la stabilité du service, la vitesse de réponse client, la lisibilité des incidents, la qualité commerciale et la capacité des équipes à mieux coordonner leurs décisions.",
            "priorities": [
                "Incidents réseau parfois mal relayés entre équipes techniques, front office et management.",
                "Relation client sous pression sur les demandes récurrentes, les réclamations et la qualité de réponse multi-canaux.",
                "Lecture commerciale encore dispersée sur churn, upsell, satisfaction et qualité de service.",
                "Coordination interéquipes à renforcer entre réseau, service client, marketing, mobile money et direction.",
            ],
            "kpis": [
                "Temps moyen de réponse et taux de résolution au premier contact.",
                "Temps de diffusion d’une synthèse d’incident fiable.",
                "Délai d’analyse des verbatims et nombre de signaux priorisés.",
                "Temps de préparation d’un comité et vitesse de décision managériale.",
            ],
            "approach": "Audit IA gratuit, cadrage des deux premiers cas d’usage prioritaires, bootcamp ciblé pour décideurs et équipes clés, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
            "results": [
                "Réponses plus rapides sur les canaux clients grâce à des scripts, synthèses et aides à la qualification.",
                "Incidents mieux pilotés côté réseau et exploitation grâce à une lecture commune des impacts.",
                "Première feuille de route sur les usages IA à lancer en télécommunications.",
            ],
            "offer_title": "Bootcamp IA Télécommunications : de l’acculturation utile au premier pilote mesurable.",
            "public": "Direction générale, service client, opérations réseau, marketing, vente B2B/B2C, qualité de service, mobile money.",
            "objectives": [
                "Identifier les priorités métier les plus crédibles à traiter par l’IA dans les télécommunications.",
                "Structurer les usages liés à la relation client, à la lecture des incidents, au churn et au pilotage commercial.",
                "Distinguer les usages activables rapidement, ceux à encadrer et ceux à piloter dans le temps.",
                "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
            ],
            "trainings": [
                "Service client augmenté par l’IA.",
                "Analyse de sentiment et feedback client.",
                "Relation client augmentée par l’IA.",
                "Pilotage réseau et qualité de service avec l’IA.",
                "Leadership data-driven avec l’IA.",
            ],
            "j1": [
                "Lecture des enjeux télécoms, cartographie des irritants prioritaires et sélection des premiers gains visés.",
                "Études de cas guidées sur la relation client augmentée et la synthèse d’incidents réseau.",
                "Méthodes, prompts, routines et gestes métier autour des verbatims, signaux de churn et synthèses opérationnelles.",
                "Cadre de confiance : confidentialité, validation managériale, supervision humaine et critères d’extension.",
            ],
            "j2": [
                "Ateliers de conception sur le pilotage commercial, la qualité de service et les routines interéquipes.",
                "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
                "Approfondissements recommandés : mobile money, distribution, marketing télécom et pilotage réseau.",
                "Plan d’action à 90 jours : gouvernance, équipes à former, priorités de déploiement et décision d’extension.",
            ],
            "deliverables": [
                "Une note de cadrage télécommunications avec priorités confirmées, risques, points de vigilance et premières décisions.",
                "Une matrice des usages prioritaires applicable au service client, au réseau, au churn et au pilotage commercial.",
                "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
                "Une recommandation de parcours de montée en compétences pour les décideurs et équipes télécoms ciblées.",
            ],
            "domains": [
                "Service client et relation omnicanale",
                "Incidents réseau et coordination opérationnelle",
                "Lecture des verbatims et signaux de churn",
                "Pilotage commercial et qualité de service",
                "Mobile money, distribution et performance",
            ],
        }
    )
    return built


def set_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def set_font(run, name: str, size: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "F1A56E", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "TransferAI Africa | www.transferai.ci | IA Assistant / WhatsApp : +225 07 16 57 39 90"
    )
    set_font(run, "Aptos", 8.8, GRAY)


def add_section_heading(document: Document, text: str, top_space: float = 10, font_size: float = 19.5, new_page: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = new_page
    paragraph.paragraph_format.space_before = Pt(top_space)
    paragraph.paragraph_format.space_after = Pt(3)
    marker = paragraph.add_run("▌ ")
    set_font(marker, "Aptos", 18, ORANGE, bold=True)
    run = paragraph.add_run(text)
    set_font(run, "Aptos Display", font_size, NAVY, bold=True)


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, "Aptos Display", 15.1, NAVY, bold=True)


def add_compact_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_font(run, "Aptos Display", 13.9, NAVY, bold=True)


def add_body(document: Document, text: str, color: RGBColor = NAVY, size: float = 12.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, "Aptos", size, color)


def add_bullets(document: Document, items: list[str], color: RGBColor = NAVY, size: float = 12.5) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(item)
        set_font(run, "Aptos", size, color)


def add_compact_bullets(document: Document, items: list[str], color: RGBColor = NAVY, size: float = 11.9) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(item)
        set_font(run, "Aptos", 11.7, color)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    remove_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell)
    cell.width = Cm(16.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    run1 = p1.add_run(title)
    set_font(run1, "Aptos Display", 16, NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    set_font(run2, "Aptos", 12.5, NAVY)


def add_meta_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, "Aptos", 12.5, GRAY)


def add_cover(document: Document, sector: dict) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("Mini-catalogue sectoriel | IA métier, formation et accompagnement")
    set_font(run, "Aptos", 11.5, GRAY)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("MINI-CATALOGUE")
    set_font(run, "Aptos Display", 18, ORANGE, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SOLUTIONS IA ET FORMATIONS STRATÉGIQUES")
    set_font(run, "Aptos Display", 26, NAVY, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(sector["cover_line"])
    set_font(run, "Aptos Display", 19, GREEN, bold=True)

    add_meta_line(document, sector["meta_1"])
    add_meta_line(document, sector["meta_2"])

    add_section_heading(document, "PROMESSE CENTRALE")
    add_body(document, sector["promise"], size=13)

    add_section_heading(document, "POSITIONNEMENT DE L’OFFRE", top_space=8)
    add_callout(document, "Proposition ciblée", sector["positioning"])

    add_section_heading(document, "BÉNÉFICES MÉTIER VISIBLES", top_space=8)
    add_bullets(document, sector["benefits"], size=12.2)


def add_exec_summary(document: Document, sector: dict) -> None:
    add_section_heading(document, f"01 - SYNTHÈSE EXÉCUTIVE {sector['title'].upper()}", top_space=0, new_page=True)
    add_body(document, sector["summary"])

    add_subheading(document, "Priorités métier les plus crédibles")
    add_bullets(document, sector["priorities"])

    add_subheading(document, "KPI de référence")
    add_bullets(document, sector["kpis"])

    add_subheading(document, "Approche recommandée")
    add_callout(document, "Séquence recommandée", sector["approach"])

    add_subheading(document, "Résultats attendus")
    add_bullets(document, sector["results"])


def add_offer(document: Document, sector: dict) -> None:
    add_section_heading(document, "02 - FORMULE PRIORITAIRE", top_space=0, new_page=True)
    add_compact_subheading(document, "Intitulé")
    add_body(document, sector["offer_title"], size=12.1)

    add_compact_subheading(document, "Public cible")
    add_body(document, sector["public"], size=12.1)

    add_compact_subheading(document, "Formats possibles")
    add_compact_bullets(document, ["Présentiel", "Hybride", "En ligne (style webinaire)"])

    add_compact_subheading(document, "Niveaux")
    add_compact_bullets(document, ["Débutant", "Intermédiaire", "Avancé"])

    add_compact_subheading(document, "Durée")
    add_compact_bullets(document, ["1 jour : version décideurs et managers.", "2 jours : version équipes, managers et mise en pratique."])

    add_compact_subheading(document, "Tarif")
    add_body(document, "Prix sur demande.", size=12.1)

    add_compact_subheading(document, "Objectifs pédagogiques")
    add_compact_bullets(document, sector["objectives"])

    add_compact_subheading(document, "Formations proposées")
    add_compact_bullets(document, sector["trainings"])


def add_programme(document: Document, sector: dict) -> None:
    add_section_heading(document, "03 - PROGRAMME PROPOSÉ", top_space=0, new_page=True)
    add_subheading(document, "Jour 1")
    add_compact_bullets(document, sector["j1"])

    add_subheading(document, "Jour 2")
    add_compact_bullets(document, sector["j2"])


def add_deliverables(document: Document, sector: dict) -> None:
    add_section_heading(document, "04 - LIVRABLES ATTENDUS", top_space=8)
    add_compact_bullets(document, sector["deliverables"])

    add_section_heading(document, "05 - DOMAINES DE RÉFÉRENCE", top_space=8)
    add_compact_bullets(document, sector["domains"])


def build_document(sector: dict) -> Document:
    document = Document()
    set_page_geometry(document)
    add_footer(document.sections[0])
    add_cover(document, sector)
    add_exec_summary(document, sector)
    add_offer(document, sector)
    add_programme(document, sector)
    add_deliverables(document, sector)
    return document


def main() -> None:
    for sector in build_all_sectors():
        output_dir = BASE_OUTPUT_DIR / sector["slug"]
        output_dir.mkdir(parents=True, exist_ok=True)
        docx_path = output_dir / standard_docx_filename(sector)
        document = build_document(sector)
        document.save(docx_path)
        print(docx_path)


if __name__ == "__main__":
    main()
