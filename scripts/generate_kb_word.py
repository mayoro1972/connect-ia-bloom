"""
Génère le document Word de la Base de Connaissance Complète TransferAI Africa
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

OUTPUT_DIR = "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin"


def margins(doc):
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.8)
        s.right_margin = Cm(2.8)


def titre(doc, text):
    p = doc.add_heading(text, 0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        r.font.size = Pt(22)
    return p


def h1(doc, text):
    p = doc.add_heading(text, 1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    return p


def h2(doc, text):
    return doc.add_heading(text, 2)


def h3(doc, text):
    return doc.add_heading(text, 3)


def corps(doc, text):
    return doc.add_paragraph(text)


def puce(doc, text, niveau=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 * (niveau + 1))
    return p


def num(doc, text):
    return doc.add_paragraph(text, style="List Number")


def code_inline(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.left_indent = Inches(0.35)
    return p


def alerte(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("⚠  " + text)
    r.bold = True
    r.font.color.rgb = RGBColor(0xB4, 0x5D, 0x09)
    return p


def ok(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("✔  " + text)
    r.font.color.rgb = RGBColor(0x16, 0x6B, 0x34)
    return p


def separateur(doc):
    doc.add_paragraph("─" * 88)


def tableau(doc, entetes, lignes, largeurs_cm=None):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.style = "Table Grid"
    r = t.add_row()
    for i, e in enumerate(entetes):
        r.cells[i].text = e
        for run in r.cells[i].paragraphs[0].runs:
            run.bold = True
    for ligne in lignes:
        row = t.add_row()
        for i, val in enumerate(ligne):
            row.cells[i].text = str(val)
    if largeurs_cm:
        for row in t.rows:
            for i, w in enumerate(largeurs_cm):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ─────────────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    margins(doc)

    # ── COUVERTURE ────────────────────────────────────────────────────────────
    titre(doc, "TransferAI Africa")
    titre(doc, "Base de Connaissance Complète du Projet")
    p = doc.add_paragraph(f"Version 1.0 · Mise à jour le {datetime.today().strftime('%d/%m/%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corps(doc, "")
    corps(doc,
        "Ce document constitue la mémoire opérationnelle complète du projet TransferAI Africa. "
        "Il couvre l'ensemble du cycle de vie du projet, depuis les phases fondatrices de 2026 "
        "jusqu'à l'état actuel du pipeline de prospection, des automatisations et de "
        "l'infrastructure technique."
    )
    separateur(doc)
    doc.add_page_break()

    # ── SOMMAIRE ─────────────────────────────────────────────────────────────
    h1(doc, "Sommaire")
    sommaire = [
        "1.  Identité du projet",
        "2.  Historique chronologique du projet",
        "3.  Architecture système actuelle",
        "4.  Base de données Supabase",
        "5.  Inventaire des 12 workflows n8n",
        "6.  Formulaire d'audit — fonctionnement et corrections",
        "7.  Données commerciales et règles métier",
        "8.  Infrastructure et déploiement",
        "9.  Problèmes résolus et points d'attention permanents",
        "10. Offre commerciale TransferAI Africa",
        "11. Index des documents de référence",
        "12. Actions en attente au 28 juin 2026",
    ]
    for s in sommaire:
        puce(doc, s)
    doc.add_page_break()

    # ── 1. IDENTITÉ ───────────────────────────────────────────────────────────
    h1(doc, "1. Identité du projet")

    h2(doc, "Nom et positionnement")
    corps(doc,
        "TransferAI Africa est une plateforme de formation, de conseil, de contenus et de solutions "
        "IA conçue pour les professionnels, les entreprises et les institutions souhaitant intégrer "
        "l'intelligence artificielle avec méthode et impact concret, avec un ancrage prioritaire "
        "en Côte d'Ivoire et en Afrique de l'Ouest."
    )

    h2(doc, "URLs de production")
    tableau(doc,
        ["Ressource", "URL"],
        [
            ["Site principal", "https://www.transferai.ci"],
            ["Formulaire d'audit IA (gratuit)", "https://www.transferai.ci/questionnaire-audit"],
            ["Back-office administrateur", "https://www.transferai.ci/admin"],
            ["Email de contact projet", "contact@transferai.ci"],
            ["Email de l'administrateur", "marius.ayoro70@gmail.com"],
        ],
        largeurs_cm=[5, 11]
    )

    h2(doc, "Stack technologique")
    tableau(doc,
        ["Composant", "Technologie"],
        [
            ["Interface utilisateur", "React + TypeScript + Vite + Tailwind CSS"],
            ["Base de données et backend", "Supabase (PostgreSQL + Edge Functions Deno)"],
            ["Déploiement", "Cloudflare Pages (site) + Cloudflare Workers"],
            ["Automatisation", "n8n (12 workflows actifs)"],
            ["Intelligence artificielle", "GPT-4 / GPT-4.1-mini via API OpenAI"],
            ["Envoi d'e-mails transactionnels", "Resend API"],
            ["Chat web", "Chatwoot"],
            ["Prise de rendez-vous", "Calendly"],
            ["CRM e-mail entrant", "Zoho Mail"],
            ["Référentiel de code", "GitHub — mayoro1972/connect-ia-bloom"],
            ["Branche de production", "main"],
        ],
        largeurs_cm=[6, 10]
    )
    doc.add_page_break()

    # ── 2. HISTORIQUE ─────────────────────────────────────────────────────────
    h1(doc, "2. Historique chronologique du projet")

    h2(doc, "Phases 1 à 20 — Construction du socle (avant mai 2026)")
    corps(doc,
        "Le projet a démarré comme un site vitrine institutionnel orienté formation. "
        "Les vingt premières phases ont progressivement construit l'ensemble de la plateforme :"
    )
    phases_initiales = [
        ("Phases 1 à 3", "Socle front : accueil, catalogue, parcours, certification, blog"),
        ("Phases 4 à 6", "Pipeline éditorial IA : blog dynamique, veille, brouillons, matrice outils IA"),
        ("Phases 7 à 9", "Newsletter par domaine : inscription, livraison, automatisation hebdomadaire"),
        ("Phases 10 à 14", "Fiabilisation des flux admin, simplification UX, pipeline partenaires, emails transactionnels"),
        ("Phases 15 à 16", "Newsletter fondatrice envoyée, page média, capsules vidéo TikTok pilotées par backend"),
        ("Phases 17 à 18", "Pipeline WhatsApp complet (Twilio → Supabase → BackOffice → notifications e-mail internes)"),
        ("Phases 19 à 20", "Documentation BackOffice, gouvernance documentaire, intégration Chatwoot website"),
    ]
    for phase, desc in phases_initiales:
        puce(doc, f"{phase} — {desc}")
    corps(doc, "")
    corps(doc,
        "Au 2 mai 2026, le projet n'est plus un simple site institutionnel : c'est une vitrine "
        "commerciale crédible, une base éditoriale structurée, un cockpit d'administration "
        "et un socle concret pour les automatisations IA avancées."
    )

    h2(doc, "Phase 21 — Pipeline de prospection IA (mai 2026)")
    corps(doc,
        "Démarrage le 17 mai 2026. Objectif : construire un pipeline de prospection B2B "
        "automatisée avec n8n et GPT-4 pour générer des packs commerciaux personnalisés."
    )
    tableau(doc,
        ["Version", "Date", "Description"],
        [
            ["V1", "22 mai 2026", "Prospect unique, manuel. Scraping → anonymisation (ORG_TARGET) → GPT-4 → pack commercial. Testé et validé."],
            ["V2", "22 mai 2026", "Ajout validation humaine avant envoi + stockage Supabase (ai_prospecting_packs)."],
            ["V3", "22–29 mai 2026", "Pipeline CRM complet (51 nœuds). Génération 4 assets : Executive Letter, Catalogue, Formulaire audit, Deck Brief. 2 variantes : [FINAL] et [V20-FIXED]."],
            ["V4", "22–29 mai 2026", "Batch orchestrator quotidien (3–5 prospects/jour). Sources : Supabase, Airtable, Google Sheets. Routage dynamique vers V3 ou Follow-Up."],
        ],
        largeurs_cm=[2.5, 3, 10.5]
    )

    h2(doc, "Phase 22 — Consolidation CRM (fin mai – début juin 2026)")
    etapes_22 = [
        ("29 mai – 1er juin", "Debug intensif V3 : fix scraping, fix champs Supabase, fix format conditions n8n. Résolution erreur filtre quota (delivery_status = submitted)."),
        ("2 juin", "Génération des guides Word complets V4/V5. Configuration du workflow Zoho Reply à partir d'un audio m4a."),
        ("8 juin", "V6 Post-Audit Expert Routing configuré et branché. Schéma architecture V1→V6. Plan dashboard BackOffice web."),
        ("13 juin", "Audit complet CRM : 4 écarts majeurs identifiés et corrigés. Ajout du workflow V89 Follow-Up automatique."),
    ]
    for date, desc in etapes_22:
        puce(doc, f"{date} : {desc}")

    h2(doc, "Phase 23 — Canaux entrants et social (15–23 juin 2026)")
    etapes_23 = [
        ("15 juin", "Google Forms Social Lead Sequence V1 mise en service."),
        ("17 juin", "Admin Controller V1 configuré (77 nœuds). Debug des conditions If (format booléen v1 cassé → migration vers format v2)."),
        ("18 juin", "Session troubleshooting : fix Edge Functions et Cloudflare. Chatwoot WhatsApp IA en direct."),
        ("19 juin", "Plan intégration Zoho Inbound + BackOffice."),
        ("20 juin", "Google Forms Social Lead Sequence V2 améliorée."),
        ("23 juin", "Google Forms V2 séquence sociale complète. Recueil Markdown complet exporté."),
    ]
    for date, desc in etapes_23:
        puce(doc, f"{date} : {desc}")

    h2(doc, "Phase 24 — Pipeline CRM complet et formulaire d'audit (24–28 juin 2026)")
    etapes_24 = [
        ("24 juin", "Guides utilisateur et troubleshooting pipeline V1 (docs 103/104). Checklist mise en production CRM."),
        ("24–26 juin", "Fix resolve-invitation (colonne organization_type inexistante). Ajout de deAnonymize() dans prospect-audit-context.ts pour corriger l'affichage ORG_TARGET. Robustesse formulaire audit : timeout AbortController 12 s, gestion erreurs 404/410, fix zone de texte."),
        ("26–27 juin", "Debug V6 complet (8 nœuds). Validation bout en bout : 2 e-mails reçus à marius.ayoro70@gmail.com."),
        ("27–28 juin", "Guides Word session 28/06 (105/106/107). Guide Maître des 12 workflows. Reconstitution de la base de connaissance (ce document)."),
    ]
    for date, desc in etapes_24:
        puce(doc, f"{date} : {desc}")

    doc.add_page_break()

    # ── 3. ARCHITECTURE ────────────────────────────────────────────────────────
    h1(doc, "3. Architecture système actuelle")

    h2(doc, "Schéma du pipeline de prospection")
    schema = (
        "ENTRÉES PROSPECTS\n"
        "  ├── V4 Batch Orchestrator (quotidien)  ← Supabase / Airtable / Google Sheets\n"
        "  ├── Google Forms Social V2 (webhook)   ← campagnes réseaux sociaux\n"
        "  └── Manuel via Admin Controller\n\n"
        "           ↓\n"
        "[V3 CRM Enhanced FINAL]\n"
        "  Scraping 5 URL → Anonymisation → GPT-4 → Pack commercial\n"
        "  Produits : Executive Letter · Mini-Catalogue · URL Audit · Deck Brief\n"
        "  Stockage : table ai_prospecting_packs (Supabase)\n\n"
        "           ↓\n"
        "[Admin Controller V1]  ← validation humaine (approuver / rejeter / envoyer)\n\n"
        "           ↓ (approuvé)\n"
        "  E-mail prospect via Resend\n\n"
        "           ↓ (si pas de réponse)\n"
        "[V3 Follow-Up Sequence]  — J+1, J+3, J+7\n\n"
        "           ↓ (prospect remplit le formulaire)\n"
        "[Edge Function resolve-invitation]  +  [Formulaire transferai.ci/questionnaire-audit]\n\n"
        "           ↓ (soumission validée)\n"
        "[V6 Post-Audit Expert Routing]  — toutes les 30 minutes\n"
        "  → Fiche pré-RDV générée par GPT-4\n"
        "  → E-mail expert [PRIORITÉ HAUTE] si priorité HIGH ou MEDIUM\n"
        "  → E-mail brief pré-RDV à contact@transferai.ci\n"
        "  → Mise à jour CRM Supabase\n\n"
        "CANAUX ENTRANTS PARALLÈLES\n"
        "  ├── Chatwoot V5.5.2  → messages web → réponse automatique GPT-4\n"
        "  ├── Zoho Reply V1    → e-mails entrants → classification GPT-4\n"
        "  └── Calendly Webhook → RDV pris → mise à jour CRM + alerte équipe"
    )
    p = doc.add_paragraph()
    r = p.add_run(schema)
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)
    doc.add_paragraph()

    doc.add_page_break()

    # ── 4. SUPABASE ────────────────────────────────────────────────────────────
    h1(doc, "4. Base de données Supabase")

    h2(doc, "Tables principales")
    tableau(doc,
        ["Table", "Rôle"],
        [
            ["ai_prospecting_packs", "Packs commerciaux générés par V3 : pack_id, organization_name, target_email, payload (JSON), commercial_priority_tier"],
            ["prospect_targets", "CRM prospects : statut, tier, score, next_action_at, suivi relances"],
            ["form_invitations", "Invitations formulaire audit : invite_token, pack_id, contact_request_id, expires_at, draft_form_data, access_context"],
            ["form_responses", "Réponses au formulaire : invitation_token (⚠ PAS pack_id), form_data, is_completed, completion_percentage"],
            ["contact_requests", "Leads entrants via formulaire de contact ou d'audit"],
            ["outreach_attempts", "Journal de toutes les tentatives d'envoi"],
            ["follow_up_tracking", "Suivi des relances J+1, J+3, J+7"],
            ["newsletter_subscribers", "Abonnés à la newsletter par domaine"],
            ["newsletter_issues", "Numéros de newsletter générés"],
            ["newsletter_delivery_logs", "Journaux d'envoi de la newsletter"],
            ["social_prospects", "Leads issus des campagnes sur les réseaux sociaux"],
            ["whatsapp_inbound_messages", "Messages WhatsApp entrants (Twilio)"],
            ["whatsapp_email_notification_logs", "Journaux des notifications e-mail WhatsApp"],
        ],
        largeurs_cm=[5.5, 10.5]
    )

    h2(doc, "Points critiques à ne pas oublier")
    alerte(doc, "form_responses n'a PAS de colonne pack_id. Pour relier une réponse à un pack, il faut passer par form_invitations via invitation_token.")
    alerte(doc, "ai_prospecting_packs n'a PAS de colonne organization_type. Cette valeur se trouve dans payload.organization_type (JSON imbriqué).")
    ok(doc, "form_invitations.access_context contient le contexte déjà dé-anonymisé (ORG_TARGET remplacé par le vrai nom).")

    h2(doc, "Edge Functions Supabase actives")
    tableau(doc,
        ["Fonction", "Rôle"],
        [
            ["resolve-invitation", "Résout un token d'invitation et retourne le draft_form_data pré-rempli. Crée l'invitation si elle n'existe pas (via pack_id)."],
            ["save-form-response", "Sauvegarde les réponses du formulaire d'audit (autosave + soumission finale)."],
            ["newsletter-drafter", "Génère les brouillons de newsletter via GPT-4."],
            ["newsletter-send", "Envoie la newsletter aux abonnés."],
            ["newsletter-scheduler", "Planifie les envois de newsletter."],
            ["chatwoot-inbound", "Reçoit les webhooks Chatwoot et les transmet à n8n."],
            ["twilio-whatsapp-webhook", "Reçoit et stocke les messages WhatsApp entrants."],
        ],
        largeurs_cm=[5, 11]
    )

    doc.add_page_break()

    # ── 5. WORKFLOWS ──────────────────────────────────────────────────────────
    h1(doc, "5. Inventaire des 12 workflows n8n")

    h2(doc, "Tableau de synthèse")
    tableau(doc,
        ["#", "Workflow", "Nœuds", "Déclencheur", "Statut"],
        [
            ["1", "Admin Prospection Controller V1", "77", "Webhook + Manuel", "✔ Production"],
            ["2", "V3 CRM Enhanced [FINAL]", "51", "Manuel + Execute Workflow", "✔ Production"],
            ["3", "V3 CRM Enhanced [V20-FIXED]", "52", "Manuel + Execute Workflow", "✔ Sauvegarde"],
            ["4", "Post-Audit Expert Routing V6 MVP", "37", "Schedule 30 min + Webhook", "✔ Production"],
            ["5", "Post-Audit Expert Routing V2", "33", "Schedule 30 min + Webhook", "⚠ À DÉSACTIVER"],
            ["6", "Google Forms Social Lead Seq V2", "19", "Webhook + Schedule horaire", "✔ Actif"],
            ["7", "Multi-Prospect V4 Batch Orchestrator", "23", "Schedule quotidien + Manuel", "✔ Actif"],
            ["8", "Zoho Reply Intelligence V1", "15", "Schedule 15 min", "✔ Actif"],
            ["9", "Chatwoot Auto Reply V5.5", "12", "Webhook Chatwoot", "⚠ À DÉSACTIVER"],
            ["10", "Chatwoot Auto Reply V5.5.2", "12", "Webhook Chatwoot", "✔ Production"],
            ["11", "V3 Follow-Up Sequence V1", "12", "Schedule quotidien + Manuel", "✔ Actif"],
            ["12", "Calendly Meeting Booked Webhook V1", "6", "Webhook Calendly", "✔ Actif"],
        ],
        largeurs_cm=[0.8, 6.2, 1.5, 4.5, 3]
    )

    h2(doc, "Actions urgentes sur les workflows")
    alerte(doc,
        "Désactiver Post-Audit V2 (workflow n° 5) dans n8n. "
        "Il partage le même schedule (toutes les 30 min) et le même webhook que V6, "
        "ce qui entraîne un double traitement de chaque soumission de formulaire."
    )
    alerte(doc,
        "Désactiver Chatwoot V5.5 (workflow n° 9) dans n8n. "
        "Il partage le même webhook Chatwoot que V5.5.2, "
        "ce qui produit deux réponses automatiques pour chaque message entrant."
    )
    alerte(doc,
        "Vérifier et activer V4 Batch Orchestrator (workflow n° 7). "
        "53 prospects de catégorie tier1 présents dans Supabase n'ont encore jamais été contactés."
    )

    h2(doc, "Description des workflows clés")

    h3(doc, "V3 CRM Enhanced [FINAL] — Générateur de pack commercial")
    corps(doc,
        "V3 est le moteur de génération de contenu. Pour chaque prospect, il scrape jusqu'à "
        "5 pages web publiques, anonymise les données avant tout appel au LLM (remplacement "
        "du nom réel par le jeton ORG_TARGET), analyse le contexte avec GPT-4, puis génère "
        "quatre assets commerciaux complets."
    )
    tableau(doc,
        ["Nœud clé", "Action"],
        [
            ["Set Target", "Définit le prospect cible (URL, nom, e-mail, commercial_priority_default)"],
            ["Fetch Public Page 1 à 5", "Scrape jusqu'à 5 pages web du prospect"],
            ["Sanitize Prospect Data For LLM", "Anonymise les données (ORG_TARGET, DECISION_MAKER_TARGET)"],
            ["Call OpenAI Pre-Audit", "GPT-4 analyse le secteur, détecte les signaux IA, identifie les besoins"],
            ["Call OpenAI Problems Solutions", "GPT-4 génère les problèmes probables et les solutions recommandées"],
            ["Call OpenAI ROI", "GPT-4 calcule une hypothèse de retour sur investissement"],
            ["Assemble Prospect Pack", "Assemble le pack final. Propagation du commercial_priority_tier corrigée en juin 2026."],
            ["Upsert To Supabase", "Stocke ou met à jour le pack dans ai_prospecting_packs"],
            ["Approval Webhook", "Notifie l'Admin Controller que le pack est prêt pour validation"],
        ],
        largeurs_cm=[5, 11]
    )

    h3(doc, "V6 Post-Audit Expert Routing MVP — Traitement post-soumission")
    corps(doc,
        "V6 s'exécute toutes les 30 minutes et détecte les nouvelles soumissions du formulaire "
        "d'audit. Il génère une fiche pré-RDV complète avec GPT-4, évalue la priorité de "
        "routage et envoie deux e-mails de notification à l'équipe TransferAI."
    )
    tableau(doc,
        ["Priorité calculée", "Conditions", "E-mail expert envoyé ?"],
        [
            ["HIGH", "Complétion ≥ 95 % · Maturité avancée · Tier = tier1 ou A", "Oui"],
            ["MEDIUM", "Tier = tier2 ou B", "Oui"],
            ["NORMAL", "Tous les autres cas", "Non"],
        ],
        largeurs_cm=[3, 9, 4]
    )

    h3(doc, "Admin Controller V1 — Orchestration administrative")
    corps(doc,
        "Ce workflow est le cerveau administratif du pipeline. Il reçoit toutes les commandes "
        "de validation via un webhook sécurisé (en-tête Authorization requis) et les dispatche "
        "vers les sous-workflows appropriés."
    )
    tableau(doc,
        ["Action", "Effet"],
        [
            ["approve", "Approuver un pack → déclenche l'envoi de l'e-mail prospect"],
            ["reject", "Rejeter un pack → met à jour le statut dans Supabase"],
            ["send", "Forcer l'envoi d'un pack déjà approuvé"],
            ["re-run", "Relancer la génération V3 pour un prospect"],
            ["health_check", "Vérifier l'état général du pipeline"],
            ["batch_run", "Déclencher V4 Batch Orchestrator"],
        ],
        largeurs_cm=[4, 12]
    )

    doc.add_page_break()

    # ── 6. FORMULAIRE D'AUDIT ─────────────────────────────────────────────────
    h1(doc, "6. Formulaire d'audit — fonctionnement et corrections")

    h2(doc, "Accès et paramètres")
    tableau(doc,
        ["Paramètre", "Valeur"],
        [
            ["URL publique", "https://www.transferai.ci/questionnaire-audit"],
            ["Composant React", "src/pages/ProspectAuditFormPage.tsx"],
            ["Accès par token d'invitation", "?token=xxx  (lien personnalisé dans l'e-mail)"],
            ["Accès par identifiant de pack", "?pack_id=yyy  (crée automatiquement une invitation)"],
        ],
        largeurs_cm=[5, 11]
    )

    h2(doc, "Flux complet")
    num(doc, "Le prospect reçoit un e-mail contenant un lien avec son token d'invitation.")
    num(doc, "La page React charge et appelle l'Edge Function resolve-invitation (timeout 12 secondes).")
    num(doc, "L'Edge Function résout le token et retourne les données pré-remplies (draft_form_data), entièrement dé-anonymisées.")
    num(doc, "Le prospect complète le formulaire. La sauvegarde automatique s'effectue à chaque modification via save-form-response.")
    num(doc, "À la soumission, le champ is_completed est mis à true dans form_responses.")
    num(doc, "V6 détecte la soumission dans les 30 minutes suivantes et déclenche le traitement post-audit.")

    h2(doc, "Corrections appliquées lors des sessions de juin 2026")
    tableau(doc,
        ["Problème", "Cause racine", "Correction appliquée"],
        [
            ["ORG_TARGET visible dans le formulaire",
             "V3 anonymise les données pour le LLM. Le résultat stocké conservait ce jeton.",
             "Ajout de la fonction deAnonymize() dans prospect-audit-context.ts (ligne 236)"],
            ["Timeout réseau sans message d'erreur",
             "Aucune gestion du délai sur le fetch resolve-invitation.",
             "AbortController avec délai de 12 secondes (~lignes 241–255 de ProspectAuditFormPage.tsx)"],
            ["Erreur 404 ou 410 non explicite",
             "Le message d'erreur générique ne distinguait pas l'invitation expirée.",
             "Messages spécifiques : « cette invitation a expiré ou n'est plus disponible »"],
            ["Sauvegarde silencieuse sans retour",
             "Aucun message affiché lorsque le token d'invitation était absent.",
             "Ajout d'un message d'erreur à la ligne ~369 de ProspectAuditFormPage.tsx"],
            ["Zone de texte affichant undefined",
             "La valeur du champ n'était pas vérifiée comme chaîne de caractères.",
             "Correction : typeof formData.prospect_context === 'string' (~ligne 1045)"],
        ],
        largeurs_cm=[4, 5, 7]
    )

    h2(doc, "Logique de dé-anonymisation (prospect-audit-context.ts)")
    code_inline(doc, "const realOrgName = trimText(pack?.organization_name) || trimText(request?.company) || ...;")
    code_inline(doc, "const deAnonymize = (text) => text.replace(/ORG_TARGET/g, realOrgName).replace(/DECISION_MAKER_TARGET/g, ...);")
    code_inline(doc, "const prospectContext = deAnonymize(rawProspectContext);")

    doc.add_page_break()

    # ── 7. DONNÉES COMMERCIALES ───────────────────────────────────────────────
    h1(doc, "7. Données commerciales et règles métier")

    h2(doc, "Segmentation des prospects — Tiers commerciaux")
    tableau(doc,
        ["Tier", "Critère d'attribution", "Traitement dans le pipeline"],
        [
            ["tier1", "Lead très qualifié, intention forte, e-mail valide, secteur prioritaire", "Priorité absolue dans V4 · Routing HIGH dans V6 · E-mail expert envoyé"],
            ["tier2", "Lead exploitable, moins qualifié", "Traitement normal dans V4 · Routing MEDIUM dans V6 · E-mail expert envoyé"],
            ["tier3", "Lead incomplet ou sans e-mail valide", "Reste en statut draft · Non envoyé par V4"],
        ],
        largeurs_cm=[2, 6, 8]
    )

    h2(doc, "Cycle de vie d'un prospect (champ status)")
    tableau(doc,
        ["Statut", "Signification"],
        [
            ["draft", "Lead reçu, données incomplètes"],
            ["ready", "Prêt à être envoyé (e-mail valide confirmé)"],
            ["sent", "Pack commercial envoyé au prospect"],
            ["completed", "Formulaire d'audit complété par le prospect"],
            ["rejected", "Rejeté par l'administrateur lors de la validation"],
            ["do_not_contact", "Inscrit sur la liste noire — aucun envoi autorisé"],
        ],
        largeurs_cm=[4, 12]
    )

    h2(doc, "Séquence de relance automatique (V3 Follow-Up)")
    tableau(doc,
        ["Délai", "Contenu de l'e-mail"],
        [
            ["J+1", "Premier suivi — confirmation de réception du pack + invitation à remplir le formulaire d'audit"],
            ["J+3", "Deuxième suivi — partage d'un cas d'usage concret dans le secteur du prospect"],
            ["J+7", "Troisième et dernier suivi — offre de prise de rendez-vous direct"],
        ],
        largeurs_cm=[3, 13]
    )

    doc.add_page_break()

    # ── 8. INFRASTRUCTURE ─────────────────────────────────────────────────────
    h1(doc, "8. Infrastructure et déploiement")

    h2(doc, "Branches Git actives")
    tableau(doc,
        ["Branche", "Rôle"],
        [
            ["main", "Branche de production — déployée automatiquement sur Cloudflare Pages"],
            ["codex/newsletter-bilingual-editorial", "Branche de travail principale actuelle"],
            ["codex/prepare-main-newsletter-merge", "Préparation du merge newsletter"],
            ["publish-main", "Déclencheur de publication Cloudflare"],
            ["codex/webinar-july-2026", "Préparation du webinaire de juillet 2026"],
        ],
        largeurs_cm=[6, 10]
    )

    h2(doc, "Processus de déploiement")
    num(doc, "Développement et corrections sur la branche codex/newsletter-bilingual-editorial.")
    num(doc, "Commit des fichiers modifiés et push vers GitHub.")
    num(doc, "Merge vers main : git merge codex/newsletter-bilingual-editorial.")
    num(doc, "Push de main → Cloudflare Pages déclenche automatiquement un build.")
    num(doc, "Le site est mis à jour sur www.transferai.ci en 1 à 2 minutes.")
    num(doc, "Les Edge Functions Supabase sont déployées séparément via : supabase functions deploy <nom>")

    h2(doc, "Fichiers de code critiques")
    tableau(doc,
        ["Fichier", "Rôle"],
        [
            ["src/pages/ProspectAuditFormPage.tsx", "Formulaire d'audit — composant React principal"],
            ["supabase/functions/resolve-invitation/index.ts", "Résolution des tokens d'invitation"],
            ["supabase/functions/save-form-response/index.ts", "Sauvegarde des réponses formulaire"],
            ["supabase/functions/_shared/prospect-audit-context.ts", "Logique de contexte prospect, dé-anonymisation, construction du draft"],
        ],
        largeurs_cm=[8, 8]
    )

    doc.add_page_break()

    # ── 9. BUGS ET POINTS D'ATTENTION ─────────────────────────────────────────
    h1(doc, "9. Problèmes résolus et points d'attention permanents")

    h2(doc, "Bugs résolus lors des sessions de mai–juin 2026")
    tableau(doc,
        ["Problème", "Cause", "Correction"],
        [
            ["ORG_TARGET affiché dans le formulaire",
             "Anonymisation LLM non inversée à l'affichage",
             "Fonction deAnonymize() dans prospect-audit-context.ts"],
            ["Condition 'If Pack Found' toujours fausse",
             "Ancien format booléen n8n v1 cassé dans la version actuelle",
             "Migration vers le format v2 avec options:{version:2}"],
            ["'Get Latest Form Response' retourne vide",
             "Filtre sur pack_id (colonne inexistante dans form_responses)",
             "Filtre sur id=eq.{response_id}"],
            ["commercial_priority_tier absent du pack",
             "Nœud Assemble Prospect Pack ne propageait pas la valeur",
             "Ajout du fallback ctx.commercial_priority_default"],
            ["Double traitement des soumissions post-audit",
             "V2 et V6 actifs simultanément avec le même schedule",
             "À corriger : désactiver V2 dans n8n"],
            ["Lien Calendly doublé dans l'e-mail",
             "Format Markdown [url](url) dans un champ HTML de n8n",
             "À corriger : remplacer par <a href='...'>Prendre rendez-vous</a>"],
            ["Erreur 500 sur resolve-invitation",
             "SELECT incluait organization_type (colonne inexistante dans ai_prospecting_packs)",
             "Suppression du champ du SELECT"],
        ],
        largeurs_cm=[4.5, 5.5, 6]
    )

    h2(doc, "Points d'attention permanents")
    tableau(doc,
        ["Point", "Détail"],
        [
            ["Délai V6", "V6 tourne toutes les 30 minutes. Une soumission peut attendre jusqu'à 30 minutes avant traitement."],
            ["Format conditions n8n", "L'ancien format booléen (value1 equals value2) est cassé dans la version actuelle de n8n. Toujours utiliser le format v2."],
            ["Colonne pack_id dans form_responses", "Cette colonne n'existe pas. Toujours filtrer via invitation_token ou id."],
            ["Quota V4", "Le filtre quota utilisait delivery_status = submitted. V3 n'écrit pas ce champ. Corrigé en juin 2026."],
            ["53 prospects tier1 en attente", "Présents dans Supabase, jamais contactés. Nécessite un déclenchement manuel de V4."],
        ],
        largeurs_cm=[4, 12]
    )

    doc.add_page_break()

    # ── 10. OFFRE COMMERCIALE ─────────────────────────────────────────────────
    h1(doc, "10. Offre commerciale TransferAI Africa")

    h2(doc, "Services proposés")
    tableau(doc,
        ["Service", "Description", "Point d'entrée"],
        [
            ["Audit IA gratuit", "Évaluation de la maturité IA de l'organisation. RDV de 45 min post-soumission.", "Formulaire en ligne"],
            ["Formation et certification", "Catalogue sectoriel, parcours métier, certification par domaine.", "Catalogue en ligne / PDF"],
            ["Accompagnement entreprise", "Adoption de l'IA, automatisation des processus, conseil stratégique.", "Prise de contact"],
            ["Solutions sur mesure", "LLM privé, workflows personnalisés, CRM IA, intégrations métier.", "Devis sur mesure"],
        ],
        largeurs_cm=[4, 8, 4]
    )

    h2(doc, "Secteurs cibles prioritaires")
    secteurs = [
        "Assistanat et secrétariat", "Ressources humaines", "Marketing et communication",
        "Finance et comptabilité", "Juridique et conformité", "Service client",
        "Data et analyse", "IT et transformation digitale", "Management et leadership",
        "Formation et pédagogie", "Santé et bien-être", "Diplomatie et affaires internationales",
    ]
    for i in range(0, len(secteurs), 2):
        ligne = secteurs[i]
        if i + 1 < len(secteurs):
            ligne += "  ·  " + secteurs[i + 1]
        puce(doc, ligne)

    h2(doc, "Géographie cible")
    puce(doc, "Priorité : Côte d'Ivoire")
    puce(doc, "Extension : Afrique de l'Ouest")
    puce(doc, "Horizon : Afrique francophone")

    doc.add_page_break()

    # ── 11. INDEX DOCUMENTS ───────────────────────────────────────────────────
    h1(doc, "11. Index des documents de référence")
    corps(doc,
        "Tous les documents sont dans le dossier :\n"
        "docs/transferai-admin/  (dans le référentiel GitHub connect-ia-bloom)"
    )

    h2(doc, "Documents fondateurs")
    tableau(doc,
        ["Fichier", "Contenu"],
        [
            ["00_INDEX.md", "Index complet de tous les documents du projet"],
            ["01_Guide_Administrateur_No1.md", "Guide de l'administrateur principal"],
            ["02_Base_Connaissance_Site.md", "Base de connaissance du site (version antérieure)"],
            ["03_FAQ_Site_Assistant_IA.md", "FAQ opérationnelle"],
            ["04_Plan_Deploiement_Exploitation.md", "Plan d'exploitation"],
            ["05_Roadmap_Historique_Projet.md", "Historique des phases 1 à 20"],
            ["KB_TransferAI_Base_Connaissance_Complete_2026-06-28.md", "CE DOCUMENT — base de connaissance maître (Markdown)"],
        ],
        largeurs_cm=[7, 9]
    )

    h2(doc, "Documents pipeline de prospection")
    tableau(doc,
        ["Fichier(s)", "Contenu"],
        [
            ["57", "Présentation du pipeline V1 → V2 → V3 → V4"],
            ["68 / 69", "Guide utilisateur et troubleshooting V3 CRM"],
            ["76 / 77", "Guide V4 Batch nœud par nœud et troubleshooting"],
            ["78", "Schéma architecture fonctionnelle V1 → V6"],
            ["79 à 82", "Plan et prompts du dashboard BackOffice web"],
            ["90", "Audit expert CRM — 13 juin 2026"],
            ["103 / 104", "Guide utilisateur et troubleshooting pipeline V1 (MD + DOCX)"],
        ],
        largeurs_cm=[4, 12]
    )

    h2(doc, "Guides sessions récentes (juin 2026)")
    tableau(doc,
        ["Fichier(s)", "Contenu"],
        [
            ["100 / 101 / 102", "Google Forms Social Lead Sequence V2 — guide complet"],
            ["104_Checklist_Mise_En_Production_CRM_2026-06-24.md", "Checklist de mise en production CRM"],
            ["105_Guide_Utilisateur_Session_2026-06-28_Audit_V6.docx", "Guide utilisateur session 28 juin — Audit V6"],
            ["106_Guide_Troubleshooting_Session_2026-06-28_Audit_V6.docx", "Guide troubleshooting session 28 juin — Audit V6"],
            ["107_Guide_Maitre_Workflows_TransferAI_V1.docx", "Guide Maître des 12 workflows TransferAI (RÉFÉRENCE PRINCIPALE)"],
        ],
        largeurs_cm=[6, 10]
    )

    doc.add_page_break()

    # ── 12. ACTIONS EN ATTENTE ────────────────────────────────────────────────
    h1(doc, "12. Actions en attente au 28 juin 2026")

    h2(doc, "Urgences — bugs actifs en production")
    alerte(doc, "Désactiver le workflow Post-Audit V2 dans n8n (conflit avec V6 — double traitement des soumissions).")
    alerte(doc, "Désactiver le workflow Chatwoot V5.5 dans n8n (conflit avec V5.5.2 — double réponse aux messages entrants).")
    alerte(doc, "Corriger le lien Calendly doublé dans le nœud 'Send Internal Brief Email' de V6 : remplacer le format Markdown par une balise HTML <a href>.")

    h2(doc, "Actions importantes à planifier")
    puce(doc, "Lancer V4 Batch Orchestrator manuellement pour les 53 prospects tier1 en attente dans Supabase.")
    puce(doc, "Effectuer un test complet de V3 avec un nouveau prospect réel pour valider l'Executive Letter, le Deck Brief et le Mini-Catalogue.")
    puce(doc, "Corriger les anciens packs qui contiennent encore ORG_TARGET dans le champ access_context (invitations créées avant le fix de juin 2026).")
    puce(doc, "Confirmer si Chatwoot V5.5 et V5.5.2 pointent vers le même compte ou des comptes différents.")
    puce(doc, "Confirmer si V6 remplace entièrement V2 (désactivation définitive) ou s'ils coexistent.")

    h2(doc, "Backlog moyen terme")
    puce(doc, "Dashboard BackOffice web prospection (plan documenté dans les fichiers 79 à 82).")
    puce(doc, "Webinaire juillet 2026 (branche codex/webinar-july-2026).")
    puce(doc, "Newsletter bilingue français / anglais (branche codex/newsletter-bilingual-editorial).")

    separateur(doc)
    corps(doc, f"Document généré automatiquement le {datetime.today().strftime('%d/%m/%Y à %H:%M')}.")
    corps(doc, "TransferAI Africa — contact@transferai.ci — www.transferai.ci")

    # ── SAUVEGARDE ────────────────────────────────────────────────────────────
    out = os.path.join(OUTPUT_DIR, "108_Base_Connaissance_Complete_TransferAI_2026-06-28.docx")
    doc.save(out)
    print(f"✓ Document sauvegardé : {out}")
    return out


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    build()
