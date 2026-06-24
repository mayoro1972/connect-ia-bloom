from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
TODAY = date.today().isoformat()

ATTACHED_V3 = Path("/Users/marius_ayoro/Downloads/TransferAI Prospecting V3 CRM Enhanced [FINAL]-10.json")
FALLBACK_V3 = DOCS_DIR / "62_n8n_Prospection_Multi_Prospect_V3_CRM_Ready_Export.json"

OUTPUT = WORD_DIR / f"TransferAI_Africa_Guide_Maitre_Prospecting_V3_Fusionne_{TODAY}.docx"


def load_v3() -> tuple[dict, Path]:
    path = ATTACHED_V3 if ATTACHED_V3.exists() else FALLBACK_V3
    return json.loads(path.read_text(encoding="utf-8")), path


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.8)

    title = document.styles["Title"]
    title.font.name = "Georgia"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Georgia"
        style.font.bold = True
    document.styles["Heading 1"].font.size = Pt(17)
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x10, 0x2A, 0x43)
    document.styles["Heading 2"].font.size = Pt(12.5)
    document.styles["Heading 2"].font.color.rgb = RGBColor(0x17, 0x4B, 0x6B)
    document.styles["Heading 3"].font.size = Pt(11)
    document.styles["Heading 3"].font.color.rgb = RGBColor(0x2A, 0x5D, 0x7B)


def add_title_block(document: Document, workflow_name: str, source_name: str, node_count: int) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("Guide d'installation V3\nNœud par nœud")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Workflow n8n Prospection Intelligente CRM Enhanced")
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    document.add_paragraph("")
    document.add_paragraph("Guide d’installation V3", style="Heading 1")
    document.add_paragraph("Nœud par nœud", style="Heading 1")

    meta_lines = [
        f"Workflow : {workflow_name}",
        f"Fichier de référence : {source_name}",
        f"Version consolidée : {TODAY}",
        f"État réel du workflow : {node_count} nœuds",
        "Référentiel utilisé : guide nœud par nœud V3, guide utilisateur validation / envoi prospect, guide d’installation CRM ready, et export n8n FINAL.",
    ]
    for line in meta_lines:
        document.add_paragraph(line)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(item)


def add_plain_dash_lines(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(f"- {item}")


def role_map() -> dict[str, str]:
    return {
        "Manual Trigger": "déclenchement manuel depuis l’interface n8n.",
        "Execute Workflow Trigger": "point d’entrée lorsqu’un autre workflow, par exemple V4, appelle la V3.",
        "Set Target": "initialise les champs prospect et conserve les métadonnées CRM.",
        "Build Source URLs": "détermine les pages publiques à explorer pour le prospect.",
        "Fetch Public Page 1": "récupère la première page publique cible.",
        "Fetch Public Page 2": "récupère la deuxième page publique cible.",
        "Fetch Public Page 3": "récupère la troisième page publique cible.",
        "Fetch Public Page 4": "récupère la quatrième page publique cible.",
        "Fetch Public Page 5": "récupère la cinquième page publique cible.",
        "Normalize Public Signals": "transforme le contenu brut en signaux publics exploitables.",
        "Sanitize Prospect Data For LLM": "protège les données sensibles avant exposition au LLM.",
        "Call OpenAI Pre-Audit": "produit le pré-audit et les hypothèses initiales.",
        "Call OpenAI Problems Solutions": "déduit les problèmes probables, solutions et angle commercial.",
        "Call OpenAI ROI": "formule une hypothèse de valeur et de ROI.",
        "Assemble Prospect Context": "fusionne les analyses pour constituer le contexte maître du prospect.",
        "Generate Executive Letter": "génère le courrier de premier contact.",
        "Generate Tailored Catalogue": "génère le mini-catalogue commercial ciblé.",
        "Generate Tailored Audit Form": "génère la forme d’audit pré-appel.",
        "Generate Deck Brief": "génère la structure du deck commercial.",
        "Assemble Prospect Pack": "assemble le pack prospect final avant rendu et stockage.",
        "Store Pack In Supabase": "enregistre le pack dans ai_prospecting_packs.",
        "Build Approval Email": "prépare l’e-mail interne avec liens Approuver / Rejeter.",
        "Send Internal Approval Email": "envoie la demande d’approbation interne.",
        "Approval Webhook": "reçoit la décision approuver ou rejeter depuis l’e-mail.",
        "Parse Approval Query": "lit pack_id et decision dans la query string.",
        "Get Pack From Supabase": "récupère le pack stocké en base après le clic d’approbation.",
        "Extract Pack Payload": "reconstruit les données nécessaires à l’envoi.",
        "If Approved": "sépare les branches approbation et rejet.",
        "Build Send Context": "vérifie qu’un envoi est réellement possible.",
        "If Ready To Send": "autorise ou bloque l’envoi selon can_send.",
        "Mark Pack Approved": "marque le pack comme approuvé.",
        "Send External Prospect Email": "envoie le courrier prospect via Resend.",
        "Parse Send Result": "capture l’identifiant de message et la date d’envoi.",
        "Mark Pack Sent": "met à jour le statut sent du pack.",
        "Log Outreach Attempt": "journalise l’envoi dans outreach_attempts.",
        "Send Internal Sent Confirmation": "notifie l’équipe qu’un envoi effectif a eu lieu.",
        "Mark Pack Approval Error": "trace un blocage fonctionnel après approbation.",
        "Mark Pack Rejected": "met le pack au statut rejected.",
        "Update Prospect Target Sent": "réécrit le statut CRM du prospect après envoi.",
        "Update Prospect Target Approval Error": "réécrit le statut CRM du prospect après blocage.",
        "Update Prospect Target Rejected": "réécrit le statut CRM du prospect après rejet.",
        "Respond to Webhook": "retourne la réponse navigateur pour la branche approuvée.",
        "Respond Rejected": "retourne la réponse navigateur pour la branche rejetée.",
        "Respond Approval Error": "retourne la réponse navigateur pour la branche erreur.",
        "Resolve Domain Catalogue": "maintient la compatibilité avec l’ancienne logique catalogue statique.",
        "Download Catalogue PDF": "télécharge l’ancien PDF catalogue public si nécessaire.",
        "Assemble Mail Attachments": "agrège les pièces jointes utilisables pour l’e-mail.",
        "Build Catalogue Render Payload": "prépare le payload du renderer catalogue.",
        "render Catalogue Artifact": "génère le catalogue PDF / DOCX dynamique côté backend.",
        "Merge Catalogue Artifact": "réinjecte les URLs réelles du catalogue dans le pack.",
        "Build Deck Render Payload": "prépare le payload du renderer deck PPTX.",
        "Render Deck Artifact": "génère le deck PPTX dynamique côté backend.",
        "Merge Deck Artifact": "fusionne le deck et les pièces jointes finales dans le pack.",
    }


def group_sections() -> list[tuple[str, list[str]]]:
    return [
        ("Phase 1 — Déclencheurs", ["Manual Trigger", "Execute Workflow Trigger"]),
        ("Phase 2 — Initialisation", ["Set Target", "Build Source URLs"]),
        ("Phase 3 — Scraping web", [
            "Fetch Public Page 1", "Fetch Public Page 2", "Fetch Public Page 3", "Fetch Public Page 4", "Fetch Public Page 5"
        ]),
        ("Phase 4 — Normalisation et protection LLM", ["Normalize Public Signals", "Sanitize Prospect Data For LLM"]),
        ("Phase 5 — Analyse IA", ["Call OpenAI Pre-Audit", "Call OpenAI Problems Solutions", "Call OpenAI ROI", "Assemble Prospect Context"]),
        ("Phase 6 — Génération de contenu", [
            "Generate Executive Letter", "Generate Tailored Catalogue", "Generate Tailored Audit Form", "Generate Deck Brief", "Assemble Prospect Pack"
        ]),
        ("Phase 7 — Rendu des artefacts", [
            "Resolve Domain Catalogue", "Download Catalogue PDF", "Build Catalogue Render Payload", "render Catalogue Artifact",
            "Merge Catalogue Artifact", "Build Deck Render Payload", "Render Deck Artifact", "Merge Deck Artifact", "Assemble Mail Attachments"
        ]),
        ("Phase 8 — Stockage et validation interne", ["Store Pack In Supabase", "Build Approval Email", "Send Internal Approval Email"]),
        ("Phase 9 — Approbation webhook", ["Approval Webhook", "Parse Approval Query", "Get Pack From Supabase", "Extract Pack Payload", "If Approved"]),
        ("Phase 10 — Préparation d’envoi prospect", ["Build Send Context", "If Ready To Send", "Mark Pack Approved", "Send External Prospect Email", "Parse Send Result", "Mark Pack Sent"]),
        ("Phase 11 — Post-envoi, logs et statuts", ["Log Outreach Attempt", "Send Internal Sent Confirmation", "Update Prospect Target Sent", "Respond to Webhook"]),
        ("Phase 12 — Rejet et erreurs", ["Mark Pack Approval Error", "Mark Pack Rejected", "Update Prospect Target Approval Error", "Update Prospect Target Rejected", "Respond Approval Error", "Respond Rejected"]),
    ]


def main() -> None:
    data, source_path = load_v3()
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_block(doc, data.get("name", "Workflow V3"), source_path.name, len(data.get("nodes", [])))

    add_heading(doc, "1. Ce qui change dans ce guide fusionné", 1)
    add_bullets(doc, [
        "Le guide 63 reste la base formelle : même logique de lecture, mêmes grandes sections, même ton opératoire.",
        "Les apports du guide utilisateur prospect ont été intégrés pour mieux expliquer l’usage quotidien, les secrets, les flux réels et les écarts entre théorie et export FINAL.",
        "Les apports du guide d’installation CRM ready ont été intégrés pour aligner les variables, la logique CRM, les tests et les résultats attendus.",
        f"La version de référence prise en compte ici est {source_path.name}, qui contient {len(data.get('nodes', []))} nœuds.",
    ])

    add_heading(doc, "2. Alignement des versions du projet prospecting", 1)
    add_paragraph(doc, "Lecture simple : V5 fait entrer les sociétés dans le CRM, V4 choisit lesquelles traiter, V3 fabrique le pack prospect et l’envoie après approbation.")
    add_plain_dash_lines(doc, [
        "V1 : génération manuelle d’un pack prospect à partir du scraping public et des appels LLM.",
        "V2 : première boucle complète avec stockage Supabase, approbation interne et envoi prospect.",
        "V3 : version CRM Enhanced, devenue la version opérable de référence.",
        "V4 : batch et quotas d’envoi, règles d’arrêt et gouvernance commerciale.",
        "V5 : growth loop CRM pour l’ingestion continue des leads publics.",
    ])

    add_heading(doc, "3. Ce qui a changé depuis l’ancienne version du guide", 1)
    add_bullets(doc, [
        f"Le workflow n’est plus à 44 ou 50 nœuds, mais à {len(data.get('nodes', []))} nœuds dans la version de référence actuelle.",
        "Le pipeline V3 inclut une chaîne de rendu catalogue complète : Build Catalogue Render Payload -> render Catalogue Artifact -> Merge Catalogue Artifact.",
        "Le pipeline V3 inclut une chaîne de rendu deck PPTX complète : Build Deck Render Payload -> Render Deck Artifact -> Merge Deck Artifact.",
        "Build Send Context doit prendre en compte des pièces jointes réelles et pas seulement un courrier texte.",
        "Le flux cible impose 2 pièces jointes finales : 1 catalogue PDF et 1 deck PPTX.",
        "Le fallback Deck_Brief_[Prospect].json ne fait plus partie du flux d’envoi cible.",
        "Les nœuds OpenAI et les accès Supabase / Resend doivent être externalisés en variables d’environnement avant production.",
    ])

    add_heading(doc, "4. Prérequis d’installation et d’exploitation", 1)
    add_bullets(doc, [
        "Une instance n8n accessible et opérationnelle.",
        "Le workflow JSON V3 final.",
        "Un projet Supabase opérationnel.",
        "Le bucket ou mécanisme prospecting-artifacts pour les rendus catalogue / deck.",
        "Les tables Supabase minimales : ai_prospecting_packs, outreach_attempts, prospect_targets.",
        "Les fonctions ou services de rendu catalogue-renderer et deck-renderer si elles sont déployées côté backend.",
        "Un compte Resend opérationnel.",
    ])

    add_heading(doc, "5. Variables et secrets recommandés", 1)
    add_heading(doc, "Variables n8n à définir", 2)
    add_bullets(doc, [
        "OPENAI_API_KEY",
        "OPENAI_MODEL",
        "CONTENT_ADMIN_TOKEN",
        "SUPABASE_URL",
        "SUPABASE_SERVICE_ROLE_KEY",
        "RESEND_API_KEY",
        "OUTREACH_FROM_EMAIL",
        "INTERNAL_REVIEW_EMAIL",
        "N8N_BASE_URL",
        "BOOKING_LINK_45MIN",
    ])
    add_heading(doc, "Secrets à externaliser avant production", 2)
    add_paragraph(
        doc,
        "Le workflow de référence peut encore contenir des valeurs sensibles écrites en dur dans certains nœuds REST Supabase, "
        "renderers ou Resend. Avant production, il faut les remplacer par des variables d’environnement ou des credentials n8n."
    )
    add_bullets(doc, [
        "apikey Supabase dans les nœuds REST",
        "Authorization: Bearer ... Supabase dans les nœuds REST",
        "Authorization: Bearer ... Resend dans les nœuds email",
        "x-admin-token dans render Catalogue Artifact et Render Deck Artifact",
    ])

    add_heading(doc, "6. Architecture actuelle", 1)
    add_heading(doc, "Phases actuelles", 2)
    add_plain_dash_lines(doc, [
        "1. Déclencheurs",
        "2. Initialisation",
        "3. Scraping web",
        "4. Normalisation et protection LLM",
        "5. Analyse IA",
        "6. Génération de contenu",
        "7. Assemblage et rendu d’artefacts",
        "8. Stockage et validation interne",
        "9. Approbation webhook",
        "10. Préparation d’envoi prospect",
        "11. Post-envoi, logs et statuts",
        "12. Rejet et erreurs",
    ])
    add_heading(doc, "Chaîne principale actuelle", 2)
    add_paragraph(
        doc,
        "`Assemble Prospect Pack -> Build Catalogue Render Payload -> render Catalogue Artifact -> Merge Catalogue Artifact -> "
        "Build Deck Render Payload -> Render Deck Artifact -> Merge Deck Artifact -> Store Pack In Supabase -> "
        "Build Approval Email -> Send Internal Approval Email`"
    )
    add_heading(doc, "Chaîne d’approbation et d’envoi", 2)
    add_paragraph(
        doc,
        "`Approval Webhook -> Parse Approval Query -> Get Pack From Supabase -> Extract Pack Payload -> If Approved -> "
        "Build Send Context -> If Ready To Send -> Mark Pack Approved -> Send External Prospect Email -> Parse Send Result -> Mark Pack Sent`"
    )

    add_heading(doc, "7. Règles critiques à connaître", 1)
    add_bullets(doc, [
        "Render Catalogue Artifact et Render Deck Artifact utilisent idéalement CONTENT_ADMIN_TOKEN via le header x-admin-token.",
        "Store Pack In Supabase doit recevoir la sortie la plus riche du pipeline, après fusion des artefacts.",
        "Build Send Context doit retourner can_send = true seulement si le courrier et l’e-mail cible existent.",
        "Le flux cible final doit transporter 2 pièces jointes : un catalogue PDF et un deck PPTX.",
        "La V3 réécrit l’état métier du prospect dans prospect_targets après envoi, rejet ou approval_error.",
        "Le LLM ne doit voir que des signaux publics assainis, jamais les données sensibles brutes.",
    ])

    add_heading(doc, "8. Installation étape par étape", 1)
    add_numbered(doc, [
        "Importer le fichier JSON de référence dans n8n.",
        "Vérifier que les variables d’environnement OpenAI, Supabase, Resend, CONTENT_ADMIN_TOKEN et N8N_BASE_URL existent.",
        "Contrôler Set Target : includeOtherFields doit rester activé pour accepter les champs CRM entrants.",
        "Configurer Supabase pour Store Pack In Supabase, Get Pack From Supabase, Mark Pack Approved, Mark Pack Sent, Mark Pack Approval Error, Mark Pack Rejected et les nœuds Update Prospect Target.",
        "Configurer Resend pour Send Internal Approval Email, Send External Prospect Email et Send Internal Sent Confirmation.",
        "Configurer les renderers catalogue et deck et vérifier le mécanisme de stockage des fichiers générés.",
        "Activer le workflow pour rendre le webhook d’approbation accessible publiquement.",
        "Lancer un test avec un prospect de démonstration, puis vérifier l’approbation, l’envoi, les logs et la mise à jour du CRM.",
    ])

    add_heading(doc, "9. Guide utilisateur quotidien", 1)
    add_bullets(doc, [
        "Vérifier d’abord le prospect entrant : organization_name, website et target_email.",
        "Contrôler que le pack a bien été généré et stocké avec un pack_id.",
        "Recevoir puis ouvrir l’e-mail d’approbation interne.",
        "Cliquer sur Approuver ou Rejeter selon la qualité du pack.",
        "En cas d’approbation, confirmer que le courrier, le catalogue PDF et le deck PPTX sont prêts.",
        "Contrôler l’état final dans ai_prospecting_packs, outreach_attempts et prospect_targets.",
    ])

    add_heading(doc, "10. Guide nœud par nœud", 1)
    roles = role_map()
    for section_title, node_names in group_sections():
        add_heading(doc, section_title, 2)
        for name in node_names:
            node = next((n for n in data.get("nodes", []) if n.get("name") == name), None)
            if not node:
                continue
            add_paragraph(doc, f"{name}")
            add_paragraph(doc, f"Type : `{node['type'].split('.')[-1]}`")
            add_paragraph(doc, f"Rôle : {roles.get(name, 'nœud à vérifier dans le canvas n8n.')}")
            add_paragraph(doc, "Action : vérifier la configuration et la connexion entrante / sortante dans le canvas.")

    add_heading(doc, "11. Logique CRM à retenir", 1)
    add_bullets(doc, [
        "V5 fait entrer les leads publics dans prospect_targets.",
        "V4 lit prospect_targets et décide qui traiter aujourd’hui.",
        "Cette V3 traite un prospect puis réécrit son état CRM.",
        "ai_prospecting_packs conserve le snapshot du pack au moment de la validation / de l’envoi.",
        "outreach_attempts conserve l’historique des courriers réellement envoyés.",
    ])

    add_heading(doc, "12. Troubleshooting rapide", 1)
    add_bullets(doc, [
        "Si Get Pack From Supabase renvoie [] : vérifier select=*, apikey, Authorization et pack_id exact.",
        "Si If Ready To Send bloque : ouvrir Build Send Context et vérifier can_send, target_email et executive_letter.",
        "Si le webhook ne répond pas : activer le workflow et tester N8N_BASE_URL.",
        "Si aucune pièce jointe n’apparaît : tester séparément le renderer catalogue puis le renderer deck.",
        "Si le CRM ne se met pas à jour : vérifier les PATCH vers prospect_targets et la clé service_role.",
    ])

    add_heading(doc, "13. Résultat attendu", 1)
    add_paragraph(
        doc,
        "À la fin, tu disposes d’un guide maître unique pour le projet prospecting. Il conserve la structure lisible du guide 63, "
        "mais aligne les contenus utilisateur, installation, CRM, validation, envoi et rendu d’artefacts sur la version V3 CRM Enhanced actuelle."
    )

    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TransferAI - Guide maître Prospecting V3 fusionné")
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    WORD_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
