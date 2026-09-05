from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-prospection/catalogues-sectoriels-premium-2026-06-12/orange-cote-divoire-clean"
)
DOCX_PATH = OUTPUT_DIR / "Mini_Catalogue_TransferAI_Orange_Cote_d_Ivoire_Clean_2026-06-12.docx"

NAVY = RGBColor(0x1C, 0x39, 0x63)
ORANGE = RGBColor(0xEF, 0x6C, 0x12)
GREEN = RGBColor(0x1C, 0x8A, 0x78)
GRAY = RGBColor(0x6B, 0x7C, 0x93)
LIGHT_GRAY = "DDE5F1"
LIGHT_FILL = "F8F3EC"


def set_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)


def set_font(run, name: str, size: float, color: RGBColor, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "F1A56E", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "TransferAI Africa | NettelecomCI | contact@transferai.ci | www.transferai.ci | WhatsApp +225 07 16 57 39 90"
    )
    set_font(run, "Aptos", 9.2, GRAY)


def add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LIGHT_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_heading(document: Document, text: str, top_space: float = 10) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(top_space)
    paragraph.paragraph_format.space_after = Pt(4)
    marker = paragraph.add_run("▌ ")
    set_font(marker, "Aptos", 18, ORANGE, bold=True)
    run = paragraph.add_run(text)
    set_font(run, "Aptos Display", 21, NAVY, bold=True)


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, "Aptos Display", 16.5, NAVY, bold=True)


def add_body(document: Document, text: str, color: RGBColor = NAVY, size: float = 12.5, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_font(run, "Aptos", size, color, italic=italic)


def add_bullets(document: Document, items: list[str], color: RGBColor = NAVY, size: float = 12.5) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_font(run, "Aptos", size, color)


def add_numbered(document: Document, items: list[str], color: RGBColor = NAVY, size: float = 12.3) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(item)
        set_font(run, "Aptos", size, color)


def add_callout(document: Document, title: str, body: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    remove_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, LIGHT_FILL)
    set_cell_border(cell)
    cell.width = Cm(16.5)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(3)
    run1 = p1.add_run(title)
    set_font(run1, "Aptos Display", 16, NAVY, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    set_font(run2, "Aptos", 12.5, NAVY)


def add_meta_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, "Aptos", 12.5, GRAY)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(36)
    run = p.add_run("Mini-catalogue ciblé | Formation, accompagnement et gouvernance IA")
    set_font(run, "Aptos", 11.5, GRAY)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("MINI-CATALOGUE")
    set_font(run, "Aptos Display", 18, ORANGE, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("SOLUTIONS IA ET FORMATIONS STRATÉGIQUES")
    set_font(run, "Aptos Display", 26, NAVY, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("pour Orange Côte d'Ivoire")
    set_font(run, "Aptos Display", 22, GREEN, bold=True)

    add_meta_line(document, "Télécommunications · Relation client · Opérations réseau · Côte d’Ivoire")
    add_meta_line(document, "Version ciblée sans date | www.orange.ci")

    add_section_heading(document, "QUI SOMMES-NOUS")
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("TransferAI Africa")
    set_font(run, "Aptos Display", 23, NAVY, bold=True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Hub IA de NettelecomCI en Côte d’Ivoire")
    set_font(run, "Aptos", 18, GREEN)

    add_bullets(
        document,
        [
            "13 experts ivoiriens et internationaux mobilisables.",
            "Bases d’intervention : Côte d’Ivoire, Royaume-Uni, États-Unis et Inde.",
            "Cas d’usage concrets et transposables dans les télécommunications, la finance, l’énergie et les services.",
            "Accompagnement pragmatique combiné à la formation, au pilotage et à la gouvernance.",
        ],
        size=12.3,
    )

    add_section_heading(document, "NOTRE MISSION")
    add_body(
        document,
        "Transformer l’IA en solutions concrètes, utiles et immédiatement activables pour les organisations africaines.",
        size=13,
    )

    add_section_heading(document, "NOTRE ENGAGEMENT", top_space=6)
    add_body(
        document,
        "Pas de l’IA pour l’IA : des solutions de performance, de formation et de transformation adaptées aux réalités terrain de chaque prospect.",
        size=13,
    )

    add_section_heading(document, "MESSAGE CENTRAL POUR ORANGE CÔTE D’IVOIRE", top_space=8)
    add_callout(
        document,
        "Proposition ciblée",
        "La recommandation n’est pas un catalogue large, mais une séquence ciblée entre audit, cas d’usage priorisés, formation utile, accompagnement sur 90 jours et décision d’extension sur la base de KPI concrets.",
    )

    add_section_heading(document, "HYPOTHÈSE DE GAINS ATTENDUS", top_space=8)
    add_bullets(
        document,
        [
            "Réponses plus rapides et plus homogènes sur les canaux clients.",
            "Incidents réseau mieux relayés entre exploitation, relation client et management.",
            "Lecture plus rapide des verbatims, signaux de churn et irritants service.",
            "Meilleure qualité de préparation des comités commerciaux et opérationnels.",
        ],
        size=12.2,
    )


def add_exec_summary(document: Document) -> None:
    document.add_page_break()
    add_section_heading(document, "01 - SYNTHÈSE EXÉCUTIVE POUR ORANGE CÔTE D’IVOIRE", top_space=0)
    add_body(
        document,
        "Dans les télécommunications, la valeur de l’IA se lit rapidement dans la stabilité du service, la vitesse de réponse client, la lisibilité des incidents, la qualité commerciale et la coordination entre les équipes réseau, relation client et management.",
    )

    add_subheading(document, "Priorités métier les plus crédibles")
    add_bullets(
        document,
        [
            "Incidents réseau mal relayés entre équipes techniques, front office et management.",
            "Relation client sous pression sur les demandes récurrentes, les réclamations et la qualité de réponse multi-canaux.",
            "Lecture commerciale dispersée sur churn, upsell, satisfaction et qualité de service.",
            "Coordination interéquipes perfectible entre réseau, service client, marketing, mobile money et direction.",
        ],
    )

    add_subheading(document, "Proposition ciblée TransferAI")
    add_callout(
        document,
        "Séquence recommandée",
        "Audit IA gratuit, cadrage des deux premiers cas d’usage, bootcamp ciblé décideurs + équipes clés, puis accompagnement sur 90 jours avec KPI simples, supervision humaine et décision d’extension.",
    )

    add_subheading(document, "Résultats attendus")
    add_bullets(
        document,
        [
            "Temps moyen de réponse réduit sur les canaux clients.",
            "Meilleure qualité de synthèse lors des incidents réseau et des perturbations majeures.",
            "Lecture plus rapide des signaux d’insatisfaction et des motifs de churn.",
            "Décisions managériales mieux préparées grâce à des KPI commentés et plus actionnables.",
        ],
    )


def add_use_cases(document: Document) -> None:
    document.add_page_break()
    add_section_heading(document, "02 - CAS D’USAGE PRIORITAIRES", top_space=0)
    cases = [
        (
            "Centre de contact augmenté par l’IA",
            "Scripts, réponses assistées, résumés d’échanges et routage plus homogènes pour améliorer la résolution au premier contact.",
        ),
        (
            "Synthèse d’incidents réseau et communication interne",
            "Production plus rapide d’une synthèse fiable pour coordonner exploitation, service client, direction et communication de crise.",
        ),
        (
            "Lecture des verbatims clients et signaux de churn",
            "Regroupement plus exploitable des motifs d’insatisfaction, incidents récurrents et attentes clients à forte valeur d’action.",
        ),
        (
            "Pilotage commercial et qualité de service",
            "Préparation accélérée des comités, lecture plus claire des arbitrages et commentaires plus utiles sur les KPI prioritaires.",
        ),
    ]
    for title, body in cases:
        add_subheading(document, title)
        add_body(document, body)

    add_subheading(document, "KPI de référence")
    add_numbered(
        document,
        [
            "Temps moyen de réponse et taux de résolution au premier contact.",
            "Temps de diffusion d’une synthèse d’incident fiable.",
            "Délai d’analyse des verbatims et nombre de signaux priorisés.",
            "Temps de préparation d’un comité et vitesse de décision managériale.",
        ],
    )


def add_offer(document: Document) -> None:
    document.add_page_break()
    add_section_heading(document, "03 - FORMULE PRIORITAIRE RECOMMANDÉE", top_space=0)
    add_subheading(document, "Intitulé")
    add_body(
        document,
        "Bootcamp IA sectoriel Télécommunications : de l’acculturation utile au premier pilote mesurable.",
    )

    add_subheading(document, "Public cible")
    add_body(
        document,
        "Direction générale, service client, opérations réseau, marketing, qualité de service, vente B2B/B2C, mobile money et équipes de pilotage.",
    )

    add_subheading(document, "Formats possibles")
    add_bullets(document, ["Présentiel", "Hybride", "En ligne (style webinaire)"])

    add_subheading(document, "Niveaux")
    add_bullets(document, ["Débutant", "Intermédiaire", "Avancé"])

    add_subheading(document, "Durée")
    add_bullets(document, ["1 jour : version décideurs et managers.", "2 jours : version équipes, managers et mise en pratique."])

    add_subheading(document, "Tarif")
    add_body(document, "Prix sur demande.")

    add_subheading(document, "Objectifs pédagogiques")
    add_bullets(
        document,
        [
            "Identifier les priorités métier les plus crédibles à traiter par l’IA dans les télécommunications.",
            "Structurer et piloter les cas d’usage relation client, incidents réseau, signaux de churn et pilotage qualité.",
            "Distinguer les usages immédiatement activables, les usages à encadrer et les usages à piloter dans le temps.",
            "Définir des KPI simples, un cadre de confiance et un premier plan d’action à 90 jours.",
        ],
    )


def add_programme(document: Document) -> None:
    document.add_page_break()
    add_section_heading(document, "04 - PROGRAMME PROPOSÉ", top_space=0)
    add_subheading(document, "Jour 1")
    add_numbered(
        document,
        [
            "Lecture des enjeux Orange Côte d’Ivoire et cartographie des irritants prioritaires.",
            "Études de cas guidées sur le centre de contact augmenté et la synthèse d’incidents réseau.",
            "Méthodes, prompts, routines et gestes métier autour de la relation client augmentée par l’IA.",
            "Cadre de confiance : confidentialité, supervision humaine, validation managériale et critères d’extension.",
        ],
    )

    add_subheading(document, "Jour 2")
    add_numbered(
        document,
        [
            "Ateliers de conception sur les verbatims clients, les signaux de churn et le pilotage qualité de service.",
            "Construction d’un premier canevas de pilote, des KPI de suivi et des livrables attendus.",
            "Approfondissements recommandés : pilotage réseau et qualité de service, marketing télécom, leadership data-driven.",
            "Plan d’action à 90 jours : gouvernance, équipes à former, décision d’extension ou de recentrage.",
        ],
    )


def add_deliverables(document: Document) -> None:
    add_section_heading(document, "05 - LIVRABLES ATTENDUS", top_space=8)
    add_bullets(
        document,
        [
            "Une note de cadrage sectorielle télécommunications avec priorités confirmées, risques et premières décisions.",
            "Une matrice des cas d’usage prioritaires adaptée au contexte Orange Côte d’Ivoire.",
            "Un canevas de pilote à 90 jours avec KPI, rôles, supervision et critères de succès.",
            "Une recommandation de parcours de montée en compétences pour les décideurs et les équipes ciblées.",
        ],
    )

    add_section_heading(document, "06 - CONTACT ET PROCHAINE ÉTAPE", top_space=8)
    add_callout(
        document,
        "Appel à l’action",
        "Planifier un audit IA gratuit suivi d’un échange expert de 45 minutes pour confirmer les deux premiers cas d’usage, les populations à former et le cadre de déploiement pertinent pour Orange Côte d’Ivoire.",
    )


def build_document() -> Document:
    document = Document()
    set_page_geometry(document)
    add_footer(document.sections[0])
    add_cover(document)
    add_exec_summary(document)
    add_use_cases(document)
    add_offer(document)
    add_programme(document)
    add_deliverables(document)
    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
