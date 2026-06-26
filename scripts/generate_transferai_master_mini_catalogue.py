from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUTPUT_DIR = ROOT / "outputs" / "master-templates" / "final"
OUTPUT_DOCX = OUTPUT_DIR / "TransferAI_Master_Mini_Catalogue_Prospect_Sectoriel.docx"

NAVY = RGBColor(0x1A, 0x35, 0x54)
ORANGE = RGBColor(0xE7, 0x6F, 0x1D)
SOFT = RGBColor(0x5F, 0x6F, 0x84)
BLACK = RGBColor(0x1F, 0x29, 0x34)
LIGHT = "F8F3EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_page_margins(section) -> None:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)


def set_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TransferAI Africa | www.transferai.ci | IA Assistant / WhatsApp : +225 07 16 57 39 90")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = SOFT


def style_normal(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK


def add_band_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"▌ {text}")
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.bold = True
    run.font.color.rgb = NAVY


def add_heading(document: Document, text: str, size: float = 20) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = NAVY


def add_body(document: Document, text: str, space_after: float = 6) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK


def add_soft_label(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = NAVY


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.3)
        run.font.color.rgb = BLACK


def add_meta_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(10.8)
    for label, value in rows:
        row = table.add_row()
        left = row.cells[0]
        right = row.cells[1]
        left.width = Cm(4.2)
        right.width = Cm(10.8)
        set_cell_shading(left, LIGHT)
        for cell in (left, right):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int((4.2 if cell is left else 10.8) * 567)))
            tc_w.set(qn("w:type"), "dxa")
        p1 = left.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label)
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r1.bold = True
        r1.font.color.rgb = NAVY
        p2 = right.paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(value)
        r2.font.name = "Arial"
        r2.font.size = Pt(10.2)
        r2.font.color.rgb = BLACK


def build_document() -> Document:
    document = Document()
    style_normal(document)
    for section in document.sections:
        set_page_margins(section)
        set_footer(section)

    # Page 1
    add_body(document, "TransferAI Africa | IA metier, formation et accompagnement", 2)
    add_heading(document, "MASTER MINI-CATALOGUE SOLUTIONS IA ET FORMATIONS STRATEGIQUES", 18)
    add_body(document, "pour {{SECTOR_NAME}}", 1)
    add_body(document, "{{FUNCTION_SCOPE}} | {{ORGANIZATION_TYPE}} | {{COUNTRY}}", 2)
    add_body(document, "Document de presentation sectoriel | Presentiel · Hybride · En ligne", 8)

    add_band_title(document, "PROMESSE CENTRALE")
    add_body(document, "{{CATALOGUE_PROMISE}}", 6)

    add_band_title(document, "POSITIONNEMENT DE L'OFFRE")
    add_body(document, "{{CATALOGUE_POSITIONING}}", 6)

    add_band_title(document, "BENEFICES METIER VISIBLES")
    add_bullets(
        document,
        [
            "{{BENEFIT_1}}",
            "{{BENEFIT_2}}",
            "{{BENEFIT_3}}",
            "{{BENEFIT_4}}",
        ],
    )

    document.add_page_break()

    # Page 2
    add_band_title(document, "01 - SYNTHESE EXECUTIVE {{SECTOR_NAME_UPPER}}")
    add_body(document, "{{EXECUTIVE_SUMMARY}}", 8)

    add_soft_label(document, "Priorites metier les plus credibles")
    add_bullets(
        document,
        [
            "{{PRIORITY_BULLET_1}}",
            "{{PRIORITY_BULLET_2}}",
            "{{PRIORITY_BULLET_3}}",
            "{{PRIORITY_BULLET_4}}",
        ],
    )

    add_soft_label(document, "KPI de reference")
    add_bullets(
        document,
        [
            "{{KPI_BULLET_1}}",
            "{{KPI_BULLET_2}}",
            "{{KPI_BULLET_3}}",
            "{{KPI_BULLET_4}}",
        ],
    )

    add_soft_label(document, "Approche recommandee")
    add_body(document, "{{RECOMMENDED_APPROACH}}", 4)

    add_soft_label(document, "Resultats attendus")
    add_bullets(
        document,
        [
            "{{EXPECTED_RESULT_1}}",
            "{{EXPECTED_RESULT_2}}",
            "{{EXPECTED_RESULT_3}}",
            "{{EXPECTED_RESULT_4}}",
        ],
    )

    document.add_page_break()

    # Page 3
    add_band_title(document, "02 - FORMULE PRIORITAIRE")
    add_meta_table(
        document,
        [
            ("Intitule", "{{PRIORITY_OFFER_TITLE}}"),
            ("Public cible", "{{TARGET_AUDIENCE}}"),
            ("Formats possibles", "{{FORMAT_1}} | {{FORMAT_2}} | {{FORMAT_3}}"),
            ("Niveaux", "{{LEVEL_1}} | {{LEVEL_2}} | {{LEVEL_3}}"),
            ("Duree", "{{DURATION_1}} | {{DURATION_2}}"),
            ("Tarif", "{{PRICE_BLOCK}}"),
        ],
    )

    add_soft_label(document, "Objectifs pedagogiques")
    add_bullets(
        document,
        [
            "{{OBJECTIVE_1}}",
            "{{OBJECTIVE_2}}",
            "{{OBJECTIVE_3}}",
            "{{OBJECTIVE_4}}",
        ],
    )

    add_soft_label(document, "Formations proposees")
    add_bullets(
        document,
        [
            "{{TRAINING_LINE_1}}",
            "{{TRAINING_LINE_2}}",
            "{{TRAINING_LINE_3}}",
            "{{TRAINING_LINE_4}}",
            "{{TRAINING_LINE_5}}",
        ],
    )

    document.add_page_break()

    # Page 4
    add_band_title(document, "03 - PROGRAMME PROPOSE")
    add_soft_label(document, "Jour 1")
    add_bullets(
        document,
        [
            "{{DAY1_ITEM_1}}",
            "{{DAY1_ITEM_2}}",
            "{{DAY1_ITEM_3}}",
            "{{DAY1_ITEM_4}}",
        ],
    )

    add_soft_label(document, "Jour 2")
    add_bullets(
        document,
        [
            "{{DAY2_ITEM_1}}",
            "{{DAY2_ITEM_2}}",
            "{{DAY2_ITEM_3}}",
            "{{DAY2_ITEM_4}}",
        ],
    )

    add_band_title(document, "04 - LIVRABLES ATTENDUS")
    add_bullets(
        document,
        [
            "{{DELIVERABLE_1}}",
            "{{DELIVERABLE_2}}",
            "{{DELIVERABLE_3}}",
            "{{DELIVERABLE_4}}",
        ],
    )

    add_band_title(document, "05 - CHAMPS VARIABLES A PERSONNALISER")
    add_bullets(
        document,
        [
            "{{CUSTOM_FIELD_1}}",
            "{{CUSTOM_FIELD_2}}",
            "{{CUSTOM_FIELD_3}}",
            "{{CUSTOM_FIELD_4}}",
            "{{CUSTOM_FIELD_5}}",
            "{{CUSTOM_FIELD_6}}",
        ],
    )

    add_body(
        document,
        "Bloc de cloture recommande : audit IA gratuit + echange expert + confirmation des deux premiers cas d'usage a lancer pour {{ORGANIZATION_NAME}}.",
        0,
    )

    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
