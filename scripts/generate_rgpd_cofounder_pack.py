from __future__ import annotations

import json
import math
import re
from collections import Counter
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
EXCEL_DIR = DOCS_DIR / "excel"

INVENTORY_MD = DOCS_DIR / "32_Inventory_Form_Fields_RGPD.md"
INVENTORY_JSON = DOCS_DIR / "32_form_fields_inventory_rgpd.json"
MATRIX_MD = DOCS_DIR / "33_Matrice_Protection_LLM_Prospects_Clients.md"
MATRIX_JSON = DOCS_DIR / "33_llm_protection_matrix_prospects_clients.json"

TODAY = date.today().isoformat()

WORD_FILES = {
    "inventory": WORD_DIR / f"TransferAI_Africa_RGPD_Inventory_Prospects_Clients_{TODAY}.docx",
    "matrix": WORD_DIR / f"TransferAI_Africa_RGPD_LLM_Protection_Matrix_{TODAY}.docx",
    "pack": WORD_DIR / f"TransferAI_Africa_Cofounders_RGPD_LLM_Data_Governance_Pack_{TODAY}.docx",
}
EXCEL_FILE = EXCEL_DIR / f"TransferAI_Africa_Cofounders_RGPD_LLM_Data_Governance_Pack_{TODAY}.xlsx"


def ensure_dirs() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    EXCEL_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def normalize_markdown_text(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(document: Document) -> None:
    section = document.sections[0]
    apply_section_layout(section)
    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.5)


def apply_section_layout(section) -> None:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)


def start_new_section(document: Document) -> None:
    paragraph = document.add_paragraph(style="Normal")
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_title(document: Document, title: str, subtitle: str | None = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x4C)
    if subtitle:
        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_kpi_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Indicateur"
    hdr[1].text = "Valeur"
    set_cell_shading(hdr[0], "DCE6F1")
    set_cell_shading(hdr[1], "DCE6F1")
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def render_markdown_to_docx(md_path: Path, output_path: Path, title: str) -> None:
    document = Document()
    style_doc(document)
    add_title(document, title, f"Version Word generee le {TODAY}")
    document.add_paragraph("")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_table = False
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buffer, in_table
        if not table_buffer:
            return
        widths = max(len(row) for row in table_buffer)
        table = document.add_table(rows=1, cols=widths)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_cells = table.rows[0].cells
        for idx, value in enumerate(table_buffer[0]):
            header_cells[idx].text = normalize_markdown_text(value)
            set_cell_shading(header_cells[idx], "DCE6F1")
        for row in table_buffer[1:]:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = normalize_markdown_text(value)
        table_buffer = []
        in_table = False

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            flush_table()
            document.add_paragraph("")
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if set(cells) <= {"---", ":---", "---:", ":---:"}:
                continue
            table_buffer.append(cells)
            in_table = True
            continue

        flush_table()

        if stripped.startswith("# "):
            p = document.add_paragraph()
            p.style = "Title"
            p.add_run(normalize_markdown_text(stripped[2:])).bold = True
            continue
        if stripped.startswith("## "):
            p = document.add_paragraph()
            p.style = "Heading 1"
            run = p.add_run(normalize_markdown_text(stripped[3:]))
            run.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x3D, 0x4C)
            continue
        if stripped.startswith("### "):
            p = document.add_paragraph()
            p.style = "Heading 2"
            run = p.add_run(normalize_markdown_text(stripped[4:]))
            run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x6E)
            continue
        if stripped.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            p.add_run(normalize_markdown_text(stripped[2:]))
            continue

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(normalize_markdown_text(stripped))

    flush_table()
    document.save(output_path)


def flatten_collection_points(inventory: dict) -> list[dict]:
    rows: list[dict] = []
    for point in inventory["collectionPoints"]:
        rows.append(
            {
                "id": point["id"],
                "route": point["route"],
                "storage": ", ".join(point.get("storage", [])),
                "input_fields_count": len(point.get("inputFields", [])),
                "stored_fields_count": len(point.get("storedFields", [])),
                "llm_excluded_fields_count": len(point.get("llmExcludeFields", [])),
            }
        )
    return rows


def classify_field(field: str) -> str:
    lowered = field.lower()
    if any(token in lowered for token in ["email", "name", "phone", "company", "city", "country", "position", "profession", "organization", "website", "wa_id"]):
        return "Direct identifier"
    if any(token in lowered for token in ["password", "hash", "token", "session", "sid", "response_id", "provider_message_id"]):
        return "Credential or token"
    if any(token in lowered for token in ["message", "notes", "body", "payload", "motivation", "form_data", "draft_form_data", "error"]):
        return "Free text or raw payload"
    if any(token in lowered for token in ["sector", "domain", "budget", "horizon", "participants", "status", "language", "source_page", "category"]):
        return "Qualification or operational metadata"
    return "Operational field"


def flatten_field_inventory(inventory: dict) -> list[dict]:
    rows: list[dict] = []
    for point in inventory["collectionPoints"]:
        excluded = set(point.get("llmExcludeFields", []))
        for field in point.get("storedFields", []):
            rows.append(
                {
                    "collection_point": point["id"],
                    "route": point["route"],
                    "storage": ", ".join(point.get("storage", [])),
                    "stored_field": field,
                    "category": classify_field(field),
                    "llm_default_status": "BLOCKED" if field in excluded else "REVIEW_REQUIRED",
                }
            )
    return rows


def summarize_inventory(inventory: dict, matrix: dict) -> dict:
    collection_points = inventory["collectionPoints"]
    field_rows = flatten_field_inventory(inventory)
    categories = Counter(row["category"] for row in field_rows)
    blocked_tables = sum(1 for table in matrix["tables"] if not table.get("pseudonymizeFields") and table.get("blockedFields"))
    return {
        "collection_points": len(collection_points),
        "tables_or_logs": len({storage for point in collection_points for storage in point.get("storage", [])}),
        "stored_fields": len(field_rows),
        "blocked_by_default": sum(1 for row in field_rows if row["llm_default_status"] == "BLOCKED"),
        "critical_tables": blocked_tables,
        "direct_identifier_fields": categories["Direct identifier"],
        "credential_fields": categories["Credential or token"],
        "free_text_fields": categories["Free text or raw payload"],
    }


def add_exec_summary(document: Document, inventory: dict, matrix: dict) -> None:
    summary = summarize_inventory(inventory, matrix)
    add_title(
        document,
        "Data Governance & Conformite RGPD",
        "Pack partageable co-fondateurs - donnees prospects / clients et usages LLM",
    )
    document.add_paragraph("")
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        "Ce pack synthetise l'ensemble des informations personnelles, metier et potentiellement sensibles collectees "
        "par le site web TransferAI Africa pour les prospects et clients. Il sert de base commune pour arbitrer "
        "la gouvernance des donnees, le perimetre LLM autorise et les priorites de securisation."
    )
    document.add_paragraph("")
    add_kpi_table(
        document,
        [
            ("Points de collecte analyses", str(summary["collection_points"])),
            ("Tables / logs concernes", str(summary["tables_or_logs"])),
            ("Champs stockes inventories", str(summary["stored_fields"])),
            ("Champs bloques par defaut pour les LLM", str(summary["blocked_by_default"])),
            ("Tables a risque maximal", str(summary["critical_tables"])),
        ],
    )
    document.add_paragraph("")
    document.add_paragraph("Actions prioritaires recommandees", style="Heading 1")
    for item in [
        "Auditer et valider chaque point de collecte prospect / client, y compris newsletters, webinaires, inscriptions, audit, WhatsApp et logs de diffusion.",
        "Bloquer par defaut tout envoi LLM contenant identifiants directs, secrets, texte libre ou donnees d'audit brut.",
        "Mettre en place un module unique de redaction, pseudonymisation et journalisation avant tout appel LLM.",
        "Definir une allowlist de champs par workflow metier, avec revue humaine sur les usages sensibles.",
        "Traiter en priorite le formulaire d'audit, les messages WhatsApp et les journaux de diffusion nominative.",
    ]:
        document.add_paragraph(item, style="List Bullet")


def add_collection_point_section(document: Document, inventory: dict) -> None:
    start_new_section(document)
    document.add_paragraph("Perimetre de collecte", style="Heading 1")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Point", "Stockage principal", "Champs stockes", "Risque LLM", "Observation cle"]
    for idx, label in enumerate(headers):
        table.rows[0].cells[idx].text = label
        set_cell_shading(table.rows[0].cells[idx], "DCE6F1")
    risk_notes = {
        "audit_questionnaire": (
            "Tres eleve",
            "Questionnaire le plus sensible; form_data doit rester hors LLM par defaut.",
        ),
        "whatsapp_inbound_messages": (
            "Tres eleve",
            "Messages libres et payloads bruts; redaction locale obligatoire.",
        ),
        "audit_request": (
            "Tres eleve",
            "Inclut credentials, hashes et tokens du portail prospect.",
        ),
        "partner_listing": (
            "Eleve",
            "Combine identite, entreprise, besoins et contenus de reponse.",
        ),
        "contact_requests": (
            "Eleve",
            "Donnees de contact et qualification commerciale nominative.",
        ),
        "webinar_registration": (
            "Eleve",
            "Donnees d'inscription nominatives avec motivation libre.",
        ),
        "registration_requests": (
            "Eleve",
            "Inscriptions formation avec message libre et coordonnees.",
        ),
        "newsletter_subscription": (
            "Modere",
            "Email direct bloque; preferences exploitables seulement apres pseudonymisation.",
        ),
        "prospect_email_delivery_logs": (
            "Eleve",
            "Logs d'envoi contenant destinataires, erreurs et meta.",
        ),
        "newsletter_delivery_logs": (
            "Eleve",
            "Logs de campagne nominative a restreindre aux agregats.",
        ),
        "catalogue_delivery_logs": (
            "Modere",
            "Suivi d'envoi individuel a convertir en statistiques anonymes.",
        ),
        "whatsapp_email_notification_logs": (
            "Modere",
            "Meta de notification a exploiter uniquement en taux agreges.",
        ),
        "catalogue_download_modal": (
            "Faible a modere",
            "Collecte limitee, mais tout identifiant doit rester hors LLM.",
        ),
    }
    for row in flatten_collection_points(inventory):
        risk_label, note = risk_notes.get(
            row["id"],
            ("A confirmer", "A qualifier selon les champs reels et les usages cibles."),
        )
        cells = table.add_row().cells
        cells[0].text = row["id"]
        cells[1].text = row["storage"]
        cells[2].text = str(row["stored_fields_count"])
        cells[3].text = risk_label
        cells[4].text = note


def add_risk_focus_section(document: Document, inventory: dict, matrix: dict) -> None:
    start_new_section(document)
    document.add_paragraph("Points de vigilance critiques", style="Heading 1")
    critical_points = [
        (
            "Formulaire d'audit",
            "Le champ `form_data` concentre des reponses riches sur gouvernance, outils, securite, systemes et conformite. Il doit rester hors LLM par defaut.",
        ),
        (
            "Portail prospect et invitations",
            "Les identifiants, hashes et tokens d'invitation relevent du secret d'authentification et ne doivent jamais etre exposes dans un prompt.",
        ),
        (
            "WhatsApp et messages libres",
            "Les messages entrants, notes internes et charges brutes doivent etre rediges localement avant toute analyse IA.",
        ),
        (
            "Logs de diffusion",
            "Les journaux d'envoi email et notification contiennent des destinataires nominaux, des erreurs et des payloads qui doivent rester hors LLM.",
        ),
    ]
    for title, body in critical_points:
        p = document.add_paragraph(style="Heading 2")
        p.add_run(title).bold = True
        document.add_paragraph(body)

    document.add_paragraph("Regles d'usage LLM", style="Heading 1")
    rules_table = document.add_table(rows=1, cols=4)
    rules_table.style = "Table Grid"
    rules_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(["Famille", "Exemples", "Statut", "Regle"]):
        rules_table.rows[0].cells[idx].text = label
        set_cell_shading(rules_table.rows[0].cells[idx], "DCE6F1")
    for rule in matrix["rules"]:
        cells = rules_table.add_row().cells
        cells[0].text = rule["dataFamily"]
        cells[1].text = ", ".join(rule["examples"])
        cells[2].text = rule["llmStatus"]
        cells[3].text = rule["rule"]


def add_table_matrix_section(document: Document, matrix: dict) -> None:
    start_new_section(document)
    document.add_paragraph("Matrice par table", style="Heading 1")
    for table_rule in matrix["tables"]:
        p = document.add_paragraph(style="Heading 2")
        p.add_run(table_rule["table"]).bold = True
        document.add_paragraph("Champs bloques: " + ", ".join(table_rule.get("blockedFields", []) or ["aucun"]))
        document.add_paragraph(
            "Champs a pseudonymiser: "
            + ", ".join(table_rule.get("pseudonymizeFields", []) or ["aucun"])
        )
        document.add_paragraph(
            "Agregats autorises: "
            + ", ".join(table_rule.get("allowedAggregates", []) or ["aucun"])
        )


def add_next_steps(document: Document) -> None:
    start_new_section(document)
    document.add_paragraph("Decision et prochaines etapes", style="Heading 1")
    for step in [
        "Valider en comite fondateur le principe `deny-by-default` pour tous les usages LLM externes.",
        "Nommer un proprietaire de la gouvernance des donnees prospects / clients et un proprietaire technique des garde-fous applicatifs.",
        "Traduire la matrice en spec technique: allowlists par workflow, regles de pseudonymisation, journalisation et revue humaine.",
        "Lancer un PIA cible sur le formulaire d'audit, les logs nominaux et les conversations WhatsApp.",
        "Verifier les durées de conservation, droits d'acces et chiffrement pour chaque table identifiee.",
    ]:
        document.add_paragraph(step, style="List Bullet")
    note = document.add_paragraph()
    note.add_run("Base d'analyse: ").bold = True
    note.add_run("code du site et schemas presents dans le depot au 2026-05-22; pas de lecture directe des donnees de production.")


def build_pack_docx(inventory: dict, matrix: dict, output_path: Path) -> None:
    document = Document()
    style_doc(document)
    add_exec_summary(document, inventory, matrix)
    add_collection_point_section(document, inventory)
    add_risk_focus_section(document, inventory, matrix)
    add_table_matrix_section(document, matrix)
    add_next_steps(document)
    document.save(output_path)


def make_sheet(ws, title: str, rows: list[list[object]], widths: list[int], freeze: str = "A2") -> None:
    ws.title = title
    for row in rows:
        ws.append(row)
    header_fill = PatternFill("solid", fgColor="0F3D4C")
    header_font = Font(color="FFFFFF", bold=True)
    body_border = Border(
        left=Side(style="thin", color="D9D9D9"),
        right=Side(style="thin", color="D9D9D9"),
        top=Side(style="thin", color="D9D9D9"),
        bottom=Side(style="thin", color="D9D9D9"),
    )
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row in ws.iter_rows():
        for cell in row:
            cell.border = body_border
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for idx, width in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(idx)].width = width
    ws.freeze_panes = freeze
    ws.auto_filter.ref = ws.dimensions
    end_row = ws.max_row
    end_col = ws.max_column
    if end_row >= 2 and end_col >= 1:
        table = Table(displayName=f"{title.replace(' ', '_')}_table", ref=f"A1:{get_column_letter(end_col)}{end_row}")
        table.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        ws.add_table(table)


def build_excel(inventory: dict, matrix: dict, output_path: Path) -> None:
    wb = Workbook()
    ws = wb.active

    summary = summarize_inventory(inventory, matrix)
    overview_rows = [
        ["Pack", "TransferAI Africa - Gouvernance des donnees prospects / clients"],
        ["Date d'analyse", inventory["generatedAt"]],
        ["Repository", inventory["repository"]],
        ["Politique LLM", inventory["defaultLlmPolicy"]["mode"]],
        ["Points de collecte", summary["collection_points"]],
        ["Tables / logs", summary["tables_or_logs"]],
        ["Champs stockes", summary["stored_fields"]],
        ["Champs bloques par defaut", summary["blocked_by_default"]],
        ["Tables a risque maximal", summary["critical_tables"]],
        ["Champ critique principal", "form_responses.form_data"],
        ["Message directeur", "Aucune donnee prospect / client ne part vers un LLM externe par defaut"],
    ]
    make_sheet(ws, "Overview", overview_rows, [28, 110], freeze="A1")

    collection_rows = [["Collection point", "Route", "Storage", "Input fields", "Stored fields", "LLM excluded fields"]]
    for row in flatten_collection_points(inventory):
        collection_rows.append(
            [
                row["id"],
                row["route"],
                row["storage"],
                row["input_fields_count"],
                row["stored_fields_count"],
                row["llm_excluded_fields_count"],
            ]
        )
    make_sheet(wb.create_sheet(), "Collection_Points", collection_rows, [28, 28, 32, 14, 14, 18])

    field_rows = [["Collection point", "Route", "Storage", "Stored field", "Category", "Default LLM status"]]
    for row in flatten_field_inventory(inventory):
        field_rows.append(
            [
                row["collection_point"],
                row["route"],
                row["storage"],
                row["stored_field"],
                row["category"],
                row["llm_default_status"],
            ]
        )
    make_sheet(wb.create_sheet(), "Field_Inventory", field_rows, [24, 24, 30, 28, 32, 20])

    rule_rows = [["Data family", "Examples", "LLM status", "Rule"]]
    for rule in matrix["rules"]:
        rule_rows.append(
            [
                rule["dataFamily"],
                ", ".join(rule["examples"]),
                rule["llmStatus"],
                rule["rule"],
            ]
        )
    make_sheet(wb.create_sheet(), "LLM_Rules", rule_rows, [28, 56, 24, 68])

    table_rows = [["Table", "Blocked fields", "Pseudonymize fields", "Allowed aggregates"]]
    for table in matrix["tables"]:
        table_rows.append(
            [
                table["table"],
                ", ".join(table.get("blockedFields", [])),
                ", ".join(table.get("pseudonymizeFields", [])),
                ", ".join(table.get("allowedAggregates", [])),
            ]
        )
    make_sheet(wb.create_sheet(), "Table_Matrix", table_rows, [30, 62, 52, 40])

    critical_rows = [["Priority", "Topic", "Why it matters", "Immediate action"]]
    for item in [
        [
            "P1",
            "Audit questionnaire data",
            "Contains rich business, governance and compliance details in form_data",
            "Keep outside LLMs by default and design a sanitized extract only if needed",
        ],
        [
            "P1",
            "Credentials and invite tokens",
            "Passwords, hashes and invitation tokens are authentication secrets",
            "Prohibit prompt exposure and review storage controls",
        ],
        [
            "P1",
            "WhatsApp raw conversations",
            "Free text and raw payloads can reveal identities, intent and operational context",
            "Use local redaction and human review only",
        ],
        [
            "P2",
            "Delivery logs",
            "Email recipients, errors and payloads create re-identification risk",
            "Limit LLM use to anonymized performance aggregates",
        ],
        [
            "P2",
            "Commercial qualification data",
            "Sector, budget, horizon and format become identifying when combined",
            "Allow only minimized pseudonymized use cases",
        ],
    ]:
        critical_rows.append(item)
    make_sheet(wb.create_sheet(), "Critical_Focus", critical_rows, [12, 28, 68, 68])

    for sheet in wb.worksheets:
        if sheet.title == "Overview":
            continue
        for row in sheet.iter_rows(min_row=2):
            if any(cell.value == "BLOCKED" or cell.value == "blocked" for cell in row):
                for cell in row:
                    cell.fill = PatternFill("solid", fgColor="FDE9D9")
            if any(cell.value == "pseudonymize_required" for cell in row):
                for cell in row:
                    if cell.fill.patternType is None:
                        cell.fill = PatternFill("solid", fgColor="FFF2CC")

    wb.save(output_path)


def main() -> None:
    ensure_dirs()
    inventory = load_json(INVENTORY_JSON)
    matrix = load_json(MATRIX_JSON)

    render_markdown_to_docx(
        INVENTORY_MD,
        WORD_FILES["inventory"],
        "Inventaire des donnees prospects / clients - Version Word",
    )
    render_markdown_to_docx(
        MATRIX_MD,
        WORD_FILES["matrix"],
        "Matrice de protection LLM - Version Word",
    )
    build_pack_docx(inventory, matrix, WORD_FILES["pack"])
    build_excel(inventory, matrix, EXCEL_FILE)

    manifest = {
        "word": {key: str(path.relative_to(ROOT)) for key, path in WORD_FILES.items()},
        "excel": str(EXCEL_FILE.relative_to(ROOT)),
    }
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
