"""
Generate User Guide + Troubleshooting Guide for session 2026-06-27/28
Covers: audit form fixes, V3 pack fix, V6 routing fixes, email notification chain
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

OUTPUT_DIR = "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin"


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p


def add_numbered(doc, text):
    return doc.add_paragraph(text, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


def add_table_row(table, *cols, header=False):
    row = table.add_row()
    for i, col in enumerate(cols):
        if i < len(row.cells):
            row.cells[i].text = col
    if header:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
    return row


# ─────────────────────────────────────────────
# DOCUMENT 1 — USER GUIDE
# ─────────────────────────────────────────────

def build_user_guide():
    doc = Document()

    # Title
    title = doc.add_heading("Guide Utilisateur — Session Correction & Validation\nFormulaire Audit IA + Pipeline Post-Audit V6", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"TransferAI — Pipeline Prospection Automatisé | Date : {datetime.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph("Couvre les corrections appliquées lors de la session du 27-28 juin 2026.")
    doc.add_horizontal_rule() if hasattr(doc, "add_horizontal_rule") else doc.add_paragraph("─" * 80)

    # 1. Vue d'ensemble
    add_heading(doc, "1. Vue d'ensemble des corrections", 1)
    add_body(doc, "Cette session a permis de corriger et valider l'intégralité de la chaîne de traitement post-audit :")
    for item in [
        "Formulaire d'audit (www.transferai.ci/questionnaire-audit) — robustesse et personnalisation",
        "Edge Function resolve-invitation — création automatique d'invitation depuis pack_id",
        "Workflow V3 (Assemble Prospect Pack) — propagation du commercial_priority_tier",
        "Workflow V6 (Post-Audit Expert Routing) — détection de priorité et envoi d'emails",
        "Notifications email via Resend — fiche pré-RDV + alerte expert [PRIORITÉ HAUTE]",
    ]:
        add_bullet(doc, item)

    # 2. Formulaire d'audit
    add_heading(doc, "2. Formulaire d'audit — Utilisation", 1)

    add_heading(doc, "2.1 Accès au formulaire", 2)
    add_body(doc, "Le formulaire s'ouvre via un lien pack ou un lien d'invitation :")
    add_code(doc, "https://www.transferai.ci/questionnaire-audit?pack_id=<PACK_ID>")
    add_code(doc, "https://www.transferai.ci/questionnaire-audit?invite=<TOKEN>")
    add_body(doc, "Lors du premier accès via pack_id, le système crée automatiquement une invitation dans form_invitations avec un contexte personnalisé (nom de l'organisation, secteur, cas d'usage suggérés).")

    add_heading(doc, "2.2 Remplissage du formulaire", 2)
    for item in [
        "Remplir les champs de contact (nom, email, organisation, poste)",
        "Sélectionner le secteur d'activité",
        "Répondre aux questions sur la maturité IA, les cas d'usage, les contraintes",
        "Le formulaire sauvegarde automatiquement (brouillon) à chaque modification",
        "Cliquer 'Soumettre' pour envoyer définitivement",
    ]:
        add_numbered(doc, item)

    add_heading(doc, "2.3 Ce qui se passe après soumission", 2)
    add_body(doc, "Après soumission :")
    for item in [
        "La réponse est enregistrée dans form_responses (is_completed = true)",
        "V6 détecte la soumission lors de sa prochaine exécution (toutes les 30 min)",
        "V6 génère une fiche pré-RDV via GPT-4",
        "Deux emails sont envoyés à contact@transferai.ci et marius.ayoro70@gmail.com :",
        "  → [PRIORITÉ HAUTE] Nouveau dossier post-audit à traiter",
        "  → Fiche pré-RDV complète (synthèse, cas d'usage, quick wins, contraintes, prochaine étape)",
    ]:
        add_bullet(doc, item)

    # 3. Personnalisation
    add_heading(doc, "3. Personnalisation du formulaire", 1)
    add_body(doc, "Le formulaire affiche un contexte personnalisé issu du pack :")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    add_table_row(table, "Champ affiché", "Source", header=True)
    for row in [
        ("Nom de l'organisation", "ai_prospecting_packs.organization_name"),
        ("Secteur d'activité", "payload.sector_guess via access_context.sectorLabel"),
        ("Contexte prospect", "payload.organization_summary (de-anonymisé)"),
        ("Cas d'usage suggérés", "payload.probable_needs + recommended_use_case"),
        ("Quick wins suggérés", "payload.probable_quick_wins"),
    ]:
        add_table_row(table, *row)
    doc.add_paragraph()

    add_body(doc, "Note : Le placeholder ORG_TARGET (utilisé par V3 pour anonymiser le texte envoyé au LLM) est maintenant remplacé automatiquement par le vrai nom de l'organisation avant affichage.")

    # 4. V3
    add_heading(doc, "4. Workflow V3 — Correction commercial_priority_tier", 1)

    add_heading(doc, "4.1 Problème corrigé", 2)
    add_body(doc, "Le pack généré par V3 avait commercial_priority_tier vide car le noeud 'Assemble Prospect Pack' lisait ctx.commercial_priority_tier qui n'était jamais peuplé. La valeur était stockée sous le nom commercial_priority_default dans 'Set Target'.")

    add_heading(doc, "4.2 Correction appliquée", 2)
    add_body(doc, "Dans le noeud 'Assemble Prospect Pack', la ligne suivante a été modifiée :")
    add_code(doc, "// AVANT")
    add_code(doc, "commercial_priority_tier: ctx.commercial_priority_tier || '',")
    add_code(doc, "// APRÈS")
    add_code(doc, "commercial_priority_tier: ctx.commercial_priority_tier || ctx.commercial_priority_default || '',")

    add_heading(doc, "4.3 Impact", 2)
    add_body(doc, "Les nouveaux packs générés auront désormais commercial_priority_tier = 'tier1', 'tier2' ou 'tier3' selon la valeur dans Set Target. Cette valeur est lue par V6 pour déterminer la priorité de routage.")

    # 5. V6
    add_heading(doc, "5. Workflow V6 — Corrections Post-Audit Expert Routing", 1)

    add_heading(doc, "5.1 Corrections apportées", 2)
    corrections = [
        ("Build Expert Routing", "Priorité tier1 non détectée", "Ajout support tier1/tier2 en plus de A/B legacy"),
        ("If Expert Notification Enabled", "Email uniquement pour 'high'", "Condition changée en high OR medium"),
        ("If Pack Found", "Condition value1=value2 (toujours false)", "Corrigé en $json.found is true"),
        ("If Audit Completed", "Condition value1=value2 (toujours false)", "Corrigé en $json.audit_completed is true"),
        ("If Already Post-Audit Processed", "Ancien format booléen n8n", "Mis à jour vers nouveau format v2"),
        ("If Slack Alerts Enabled", "Condition value1=value2 (toujours false)", "Corrigé en send_slack_alert is true"),
        ("Get Latest Form Response", "Filtre par pack_id inexistant", "URL changée pour filtrer par response_id"),
    ]
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    add_table_row(table2, "Noeud", "Problème", "Correction", header=True)
    for r in corrections:
        row = table2.add_row()
        row.cells[0].text = r[0]
        row.cells[1].text = r[1]
        row.cells[2].text = r[2]
    doc.add_paragraph()

    add_heading(doc, "5.2 Logique de priorité", 2)
    add_body(doc, "Le noeud 'Build Expert Routing' détermine routing_priority selon :")
    for rule in [
        "completion_percentage >= 95 → high",
        "maturity contient 'avance' ou 'advanced' → high",
        "commercial_priority_tier = 'tier1' ou 'A' → high",
        "commercial_priority_tier = 'tier2' ou 'B' → medium",
        "Sinon → normal",
    ]:
        add_bullet(doc, rule)

    add_body(doc, "L'email expert est envoyé si routing_priority = 'high' OU 'medium'.")

    # 6. Emails reçus
    add_heading(doc, "6. Emails de notification — Structure", 1)

    add_heading(doc, "Email 1 : Notification expert [PRIORITÉ HAUTE]", 2)
    for field in [
        "Objet : [PRIORITE HAUTE] Nouveau dossier post-audit à traiter — {organisation}",
        "De : TransferAI <contact@transferai.ci>",
        "À : assigned_expert_email (contact@transferai.ci par défaut)",
        "Contenu : Organisation, contact, secteur, maturité IA, service recommandé, prochaine action, Pack ID, lien Calendly",
    ]:
        add_bullet(doc, field)

    add_heading(doc, "Email 2 : Fiche pré-RDV interne", 2)
    for field in [
        "Objet : Fiche pré-RDV post-audit — {organisation}",
        "De : TransferAI <contact@transferai.ci>",
        "À : contact@transferai.ci + expert assigné",
        "Contenu : Fiche complète générée par GPT-4 (identification, synthèse, maturité IA, cas d'usage, quick wins, contraintes, orientation TransferAI, prochaine étape, documents à consulter, lien RDV)",
    ]:
        add_bullet(doc, field)

    # 7. Réinitialisation
    add_heading(doc, "7. Réinitialisation d'une invitation pour re-test", 1)
    add_body(doc, "Pour effacer une soumission et retester le formulaire depuis zéro :")
    add_code(doc, "-- Remplacer PACK_ID par le vrai pack_id")
    add_code(doc, "DELETE FROM form_responses")
    add_code(doc, "WHERE invitation_token IN (")
    add_code(doc, "  SELECT invite_token FROM form_invitations WHERE pack_id = 'PACK_ID'")
    add_code(doc, ");")
    add_code(doc, "DELETE FROM form_invitations WHERE pack_id = 'PACK_ID';")
    add_body(doc, "Après exécution, rouvrir le lien pack_id — une nouvelle invitation sera créée automatiquement.")

    # 8. Déclenchement manuel V6
    add_heading(doc, "8. Déclencher V6 manuellement", 1)
    add_body(doc, "Pour tester sans attendre l'exécution automatique (toutes les 30 min) :")
    add_numbered(doc, "Ouvrir V6 dans n8n")
    add_numbered(doc, "Aller dans 'Set Post-Audit Manual Input'")
    add_numbered(doc, "Renseigner : pack_id, response_id, internal_email_to, next_action_delay_days=1, force_rerun=true")
    add_numbered(doc, "Fermer le noeud → Execute Workflow depuis l'Editor")
    add_body(doc, "Note : response_id se trouve dans Supabase → Table form_responses → colonne id.")

    out_path = os.path.join(OUTPUT_DIR, "105_Guide_Utilisateur_Session_2026-06-28_Audit_V6.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ─────────────────────────────────────────────
# DOCUMENT 2 — TROUBLESHOOTING GUIDE
# ─────────────────────────────────────────────

def build_troubleshooting_guide():
    doc = Document()

    title = doc.add_heading("Guide Troubleshooting — Formulaire Audit + Pipeline V6\nSession 2026-06-27/28", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"TransferAI | Date : {datetime.today().strftime('%d/%m/%Y')}")
    doc.add_paragraph("─" * 80)

    add_heading(doc, "1. Formulaire d'audit — Problèmes et solutions", 1)

    issues = [
        (
            "Formulaire affiche 'Le questionnaire n'est pas encore disponible'",
            [
                "L'invitation a expiré (status=expired ou expires_at dépassé)",
                "Le pack_id n'existe pas dans ai_prospecting_packs",
                "L'ancienne version de la Edge Function resolve-invitation était déployée (v12)",
            ],
            [
                "Vérifier dans Supabase : SELECT * FROM form_invitations WHERE pack_id = '...'",
                "Si l'invitation est expirée : DELETE FROM form_invitations WHERE pack_id = '...' puis recharger",
                "Si la Edge Function est ancienne : supabase functions deploy resolve-invitation --no-verify-jwt",
                "Vérifier que la table ai_prospecting_packs contient bien le pack",
            ]
        ),
        (
            "Formulaire affiche 'Missing invite token'",
            [
                "L'ancienne Edge Function (v12) ne gérait pas le paramètre pack_id",
                "Le packId n'était pas transmis dans le body de la requête",
            ],
            [
                "Redéployer la Edge Function : supabase functions deploy resolve-invitation --no-verify-jwt",
                "La version actuelle gère pack_id ET invite_token",
            ]
        ),
        (
            "Les mots se collent sans espaces dans les textareas",
            [
                "La fonction trimText() était appliquée sur value={} du textarea",
                "trimText() supprime les espaces en fin de chaîne à chaque frappe",
            ],
            [
                "Corrigé dans ProspectAuditFormPage.tsx :",
                "  AVANT : value={trimText(formData[key])}",
                "  APRÈS : value={typeof formData[key] === 'string' ? formData[key] : ''}",
            ]
        ),
        (
            "Le texte affiche 'ORG_TARGET' au lieu du nom de l'organisation",
            [
                "V3 anonymise le texte public avec ORG_TARGET avant envoi au LLM",
                "Le texte généré par le LLM conservait ORG_TARGET dans organization_summary",
                "La Edge Function ne de-anonymisait pas avant de stocker dans access_context",
            ],
            [
                "Corrigé dans _shared/prospect-audit-context.ts :",
                "  Ajout d'une fonction deAnonymize() qui remplace ORG_TARGET par le vrai nom",
                "  Appliqué au rawProspectContext avant stockage",
                "Pour les anciennes invitations : DELETE + rechargement du lien pour recréer",
            ]
        ),
        (
            "Formulaire bloqué en chargement indéfini",
            [
                "La Edge Function ne répond pas (timeout réseau)",
                "Aucun timeout n'était configuré sur le fetch()",
            ],
            [
                "Corrigé : AbortController avec timeout 12 secondes ajouté",
                "Message d'erreur affiché après 12s si pas de réponse",
            ]
        ),
        (
            "Erreur 'column ai_prospecting_packs.organization_type does not exist'",
            [
                "La requête SELECT dans resolve-invitation incluait organization_type",
                "Cette colonne n'existe pas dans la table ai_prospecting_packs",
            ],
            [
                "Corrigé dans resolve-invitation/index.ts :",
                "  AVANT : .select('pack_id, organization_name, organization_type, target_email, payload')",
                "  APRÈS : .select('pack_id, organization_name, target_email, payload')",
            ]
        ),
    ]

    for title_text, causes, solutions in issues:
        add_heading(doc, title_text, 2)
        add_body(doc, "Causes :")
        for c in causes:
            add_bullet(doc, c, level=0)
        add_body(doc, "Solutions :")
        for s in solutions:
            add_bullet(doc, s, level=0)
        doc.add_paragraph()

    add_heading(doc, "2. Workflow V6 — Problèmes et solutions", 1)

    v6_issues = [
        (
            "V6 s'exécute mais aucun email n'est envoyé",
            [
                "routing_priority = 'normal' car commercial_priority_tier est vide dans le pack",
                "La condition 'If Expert Notification Enabled' ne passait que pour 'high'",
                "Les noeuds If Pack Found / If Audit Completed avaient des conditions cassées",
            ],
            [
                "Corriger V3 : commercial_priority_tier: ctx.commercial_priority_tier || ctx.commercial_priority_default || ''",
                "V6 Build Expert Routing : ajouter support tier1 en plus de 'A'",
                "If Expert Notification Enabled : ajouter condition OR medium",
                "Si completion_percentage >= 95, la priorité est automatiquement 'high'",
            ]
        ),
        (
            "V6 exécution succeeded en 36ms sans traiter de données",
            [
                "Aucun formulaire complété depuis la dernière exécution",
                "Fetch Recent Completed Responses ne trouve rien dans la fenêtre de temps",
            ],
            [
                "Normal si aucune soumission récente",
                "Pour forcer le traitement : utiliser Manual Trigger avec Set Post-Audit Manual Input",
            ]
        ),
        (
            "If Pack Found bloque avec 'No output data'",
            [
                "Condition cassée : value1 equals value2 (ancien format n8n non migré)",
                "Le pack_id était undefined car Get Pack Row retournait vide",
                "Get Pack Row filtre par pack_id mais form_responses n'a pas ce champ",
            ],
            [
                "Corriger la condition : $json.found is true",
                "Si manuel : renseigner response_id dans Set Post-Audit Manual Input",
                "Get Latest Form Response : changer URL pour filtrer par id=eq.{response_id}",
            ]
        ),
        (
            "Get Latest Form Response retourne vide",
            [
                "L'URL utilisait form_responses?pack_id=eq.{pack_id}",
                "La table form_responses n'a pas de colonne pack_id (elle a invitation_token)",
            ],
            [
                "Changer l'URL en :",
                "  form_responses?id=eq.{encodeURIComponent($json.response_id)}&select=*&order=submitted_at.desc.nullslast&limit=1",
            ]
        ),
        (
            "Lien Calendly apparaît deux fois dans l'email",
            [
                "Le noeud Send Internal Brief Email utilise du Markdown dans un champ HTML",
                "[url](url) est rendu littéralement au lieu d'être converti en lien HTML",
            ],
            [
                "Dans le noeud Send Internal Brief Email, remplacer :",
                "  [https://calendly.com/...](https://calendly.com/...)",
                "  par : <a href='https://calendly.com/contact-transferai/30min'>Prendre rendez-vous</a>",
            ]
        ),
        (
            "Extract Prospect Target Row bloque (No output data)",
            [
                "Le prospect_id dans le pack est 'manual-prospect-001' qui n'existe pas dans prospect_targets",
                "Get Prospect Target Row ne trouve aucune entrée",
            ],
            [
                "Activer 'Always Output Data' sur le noeud Extract Prospect Target Row",
                "Le workflow continuera avec un prospect vide (non bloquant)",
            ]
        ),
    ]

    for title_text, causes, solutions in v6_issues:
        add_heading(doc, title_text, 2)
        add_body(doc, "Causes :")
        for c in causes:
            add_bullet(doc, c, level=0)
        add_body(doc, "Solutions :")
        for s in solutions:
            add_bullet(doc, s, level=0)
        doc.add_paragraph()

    add_heading(doc, "3. Edge Functions Supabase", 1)

    add_heading(doc, "Redéployer une Edge Function", 2)
    add_code(doc, "cd /Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
    add_code(doc, "supabase functions deploy resolve-invitation --no-verify-jwt")
    add_code(doc, "supabase functions deploy save-form-response --no-verify-jwt")

    add_heading(doc, "Vérifier la version déployée", 2)
    add_body(doc, "Dans Supabase Dashboard → Edge Functions → resolve-invitation → vérifier la date de déploiement. Si la date est antérieure aux corrections, redéployer.")

    add_heading(doc, "4. Supabase SQL — Requêtes utiles", 1)

    sql_queries = [
        ("Vérifier une invitation", "SELECT * FROM form_invitations WHERE pack_id = 'PACK_ID';"),
        ("Vérifier une soumission", "SELECT fr.* FROM form_responses fr JOIN form_invitations fi ON fi.invite_token = fr.invitation_token WHERE fi.pack_id = 'PACK_ID' ORDER BY fr.submitted_at DESC LIMIT 5;"),
        ("Réinitialiser un test", "DELETE FROM form_responses WHERE invitation_token IN (SELECT invite_token FROM form_invitations WHERE pack_id = 'PACK_ID');\nDELETE FROM form_invitations WHERE pack_id = 'PACK_ID';"),
        ("Vérifier le commercial_priority_tier", "SELECT pack_id, organization_name, payload->>'commercial_priority_tier' as tier FROM ai_prospecting_packs ORDER BY created_at DESC LIMIT 10;"),
    ]

    for label, sql in sql_queries:
        add_heading(doc, label, 2)
        add_code(doc, sql)

    add_heading(doc, "5. Checklist de vérification avant test", 1)
    checklist = [
        "Edge Function resolve-invitation déployée (version récente)",
        "Pack existe dans ai_prospecting_packs avec organization_name renseigné",
        "Invitation réinitialisée si re-test (SQL DELETE)",
        "V3 publié avec correction commercial_priority_tier",
        "V6 publié avec corrections Build Expert Routing + If Expert Notification Enabled",
        "V6 noeuds If Pack Found / If Audit Completed / If Already Post-Audit Processed / If Slack Alerts Enabled corrigés et sauvegardés",
        "Set Post-Audit Manual Input renseigné pour test manuel",
    ]
    for item in checklist:
        add_bullet(doc, "☐ " + item)

    out_path = os.path.join(OUTPUT_DIR, "106_Guide_Troubleshooting_Session_2026-06-28_Audit_V6.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    p1 = build_user_guide()
    p2 = build_troubleshooting_guide()
    print("\n✓ Documents générés :")
    print(f"  {p1}")
    print(f"  {p2}")
