from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
SOURCE_JSON = DOCS_DIR / "88_n8n_Admin_Prospection_Controller_V1_Zoho_First_Exportable.json"
TODAY = date.today().isoformat()
OUTPUT = WORD_DIR / f"89_Guide_Installation_Controller_Zoho_First_Node_Par_Node_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B)),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_after = Pt(4)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.03
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.03
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].width = Inches(widths[idx])
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=95, start=110, bottom=95, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.02
                for run in p.runs:
                    run.font.size = Pt(9.0)


def add_callout(document: Document, text: str, fill: str = "EAF2F8") -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text).font.size = Pt(10.0)
    document.add_paragraph("")


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TransferAI | Guide d'installation n8n Zoho-first | 11 juin 2026")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def load_workflow() -> dict:
    return json.loads(SOURCE_JSON.read_text(encoding="utf-8"))


def env_var_rows() -> list[list[str]]:
    return [
        ["ADMIN_BACKOFFICE_TOKEN", "Oui", "Token Bearer controle par le tronc commun du controller.", "Token long unique reserve au back-office admin."],
        ["SUPABASE_URL", "Oui", "Base URL de l'API REST Supabase pour lire et mettre a jour packs, prospects et logs.", "https://<project>.supabase.co"],
        ["SUPABASE_SERVICE_ROLE_KEY", "Oui", "Cle de service pour les requetes HTTP Request sur ai_prospecting_packs, prospect_targets et outreach_attempts.", "service_role ..."],
        ["PROSPECTING_V3_WORKFLOW_ID", "Oui", "Identifiant du Workflow V3 principal appele par Execute Workflow.", "Workflow ID n8n de la V3 CRM ready."],
        ["N8N_BASE_URL", "Oui", "Base URL publique n8n utilisee pour les callbacks legacy approve/reject.", "https://votre-instance-n8n"],
        ["OUTREACH_FROM_EMAIL", "Oui", "Adresse d'expedition par defaut utilisee par les noeuds Zoho SMTP.", "contact@transferai.ci"],
        ["BOOKING_LINK_45MIN", "Oui", "Lien de rendez-vous injecte dans le corps du courrier si absent du pack.", "https://calendly.com/contact-transferai/30min"],
    ]


def credential_rows() -> list[list[str]]:
    return [
        ["TransferAI Zoho SMTP", "SMTP", "Credential utilise par Send Zoho SMTP Email et Send Zoho Test Email."],
        ["Host SMTP Zoho", "Parametre", "Configuration standard Zoho Mail pro: smtppro.zoho.com."],
        ["Port", "Parametre", "465 en SSL ou 587 en TLS selon votre tenant et votre politique de securite."],
        ["Utilisateur", "Parametre", "Boite d'envoi TransferAI, par exemple contact@transferai.ci."],
        ["Mot de passe", "Parametre", "Mot de passe de la boite ou mot de passe applicatif si la double authentification est activee."],
    ]


def action_rows() -> list[list[str]]:
    return [
        ["generate_pack", "prospect_id, triggered_by, source, force", "Declenche V3 pour generer un pack prospect neuf."],
        ["regenerate_pack", "pack_id ou prospect_id, triggered_by, reason", "Relance V3 pour produire une nouvelle version du pack."],
        ["approve_pack", "pack_id, reviewer_email, mode", "Valide le pack; en mode approve_and_send la branche Zoho part ensuite."],
        ["reject_pack", "pack_id, reviewer_email, reason", "Refuse le pack et remet le prospect dans un etat corrigible."],
        ["send_pack", "pack_id, reviewer_email, force", "Envoie un pack deja pret vers Zoho SMTP."],
        ["test_provider", "provider=zoho, test_to_email", "Teste la configuration Zoho sans toucher a un vrai pack."],
        ["health_check", "aucun", "Controle la disponibilite du controller et l'acces Supabase."],
    ]


def node_groups() -> list[tuple[str, list[list[str]]]]:
    return [
        (
            "Tronc commun et securite",
            [
                ["Admin Controller Webhook", "Webhook", "Expose POST /webhook/admin-prospect-pack-action-zoho et recoit les actions admin.", "Path = admin-prospect-pack-action-zoho, Response mode = last node."],
                ["Normalize Request", "Code", "Uniformise body, headers, actor_email, request_id et payload.", "Produit action, payload, authorization, request_id, actor_email, received_at."],
                ["Validate Auth Header", "Code", "Compare le bearer token recu avec ADMIN_BACKOFFICE_TOKEN.", "Retourne auth_ok = true/false."],
                ["If Auth Valid", "If", "Route vers le flux normal ou l'erreur d'authentification.", "Condition bool sur auth_ok."],
                ["Build Unauthorized Response", "Code", "Construit la reponse JSON standard 401 logique.", "Status = unauthorized, code = UNAUTHORIZED."],
                ["Validate Action", "Code", "Verifie que l'action demandee fait partie de la liste supportee.", "Supporte generate, regenerate, approve, reject, send, test, health."],
                ["If Action Supported", "If", "Laisse passer ou renvoie une erreur d'action non supportee.", "Condition sur action_valid."],
                ["Build Unsupported Action Response", "Code", "Reponse standard pour une action inconnue.", "Status = invalid_action."],
            ],
        ),
        (
            "Branche generate_pack",
            [
                ["If Generate Action", "If", "Identifie l'action generate_pack.", "value2 = generate_pack."],
                ["Validate Generate Payload", "Code", "Controle la presence de prospect_id et des metadonnees minimales.", "Retourne payload_valid et erreurs."],
                ["If Generate Payload Valid", "If", "Ouvre la voie vers l'appel V3 ou retourne l'erreur de payload.", "Condition bool sur payload_valid."],
                ["Build Generate Payload Error", "Code", "Retourne une reponse propre pour un payload incomplet.", "Status = invalid_payload."],
                ["Call Workflow V3 Generate", "Execute Workflow", "Reutilise le Workflow V3 principal via PROSPECTING_V3_WORKFLOW_ID.", "Le payload normalise est transmis a la V3."],
                ["Build Generate Response", "Code", "Normalise la reponse envoyee au frontend apres lancement.", "Renvoie ok, action, pack_id/prospect_id, execution_id, refresh_hint."],
            ],
        ),
        (
            "Branche regenerate_pack",
            [
                ["If Regenerate Action", "If", "Identifie l'action regenerate_pack.", "value2 = regenerate_pack."],
                ["Validate Regenerate Payload", "Code", "Accepte pack_id ou prospect_id, plus reason et actor.", "Controle les entrees minimales."],
                ["If Regenerate Payload Valid", "If", "Route vers V3 ou l'erreur de validation.", "Condition bool sur payload_valid."],
                ["Build Regenerate Payload Error", "Code", "Construit la reponse d'erreur pour une regeneration mal renseignee.", "Status = invalid_payload."],
                ["Call Workflow V3 Regenerate", "Execute Workflow", "Relance la V3 avec le contexte de regeneration.", "Appel via PROSPECTING_V3_WORKFLOW_ID."],
                ["Build Regenerate Response", "Code", "Retour standardise pour le back-office.", "Status = accepted."],
            ],
        ),
        (
            "Branche approve_pack",
            [
                ["If Approve Action", "If", "Identifie l'action approve_pack.", "value2 = approve_pack."],
                ["Validate Approve Payload", "Code", "Controle pack_id, reviewer_email et mode.", "Modes attendus: approve_only ou approve_and_send."],
                ["If Approve Payload Valid", "If", "Passe au flux d'approbation ou retourne l'erreur.", "Condition bool sur payload_valid."],
                ["Build Approve Payload Error", "Code", "Reponse standard si le payload d'approbation est incomplet.", "Status = invalid_payload."],
                ["If Approve And Send Mode", "If", "Distingue approbation simple et approbation suivie d'envoi.", "Condition sur mode = approve_and_send."],
                ["Mark Pack Approved Direct", "HTTP Request", "Met a jour le statut du pack dans Supabase sans envoi.", "PATCH sur ai_prospecting_packs status = approved."],
                ["Build Approve Only Response", "Code", "Retourne la reponse a afficher si le pack est seulement valide.", "Status = approved."],
            ],
        ),
        (
            "Branche reject_pack",
            [
                ["If Reject Action", "If", "Identifie l'action reject_pack.", "value2 = reject_pack."],
                ["Validate Reject Payload", "Code", "Controle pack_id, reviewer_email et reason.", "Retourne payload_valid."],
                ["If Reject Payload Valid", "If", "Autorise l'appel de rejet ou retourne l'erreur.", "Condition bool sur payload_valid."],
                ["Build Reject Payload Error", "Code", "Reponse si les champs de rejet sont incomplets.", "Status = invalid_payload."],
                ["Call Legacy Reject Webhook", "HTTP Request", "Reutilise l'endpoint V3 legacy approve-prospect-pack-v3 avec decision = rejected.", "Base URL issue de N8N_BASE_URL."],
                ["Build Reject Response", "Code", "Normalise la confirmation de rejet.", "Status = rejected."],
            ],
        ),
        (
            "Branche send_pack Zoho-first",
            [
                ["If Send Action", "If", "Identifie l'action send_pack.", "value2 = send_pack."],
                ["Validate Send Payload", "Code", "Controle pack_id, reviewer_email et force.", "Retourne payload_valid."],
                ["If Send Payload Valid", "If", "Oriente vers la branche d'envoi ou l'erreur de validation.", "Condition bool sur payload_valid."],
                ["Build Send Payload Error", "Code", "Retourne l'erreur standard si pack_id est absent.", "Status = invalid_payload."],
                ["Get Pack For Zoho Send", "HTTP Request", "Lit ai_prospecting_packs dans Supabase via service role.", "GET cible sur le pack demande."],
                ["Extract Pack For Zoho Send", "Code", "Isole la ligne pack_row et ses artefacts utiles.", "Normalise la lecture Supabase."],
                ["If Pack For Zoho Send Found", "If", "Verifie qu'un pack a bien ete trouve.", "Sinon branche not found."],
                ["Build Pack Not Found For Send Response", "Code", "Reponse propre si le pack n'existe pas.", "Status = not_found."],
                ["Build Zoho Send Context", "Code", "Reconstitue target_email, HTML, URLs PDF/PPTX, attachments_count, can_send et send_failure_reason.", "Utilise BOOKING_LINK_45MIN et OUTREACH_FROM_EMAIL si necessaire."],
                ["If Zoho Send Ready", "If", "Verifie can_send avant de telecharger et d'envoyer.", "Condition bool sur can_send."],
                ["Download Zoho Catalogue PDF", "HTTP Request", "Telecharge le mini catalogue depuis l'URL PDF du pack.", "ResponseFormat = file, outputPropertyName = catalogue_pdf."],
                ["Download Zoho Deck PPTX", "HTTP Request", "Telecharge le deck PPTX depuis l'URL du pack.", "ResponseFormat = file, outputPropertyName = deck_pptx."],
                ["Merge Zoho Files", "Merge", "Assemble les deux flux binaires.", "Mode = combineByPosition."],
                ["Assemble Zoho Email Payload", "Code", "Fusionne le contexte JSON et les binaires pour l'e-mail final.", "Produit json + binary."],
                ["Send Zoho SMTP Email", "Email Send", "Expedie le mail via le credential SMTP TransferAI Zoho SMTP.", "From = OUTREACH_FROM_EMAIL, To = target_email, Subject = subject."],
                ["Build Zoho Send Result", "Code", "Capture l'etat d'envoi avant ecriture en base.", "Prepare status sent et meta provider."],
                ["Mark Pack Approved For Zoho", "HTTP Request", "Passe le pack au statut approved avant la marque sent.", "PATCH Supabase."],
                ["Mark Pack Sent Zoho", "HTTP Request", "Marque ai_prospecting_packs comme sent avec sent_at.", "PATCH Supabase."],
                ["Log Outreach Attempt Zoho", "HTTP Request", "Cree la ligne outreach_attempts de l'envoi reel.", "INSERT Supabase."],
                ["Update Prospect Target Sent Zoho", "HTTP Request", "Met prospect_targets a active / outreach_started.", "PATCH Supabase."],
                ["Build Zoho Final Response", "Code", "Retourne la reponse finale au back-office apres envoi.", "Status = sent."],
                ["Mark Pack Approval Error Zoho", "HTTP Request", "Marque le pack en approval_error si can_send = false ou artefact manquant.", "PATCH Supabase."],
                ["Update Prospect Target Approval Error Zoho", "HTTP Request", "Met le prospect en paused / internal_fix_required.", "PATCH Supabase."],
                ["Build Zoho Send Error Response", "Code", "Retourne la raison detaillee qui bloque l'envoi.", "Status = approval_error."],
            ],
        ),
        (
            "Branche test_provider et health_check",
            [
                ["If Test Provider Action", "If", "Identifie l'action test_provider.", "value2 = test_provider."],
                ["Validate Test Provider Payload", "Code", "Controle provider = zoho et test_to_email.", "Retourne payload_valid."],
                ["If Test Provider Payload Valid", "If", "Oriente vers l'e-mail de test ou l'erreur.", "Condition bool sur payload_valid."],
                ["Build Test Provider Payload Error", "Code", "Reponse standard si le test provider est mal renseigne.", "Status = invalid_payload."],
                ["Prepare Zoho Test Email", "Code", "Construit un petit sujet et message HTML de verification.", "Inclut actor_email et horodatage."],
                ["Send Zoho Test Email", "Email Send", "Envoie un e-mail de test via le credential Zoho SMTP.", "Permet de valider la boite avant usage CRM."],
                ["Build Test Provider Response", "Code", "Retourne une confirmation simple au back-office.", "Status = provider_test_sent."],
                ["If Health Action", "If", "Identifie l'action health_check.", "value2 = health_check."],
                ["Check Supabase Health", "HTTP Request", "Teste l'acces a Supabase via une requete legere.", "GET REST avec service role."],
                ["Build Health Response", "Code", "Retourne l'etat du controller et de la connectivite.", "Status = healthy."],
                ["Build No Route Response", "Code", "Reponse de secours si aucune branche ne matche.", "Status = no_route."],
            ],
        ),
    ]


def supabase_rows() -> list[list[str]]:
    return [
        ["ai_prospecting_packs", "Get Pack For Zoho Send / Mark Pack Approved Direct / Mark Pack Approved For Zoho / Mark Pack Sent Zoho / Mark Pack Approval Error Zoho", "Lecture du pack, validation, marquage sent ou approval_error."],
        ["prospect_targets", "Update Prospect Target Sent Zoho / branche legacy reject", "Synchronisation du statut commercial du prospect."],
        ["outreach_attempts", "Log Outreach Attempt Zoho", "Journalisation de l'envoi effectif."],
    ]


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guide d'installation node-by-node")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Workflow n8n Controller Zoho-first pour Prospection Packs")
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4B, 0x5D, 0x6B)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version générée le {TODAY}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_callout(
        document,
        "Ce guide décrit la configuration étape par étape du controller Zoho-first livré dans le fichier "
        "88_n8n_Admin_Prospection_Controller_V1_Zoho_First_Exportable.json. "
        "Il couvre le webhook d'entrée, les branches métier, la branche d'envoi Zoho SMTP, "
        "les écritures Supabase et la liste complète des variables d'environnement n8n à renseigner.",
    )


def build_document() -> Document:
    workflow = load_workflow()
    document = Document()
    style_document(document)
    add_cover_page(document)
    add_footer(document)

    add_heading(document, "1. Objet et périmètre", 1)
    add_paragraph(
        document,
        "Ce workflow controller agit comme couche d'orchestration entre le futur back-office, la base Supabase, "
        "le Workflow V3 principal et le fournisseur d'envoi Zoho. Le frontend lit Supabase, puis appelle ce controller "
        "pour générer, régénérer, approuver, rejeter, tester ou envoyer un pack prospect.",
    )
    add_bullets(
        document,
        [
            "webhook principal : POST /webhook/admin-prospect-pack-action-zoho",
            "source JSON à importer : docs/transferai-admin/88_n8n_Admin_Prospection_Controller_V1_Zoho_First_Exportable.json",
            "moteur de génération réutilisé : Workflow V3 principal via PROSPECTING_V3_WORKFLOW_ID",
            "fournisseur d'envoi visé : Zoho SMTP via le credential TransferAI Zoho SMTP",
        ],
    )

    add_heading(document, "2. Pré-requis avant import", 1)
    add_numbered(
        document,
        [
            "Disposer d'une instance n8n accessible depuis le back-office.",
            "Avoir déjà importé ou identifié le Workflow V3 principal utilisé pour la génération des packs.",
            "Avoir un projet Supabase avec les tables prospect_targets, ai_prospecting_packs et outreach_attempts.",
            "Avoir une boîte Zoho prête à l'envoi et un credential SMTP créé dans n8n.",
            "Renseigner toutes les variables d'environnement listées dans la section suivante.",
        ],
    )

    add_heading(document, "3. Variables d'environnement n8n à renseigner", 1)
    add_table(document, ["Variable", "Obligatoire", "Utilité", "Exemple / valeur cible"], env_var_rows(), [2.0, 0.8, 2.45, 1.55])
    add_paragraph(
        document,
        "Ces sept variables sont les seules référencées directement par le workflow Zoho-first. "
        "La partie SMTP Zoho n'utilise pas d'autres variables d'environnement : ses secrets résident dans le credential n8n.",
    )

    add_heading(document, "4. Credentials et paramètres n8n à créer", 1)
    add_table(document, ["Élément", "Type", "Rôle"], credential_rows(), [1.7, 1.1, 3.7])
    add_callout(
        document,
        "Pour la configuration standard Zoho Mail pro, préparer un credential SMTP avec l'hôte smtppro.zoho.com, "
        "port 465 en SSL ou 587 en TLS, l'adresse de la boîte TransferAI et un mot de passe applicatif si la double authentification est activée.",
        fill="F4F8E8",
    )

    add_heading(document, "5. Actions supportées par le webhook", 1)
    add_table(document, ["Action", "Entrants minimum", "Résultat attendu"], action_rows(), [1.35, 2.65, 2.0])

    add_heading(document, "6. Vue d'ensemble des noeuds", 1)
    add_paragraph(
        document,
        f"Le workflow exporté contient actuellement {len(workflow['nodes'])} noeuds. "
        "La structure est volontairement séparée entre un tronc commun de contrôle, les branches generate/regenerate/approve/reject, "
        "la branche send_pack Zoho-first et les branches utilitaires test_provider et health_check.",
    )

    for idx, (group_title, rows) in enumerate(node_groups(), start=1):
        add_heading(document, f"6.{idx} {group_title}", 2)
        add_table(document, ["Nœud", "Type", "Rôle", "Configuration clé"], rows, [1.7, 1.0, 2.15, 1.95])

    add_heading(document, "7. Écritures Supabase réalisées par le controller", 1)
    add_table(document, ["Table", "Nœuds concernés", "But métier"], supabase_rows(), [1.45, 2.95, 2.25])

    add_heading(document, "8. Ordre d'installation recommandé dans n8n", 1)
    add_numbered(
        document,
        [
            "Importer le JSON 88 dans n8n.",
            "Créer le credential SMTP TransferAI Zoho SMTP et l'associer aux deux noeuds Email Send.",
            "Renseigner les sept variables d'environnement côté n8n.",
            "Vérifier que PROSPECTING_V3_WORKFLOW_ID pointe vers le bon Workflow V3 principal.",
            "Contrôler N8N_BASE_URL afin que la branche reject utilise bien l'URL publique correcte.",
            "Lancer un health_check depuis le back-office ou via Manual Trigger.",
            "Exécuter test_provider pour valider la boîte Zoho.",
            "Tester generate_pack avec un prospect de démonstration, puis send_pack sur un pack approuvé.",
            "Activer le workflow uniquement après validation de bout en bout.",
        ],
    )

    add_heading(document, "9. Payloads minimum à envoyer depuis le back-office", 1)
    add_paragraph(document, "Exemple generate_pack :")
    add_callout(
        document,
        '{ "action": "generate_pack", "payload": { "prospect_id": "fipme-001", "triggered_by": "marius@transferai.ci", "source": "backoffice", "force": false } }',
        fill="F7F7F7",
    )
    add_paragraph(document, "Exemple approve_pack en mode approve_and_send :")
    add_callout(
        document,
        '{ "action": "approve_pack", "payload": { "pack_id": "pack-abc123", "reviewer_email": "marius@transferai.ci", "mode": "approve_and_send", "notes": "ok pour envoi" } }',
        fill="F7F7F7",
    )
    add_paragraph(document, "Exemple send_pack :")
    add_callout(
        document,
        '{ "action": "send_pack", "payload": { "pack_id": "pack-abc123", "reviewer_email": "marius@transferai.ci", "force": false } }',
        fill="F7F7F7",
    )

    add_heading(document, "10. Check-list finale avant mise en production", 1)
    add_bullets(
        document,
        [
            "Le webhook admin répond avec un JSON cohérent aux actions health_check et test_provider.",
            "Le credential Zoho SMTP envoie un e-mail de test vers une boîte interne TransferAI.",
            "Le controller lit bien ai_prospecting_packs et prospect_targets via SUPABASE_SERVICE_ROLE_KEY.",
            "La branche send_pack télécharge réellement le PDF et le PPTX du pack.",
            "Après un envoi réussi, ai_prospecting_packs, outreach_attempts et prospect_targets sont mis à jour.",
            "Le back-office ne lit pas n8n pour les listes : il relit toujours Supabase après action.",
        ],
    )

    return document


def main() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
