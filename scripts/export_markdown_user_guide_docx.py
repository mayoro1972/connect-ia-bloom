from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


H1_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H2_COLOR = RGBColor(0x2E, 0x74, 0xB5)
H3_COLOR = RGBColor(0x1F, 0x4D, 0x78)
BODY_COLOR = RGBColor(0x10, 0x19, 0x28)
MUTED_COLOR = RGBColor(0x47, 0x55, 0x67)
RULE_COLOR = "D9E2F1"


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=BODY_COLOR):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE_COLOR)


def configure_document(document: Document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    nfmt = normal.paragraph_format
    nfmt.space_before = Pt(0)
    nfmt.space_after = Pt(6)
    nfmt.line_spacing = 1.25

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    h1 = document.styles["Heading 1"]
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = H1_COLOR
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.15

    h2 = document.styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = H2_COLOR
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.15

    h3 = document.styles["Heading 3"]
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = H3_COLOR
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.line_spacing = 1.15

    if "Code Block" not in document.styles:
        code_style = document.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(9.5)
        code_style.font.color.rgb = BODY_COLOR
        code_style.paragraph_format.space_before = Pt(4)
        code_style.paragraph_format.space_after = Pt(8)
        code_style.paragraph_format.line_spacing = 1.1
        code_style.paragraph_format.left_indent = Inches(0.25)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Guide utilisateur TransferAI - Prospection V3 CRM + Audit dynamique")
    set_run_font(footer_run, size=9, color=MUTED_COLOR)


def add_title_block(document: Document, title: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    set_run_font(run, size=22, bold=True, color=RGBColor(0x0B, 0x25, 0x45))
    set_paragraph_spacing(p, before=0, after=4, line=1.05)

    subtitle = document.add_paragraph()
    srun = subtitle.add_run("Référence opérationnelle n8n, Supabase, OpenAI, Resend et formulaire d’audit dynamique")
    set_run_font(srun, size=10.5, color=MUTED_COLOR)
    set_paragraph_spacing(subtitle, before=0, after=12, line=1.15)

    add_bottom_rule(subtitle)


def strip_link_markup(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def add_inline_runs(paragraph, text: str):
    text = strip_link_markup(text)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=RGBColor(0x11, 0x18, 0x27))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=11)


def add_bullet(document: Document, text: str):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph_spacing(p, before=0, after=4, line=1.25)
    bullet = p.add_run("• ")
    set_run_font(bullet, size=11)
    add_inline_runs(p, text)


def add_number(document: Document, number: str, text: str):
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    set_paragraph_spacing(p, before=0, after=4, line=1.25)
    prefix = p.add_run(f"{number} ")
    set_run_font(prefix, size=11, bold=True)
    add_inline_runs(p, text)


def add_paragraph(document: Document, text: str):
    p = document.add_paragraph(style="Normal")
    set_paragraph_spacing(p, before=0, after=6, line=1.25)
    add_inline_runs(p, text)


def add_code_block(document: Document, lines: list[str]):
    for idx, line in enumerate(lines):
        p = document.add_paragraph(style="Code Block")
        set_paragraph_spacing(p, before=4 if idx == 0 else 0, after=0 if idx < len(lines) - 1 else 8, line=1.1)
        run = p.add_run(line.rstrip("\n"))
        set_run_font(run, name="Consolas", size=9.5, color=RGBColor(0x11, 0x18, 0x27))


def export_markdown(md_path: Path, docx_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_document(document)

    first_heading_used = False
    in_code = False
    code_lines: list[str] = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(document, code_lines)
                in_code = False
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue

        if re.fullmatch(r"-{3,}", line.strip()):
            continue

        if line.startswith("# "):
            if not first_heading_used:
                add_title_block(document, line[2:].strip())
                first_heading_used = True
            else:
                p = document.add_paragraph(line[2:].strip(), style="Heading 1")
                set_paragraph_spacing(p, before=18, after=10, line=1.15)
            continue

        if line.startswith("## "):
            p = document.add_paragraph(line[3:].strip(), style="Heading 1")
            set_paragraph_spacing(p, before=18, after=10, line=1.15)
            continue

        if line.startswith("### "):
            p = document.add_paragraph(line[4:].strip(), style="Heading 2")
            set_paragraph_spacing(p, before=14, after=7, line=1.15)
            continue

        if line.startswith("#### "):
            p = document.add_paragraph(line[5:].strip(), style="Heading 3")
            set_paragraph_spacing(p, before=10, after=5, line=1.15)
            continue

        bullet_match = re.match(r"^\s*-\s+(.*)$", line)
        if bullet_match:
            add_bullet(document, bullet_match.group(1))
            continue

        number_match = re.match(r"^\s*(\d+\.)\s+(.*)$", line)
        if number_match:
            add_number(document, number_match.group(1), number_match.group(2))
            continue

        add_paragraph(document, line)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(docx_path)


def main():
    if len(sys.argv) != 3:
        raise SystemExit("Usage: export_markdown_user_guide_docx.py <input.md> <output.docx>")

    md_path = Path(sys.argv[1]).resolve()
    docx_path = Path(sys.argv[2]).resolve()
    export_markdown(md_path, docx_path)


if __name__ == "__main__":
    main()
