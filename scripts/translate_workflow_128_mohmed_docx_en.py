from pathlib import Path

from docx import Document


SOURCE = Path("/Users/marius_ayoro/Downloads/137_Guide_Utilisateur_Workflow_128_CR_Reunion_Mohmed.docx")
OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "140_User_Guide_Workflow_128_Meeting_Report_Mohmed_EN.docx"
)


PARAGRAPH_MAP = {
    "Guide utilisateur - Workflow 128": "User Guide - Workflow 128",
    "Compte rendu de réunion international, pseudonymisation locale et livrables assainis": "International meeting report, local pseudonymization, and sanitized deliverables",
    "1. Objet du document": "1. Purpose of this document",
    "Ce guide explique comment utiliser et configurer le workflow 128 pour traiter des enregistrements de réunions sensibles, produire des comptes rendus structurés et générer des livrables externes assainis.": "This guide explains how to use and configure workflow 128 to process sensitive meeting recordings, produce structured meeting reports, and generate sanitized external deliverables.",
    "2. Cas d’usage recommandé": "2. Recommended use case",
    "Le client enregistre ses réunions de cadrage, de suivi ou de restitution via Zoom, Webex, Microsoft Teams, un dictaphone, un smartphone ou un dispositif de salle.": "The client records its scoping, follow-up, or debrief meetings through Zoom, Webex, Microsoft Teams, a dictaphone, a smartphone, or a meeting room device.",
    "L’objectif est d’obtenir rapidement un compte rendu exploitable, un e-mail de suivi client et une version assainie, tout en protégeant les données sensibles avant appel à l’IA.": "The goal is to obtain quickly a usable meeting report, a client follow-up email, and a sanitized version, while protecting sensitive data before any AI call.",
    "Réunion de cadrage : Programme Formation IA Secrétariat Direction": "Scoping meeting: AI Training Programme Executive Secretariat",
    "Source d’enregistrement : Teams, Zoom, Webex, dictaphone ou smartphone": "Recording source: Teams, Zoom, Webex, dictaphone, or smartphone",
    "Sortie attendue : compte rendu, e-mail client, archive assainie": "Expected output: meeting report, client email, sanitized archive",
    "3. Architecture du workflow": "3. Workflow architecture",
    "Entrée du recording": "Recording intake",
    "Normalisation des métadonnées": "Metadata normalization",
    "Récupération ou réception de l’audio": "Audio retrieval or reception",
    "Découpage audio": "Audio splitting",
    "Transcription OpenAI": "OpenAI transcription",
    "Pseudonymisation locale": "Local pseudonymization",
    "Génération du rapport assaini": "Sanitized report generation",
    "Validation humaine puis envoi final": "Human validation and final delivery",
    "4. Modes d’entrée supportés": "4. Supported input modes",
    "4.1 Upload manuel": "4.1 Manual upload",
    "À utiliser pour les fichiers .m4a, .mp3, les exports smartphone et les exports dictaphone.": "Use this for .m4a files, .mp3 files, smartphone exports, and dictaphone exports.",
    "4.2 Google Drive": "4.2 Google Drive",
    "À utiliser si les fichiers sont déposés dans un dossier surveillé ou si un lien Drive est fourni dans le formulaire.": "Use this when files are uploaded into a watched folder or when a Drive link is provided in the form.",
    "4.3 Webhook plateforme meeting": "4.3 Meeting platform webhook",
    "À utiliser si un workflow amont Zoom, Teams ou Webex récupère le recording puis appelle le webhook d’ingestion du workflow 128.": "Use this when an upstream Zoom, Teams, or Webex workflow retrieves the recording and then calls the ingestion webhook of workflow 128.",
    "5. Pré-requis techniques": "5. Technical prerequisites",
    "6. Configuration pas à pas": "6. Step-by-step configuration",
    "7. Paramétrage conseillé": "7. Recommended settings",
    "Livrable demandé : Les deux": "Requested deliverable: Both",
    "Niveau de masquage : Strict": "Masking level: Strict",
    "Langue de sortie : Français et Anglais": "Output language: French and English",
    "Dossier source : Recordings": "Source folder: Recordings",
    "Dossier archive : Archives assainies": "Archive folder: Sanitized archives",
    "8. Configuration des sources de réunion": "8. Meeting source configuration",
    "8.1 Dictaphone": "8.1 Dictaphone",
    "Mode simple : exporter le fichier puis l’envoyer via le formulaire.": "Simple mode: export the file and upload it through the form.",
    "Mode semi-automatique : synchroniser les fichiers vers un dossier Google Drive surveillé.": "Semi-automated mode: sync the files to a watched Google Drive folder.",
    "8.2 Zoom": "8.2 Zoom",
    "Activer Cloud Recording dans Zoom.": "Enable Cloud Recording in Zoom.",
    "Utiliser un workflow amont recevant l’événement recording.completed.": "Use an upstream workflow that receives the recording.completed event.",
    "Télécharger le recording puis appeler le webhook cr-intl-source-ingest.": "Download the recording, then call the cr-intl-source-ingest webhook.",
    "8.3 Microsoft Teams": "8.3 Microsoft Teams",
    "Récupérer les enregistrements dans OneDrive ou SharePoint.": "Retrieve recordings from OneDrive or SharePoint.",
    "Passer par Microsoft Graph ou un workflow de collecte amont.": "Use Microsoft Graph or an upstream collection workflow.",
    "Transmettre ensuite le binaire ou les métadonnées au workflow 128.": "Then pass the binary file or the metadata to workflow 128.",
    "8.4 Webex": "8.4 Webex",
    "Activer les recordings Webex.": "Enable Webex recordings.",
    "Utiliser l’API Recordings pour récupérer le fichier.": "Use the Recordings API to retrieve the file.",
    "Faire suivre le recording vers le webhook d’ingestion du workflow 128.": "Forward the recording to the ingestion webhook of workflow 128.",
    "9. Structure attendue pour le webhook d’ingestion": "9. Expected ingestion webhook structure",
    "Le webhook cr-intl-source-ingest doit idéalement recevoir : source_system, source_label, meeting_title, meeting_date, participants, e-mail destinataire, e-mail validateur, langue, niveau de masquage, référence de réunion, objet source et recording_url. Si possible, transmettre directement le binaire audio.": "The cr-intl-source-ingest webhook should ideally receive: source_system, source_label, meeting_title, meeting_date, participants, recipient email, validator email, language, masking level, meeting reference, source object, and recording_url. When possible, send the audio binary directly.",
    "10. Checklist de mise en production": "10. Go-live checklist",
    "Le workflow s’importe sans erreur.": "The workflow imports without error.",
    "Google Drive est connecté.": "Google Drive is connected.",
    "Resend est connecté.": "Resend is connected.",
    "OpenAI est connecté.": "OpenAI is connected.",
    "Le service de découpage audio répond.": "The audio splitting service responds.",
    "Le validateur reçoit l’aperçu.": "The validator receives the preview.",
    "L’archive HTML est créée.": "The HTML archive is created.",
    "Aucune version réidentifiée n’est envoyée au client.": "No re-identified version is sent to the client.",
    "11. Recommandation finale": "11. Final recommendation",
    "Pour le test, il est recommandé de démarrer avec l’upload manuel et Google Drive, de valider le processus métier et la qualité des rapports, puis d’ajouter dans un second temps des workflows d’ingestion dédiés à Zoom, Teams et Webex.": "For testing, the recommended starting point is manual upload and Google Drive. Once the business process and report quality are validated, add dedicated ingestion workflows for Zoom, Teams, and Webex.",
}


TABLE_MAP = {
    "Élément": "Item",
    "Configuration": "Configuration",
    "Clé OpenAI": "OpenAI key",
    "Variable d’environnement OPENAI_API_KEY": "Environment variable OPENAI_API_KEY",
    "Clé Resend": "Resend key",
    "Variable d’environnement RESEND_API_KEY": "Environment variable RESEND_API_KEY",
    "Modèle rapport": "Report model",
    "gpt-5 recommandé ; sinon gpt-4.1 ou gpt-4o": "gpt-5 recommended; otherwise gpt-4.1 or gpt-4o",
    "Google Drive": "Google Drive",
    "Credential OAuth actif pour lecture et archivage": "Active OAuth credential for reading and archiving",
    "Découpage audio": "Audio splitting",
    "Service disponible sur http://n8n-pxlk-audio-splitter-1:8000/split": "Service available at http://n8n-pxlk-audio-splitter-1:8000/split",
    "Domaine webhook": "Webhook domain",
    "Étape 1 - Importer le workflow": "Step 1 - Import the workflow",
    "Dans n8n, ouvrir Workflows, choisir Import from file, puis sélectionner le fichier JSON du workflow 128.": "In n8n, open Workflows, choose Import from file, then select the JSON file for workflow 128.",
    "Étape 2 - Vérifier les points d’entrée": "Step 2 - Check the entry points",
    "Contrôler les nœuds Formulaire - Audio sensible, Google Drive - Nouveau fichier audio et Webhook - Ingestion plateforme meeting.": "Review the nodes Formulaire - Audio sensible, Google Drive - Nouveau fichier audio, and Webhook - Ingestion plateforme meeting.",
    "Étape 3 - Configurer le formulaire": "Step 3 - Configure the form",
    "Renseigner les champs sujet, participants, date, source d’enregistrement, mode d’ingestion, destinataire final, validateur, langue audio, langue de sortie, livrable demandé, entités sensibles à masquer et niveau de masquage.": "Fill in the fields for subject, participants, date, recording source, ingestion mode, final recipient, validator, audio language, output language, requested deliverable, sensitive entities to mask, and masking level.",
    "Étape 4 - Configurer le dossier source": "Step 4 - Configure the source folder",
    "Vérifier l’ID du dossier surveillé dans le nœud Google Drive - Nouveau fichier audio.": "Check the watched folder ID in the node Google Drive - Nouveau fichier audio.",
    "Étape 5 - Configurer le dossier d’archive": "Step 5 - Configure the archive folder",
    "Vérifier l’ID du dossier d’archivage dans le nœud Archiver version assainie.": "Check the archive folder ID in the node Archiver version assainie.",
    "Étape 6 - Configurer Resend": "Step 6 - Configure Resend",
    "Valider les nœuds Envoyer au validateur et Envoyer le livrable final assaini, ainsi que l’adresse expéditeur autorisée.": "Validate the nodes Envoyer au validateur and Envoyer le livrable final assaini, as well as the authorized sender address.",
    "Étape 7 - Configurer OpenAI": "Step 7 - Configure OpenAI",
    "Vérifier les nœuds OpenAI - Transcrire segment et OpenAI - Rapport structuré assaini.": "Check the nodes OpenAI - Transcrire segment and OpenAI - Rapport structure sanitise.",
    "Étape 8 - Vérifier les webhooks de validation": "Step 8 - Verify the validation webhooks",
    "Contrôler les nœuds Webhook - Approuver et Webhook - Rejeter, ainsi que le domaine de production.": "Review the nodes Webhook - Approuver and Webhook - Rejeter, together with the production domain.",
    "Étape 9 - Lancer un test pilote": "Step 9 - Run a pilot test",
    "Faire un test avec un petit enregistrement, approuver le livrable, puis vérifier l’archive HTML.": "Test with a short recording, approve the deliverable, then verify the archived HTML output.",
}


HEADER_MAP = {
    "TransferAI - Compte rendu de réunion sensible": "TransferAI - Sensitive meeting report",
}


FOOTER_MAP = {
    "Guide utilisateur Workflow 128 - TransferAI": "Workflow 128 User Guide - TransferAI",
}


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def translate_paragraphs(paragraphs, mapping):
    for paragraph in paragraphs:
        source = paragraph.text.strip()
        if not source:
            continue
        translated = mapping.get(source)
        if translated:
            replace_paragraph_text(paragraph, translated)


def translate_table_cells(tables, mapping):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    source = paragraph.text.strip()
                    if not source:
                        continue
                    translated = mapping.get(source)
                    if translated:
                        replace_paragraph_text(paragraph, translated)


def main():
    doc = Document(SOURCE)
    translate_paragraphs(doc.paragraphs, PARAGRAPH_MAP)
    translate_table_cells(doc.tables, TABLE_MAP)

    for section in doc.sections:
        translate_paragraphs(section.header.paragraphs, HEADER_MAP)
        translate_paragraphs(section.footer.paragraphs, FOOTER_MAP)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
