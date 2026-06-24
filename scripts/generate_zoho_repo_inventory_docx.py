from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORD_DIR = ROOT / "docs" / "transferai-admin" / "word"
TODAY = date.today().isoformat()
OUTPUT = WORD_DIR / f"90_Inventory_Zoho_Email_TransferAI_Repo_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.space_after = Pt(4)


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].width = Inches(widths[idx])
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=95, start=110, bottom=95, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].width = Inches(widths[idx])
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.0)


def add_callout(document: Document, text: str, fill: str = "EAF2F8") -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=170, bottom=140, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run(text).font.size = Pt(10.0)
    document.add_paragraph("")


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TransferAI | Inventaire Zoho Email repo | 11 juin 2026")
    run.font.name = "Aptos"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_document() -> Document:
    document = Document()
    style_document(document)
    add_footer(document)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Inventaire Zoho Email - TransferAI")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Document généré le {TODAY}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    add_callout(
        document,
        "Ce document rassemble les informations retrouvées dans le dépôt connect-ia-bloom concernant l'usage de l'adresse TransferAI, "
        "la future configuration Zoho SMTP, les composants déjà branchés sur Resend et les éléments non encore présents dans le repo.",
    )

    add_heading(document, "1. Synthèse rapide", 1)
    add_bullets(
        document,
        [
            "Le code actif du site et des fonctions Supabase utilise encore principalement Resend.",
            "L'adresse e-mail métier la plus présente dans le dépôt est contact@transferai.ci.",
            "La future intégration Zoho est déjà modélisée dans le workflow n8n Zoho-first.",
            "Le credential n8n attendu pour Zoho s'appelle TransferAI Zoho SMTP.",
            "Aucun identifiant Zoho, mot de passe SMTP, token ZeptoMail ou variable ZOHO_* n'a été trouvé dans le dépôt.",
        ],
    )

    add_heading(document, "2. Informations retrouvées et jugées fiables", 1)
    add_table(
        document,
        ["Élément", "Valeur / constat", "Source principale"],
        [
            ["Adresse e-mail principale", "contact@transferai.ci", "supabase/functions/send-prospect-emails/index.ts"],
            ["Format sender utilisé", "TransferAI Africa <contact@transferai.ci>", "supabase/functions/send-prospect-emails/index.ts"],
            ["Credential n8n attendu", "TransferAI Zoho SMTP", "docs/transferai-admin/88_n8n_Admin_Prospection_Controller_V1_Zoho_First_Exportable.json"],
            ["Hôte SMTP Zoho documenté", "smtppro.zoho.com", "docs/transferai-admin/88...json + scripts/generate_zoho_first_controller_docx.py"],
            ["Ports SMTP documentés", "465 SSL ou 587 TLS", "docs/transferai-admin/88...json + scripts/generate_zoho_first_controller_docx.py"],
            ["Adresse from par défaut côté n8n Zoho-first", "OUTREACH_FROM_EMAIL ou contact@transferai.ci", "docs/transferai-admin/88...json"],
        ],
        [1.8, 2.4, 2.2],
    )

    add_heading(document, "3. Endroits du dépôt où l'e-mail TransferAI est utilisé aujourd'hui", 1)
    add_table(
        document,
        ["Zone", "Constat", "Fichier"],
        [
            ["Fonctions Supabase", "MAIL_FROM et MAIL_TO par défaut pointent vers contact@transferai.ci.", "supabase/functions/send-prospect-emails/index.ts"],
            ["Fonctions Supabase", "FROM et ADMIN_EMAIL utilisent contact@transferai.ci.", "supabase/functions/webinar-admin-weekly-digest/index.ts"],
            ["Fonctions Supabase", "Diverses fonctions newsletters et follow-up réutilisent contact@transferai.ci.", "supabase/functions/newsletter-subscribe/index.ts, process-prospect-followups/index.ts, partner-followup-send/index.ts"],
            ["Frontend", "Les pages Privacy et ProspectRequest mentionnent contact@transferai.ci.", "src/pages/Privacy.tsx, src/pages/ProspectRequestPage.tsx"],
            ["Supports commerciaux", "Les decks, lettres et modèles de prospection affichent contact@transferai.ci.", "outputs/... et docs/transferai-prospection/..."],
        ],
        [1.45, 2.95, 2.0],
    )

    add_heading(document, "4. Ce que le dépôt dit déjà sur Zoho", 1)
    add_table(
        document,
        ["Point Zoho", "Détail", "Fichier"],
        [
            ["Workflow exportable", "Le controller Zoho-first est déjà généré et prêt à importer dans n8n.", "docs/transferai-admin/88_n8n_Admin_Prospection_Controller_V1_Zoho_First_Exportable.json"],
            ["Nœud d'envoi", "Send Zoho SMTP Email utilise le credential TransferAI Zoho SMTP.", "docs/transferai-admin/88...json"],
            ["Nœud de test", "Send Zoho Test Email permet de valider la boîte avant usage CRM.", "docs/transferai-admin/88...json"],
            ["Documentation", "Le guide node-by-node rappelle que les secrets Zoho ne vivent pas dans les envs du repo.", "scripts/generate_zoho_first_controller_docx.py"],
            ["Blueprint", "Le contrat backend et le blueprint utilisent provider_key = zoho-primary.", "docs/transferai-admin/84...md et 86...md"],
        ],
        [1.5, 3.0, 1.9],
    )

    add_heading(document, "5. Ce que je n'ai pas trouvé dans le repo", 1)
    add_bullets(
        document,
        [
            "aucune variable d'environnement ZOHO_* dans .env ou .env.example",
            "aucun mot de passe SMTP Zoho",
            "aucun token API ZeptoMail",
            "aucun login utilisateur Zoho stocké dans les sources",
            "aucune configuration MX/DNS Zoho exportée dans le dépôt",
        ],
    )

    add_heading(document, "6. Conclusion opérationnelle", 1)
    add_paragraph(
        document,
        "Le dépôt montre clairement que votre adresse métier de référence est contact@transferai.ci et que la cible d'évolution côté prospection est une branche d'envoi Zoho SMTP dans n8n. "
        "En revanche, les vrais identifiants Zoho ne sont pas présents dans le repo. La prochaine étape pratique consiste donc à récupérer les paramètres de votre boîte Zoho côté admin (email exact, mot de passe applicatif si nécessaire, politique SSL/TLS), puis à les saisir dans le credential n8n TransferAI Zoho SMTP.",
    )

    add_heading(document, "7. Prochaine étape recommandée", 1)
    add_bullets(
        document,
        [
            "confirmer l'adresse exacte de la boîte Zoho à utiliser pour l'envoi CRM",
            "créer ou vérifier le credential SMTP TransferAI Zoho SMTP dans n8n",
            "renseigner OUTREACH_FROM_EMAIL avec l'adresse d'envoi retenue",
            "exécuter test_provider dans le workflow Zoho-first",
        ],
    )

    return document


def main() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
