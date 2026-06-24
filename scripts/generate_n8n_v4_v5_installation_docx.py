from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "transferai-admin"
WORD_DIR = DOCS_DIR / "word"
TODAY = date.today().isoformat()

V4_JSON = DOCS_DIR / "47_n8n_Prospection_Multi_Prospect_V4_Batch.json"
V5_JSON = DOCS_DIR / "60_n8n_Prospection_CRM_V5_Growth_Loop.json"
ATTACHED_V3_JSON = Path("/Users/marius_ayoro/Downloads/TransferAI Prospecting V3 CRM Enhanced [FINAL]-10.json")
FALLBACK_V3_JSON = DOCS_DIR / "62_n8n_Prospection_Multi_Prospect_V3_CRM_Ready_Export.json"

V4_OUTPUT = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V4_n8n_{TODAY}.docx"
V5_OUTPUT = WORD_DIR / f"TransferAI_Africa_Installation_Complete_Workflow_V5_n8n_{TODAY}.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 1", 15, RGBColor(0x10, 0x2A, 0x43)),
        ("Heading 2", 12, RGBColor(0x17, 0x4B, 0x6B)),
        ("Heading 3", 10.5, RGBColor(0x2A, 0x5D, 0x7B)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def add_title_page(document: Document, title: str, subtitle: str, intro: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(f"Version générée le {TODAY}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.4)
    set_cell_shading(cell, "EAF2F8")
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(intro)
    run.font.size = Pt(10.5)
    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        set_cell_shading(hdr[idx], "D9E7F5")
        set_cell_margins(hdr[idx], top=90, start=110, bottom=90, end=110)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            set_cell_margins(cells[idx], top=70, start=100, bottom=70, end=100)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def list_nodes(data: dict) -> list[list[str]]:
    return [[str(i), node["name"], node["type"].split(".")[-1]] for i, node in enumerate(data["nodes"], 1)]


def choose_v3_reference() -> Path:
    return ATTACHED_V3_JSON if ATTACHED_V3_JSON.exists() else FALLBACK_V3_JSON


def add_footer(document: Document, label: str) -> None:
    footer = document.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def build_v4_doc(v4_data: dict, v3_data: dict, v3_ref: Path) -> Document:
    doc = Document()
    style_document(doc)
    add_title_page(
        doc,
        "Installation complète du workflow V4 dans n8n",
        "Batch quotidien multi-prospects aligné sur la V3 CRM Enhanced finale",
        "Ce guide explique comment installer et exploiter la V4 batch orchestrator à partir de la V3 finale jointe. "
        "La V4 lit la CRM, applique les quotas et règles d’arrêt, puis déclenche la V3 pour chaque prospect éligible.",
    )

    add_heading(doc, "1. Dépendance principale", 1)
    add_paragraph(
        doc,
        f"La V4 doit appeler la V3 enfant. La référence de travail retenue ici est le fichier {v3_ref.name}, "
        f"qui contient {len(v3_data.get('nodes', []))} nœuds et gère la génération, le stockage Supabase, "
        "l’approbation, l’envoi et la mise à jour CRM.",
    )
    add_bullets(doc, [
        "Configurer dans n8n l’identifiant du workflow V3 final dans N8N_CHILD_WORKFLOW_ID_V3.",
        "S’assurer que la V3 est active et testée avant de brancher la V4.",
        "Remplacer toute clé hardcodée dans la V3 par des variables d’environnement avant production.",
    ])

    add_heading(doc, "2. Variables d’environnement à prévoir", 1)
    add_table(
        doc,
        ["Variable", "Utilité"],
        [
            ["SUPABASE_URL", "URL du projet Supabase utilisé comme CRM maître"],
            ["SUPABASE_SERVICE_ROLE_KEY", "clé service_role pour lire prospect_targets, outreach_attempts et écrire les batch runs"],
            ["N8N_CHILD_WORKFLOW_ID_V3", "identifiant du workflow V3 CRM Enhanced final appelé par la V4"],
            ["BOOKING_LINK_45MIN", "lien de prise de rendez-vous propagé à la V3"],
            ["AIRTABLE_API_KEY", "clé Airtable si la source batch provient d’Airtable"],
            ["AIRTABLE_BASE_ID", "base Airtable contenant les prospects prêts"],
            ["AIRTABLE_TABLE_NAME", "nom de table Airtable source"],
            ["AIRTABLE_READY_VIEW", "vue Airtable prête pour la V4"],
            ["GOOGLE_SHEETS_CSV_URL", "URL publique CSV si la source batch est Google Sheets"],
        ],
        widths=[2.35, 4.0],
    )

    add_heading(doc, "3. Tables et vues Supabase attendues", 1)
    add_bullets(doc, [
        "prospect_targets : file CRM principale avec les statuts, le nombre de tentatives, les stop rules et la source du prospect.",
        "outreach_attempts : historique des courriers réellement envoyés.",
        "prospecting_batch_runs : résumé de chaque run quotidien.",
        "prospecting_batch_run_items : détail prospect par prospect du traitement batch.",
        "Vue optionnelle prospect_targets_ready_for_batch : facilite la sélection des prospects prêts.",
    ])

    add_heading(doc, "4. Paramètres métier de Set Batch Config", 1)
    add_table(
        doc,
        ["Champ", "Valeur recommandée", "Rôle"],
        [
            ["source_backend", "supabase", "backend principal pour la lecture des prospects"],
            ["batch_fetch_limit", "25", "nombre maximum de prospects chargés au démarrage"],
            ["daily_send_limit", "5", "nombre maximum d’envois autorisés par jour"],
            ["max_attempts_per_prospect", "3", "stop rule après 3 tentatives sans signal positif"],
            ["min_confidence_score", "0.45", "seuil minimal avant dispatch vers la V3"],
            ["batch_run_label", "weekday-morning-run", "étiquette métier du run"],
            ["child_workflow_id_v3", "N8N_CHILD_WORKFLOW_ID_V3", "workflow enfant réellement exécuté"],
        ],
        widths=[1.7, 1.6, 3.0],
    )

    add_heading(doc, "5. Ordre réel du workflow V4", 1)
    add_table(doc, ["#", "Nœud", "Type"], list_nodes(v4_data), widths=[0.45, 3.7, 2.0])

    add_heading(doc, "6. Logique opérationnelle", 1)
    add_numbered(doc, [
        "Lancer la V4 manuellement ou via le déclencheur planifié Daily Schedule Trigger.",
        "Set Batch Config définit le backend, les quotas et l’identifiant de la V3 enfant.",
        "Fetch Today Outreach Count calcule le volume déjà envoyé aujourd’hui.",
        "Create Batch Run ouvre une ligne de suivi dans prospecting_batch_runs.",
        "Le workflow charge les prospects depuis Supabase, Airtable ou Google Sheets selon source_backend.",
        "Build Eligible Prospect Queue calcule remaining_capacity et prépare la file éligible.",
        "If Prospect Eligible décide qui part maintenant et qui est ignoré.",
        "Execute Prospect Workflow V3 appelle la V3 finale pour chaque prospect process_now.",
        "Mark Dispatched To V3, Mark Skipped In Batch et Merge Batch Outcomes consolident les résultats.",
        "Finalize Batch Run ferme le run avec les compteurs processed_count et skipped_count.",
    ])

    add_heading(doc, "7. Nœuds à surveiller en priorité", 1)
    add_table(
        doc,
        ["Nœud", "Contrôle à faire"],
        [
            ["Fetch Prospects From Supabase", "vérifier les headers apikey et Authorization avec la service role key"],
            ["Build Eligible Prospect Queue", "confirmer le calcul de remaining_capacity et des stop rules"],
            ["Execute Prospect Workflow V3", "pointer vers l’ID réel du workflow V3 final"],
            ["Log Processed Batch Item", "contrôler l’écriture correcte dans prospecting_batch_run_items"],
            ["Finalize Batch Run", "vérifier le PATCH vers prospecting_batch_runs avec status completed"],
        ],
        widths=[2.4, 3.9],
    )

    add_heading(doc, "8. Tests recommandés", 1)
    add_numbered(doc, [
        "Tester la V3 finale seule avec un prospect unique et confirmer un cycle complet jusqu’à l’envoi.",
        "Injecter 3 prospects tests dans prospect_targets avec des statuts différents.",
        "Lancer la V4 avec daily_send_limit = 1 pour vérifier la logique de quota.",
        "Confirmer qu’un prospect do_not_contact ou paused est bien routé vers la branche skipped.",
        "Vérifier qu’un prospect eligible déclenche bien Execute Prospect Workflow V3 avec les bons champs d’entrée.",
    ])

    add_heading(doc, "9. Checklist avant production", 1)
    add_bullets(doc, [
        "La V3 finale jointe est active et référencée par N8N_CHILD_WORKFLOW_ID_V3.",
        "Aucune clé sensible n’est hardcodée dans les nœuds HTTP.",
        "Les tables Supabase et les policies RLS ont été appliquées.",
        "Le quota quotidien est volontairement bas pour la première semaine.",
        "Le suivi de batch et le suivi des réponses prospect sont visibles dans Supabase.",
    ])

    add_footer(doc, "TransferAI - Guide d'installation V4 batch")
    return doc


def build_v5_doc(v5_data: dict, v4_data: dict, v3_data: dict, v3_ref: Path) -> Document:
    doc = Document()
    style_document(doc)
    add_title_page(
        doc,
        "Installation complète du workflow V5 dans n8n",
        "Boucle de croissance CRM alignée sur la chaîne V5 → V4 → V3",
        "Ce guide décrit l’installation de la V5 CRM Growth Loop, conçue pour alimenter la CRM depuis des sources scrappées, "
        "préparer les prospects dans prospect_targets puis déclencher la V4, laquelle dispatchera la V3 finale jointe.",
    )

    add_heading(doc, "1. Positionnement de la V5", 1)
    add_bullets(doc, [
        "La V5 ne prospecte pas elle-même ; elle alimente la CRM.",
        "La V5 normalise les leads publics et les upsert dans prospect_targets.",
        "La V5 peut ensuite déclencher la V4 batch, qui elle-même appelle la V3 finale.",
        f"La V3 de référence utilisée dans la chaîne est {v3_ref.name}.",
    ])

    add_heading(doc, "2. Variables d’environnement à prévoir", 1)
    add_table(
        doc,
        ["Variable", "Utilité"],
        [
            ["SUPABASE_URL", "URL du projet Supabase qui porte la CRM"],
            ["SUPABASE_SERVICE_ROLE_KEY", "clé service_role pour écrire dans prospect_targets"],
            ["SCRAPED_PUBLIC_LEADS_CSV_URL", "URL du flux CSV de leads scrappés"],
            ["N8N_CHILD_WORKFLOW_ID_V4", "identifiant du workflow V4 batch déclenché par la V5"],
            ["BOOKING_LINK_45MIN", "lien de rendez-vous propagé aux nouvelles fiches CRM"],
        ],
        widths=[2.35, 4.0],
    )

    add_heading(doc, "3. Champs CRM enrichis par la V5", 1)
    add_table(
        doc,
        ["Champ", "Rôle"],
        [
            ["prospect_id", "identifiant unique utilisé pour l’upsert"],
            ["organization_name", "raison sociale ou nom du prospect"],
            ["website", "site public servant de base au scraping V3"],
            ["target_email", "email cible si disponible"],
            ["status", "statut d’entrée, généralement ready"],
            ["paused", "bloque le prospect sans suppression"],
            ["do_not_contact", "empêche toute prospection future"],
            ["outreach_attempt_count", "compteur alimenté ensuite par la V3"],
            ["last_response_status", "retour prospect pour stop rule V4"],
            ["niche_status", "qualité de la niche ou angle de contact"],
            ["next_action_at", "date de relance ou de retraitement"],
        ],
        widths=[2.2, 4.15],
    )

    add_heading(doc, "4. Ordre réel du workflow V5", 1)
    add_table(doc, ["#", "Nœud", "Type"], list_nodes(v5_data), widths=[0.45, 3.7, 2.0])

    add_heading(doc, "5. Logique opérationnelle", 1)
    add_numbered(doc, [
        "La V5 démarre manuellement, par schedule ou via le webhook Scraped Leads Webhook.",
        "Set CRM Growth Config définit le backend CRM, l’URL du flux scrappé et l’ID de la V4 enfant.",
        "If Direct Lead Payload choisit entre payload direct et récupération CSV.",
        "Normalize Inbound Leads transforme les entrées au format canonique prospect.",
        "Filter Valid Leads ne conserve que les lignes minimales exploitables.",
        "Prepare CRM Upserts ajoute les champs de pilotage CRM et les stop flags par défaut.",
        "Upsert Prospects Into CRM écrit ou met à jour les lignes dans prospect_targets.",
        "Build CRM Import Summary consolide le volume importé.",
        "If Dispatch To V4 décide si le run doit déclencher immédiatement la V4 batch.",
        "Execute V4 Batch Workflow transmet la main à la V4, qui dispatchera ensuite la V3 finale.",
    ])

    add_heading(doc, "6. Dépendances inter-workflows", 1)
    add_table(
        doc,
        ["Chaîne", "Ce qu’il faut configurer"],
        [
            ["V5 → V4", "N8N_CHILD_WORKFLOW_ID_V4 doit pointer vers le workflow V4 batch réellement déployé"],
            ["V4 → V3", "N8N_CHILD_WORKFLOW_ID_V3 doit pointer vers la V3 finale jointe et validée"],
            ["V3 → Supabase", "la V3 doit écrire dans ai_prospecting_packs, outreach_attempts et prospect_targets"],
            ["V3 → Resend", "la V3 doit disposer des clés d’envoi et du webhook d’approbation actifs"],
        ],
        widths=[1.35, 4.95],
    )

    add_heading(doc, "7. Nœuds à surveiller en priorité", 1)
    add_table(
        doc,
        ["Nœud", "Contrôle à faire"],
        [
            ["Fetch Scraped Leads CSV", "vérifier que l’URL CSV renvoie des colonnes cohérentes"],
            ["Normalize Inbound Leads", "contrôler le mapping vers organization_name, website, target_email et sector_guess"],
            ["Prepare CRM Upserts", "vérifier status, paused, do_not_contact et next_action_at"],
            ["Upsert Prospects Into CRM", "contrôler on_conflict=prospect_id et les headers Supabase"],
            ["Execute V4 Batch Workflow", "vérifier l’identifiant du workflow V4 réellement installé"],
        ],
        widths=[2.4, 3.9],
    )

    add_heading(doc, "8. Séquence de test recommandée", 1)
    add_numbered(doc, [
        "Injecter 3 leads fictifs via Scraped Leads Webhook.",
        "Vérifier leur apparition dans prospect_targets après Upsert Prospects Into CRM.",
        "Activer dispatch_to_v4 = false pour tester la seule phase CRM.",
        "Réactiver dispatch_to_v4 = true puis confirmer le déclenchement de la V4.",
        "Suivre la cascade V5 → V4 → V3 jusqu’à la génération d’un ai_prospecting_pack.",
    ])

    add_heading(doc, "9. Checklist avant production", 1)
    add_bullets(doc, [
        "Les flux scrappés n’apportent que des données publiques autorisées.",
        "La V4 est déjà installée et testée avec quotas prudents.",
        "La V3 finale jointe est stabilisée, avec secrets passés en variables d’environnement.",
        "Les policies RLS Supabase sont en place avant de lancer des imports réels.",
        "La file CRM est revue régulièrement pour éviter l’accumulation de prospects non qualifiés.",
    ])

    add_paragraph(
        doc,
        f"Pour mémoire, la V4 actuellement documentée comporte {len(v4_data.get('nodes', []))} nœuds, "
        f"et la V3 finale jointe en comporte {len(v3_data.get('nodes', []))}. La V5 doit être pensée comme "
        "la couche de croissance et d’hygiène CRM, pas comme un workflow d’envoi direct.",
    )

    add_footer(doc, "TransferAI - Guide d'installation V5 CRM growth loop")
    return doc


def main() -> None:
    WORD_DIR.mkdir(parents=True, exist_ok=True)
    v3_ref = choose_v3_reference()
    v3_data = load_json(v3_ref)
    v4_data = load_json(V4_JSON)
    v5_data = load_json(V5_JSON)

    build_v4_doc(v4_data, v3_data, v3_ref).save(V4_OUTPUT)
    build_v5_doc(v5_data, v4_data, v3_data, v3_ref).save(V5_OUTPUT)

    print(V4_OUTPUT)
    print(V5_OUTPUT)


if __name__ == "__main__":
    main()
