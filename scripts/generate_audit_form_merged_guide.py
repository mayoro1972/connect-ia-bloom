from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom")
OUT_DIR = ROOT / "docs" / "transferai-admin"
OUT_PATH = OUT_DIR / "Guide_Formulaire_Audit_Resume_Troubleshooting_Fusionne.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_document_defaults(document):
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 22, RGBColor(15, 46, 75)),
        ("Heading 1", 15, RGBColor(15, 46, 75)),
        ("Heading 2", 12.5, RGBColor(36, 82, 120)),
        ("Heading 3", 11, RGBColor(58, 86, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_title_block(document, title, subtitle):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 46, 75)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.name = "Arial"
    run2.font.size = Pt(10.5)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(90, 90, 90)

    rows = [
        ("Projet", "Formulaire d'audit TransferAI - synthese fonctionnelle et troubleshooting"),
        ("Date", "05 juin 2026"),
        ("Portee", "Fusion du document expert et de l'implementation reelle du workflow"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.8)
        cells[1].width = Inches(5.2)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "DDEBF7")
    document.add_paragraph()


def add_section(document, title, paragraphs=None, bullets=None, numbered=None):
    document.add_heading(title, level=1)
    for paragraph in paragraphs or []:
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(5)
    for item in bullets or []:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)
    for item in numbered or []:
        p = document.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_subsection(document, title, paragraphs=None, bullets=None):
    document.add_heading(title, level=2)
    for paragraph in paragraphs or []:
        p = document.add_paragraph(paragraph)
        p.paragraph_format.space_after = Pt(4)
    for item in bullets or []:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(2)


def add_issue(document, title, symptome, cause, resolution, statut="Resolu"):
    document.add_heading(title, level=2)
    rows = [
        ("Symptome", symptome),
        ("Cause", cause),
        ("Resolution", resolution),
        ("Statut", statut),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].width = Inches(1.5)
        cells[1].width = Inches(5.5)
        cells[0].text = left
        cells[1].text = right
        set_cell_shading(cells[0], "EAF2F8")
    document.add_paragraph()


def build_document():
    doc = Document()
    set_document_defaults(doc)
    add_title_block(
        doc,
        "Guide Fusionne - Formulaire d'Audit TransferAI",
        "Resume utilisateur, architecture reelle, corrections appliquees et troubleshooting consolide",
    )

    add_section(
        doc,
        "1. Objet du document",
        paragraphs=[
            "Ce document fusionne le guide fourni par l'expert et l'implementation reelle mise en place dans le projet connect-ia-bloom. Il sert de reference unique avant la phase suivante, consacree a la collecte et a l'exploitation des reponses du formulaire.",
            "Le document couvre le parcours complet du formulaire d'audit, les points techniques effectivement utilises dans le code, les problemes rencontres, les ecarts avec l'ancienne conception et les solutions qui ont permis de stabiliser le systeme.",
        ],
    )

    add_section(
        doc,
        "2. Objectif fonctionnel du formulaire d'audit",
        paragraphs=[
            "Le formulaire d'audit permet au prospect de recevoir un lien unique base sur son pack_id, d'ouvrir un questionnaire adapte a son contexte, de repondre avant le rendez-vous, puis de transmettre ces informations afin que les experts TransferAI preparent l'echange.",
        ],
        bullets=[
            "Meilleure qualification du prospect",
            "Preparation des experts avant le rendez-vous",
            "Priorisation des cas d'usage et des quick wins",
            "Identification des contraintes, de la maturite et des besoins d'accompagnement",
        ],
    )

    add_section(
        doc,
        "3. Parcours complet du workflow jusqu'au formulaire",
        numbered=[
            "n8n genere un pack prospect personnalise.",
            "Le pack est stocke dans Supabase avec un pack_id unique.",
            "Un mail interne de validation est envoye a l'equipe TransferAI.",
            "Apres clic sur Approuver et envoyer, le prospect recoit le courrier final.",
            "Le mail prospect contient le lien du formulaire d'audit, le mini-catalogue PDF et le deck PPTX.",
            "Le prospect ouvre l'URL /questionnaire-audit?pack_id=....",
            "La page frontend charge le contexte d'audit a travers les fonctions edge Supabase.",
            "Le prospect remplit le questionnaire, le brouillon est sauvegarde localement, puis les reponses sont transmises.",
            "Les reponses sont stockees dans Supabase pour traitement ulterieur par les experts et les commerciaux.",
        ],
    )

    add_section(
        doc,
        "4. Architecture technique reelle retenue",
        bullets=[
            "Frontend public : React + Vite sur www.transferai.ci",
            "Route frontend : /questionnaire-audit",
            "Page React utilisee : src/pages/ProspectAuditFormPage.tsx",
            "Routing SPA : public/_redirects avec /* /index.html 200",
            "Client Supabase frontend : src/integrations/supabase/client.ts",
            "Fonctions edge principales : resolve-invitation et save-form-response",
            "Stockage source du pack : ai_prospecting_packs",
            "Gestion des invitations : form_invitations",
            "Stockage des reponses : form_responses",
        ],
    )

    add_subsection(
        doc,
        "4.1 Difference importante avec le document expert initial",
        paragraphs=[
            "Le document expert d'origine decrivait une lecture et une soumission via une edge function unique get-audit-pack, ainsi qu'un stockage dans ai_prospecting_packs.questionnaire_answers. La version finale du projet est plus structuree : la resolution d'acces passe par resolve-invitation, la sauvegarde par save-form-response, et les reponses sont principalement centralisees dans form_responses, avec le contexte d'acces dans form_invitations.",
            "Cette difference est importante pour la future collecte des donnees : la table de reference pour exploiter les reponses n'est pas seulement ai_prospecting_packs, mais surtout form_responses, reliee a form_invitations et au pack d'origine.",
        ],
    )

    add_section(
        doc,
        "5. Fichiers et composants effectivement utilises",
        bullets=[
            "src/App.tsx : declaration de la route /questionnaire-audit",
            "src/pages/ProspectAuditFormPage.tsx : ecran principal du questionnaire",
            "src/integrations/supabase/client.ts : detection de la configuration Supabase frontend",
            "supabase/functions/resolve-invitation/index.ts : resolution de l'acces, du pack et du contexte",
            "supabase/functions/save-form-response/index.ts : enregistrement et mise a jour des reponses",
            "public/_redirects : fallback SPA pour routes profondes",
        ],
    )

    add_section(
        doc,
        "6. Ce qui a ete corrige dans n8n pour supporter le formulaire",
        bullets=[
            "Generation stable de audit_form_url dans le pack prospect",
            "Injection du lien du formulaire dans le courrier prospect",
            "Stabilisation de la branche d'approbation",
            "Reconstruction fiable des pieces jointes PDF et PPTX apres approbation",
            "Envoi prospect via Resend avec un expediteur valide du domaine transferai.ci",
        ],
    )

    add_subsection(
        doc,
        "6.1 Resultat fonctionnel obtenu",
        bullets=[
            "Le prospect recoit bien le mail final",
            "Le mail contient bien le lien du formulaire d'audit",
            "Le mail contient bien les deux pieces jointes",
            "Le lien /questionnaire-audit charge maintenant la bonne page publique",
            "Le frontend utilise correctement les variables Supabase du build",
        ],
    )

    add_section(
        doc,
        "7. Troubleshooting consolide",
        paragraphs=[
            "Les incidents ci-dessous reprennent les problemes reels rencontres pendant une semaine de debug, ainsi que les corrections effectivement appliquees. Cette section remplace et met a jour le troubleshooting initial fourni par l'expert.",
        ],
    )

    add_issue(
        doc,
        "7.1 La route /questionnaire-audit retournait une 404",
        symptome="Le lien envoye dans le mail prospect ouvrait une page 404 ou une page NotFound du site public.",
        cause="La route n'etait pas servie correctement en production, le build public etait ancien, et un ancien contournement statique pouvait interferer avec la route React.",
        resolution="La route React /questionnaire-audit a ete confirmee dans src/App.tsx, le fichier public/_redirects a ete ajoute pour assurer le fallback SPA, l'ancien pont statique conflictuel a ete retire, puis le frontend a ete redeploye.",
    )

    add_issue(
        doc,
        "7.2 Le frontend affichait Supabase non configure",
        symptome="La page du questionnaire chargeait mais affichait un message indiquant que la connexion a la base n'etait pas configuree.",
        cause="Les variables VITE_SUPABASE_URL et la cle publishable/anon n'etaient pas disponibles dans le build public, et le client frontend validait initialement seulement VITE_SUPABASE_PUBLISHABLE_KEY.",
        resolution="Les variables VITE_SUPABASE_* ont ete configurees cote hebergement et le fichier src/integrations/supabase/client.ts a ete ajuste pour accepter VITE_SUPABASE_PUBLISHABLE_KEY ou VITE_SUPABASE_ANON_KEY.",
    )

    add_issue(
        doc,
        "7.3 Le mail prospect ne partait pas apres approbation",
        symptome="Le clic sur Approuver et envoyer repondait correctement mais aucun mail prospect n'arrivait.",
        cause="La branche d'approbation contenait des scripts fragiles dans Extract Pack Payload et Build Send Context, ce qui empechait le passage complet jusqu'au noeud d'envoi.",
        resolution="Extract Pack Payload a ete stabilise, Build Send Context a ete reconstruit, If Ready To Send a ete configure sur can_send=true, et l'adresse d'envoi Resend a ete basculee vers TransferAI <contact@transferai.ci>.",
    )

    add_issue(
        doc,
        "7.4 Le mail prospect arrivait sans pieces jointes",
        symptome="Le prospect recevait le courrier et le lien du formulaire, mais pas le mini-catalogue ni le deck.",
        cause="Une version temporairement simplifiee de Build Send Context n'incluait plus la reconstruction des attachments.",
        resolution="Le noeud Build Send Context a ete refait pour relire payload.mail_attachments, ou reconstruire les pieces depuis payload.catalogue_artifact.pdf_url et payload.deck_artifact.pptx_url, puis controler la presence d'un PDF et d'un PPTX avant envoi.",
    )

    add_issue(
        doc,
        "7.5 Le mail de confirmation interne cassait apres l'envoi prospect",
        symptome="Send Internal Sent Confirmation renvoyait JSON parameter needs to be valid JSON.",
        cause="Le body JSON etait invalide dans le noeud HTTP de confirmation.",
        resolution="Le body a ete reecrit avec une concatentation simple et un JSON valide, ce qui a evite de casser l'execution apres le succes du mail prospect.",
    )

    add_issue(
        doc,
        "7.6 Les noeuds OpenAI et certains scripts n8n etaient fragiles",
        symptome="Erreurs de syntaxe, access to env vars denied, missing / ou invalid regular expression sur differents noeuds.",
        cause="Des expressions trop fragiles, des regex mal collees et une dependance excessive a $env dans n8n.",
        resolution="Les noeuds ont ete simplifies, les expressions reecrites, les regex fragiles retirees quand necessaire et les dependances a $env supprimees dans les noeuds concernes.",
    )

    add_section(
        doc,
        "8. Ce qu'il ne faut pas retenir du document expert sans adaptation",
        paragraphs=[
            "Le document expert contenait de bonnes intuitions fonctionnelles, mais certaines recommandations techniques ne doivent pas etre reprises telles quelles sans verification par rapport au code reel.",
        ],
        bullets=[
            "Ne pas supposer que get-audit-pack est l'unique fonction edge encore utilisee.",
            "Ne pas considerer ai_prospecting_packs comme l'unique table cible des reponses.",
            "Ne pas ouvrir trop largement la table ai_prospecting_packs via une policy de type USING (true) sans justification forte.",
            "Prendre comme reference la structure actuelle du repo : resolve-invitation, save-form-response, form_invitations et form_responses.",
        ],
    )

    add_section(
        doc,
        "9. Etat final avant la phase de collecte des donnees",
        bullets=[
            "Le prospect recoit le bon mail",
            "Le lien du formulaire fonctionne",
            "Le questionnaire se charge sur la bonne route publique",
            "Le frontend parle correctement a Supabase",
            "Les pieces jointes sont presentes dans le mail prospect",
            "Le workflow n8n de validation et d'approbation est stabilise",
        ],
    )

    add_subsection(
        doc,
        "9.1 Point de depart recommande pour la prochaine phase",
        paragraphs=[
            "La prochaine phase ne consiste plus a stabiliser le formulaire, mais a exploiter les reponses. Le point d'entree le plus propre est la table form_responses, enrichie par form_invitations et reliee au pack d'origine via le pack_id.",
            "C'est sur cette base qu'il sera possible de construire ensuite un workflow post-audit, un dashboard expert, des notifications internes ou une synchronisation CRM.",
        ],
    )

    add_section(
        doc,
        "10. Conclusion",
        paragraphs=[
            "Le formulaire d'audit est desormais operationnel de bout en bout. Le lien fonctionne, le frontend charge, Supabase est accessible, les emails partent avec les bonnes pieces jointes, et la logique d'approbation est stabilisee.",
            "Le present document constitue donc la reference fusionnee a utiliser avant d'ouvrir la prochaine etape : la collecte, la synthese et la redistribution des reponses aux experts et aux commerciaux.",
        ],
    )

    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUT_PATH)
    print(str(OUT_PATH))


if __name__ == "__main__":
    main()
