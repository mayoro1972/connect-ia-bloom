from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FRENCH_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "141_Guide_Utilisateur_Workflow_128_Reunions_Internationales_Multilingue.docx"
)

ENGLISH_OUTPUT = Path(
    "/Users/marius_ayoro/Documents/GitHub/connect-ia-bloom/docs/transferai-admin/"
    "142_User_Guide_Workflow_128_International_Multilingual_Meeting_Report.docx"
)


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(document, text, subtitle):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=8, line=1.0)
    r = p.add_run(text)
    set_font(r, size=20, bold=True, color="10263F")

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p2, before=0, after=14, line=1.0)
    r2 = p2.add_run(subtitle)
    set_font(r2, size=11, color="666666")


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=14 if level == 1 else 10, after=6 if level == 1 else 4)
    r = p.add_run(text)
    if level == 1:
        set_font(r, size=15, bold=True, color="2E74B5")
    elif level == 2:
        set_font(r, size=12.5, bold=True, color="2E74B5")
    else:
        set_font(r, size=11.5, bold=True, color="1F4D78")


def add_body(document, text):
    p = document.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.15)
    r = p.add_run(text)
    set_font(r, size=11)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def add_numbered(document, items):
    for item in items:
        p = document.add_paragraph(style="List Number")
        set_paragraph_spacing(p, before=0, after=3, line=1.15)
        r = p.add_run(item)
        set_font(r, size=11)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_two_col_table(document, rows, widths=(2.05, 4.45), header=("Élément", "Configuration")):
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(widths[0])
    table.columns[1].width = Inches(widths[1])

    hdr = table.rows[0].cells
    hdr[0].text = header[0]
    hdr[1].text = header[1]
    for cell in hdr:
        shade_cell(cell, "E8EEF5")
        for p in cell.paragraphs:
            set_paragraph_spacing(p, before=0, after=0, line=1.0)
            if p.runs:
                set_font(p.runs[0], size=10.5, bold=True, color="10263F")

    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before=0, after=0, line=1.08)
                if p.runs:
                    set_font(p.runs[0], size=10.5, bold=(idx == 0))

    document.add_paragraph()


def add_header(section, text):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run(text)
    set_font(r, size=9, bold=True, color="777777")


def add_footer(section, text):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, before=0, after=0, line=1.0)
    r = p.add_run(text)
    set_font(r, size=9, color="777777")


def build_french_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header(section, "TransferAI - Compte rendu de réunion sensible")
    add_footer(section, "Guide utilisateur Workflow 128 - TransferAI")

    add_title(
        doc,
        "Guide utilisateur - Workflow 128",
        "Compte rendu de réunion international, pseudonymisation locale et livrables multilingues assainis",
    )

    add_heading(doc, "1. Objet du document")
    add_body(
        doc,
        "Ce guide explique comment utiliser et configurer le workflow 128 pour traiter des enregistrements de réunions sensibles, produire des comptes rendus structurés et générer des livrables externes assainis."
    )
    add_body(
        doc,
        "Le workflow est conçu pour des usages récurrents dans des contextes internationaux où la confidentialité, la validation humaine et la cohérence linguistique sont indispensables."
    )

    add_heading(doc, "2. Cas d’usage recommandé")
    add_body(
        doc,
        "Le workflow 128 peut être utilisé pour plusieurs contextes sensibles : sessions de formation IA, organisations internationales présentes en Côte d’Ivoire, chancelleries, ambassades, institutions publiques ou projets multi-pays."
    )
    add_body(
        doc,
        "Les réunions peuvent être enregistrées via Zoom, Webex, Microsoft Teams, un dictaphone, un smartphone ou un dispositif de salle."
    )
    add_bullets(
        doc,
        [
            "Sorties prises en charge : compte rendu, e-mail de suivi, archive HTML assainie",
            "Langues de sortie prises en charge : Français, Anglais, Espagnol",
            "Données sensibles protégées localement avant tout appel IA",
        ],
    )

    add_heading(doc, "3. Architecture du workflow")
    add_numbered(
        doc,
        [
            "Entrée du recording",
            "Normalisation des métadonnées",
            "Récupération ou réception de l’audio",
            "Découpage audio",
            "Transcription OpenAI",
            "Pseudonymisation locale",
            "Génération du rapport assaini",
            "Validation humaine puis envoi final",
        ],
    )

    add_heading(doc, "4. Modes d’entrée supportés")
    add_heading(doc, "4.1 Upload manuel", level=2)
    add_body(doc, "À utiliser pour les fichiers .m4a, .mp3, les exports smartphone et les exports dictaphone.")
    add_heading(doc, "4.2 Google Drive", level=2)
    add_body(doc, "À utiliser si les fichiers sont déposés dans un dossier surveillé ou si un lien Drive est fourni dans le formulaire.")
    add_heading(doc, "4.3 Webhook plateforme meeting", level=2)
    add_body(doc, "À utiliser si un workflow amont Zoom, Teams ou Webex récupère le recording puis appelle le webhook d’ingestion du workflow 128.")

    add_heading(doc, "5. Pré-requis techniques")
    add_two_col_table(
        doc,
        [
            ("Clé OpenAI", "Variable d’environnement OPENAI_API_KEY"),
            ("Clé Resend", "Variable d’environnement RESEND_API_KEY"),
            ("Modèle rapport", "gpt-5 recommandé ; sinon gpt-4.1 ou gpt-4o"),
            ("Google Drive", "Credential OAuth actif pour lecture et archivage"),
            ("Découpage audio", "Service disponible sur http://n8n-pxlk-audio-splitter-1:8000/split"),
            ("Domaine webhook", "https://n8n-pxlk.srv1480638.hstgr.cloud/webhook"),
        ],
    )

    add_heading(doc, "6. Configuration pas à pas")
    add_two_col_table(
        doc,
        [
            ("Étape 1 - Importer le workflow", "Dans n8n, ouvrir Workflows, choisir Import from file, puis sélectionner le fichier JSON du workflow 128."),
            ("Étape 2 - Vérifier les entrées", "Contrôler les nœuds Formulaire - Audio sensible, Google Drive - Nouveau fichier audio et Webhook - Ingestion plateforme meeting."),
            ("Étape 3 - Configurer le formulaire", "Renseigner les champs sujet, participants, date, source d’enregistrement, mode d’ingestion, destinataire final, validateur, langue audio, langue de sortie, livrable demandé, entités sensibles à masquer et niveau de masquage."),
            ("Étape 4 - Configurer le dossier source", "Vérifier l’ID du dossier surveillé dans le nœud Google Drive - Nouveau fichier audio."),
            ("Étape 5 - Configurer le dossier d’archive", "Vérifier l’ID du dossier d’archivage dans le nœud Archiver version assainie."),
            ("Étape 6 - Configurer Resend", "Valider les nœuds Envoyer au validateur et Envoyer le livrable final assaini, ainsi que l’adresse expéditeur autorisée."),
            ("Étape 7 - Configurer OpenAI", "Vérifier les nœuds OpenAI - Transcrire segment et OpenAI - Rapport structure sanitisé."),
            ("Étape 8 - Vérifier les webhooks de validation", "Contrôler les nœuds Webhook - Approuver et Webhook - Rejeter, ainsi que le domaine de production."),
            ("Étape 9 - Lancer un test pilote", "Faire un test avec un petit enregistrement, approuver le livrable, puis vérifier l’archive HTML et la cohérence de la langue finale."),
        ],
    )

    add_heading(doc, "7. Paramétrage conseillé")
    add_bullets(
        doc,
        [
            "Livrable demandé : Les deux",
            "Niveau de masquage : Strict",
            "Langues de sortie : Français, Anglais, Espagnol",
            "Dossier source : Recordings - Meetings International",
            "Dossier archive : Sanitized Archives - Meetings International",
        ],
    )

    add_heading(doc, "8. Configuration des sources de réunion")
    add_heading(doc, "8.1 Dictaphone", level=2)
    add_bullets(doc, ["Mode simple : exporter le fichier puis l’envoyer via le formulaire.", "Mode semi-automatique : synchroniser les fichiers vers un dossier Google Drive surveillé."])
    add_heading(doc, "8.2 Zoom", level=2)
    add_bullets(doc, ["Activer Cloud Recording dans Zoom.", "Utiliser un workflow amont recevant l’événement recording.completed.", "Télécharger le recording puis appeler le webhook cr-intl-source-ingest."])
    add_heading(doc, "8.3 Microsoft Teams", level=2)
    add_bullets(doc, ["Récupérer les enregistrements dans OneDrive ou SharePoint.", "Passer par Microsoft Graph ou un workflow de collecte amont.", "Transmettre ensuite le binaire ou les métadonnées au workflow 128."])
    add_heading(doc, "8.4 Webex", level=2)
    add_bullets(doc, ["Activer les recordings Webex.", "Utiliser l’API Recordings pour récupérer le fichier.", "Faire suivre le recording vers le webhook d’ingestion du workflow 128."])

    add_heading(doc, "9. Structure attendue pour le webhook d’ingestion")
    add_body(
        doc,
        "Le webhook cr-intl-source-ingest doit idéalement recevoir : source_system, source_label, meeting_title, meeting_date, participants, e-mail destinataire, e-mail validateur, langue, langue de sortie, niveau de masquage, référence de réunion, objet source et recording_url. Si possible, transmettre directement le binaire audio."
    )

    add_heading(doc, "10. Checklist de mise en production")
    add_bullets(
        doc,
        [
            "Le workflow s’importe sans erreur.",
            "Google Drive est connecté.",
            "Resend est connecté.",
            "OpenAI est connecté.",
            "Le service de découpage audio répond.",
            "Le validateur reçoit l’aperçu.",
            "L’archive HTML est créée.",
            "Aucune version réidentifiée n’est envoyée au client.",
            "Les livrables sortent intégralement dans la langue cible choisie.",
        ],
    )

    add_heading(doc, "11. Limites actuelles et recommandation finale")
    add_body(
        doc,
        "Le workflow 128 est un cœur de traitement. Il ne gère pas encore à lui seul l’authentification native Zoom, Webex ou Microsoft Graph, ni le téléchargement automatique depuis chaque plateforme."
    )
    add_body(
        doc,
        "Pour un déploiement progressif, il est recommandé de commencer avec le mode upload manuel + Google Drive, puis d’ajouter des workflows d’ingestion dédiés pour Zoom, Teams et Webex."
    )

    doc.save(FRENCH_OUTPUT)


def build_english_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_header(section, "TransferAI - Sensitive meeting report workflow")
    add_footer(section, "Workflow 128 User Guide - TransferAI")

    add_title(
        doc,
        "User Guide - Workflow 128",
        "International meeting report, local pseudonymization, and multilingual sanitized deliverables",
    )

    add_heading(doc, "1. Purpose of this document")
    add_body(
        doc,
        "This guide explains how to use and configure workflow 128 to process sensitive meeting recordings, produce structured meeting reports, and generate sanitized external deliverables."
    )
    add_body(
        doc,
        "The workflow is intended for recurring international use cases where confidentiality, human validation, and language consistency are critical."
    )

    add_heading(doc, "2. Recommended use case")
    add_body(
        doc,
        "Workflow 128 can support several sensitive contexts: AI training sessions, international organizations operating in Côte d’Ivoire, chancelleries, embassies, public institutions, or multi-country programmes."
    )
    add_body(
        doc,
        "Meetings may be recorded through Zoom, Webex, Microsoft Teams, a dictaphone, a smartphone, or a meeting room device."
    )
    add_bullets(
        doc,
        [
            "Supported outputs: meeting report, follow-up email, sanitized HTML archive",
            "Supported output languages: French, English, Spanish",
            "Sensitive data is protected locally before any AI call",
        ],
    )

    add_heading(doc, "3. Workflow architecture")
    add_numbered(
        doc,
        [
            "Recording intake",
            "Metadata normalization",
            "Audio retrieval or reception",
            "Audio splitting",
            "OpenAI transcription",
            "Local pseudonymization",
            "Sanitized report generation",
            "Human validation and final delivery",
        ],
    )

    add_heading(doc, "4. Supported input modes")
    add_heading(doc, "4.1 Manual upload", level=2)
    add_body(doc, "Use this for .m4a files, .mp3 files, smartphone exports, and dictaphone exports.")
    add_heading(doc, "4.2 Google Drive", level=2)
    add_body(doc, "Use this when files are uploaded into a watched folder or when a Drive link is provided in the form.")
    add_heading(doc, "4.3 Meeting platform webhook", level=2)
    add_body(doc, "Use this when an upstream Zoom, Teams, or Webex workflow retrieves the recording and then calls the ingestion webhook of workflow 128.")

    add_heading(doc, "5. Technical prerequisites")
    add_two_col_table(
        doc,
        [
            ("OpenAI key", "Environment variable OPENAI_API_KEY"),
            ("Resend key", "Environment variable RESEND_API_KEY"),
            ("Report model", "gpt-5 recommended; otherwise gpt-4.1 or gpt-4o"),
            ("Google Drive", "Active OAuth credential for reading and archiving"),
            ("Audio splitting", "Service available at http://n8n-pxlk-audio-splitter-1:8000/split"),
            ("Webhook domain", "https://n8n-pxlk.srv1480638.hstgr.cloud/webhook"),
        ],
        header=("Item", "Configuration"),
    )

    add_heading(doc, "6. Step-by-step configuration")
    add_two_col_table(
        doc,
        [
            ("Step 1 - Import the workflow", "In n8n, open Workflows, choose Import from file, then select the JSON file for workflow 128."),
            ("Step 2 - Check the entry points", "Review the nodes Formulaire - Audio sensible, Google Drive - Nouveau fichier audio, and Webhook - Ingestion plateforme meeting."),
            ("Step 3 - Configure the form", "Fill in the fields for subject, participants, date, recording source, ingestion mode, final recipient, validator, audio language, output language, requested deliverable, sensitive entities to mask, and masking level."),
            ("Step 4 - Configure the source folder", "Check the watched folder ID in the node Google Drive - Nouveau fichier audio."),
            ("Step 5 - Configure the archive folder", "Check the archive folder ID in the node Archiver version sanitisee."),
            ("Step 6 - Configure Resend", "Validate the nodes Envoyer au validateur and Envoyer livrable final sanitise, as well as the authorized sender address."),
            ("Step 7 - Configure OpenAI", "Check the nodes OpenAI - Transcrire segment and OpenAI - Rapport structure sanitise."),
            ("Step 8 - Verify the validation webhooks", "Review the nodes Webhook - Approuver and Webhook - Rejeter, together with the production domain."),
            ("Step 9 - Run a pilot test", "Test with a short recording, approve the deliverable, then verify the HTML archive and target-language consistency."),
        ],
        header=("Item", "Configuration"),
    )

    add_heading(doc, "7. Recommended settings")
    add_bullets(
        doc,
        [
            "Requested deliverable: Both",
            "Masking level: Strict",
            "Output languages: French, English, Spanish",
            "Source folder: Recordings - International Meetings",
            "Archive folder: Sanitized Archives - International Meetings",
        ],
    )

    add_heading(doc, "8. Meeting source configuration")
    add_heading(doc, "8.1 Dictaphone", level=2)
    add_bullets(doc, ["Simple mode: export the file and upload it through the form.", "Semi-automated mode: sync the files to a watched Google Drive folder."])
    add_heading(doc, "8.2 Zoom", level=2)
    add_bullets(doc, ["Enable Cloud Recording in Zoom.", "Use an upstream workflow that receives the recording.completed event.", "Download the recording, then call the cr-intl-source-ingest webhook."])
    add_heading(doc, "8.3 Microsoft Teams", level=2)
    add_bullets(doc, ["Retrieve recordings from OneDrive or SharePoint.", "Use Microsoft Graph or an upstream collection workflow.", "Then pass the binary file or the metadata to workflow 128."])
    add_heading(doc, "8.4 Webex", level=2)
    add_bullets(doc, ["Enable Webex recordings.", "Use the Recordings API to retrieve the file.", "Forward the recording to the ingestion webhook of workflow 128."])

    add_heading(doc, "9. Expected ingestion webhook structure")
    add_body(
        doc,
        "The cr-intl-source-ingest webhook should ideally receive: source_system, source_label, meeting_title, meeting_date, participants, recipient email, validator email, language, output language, masking level, meeting reference, source object, and recording_url. When possible, send the audio binary directly."
    )

    add_heading(doc, "10. Go-live checklist")
    add_bullets(
        doc,
        [
            "The workflow imports without error.",
            "Google Drive is connected.",
            "Resend is connected.",
            "OpenAI is connected.",
            "The audio splitting service responds.",
            "The validator receives the preview.",
            "The HTML archive is created.",
            "No re-identified version is sent to the client.",
            "Deliverables are fully generated in the selected target language.",
        ],
    )

    add_heading(doc, "11. Current limitations and final recommendation")
    add_body(
        doc,
        "Workflow 128 is a core processing engine. It does not yet handle on its own native Zoom, Webex, or Microsoft Graph authentication, nor automatic download from each source platform."
    )
    add_body(
        doc,
        "For progressive deployment, the recommended starting point is manual upload + Google Drive, then dedicated ingestion workflows for Zoom, Teams, and Webex."
    )

    doc.save(ENGLISH_OUTPUT)


if __name__ == "__main__":
    FRENCH_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_french_doc()
    build_english_doc()
